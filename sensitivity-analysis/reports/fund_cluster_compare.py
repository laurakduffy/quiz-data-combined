"""
fund_cluster_compare.py
Shared helper for the fund-level vs cluster-level SI comparison section
used in generate_report.py and generate_cluster_report.py.

For each sensitivity analysis it joins the fund-level and cause-area SI
CSVs on the appropriate keys, computes
    cross_cluster_share = cluster_SI / fund_SI
and produces the per-analysis combined_si.csv plus a scatter plot.
"""

import io
import os
from dataclasses import dataclass, field
from typing import Callable, Optional

import numpy as np
import pandas as pd

# NOTE: matplotlib and python-docx are imported lazily inside the plotting and
# docx-section functions below. Keeping them out of module scope means the CSV
# merge/writer (merge_pair / write_combined_csv / build_all) — and callers like
# regen_combined_si.py and run_all.js — only need pandas + numpy, not the full
# Word-report toolchain.


BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # parent sensitivity-analysis/ — input CSVs live here


def _p(*parts):
    return os.path.join(BASE, *parts)


@dataclass
class AnalysisSpec:
    id: str
    label: str
    output_dir: str          # e.g. "across-the-board/outputs"
    fund_csv: str            # path relative to BASE
    cause_csv: str           # path relative to BASE
    join_keys: list          # columns to merge on
    label_fn: Callable[[pd.Series], str]   # produces a scenario label from a merged row
    # Filename for the combined SI CSV written into output_dir. Defaults to
    # "combined_si.csv"; override when multiple specs share an output_dir
    # (e.g. the three diminishing-returns variants).
    combined_csv_name: str = "combined_si.csv"
    # Optional filter applied to the merged frame (e.g. drop bound=="single" rows).
    pre_filter: Optional[Callable[[pd.DataFrame], pd.DataFrame]] = None


# ── label helpers ─────────────────────────────────────────────────────────────

def _fmt_num(v):
    try:
        f = float(v)
        return f"{f:g}"
    except (TypeError, ValueError):
        return str(v)


def _drop_single_bound(df):
    if "bound" in df.columns:
        return df[~df["bound"].astype(str).isin(["single"])].copy()
    return df


# ── ANALYSIS_SPECS ────────────────────────────────────────────────────────────

ANALYSIS_SPECS = [
    AnalysisSpec(
        id="across-the-board",
        label="CE Multipliers",
        output_dir="across-the-board/outputs",
        fund_csv="across-the-board/outputs/fund/ce_multiplier_si.csv",
        cause_csv="across-the-board/outputs/cause/cause_area_si.csv",
        join_keys=["fund_varied", "multiplier"],
        label_fn=lambda r: f"{r['fund_varied']} ×{_fmt_num(r['multiplier'])}",
    ),
    AnalysisSpec(
        id="time-discounts",
        label="Time Discounts",
        output_dir="time-discounts/outputs",
        fund_csv="time-discounts/outputs/fund/discount_fund_si.csv",
        cause_csv="time-discounts/outputs/cause/discount_cause_area_si.csv",
        join_keys=["scenario_group", "multiplier"],
        label_fn=lambda r: f"{r['scenario_group']} ×{_fmt_num(r['multiplier'])}",
    ),
    AnalysisSpec(
        id="worldview-sensitivity",
        label="Worldview Credences",
        output_dir="worldview-sensitivity/outputs",
        fund_csv="worldview-sensitivity/outputs/fund/split_credences_index.csv",
        cause_csv="worldview-sensitivity/outputs/cause/cause_area_index.csv",
        join_keys=["scenario", "worldview", "bound", "credence_base", "credence_scenario"],
        label_fn=lambda r: str(r["scenario"]),
        pre_filter=_drop_single_bound,
    ),
    AnalysisSpec(
        id="aggregation-methods",
        label="Aggregation Methods",
        output_dir="aggregation-methods/outputs",
        fund_csv="aggregation-methods/outputs/fund/split_credences_index.csv",
        cause_csv="aggregation-methods/outputs/cause/cause_area_index.csv",
        join_keys=["scenario", "method", "bound", "credence_base", "credence_scenario"],
        label_fn=lambda r: str(r["scenario"]),
        pre_filter=_drop_single_bound,
    ),
    AnalysisSpec(
        id="risk-aversion",
        label="Risk Aversion",
        output_dir="risk-aversion/outputs",
        fund_csv="risk-aversion/outputs/fund/risk_aversion_summary.csv",
        cause_csv="risk-aversion/outputs/cause/risk_aversion_cause_area_summary.csv",
        join_keys=["test"],
        label_fn=lambda r: str(r["test"]),
    ),
    AnalysisSpec(
        id="moral-weights",
        label="Moral Weights",
        output_dir="moral-weights/outputs",
        fund_csv="moral-weights/outputs/fund/moral_weights_overall_si.csv",
        cause_csv="moral-weights/outputs/cause/moral_weights_overall_cause_area_si.csv",
        join_keys=["multiplier"],
        label_fn=lambda r: f"weight ×{_fmt_num(r['multiplier'])}",
    ),
]


# ── core ──────────────────────────────────────────────────────────────────────

def merge_pair(spec: AnalysisSpec, apply_filter: bool = False) -> Optional[pd.DataFrame]:
    """Join fund-level and cluster-level SI CSVs for one analysis.

    Returns a DataFrame with columns [*join_keys, fund_si, cluster_si,
    cross_cluster_share, gap, label].  Returns None if either CSV is missing.

    When apply_filter is True, the spec's pre_filter (if any) is applied —
    used for the report's section to drop bound=="single" rows. The CSV
    writer leaves apply_filter=False so the spreadsheet stays comprehensive.
    """
    fund_path = _p(spec.fund_csv)
    cause_path = _p(spec.cause_csv)
    if not (os.path.exists(fund_path) and os.path.exists(cause_path)):
        return None

    fund_df = pd.read_csv(fund_path)
    cause_df = pd.read_csv(cause_path)

    fund_sub = fund_df[spec.join_keys + ["sensitivity_index"]].rename(
        columns={"sensitivity_index": "fund_si"}
    )
    cause_sub = cause_df[spec.join_keys + ["sensitivity_index"]].rename(
        columns={"sensitivity_index": "cluster_si"}
    )

    merged = pd.merge(fund_sub, cause_sub, on=spec.join_keys, how="inner")

    if apply_filter and spec.pre_filter is not None:
        merged = spec.pre_filter(merged)

    merged["fund_si"] = merged["fund_si"].astype(float)
    merged["cluster_si"] = merged["cluster_si"].astype(float)
    merged["gap"] = merged["fund_si"] - merged["cluster_si"]
    # cross_cluster_share = cluster_si / fund_si; NaN when fund_si == 0
    with np.errstate(divide="ignore", invalid="ignore"):
        merged["cross_cluster_share"] = np.where(
            merged["fund_si"] > 0,
            merged["cluster_si"] / merged["fund_si"],
            np.nan,
        )

    merged["label"] = merged.apply(spec.label_fn, axis=1)
    return merged.reset_index(drop=True)


def write_combined_csv(spec: AnalysisSpec, merged_df: pd.DataFrame) -> str:
    """Write the per-analysis combined SI CSV. Returns the absolute path."""
    out_dir = _p(spec.output_dir)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, spec.combined_csv_name)
    cols = spec.join_keys + ["fund_si", "cluster_si", "cross_cluster_share"]
    out = merged_df[cols].copy()
    out["fund_si"] = out["fund_si"].round(4)
    out["cluster_si"] = out["cluster_si"].round(4)
    out["cross_cluster_share"] = out["cross_cluster_share"].round(4)
    out.to_csv(out_path, index=False)
    return out_path


def summary_stats_row(spec: AnalysisSpec, merged_df: pd.DataFrame) -> dict:
    """One row for the section's overview table."""
    active = merged_df[merged_df["fund_si"] > 0]
    if active.empty:
        return {
            "label": spec.label,
            "n_tests": 0,
            "max_fund_si": float("nan"),
            "max_cluster_si": float("nan"),
            "median_share": float("nan"),
            "min_share": float("nan"),
        }
    return {
        "label": spec.label,
        "n_tests": len(active),
        "max_fund_si": active["fund_si"].max(),
        "max_cluster_si": active["cluster_si"].max(),
        "median_share": active["cross_cluster_share"].median(),
        "min_share": active["cross_cluster_share"].min(),
    }


# ── plotting ──────────────────────────────────────────────────────────────────

def _fig_to_stream(fig):
    import matplotlib.pyplot as plt
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def scatter_fund_vs_cluster(merged_df: pd.DataFrame, title: str):
    """Scatter with cluster SI on x-axis, fund SI on y-axis, y=x reference line.

    Marker color encodes cross_cluster_share (1 = on the diagonal = all
    sensitivity is cross-cluster; 0 = far above the diagonal = pure
    within-cluster reshuffling).
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    active = merged_df[merged_df["fund_si"] > 0].copy()
    if active.empty:
        return None

    fig, ax = plt.subplots(figsize=(6.5, 5.5))
    max_v = max(active["fund_si"].max(), active["cluster_si"].max()) * 1.08 + 1
    xs = np.linspace(0, max_v, 100)
    ax.fill_between(xs, xs, max_v, alpha=0.06, color="orange",
                    label="Within-cluster reshuffling region")
    ax.plot([0, max_v], [0, max_v], "k--", lw=0.9, alpha=0.5,
            label="y = x (all sensitivity is cross-cluster)")

    shares = active["cross_cluster_share"].clip(0, 1)
    sc = ax.scatter(
        active["cluster_si"], active["fund_si"],
        c=shares, cmap="viridis", vmin=0, vmax=1,
        s=42, alpha=0.85, edgecolor="white", linewidth=0.5, zorder=3,
    )

    cbar = plt.colorbar(sc, ax=ax, pad=0.02)
    cbar.set_label("cross_cluster_share = cluster_SI / fund_SI",
                   fontsize=8)

    ax.set_xlim(0, max_v)
    ax.set_ylim(0, max_v)
    ax.set_xlabel("Cluster-Level SI (pp)", fontsize=10)
    ax.set_ylabel("Fund-Level SI (pp)", fontsize=10)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=7, loc="lower right")
    fig.tight_layout()
    return _fig_to_stream(fig)


# ── high-level convenience ────────────────────────────────────────────────────

def build_all(write_csvs: bool = True):
    """Merge every analysis pair. Optionally writes combined_si.csv files.

    Returns a list of (spec, merged_df) pairs (unfiltered). Specs whose CSVs
    are missing are skipped with a printed warning.
    """
    out = []
    for spec in ANALYSIS_SPECS:
        merged = merge_pair(spec, apply_filter=False)
        if merged is None:
            print(f"  [fund_cluster_compare] Skipping {spec.id}: "
                  f"missing fund or cause CSV.")
            continue
        if write_csvs:
            path = write_combined_csv(spec, merged)
            print(f"  [fund_cluster_compare] Wrote {path}")
        out.append((spec, merged))
    return out


# ── docx section builder ──────────────────────────────────────────────────────

def _shade_row(row, hex_color="D9D9D9"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _add_table(doc, headers, rows, col_widths=None):
    from docx.shared import Inches, Pt
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    _shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def _add_picture(doc, stream, width=None):
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if stream is None:
        return
    doc.add_picture(stream, width=width if width is not None else Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fmt_si(v):
    return "—" if v is None or (isinstance(v, float) and np.isnan(v)) else f"{v:.2f}"


def _fmt_share(v):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    return f"{v:.0%}"


def build_fund_vs_cluster_section(doc, section_number=8, write_csvs=True,
                                  top_n_per_analysis=8):
    """Append the fund-level vs cluster-level SI comparison section.

    Builds:
      • Section heading + explanation paragraph.
      • Overview table across all analyses.
      • One subsection per analysis with a per-scenario detail table
        (sorted by cross_cluster_share ascending — most within-cluster first)
        and a scatter plot of fund SI vs cluster SI.

    When write_csvs=True, also writes the per-analysis combined_si.csv files.
    """
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h = doc.add_heading(
        f"{section_number}. Fund-Level vs Cluster-Level Sensitivity", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "Every test in this report has two sensitivity indices: a fund-level SI "
        "(how much money moves across the eight individual funds) and a "
        "cluster-level SI (how much moves across the three cause-area clusters: "
        "GHD, GCR, and AW). The fund-level SI is always at least as large as the "
        "cluster-level SI — aggregating fund deltas by cluster can only cancel out "
        "offsetting movements, never amplify them."
    )
    doc.add_paragraph(
        "We define the cross-cluster share as cluster_SI / fund_SI. This ratio is "
        "bounded in [0, 1]:"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("share ≈ 1: ").bold = True
    p.add_run(
        "the parameter genuinely redistributes money across cause areas; "
        "fund-level movement is entirely cross-cluster.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("share ≈ 0: ").bold = True
    p.add_run(
        "the parameter only reshuffles money between funds inside the same "
        "cause area — the fund-level SI overstates how much the cause-area "
        "mix actually changes.")
    doc.add_paragraph(
        "A high fund-level SI combined with a low cross-cluster share is a "
        "signal that an apparent sensitivity is mostly within-cluster churn "
        "rather than a real reallocation between cause areas."
    )

    # ── build merged frames once ──
    # The CSV gets the unfiltered merge (comprehensive data). The section's
    # tables, plots, and summary stats use the filtered merge (drops
    # bound=="single" rows etc.) so the visualizations focus on real
    # parameter-shift tests.
    pairs = []
    for spec in ANALYSIS_SPECS:
        raw = merge_pair(spec, apply_filter=False)
        if raw is None:
            print(f"  [fund_cluster_compare] Skipping {spec.id}: "
                  f"missing fund or cause CSV.")
            continue
        if write_csvs:
            path = write_combined_csv(spec, raw)
            print(f"  [fund_cluster_compare] Wrote {path}")
        filtered = (spec.pre_filter(raw) if spec.pre_filter is not None
                    else raw).reset_index(drop=True)
        pairs.append((spec, filtered))

    # ── overview table ──
    doc.add_paragraph()
    h2 = doc.add_heading("Overview across analyses", level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    overview_headers = [
        "Analysis", "# tests",
        "Max fund SI", "Max cluster SI",
        "Median share", "Min share",
    ]
    overview_rows = []
    for spec, merged in pairs:
        s = summary_stats_row(spec, merged)
        overview_rows.append([
            s["label"],
            str(s["n_tests"]),
            _fmt_si(s["max_fund_si"]),
            _fmt_si(s["max_cluster_si"]),
            _fmt_share(s["median_share"]),
            _fmt_share(s["min_share"]),
        ])
    _add_table(doc, overview_headers, overview_rows,
               col_widths=[2.2, 0.8, 1.0, 1.0, 1.0, 0.9])

    doc.add_paragraph(
        "Analyses with median share near 1 (worldview credences, aggregation "
        "methods, risk aversion, moral weights, time discounts) are dominated "
        "by genuine cross-cluster reallocation. Analyses with lower median "
        "share (diminishing returns, CE multipliers on intra-cluster groups) "
        "include a substantial portion of within-cluster reshuffling."
    )

    # ── per-analysis subsections ──
    for spec, merged in pairs:
        doc.add_paragraph()
        h2 = doc.add_heading(spec.label, level=2)
        h2.alignment = WD_ALIGN_PARAGRAPH.LEFT

        active = merged[merged["fund_si"] > 0].copy()
        if active.empty:
            doc.add_paragraph("No non-baseline tests with fund SI > 0.")
            continue

        # Detail table — sorted by cross_cluster_share ascending
        active_sorted = active.sort_values(
            "cross_cluster_share", ascending=True, na_position="last"
        ).head(top_n_per_analysis)
        detail_headers = ["Scenario", "Fund SI", "Cluster SI", "Cross-cluster share"]
        detail_rows = [
            [
                str(r["label"])[:60],
                _fmt_si(r["fund_si"]),
                _fmt_si(r["cluster_si"]),
                _fmt_share(r["cross_cluster_share"]),
            ]
            for _, r in active_sorted.iterrows()
        ]
        doc.add_paragraph(
            f"Top {len(detail_rows)} scenarios by lowest cross-cluster share "
            f"(most within-cluster reshuffling first):"
        )
        _add_table(doc, detail_headers, detail_rows,
                   col_widths=[3.0, 0.9, 0.9, 1.3])

        _add_picture(
            doc,
            scatter_fund_vs_cluster(merged, f"{spec.label}: fund SI vs cluster SI"),
            width=Inches(5.8),
        )
