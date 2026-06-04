"""
update_report_tables.py
=======================

Turn the raw sensitivity-analysis CSVs into the presentation tables used in the
Fund-Level Sensitivity Analysis report — automatically.

THE PROBLEM THIS SOLVES
-----------------------
Every time a model is re-run, the SI CSVs change, and each report table
(Scenario · Total shift · Between-cause shift · Share between-causes · Largest
gain · Largest loss) has to be hand-rebuilt from the raw output. This script
does that mechanical transform from a single dictionary (`TABLE_SPECS`) that
maps each report table to its source CSV(s).

WHAT IT PRODUCES
----------------
For every table in `TABLE_SPECS` it can:
  • write a Markdown preview (always) — easy to eyeball / diff in git, and
  • either UPDATE the table in place inside a *copy* of the report .docx
    (mode="inplace"), or EMIT a standalone .docx table to paste in
    (mode="emit", for analyses that don't have a table in the report yet).

The original report .docx is never modified — output goes to a new file
(default: "<report> (auto).docx").

THE DICTIONARY (how to add / change a table)
--------------------------------------------
Each entry in `TABLE_SPECS` is a `TableSpec`. The fields you set:

    TableSpec(
        name="time-discounts",            # id you pass on the CLI
        mode="inplace",                    # "inplace" | "emit"
        fund_csv="time-discounts/outputs/fund/discount_fund_si.csv",
        key_cols=["scenario_group", "multiplier"],   # cols that identify a row
        # between-cause SI: auto-detected from the CSV (si_cluster / cluster_si /
        # ca_sensitivity_index) unless you set between_col=..., or point at a
        # separate cause CSV with cause_csv=... (see across-the-board).
        locate=Locate(header_contains=["Total portfolio shift",
                                       "Largest gain"]),   # find the docx table
        columns=[                          # the table's columns, left→right
            Col.label("Scenario"),
            Col.label("Time discount multiplier"),
            Col.total(),
            Col.between(),
            Col.share(),
            Col.gain(),
            Col.loss(),
        ],
        rows=[                             # one entry per output row
            Row(labels=["Discount 500+", "1E-07"], keys=[["discount 500+", "1e-7"]]),
            ...
        ],
        sort="total_desc",                 # or "none" to keep `rows` order
    )

Notes:
  • `rows` labels are EDITORIAL (they don't change on re-run); the numbers are
    recomputed from the CSV every time. A row's `keys` is a list of CSV key
    tuples — give it more than one to collapse several identical scenarios into
    one labelled row (e.g. "All GCR ×100, 1000, 10000"); their deltas are
    averaged.
  • For 1-row-per-scenario tables (like GCR), use `rows="auto"` with a
    `label_map`/`exclude` instead of listing rows by hand.
  • Per-fund delta columns are auto-detected whether they are named
    `diff_<fund>` or `<fund>_delta`.

CLI
---
    python update_report_tables.py --list-csv NAME      # show a CSV's columns
    python update_report_tables.py --list-doc-tables    # index+header of every docx table
    python update_report_tables.py                      # all specs → md (+ docx)
    python update_report_tables.py time-discounts gcr-params   # just these
    python update_report_tables.py --no-docx            # markdown only
    python update_report_tables.py --in REPORT.docx --out OUT.docx

Run from anywhere; paths in specs are relative to sensitivity-analysis/.
"""

from __future__ import annotations

import argparse
import math
import os
import sys
from dataclasses import dataclass, field
from typing import Optional, Union

import pandas as pd

# sensitivity-analysis/ — all spec CSV paths are relative to this
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORTS = os.path.join(BASE, "reports")
DEFAULT_REPORT = os.path.join(REPORTS, "Fund-Level Sensitivity Analysis Draft 2.docx")
MD_OUT_DIR = os.path.join(REPORTS, "auto_tables")


# ── fund / cause display names ────────────────────────────────────────────────
# Internal slug → the label used in the report. Edit here if the report's
# wording changes.
FUND_DISPLAY = {
    "ea_awf": "EA AWF",
    "givewell": "GiveWell",
    "leaf": "LEAF",
    "longview_ai": "Longview AI",
    "longview_nuclear": "Longview Nuclear",
    "navigation_fund_cagefree": "TNF - Cage-free",
    "navigation_fund_general": "TNF - General",
    "sentinel_bio": "Sentinel Bio",
}

# Candidate column names for the between-cause (cluster) SI, in priority order.
BETWEEN_CANDIDATES = ["si_cluster", "cluster_si", "ca_sensitivity_index"]

FUND_DISPLAY_INV = {v: k for k, v in FUND_DISPLAY.items()}

# ── cell-shading scheme (largest gain / loss / out-of-cause columns) ───────────
# Each fund's cells are shaded by its cause cluster. The report used several
# shades within each family; these are one consistent shade per fund (the
# dominant colour from the original report). Edit freely — GHD = blues,
# AW = greens, GCR = purples/reds.
FUND_FILL = {
    # Global Health & Development — blues
    "givewell": "a4c2f4",
    "leaf": "cfe2f3",
    # Animal Welfare — greens
    "ea_awf": "b6d7a8",
    "navigation_fund_general": "b6d7a8",
    "navigation_fund_cagefree": "b6d7a8",
    # Global Catastrophic Risk — purples / reds
    "longview_ai": "d5a6bd",
    "sentinel_bio": "d5a6bd",
    "longview_nuclear": "d5a6bd",
}

# Light grey for risk-neutral worldview rows (worldview credences table).
NEUTRAL_GREY = "d9d9d9"

# Risk-neutral worldviews (exact `worldview` CSV strings) — shaded grey; all
# other (risk-averse) worldviews are left clear. Taken from the original report;
# add a worldview here if a new neutral one starts appearing in the table.
RISK_NEUTRAL_WORLDVIEWS = {
    "Total Utilitarianism — Default",
    "Non-Utilitarian Consequentialism — Default",
    "Non-Utilitarian Consequentialism — Person-Affecting/Cluelessness",
    "Contractualism — Person-Affecting/Cluelessness and Animals Baseline - Risk Neutral",
    "Non-Utilitarian Consequentialism — High Life Value / Cluelessness / Low animals",
}


def _fund_fill_from_text(text):
    """Cluster fill colour for a gain/loss cell, keyed on the first fund named."""
    if not text or text.strip() == "NA":
        return None
    first = text.split("(")[0].strip().rstrip(",").strip()
    slug = FUND_DISPLAY_INV.get(first)
    return FUND_FILL.get(slug)


# ── number formatting (matches the report's dominant style) ───────────────────
# Main columns: integer when |v| >= 10, one decimal below. Gains/losses carry an
# explicit sign and a "pp" suffix. Share is an integer percent, "NA" when the
# total shift is ~0.

def fmt_num(v: float) -> str:
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return "NA"
    return f"{v:.0f}" if abs(v) >= 9.95 else f"{v:.1f}"


def fmt_signed(v: float) -> str:
    """e.g. 27.0 -> '+27pp', -9.2 -> '-9.2pp', 0.18 -> '+0.2pp'."""
    sign = "-" if v < 0 else "+"
    a = abs(v)
    body = f"{a:.0f}" if a >= 9.95 else f"{a:.1f}"
    return f"{sign}{body}pp"


def fmt_share(total: float, between: float) -> str:
    if total is None or abs(total) < 0.05:
        return "NA"
    return f"{between / total:.0%}"


def _disp_key(v: float) -> str:
    """Rounded-display key, used to group ties (funds that print the same)."""
    a = abs(v)
    return f"{a:.0f}" if a >= 9.95 else f"{a:.1f}"


# ── delta-column handling ─────────────────────────────────────────────────────

def detect_delta_cols(columns) -> dict:
    """Map fund slug -> delta column name, handling diff_<fund> and <fund>_delta."""
    out = {}
    for c in columns:
        if c.startswith("diff_"):
            slug = c[len("diff_"):]
        elif c.endswith("_delta"):
            slug = c[: -len("_delta")]
        else:
            continue
        # only keep columns that correspond to a known fund
        if slug in FUND_DISPLAY:
            out[slug] = c
    return out


def detect_between_col(columns) -> Optional[str]:
    for cand in BETWEEN_CANDIDATES:
        if cand in columns:
            return cand
    return None


# ── metric computation ────────────────────────────────────────────────────────

@dataclass
class Metrics:
    total: float
    between: float
    deltas: dict          # fund slug -> pp delta

    def gain_str(self) -> str:
        return self._extreme(positive=True)

    def loss_str(self) -> str:
        return self._extreme(positive=False)

    def _extreme(self, positive: bool) -> str:
        if not self.deltas:
            return "NA"
        items = list(self.deltas.items())
        target = max(items, key=lambda kv: kv[1]) if positive \
            else min(items, key=lambda kv: kv[1])
        # No real movement in this direction -> the report shows "NA".
        if abs(target[1]) < 0.05:
            return "NA"
        target_disp = _disp_key(target[1])
        # ties = funds whose value prints identically AND shares the sign side
        winners = [
            (slug, v) for slug, v in items
            if _disp_key(v) == target_disp and ((v >= 0) == (target[1] >= 0))
        ]
        winners.sort(key=lambda kv: -abs(kv[1]))
        return ", ".join(
            f"{FUND_DISPLAY[slug]} ({fmt_signed(v)})" for slug, v in winners
        )


def compute_metrics(group: pd.DataFrame, delta_cols: dict,
                    between_col: Optional[str]) -> Metrics:
    """Compute metrics for one output row (group = 1+ source rows, averaged)."""
    total = float(group["sensitivity_index"].astype(float).mean())
    if between_col and between_col in group.columns:
        between = float(group[between_col].astype(float).mean())
    else:
        between = float("nan")
    deltas = {
        slug: float(group[col].astype(float).mean())
        for slug, col in delta_cols.items()
    }
    return Metrics(total=total, between=between, deltas=deltas)


# ── table specification ───────────────────────────────────────────────────────

@dataclass
class Locate:
    """How to find the table inside the .docx (inplace mode only)."""
    header_contains: list = field(default_factory=list)   # substrings in header row
    header_excludes: list = field(default_factory=list)   # substrings that must NOT appear
    first_cell: Optional[str] = None      # substring in first body row's first cell
                                          # (disambiguates tables with identical headers)
    index: Optional[int] = None           # or an explicit table index


@dataclass
class Col:
    kind: str                 # label | total | between | share | gain | loss | custom
    header: str = ""
    fn: Optional[object] = None   # for kind="custom": ctx dict -> str

    @staticmethod
    def label(header):  return Col("label", header)
    @staticmethod
    def total(header="Total portfolio shift (pp)"):    return Col("total", header)
    @staticmethod
    def between(header="Between-cause shift (pp)"):     return Col("between", header)
    @staticmethod
    def share(header="Share between-causes (%)"):       return Col("share", header)
    @staticmethod
    def gain(header="Largest gain"):                    return Col("gain", header)
    @staticmethod
    def loss(header="Largest loss"):                    return Col("loss", header)
    @staticmethod
    def custom(header, fn):                             return Col("custom", header, fn)


@dataclass
class Row:
    labels: list          # values for the label columns, left→right
    keys: list            # list of key tuples (lists) into the CSV; >1 = averaged


@dataclass
class TableSpec:
    name: str
    mode: str                                   # "inplace" | "emit"
    fund_csv: str
    key_cols: list
    columns: list                               # list[Col]
    # Row-selection mode:
    #   "editorial"  -> use the explicit `rows` list (fixed labels/grouping).
    #   "auto_top"   -> one row per CSV scenario, ranked by total SI descending,
    #                   keeping max(min_rows, #scenarios with total > threshold).
    select: str = "editorial"
    rows: Union[list, str] = field(default_factory=list)   # list[Row] for editorial
    threshold: float = 10.0                      # SI cutoff for auto_top inclusion
    min_rows: Optional[int] = None               # floor for auto_top row count;
                                                 # for inplace tables, None means
                                                 # "use the table's current row count"
    min_total: Optional[float] = None            # if set, auto_top keeps ALL rows with
                                                 # total >= this (hard cutoff; bypasses
                                                 # the threshold/ratchet count rule)
    row_filter: Optional[object] = None          # row_dict -> bool; keep only True rows
    label_fn: Optional[object] = None            # row_dict -> list[str] label cells
    row_shade_fn: Optional[object] = None        # meta -> hex|None; shades the first
                                                 # label cell (e.g. grey for neutral)
    locate: Optional[Locate] = None
    between_col: Optional[str] = None           # override auto-detect
    cause_csv: Optional[str] = None             # alt source for between-cause SI
    cause_si_col: str = "sensitivity_index"     # SI column within cause_csv
    sort: str = "total_desc"                    # "total_desc" | "none" (editorial)
    title: str = ""                             # heading for emit mode
    # convenience for single-key auto_top instead of a label_fn:
    label_map: dict = field(default_factory=dict)   # key value -> pretty label
    exclude: list = field(default_factory=list)     # key values to skip (e.g. baseline)

    def __post_init__(self):
        # Back-compat: rows="auto" is shorthand for select="auto_top".
        if self.rows == "auto":
            self.select = "auto_top"
            self.rows = []


# ── CSV loading ───────────────────────────────────────────────────────────────

def _load_fund_csv(spec: TableSpec) -> pd.DataFrame:
    path = os.path.join(BASE, spec.fund_csv)
    if not os.path.exists(path):
        raise FileNotFoundError(f"[{spec.name}] fund CSV not found: {path}")
    # key columns read as strings so keys match exactly (e.g. '1e-7', '0').
    conv = {k: str for k in spec.key_cols}
    df = pd.read_csv(path, converters=conv)
    df.columns = [c.strip() for c in df.columns]
    for k in spec.key_cols:
        df[k] = df[k].astype(str).str.strip()
    return df


def _attach_between_from_cause(spec: TableSpec, fund_df: pd.DataFrame) -> str:
    """Join a separate cause CSV's SI onto fund_df as '_between'. Returns col name."""
    cpath = os.path.join(BASE, spec.cause_csv)
    if not os.path.exists(cpath):
        raise FileNotFoundError(f"[{spec.name}] cause CSV not found: {cpath}")
    conv = {k: str for k in spec.key_cols}
    cdf = pd.read_csv(cpath, converters=conv)
    cdf.columns = [c.strip() for c in cdf.columns]
    for k in spec.key_cols:
        cdf[k] = cdf[k].astype(str).str.strip()
    sub = cdf[spec.key_cols + [spec.cause_si_col]].rename(
        columns={spec.cause_si_col: "_between"})
    merged = fund_df.merge(sub, on=spec.key_cols, how="left")
    return "_between", merged


def _auto_label_cells(spec: TableSpec, row_dict: dict) -> list:
    """Produce the label cells for one CSV row in auto_top mode."""
    if spec.label_fn is not None:
        return list(spec.label_fn(row_dict))
    if len(spec.key_cols) == 1:
        val = row_dict[spec.key_cols[0]]
        return [spec.label_map.get(val, val)]
    raise ValueError(
        f"[{spec.name}] select='auto_top' with multiple key_cols needs a label_fn")


def _computed_editorial(spec, fund_df, delta_cols, between_col, n_label_cols):
    """List of (label_cells, metrics, meta) from the explicit `rows` list."""
    computed = []
    for row in spec.rows:
        if len(row.labels) != n_label_cols:
            raise ValueError(
                f"[{spec.name}] row {row.labels!r} has {len(row.labels)} labels "
                f"but spec has {n_label_cols} label columns")
        frames = []
        for key in row.keys:
            mask = pd.Series(True, index=fund_df.index)
            for kc, kv in zip(spec.key_cols, key):
                mask &= fund_df[kc] == str(kv).strip()
            sub = fund_df[mask]
            if sub.empty:
                raise ValueError(
                    f"[{spec.name}] no CSV row for key "
                    f"{dict(zip(spec.key_cols, key))}")
            frames.append(sub)
        group = pd.concat(frames)
        meta = {c: group.iloc[0][c] for c in group.columns}
        computed.append((row.labels, compute_metrics(group, delta_cols, between_col),
                         meta))
    if spec.sort == "total_desc":
        computed.sort(key=lambda lmt: -lmt[1].total)
    return computed


def _computed_auto_top(spec, fund_df, delta_cols, between_col, current_count):
    """List of (label_cells, metrics, meta), ranked by total SI and trimmed.

    By default keeps max(min_rows, #scenarios with total > threshold) scenarios;
    for in-place tables min_rows defaults to the table's current row count, so
    the table never shrinks below what's published but grows when more scenarios
    cross the threshold (and always re-ranks by current SI). If `min_total` is
    set, instead keeps every row with total >= min_total (a hard cutoff).
    """
    computed = []
    for _, r in fund_df.iterrows():
        row_dict = {c: r[c] for c in fund_df.columns}
        # skip excluded scenarios (match on any key column)
        if any(str(r[kc]).strip() in spec.exclude for kc in spec.key_cols):
            continue
        if spec.row_filter is not None and not spec.row_filter(row_dict):
            continue
        sub = fund_df[fund_df.index == r.name]
        computed.append((_auto_label_cells(spec, row_dict),
                         compute_metrics(sub, delta_cols, between_col),
                         row_dict))

    computed.sort(key=lambda lmt: -lmt[1].total)

    if spec.min_total is not None:
        return [c for c in computed if c[1].total >= spec.min_total]

    floor = spec.min_rows if spec.min_rows is not None else (current_count or 0)
    n_above = sum(1 for _, m, _ in computed if m.total > spec.threshold)
    keep = max(floor, n_above)
    return computed[:keep]


def build_rows(spec: TableSpec, current_count: Optional[int] = None):
    """Return (header, data_rows) where data_rows are lists of formatted strings.

    `current_count` is the number of body rows currently in the matched .docx
    table; used as the auto_top floor when spec.min_rows is None.
    """
    fund_df = _load_fund_csv(spec)

    between_col = spec.between_col or detect_between_col(fund_df.columns)
    if spec.cause_csv:
        between_col, fund_df = _attach_between_from_cause(spec, fund_df)

    delta_cols = detect_delta_cols(fund_df.columns)
    if not delta_cols:
        raise ValueError(f"[{spec.name}] no per-fund delta columns found in "
                         f"{spec.fund_csv}")

    n_label_cols = sum(1 for c in spec.columns if c.kind == "label")

    if spec.select == "auto_top":
        computed = _computed_auto_top(spec, fund_df, delta_cols, between_col,
                                      current_count)
    else:
        computed = _computed_editorial(spec, fund_df, delta_cols, between_col,
                                       n_label_cols)

    header = [c.header for c in spec.columns]
    data_rows = []
    fills = []
    for label_cells, m, meta in computed:
        label_iter = iter(label_cells)
        cells = []
        row_fill = []
        seen_label = False
        for c in spec.columns:
            f = None
            if c.kind == "label":
                cells.append(str(next(label_iter)))
                # shade the FIRST label cell only (e.g. grey for neutral rows)
                if not seen_label and spec.row_shade_fn is not None:
                    f = spec.row_shade_fn(meta)
                seen_label = True
            elif c.kind == "total":
                cells.append(fmt_num(m.total))
            elif c.kind == "between":
                cells.append(fmt_num(m.between))
            elif c.kind == "share":
                cells.append(fmt_share(m.total, m.between))
            elif c.kind == "gain":
                cells.append(m.gain_str())
                f = _fund_fill_from_text(cells[-1])
            elif c.kind == "loss":
                cells.append(m.loss_str())
                f = _fund_fill_from_text(cells[-1])
            elif c.kind == "custom":
                cells.append(c.fn({"metrics": m, "meta": meta}))
                f = _fund_fill_from_text(cells[-1])
            else:
                raise ValueError(f"unknown column kind {c.kind!r}")
            row_fill.append(f)
        data_rows.append(cells)
        fills.append(row_fill)
    return header, data_rows, fills


# ── baseline allocation table (different shape: cluster subtotals + funds) ─────

@dataclass
class AllocSpec:
    """The baseline portfolio table: cluster subtotal rows + per-fund rows + a
    Total, with $M and % columns. Structurally unlike the SI tables, so it has
    its own builder."""
    name: str
    mode: str                       # "inplace"
    title: str
    alloc_csv: str                  # CSV with per-fund baseline allocation %
    fund_col: str                   # column holding the fund slug
    pct_col: str                    # column holding the allocation %
    baseline_filter: dict           # rows to keep, e.g. {"fund_varied": "baseline"}
    clusters: list                  # [(cluster_display, [fund_slug, ...]), ...]
    fund_display: dict              # fund slug -> full display name
    columns: list                   # 3 header strings
    locate: Optional[Locate] = None
    budget_m: Optional[float] = None    # total $M; None -> sum baseline.json stages
    select: str = "n/a"             # (so main() can read .select uniformly)


def _round_half_up(x):
    return int(math.floor(float(x) + 0.5))


def _baseline_budget_m():
    """Total budget = sum of stage budgets in sensitivity-analysis/baseline.json."""
    import json
    with open(os.path.join(BASE, "baseline.json"), encoding="utf-8") as f:
        stages = json.load(f)["stages"]
    return sum(s["budget"] for s in stages)


def build_alloc_rows(spec: AllocSpec, current_count=None):
    df = pd.read_csv(os.path.join(BASE, spec.alloc_csv))
    df.columns = [c.strip() for c in df.columns]
    mask = pd.Series(True, index=df.index)
    for k, v in spec.baseline_filter.items():
        mask &= df[k].astype(str).str.strip() == str(v)
    base = df[mask]
    pct = {str(r[spec.fund_col]).strip(): float(r[spec.pct_col])
           for _, r in base.iterrows()}

    budget = spec.budget_m if spec.budget_m is not None else _baseline_budget_m()

    def row(name, p):
        return [name, str(_round_half_up(p / 100.0 * budget)), f"{_round_half_up(p)}%"]

    data_rows = []
    for cluster_name, slugs in spec.clusters:
        missing = [s for s in slugs if s not in pct]
        if missing:
            raise ValueError(f"[{spec.name}] funds missing from {spec.alloc_csv}: "
                             f"{missing}")
        cpct = sum(pct[s] for s in slugs)
        data_rows.append(row(cluster_name, cpct))
        for s in slugs:
            data_rows.append(row(spec.fund_display[s], pct[s]))
    data_rows.append(["Total", str(_round_half_up(budget)), "100%"])
    fills = [[None] * len(spec.columns) for _ in data_rows]
    return list(spec.columns), data_rows, fills


# ── markdown ──────────────────────────────────────────────────────────────────

def to_markdown(header, rows) -> str:
    out = ["| " + " | ".join(header) + " |",
           "| " + " | ".join("---" for _ in header) + " |"]
    for r in rows:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


# ── docx helpers (styling matches fund_cluster_compare._add_table) ────────────

def _shade_header(row, hex_color="D9D9D9"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _set_cell(cell, text, bold=False, size_pt=9):
    from docx.shared import Pt
    cell.text = str(text)
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(size_pt)


def _set_cell_fill(cell, hex_color):
    """Set (or clear) a cell's background shading. Removes any existing fill so
    stale colours from a cloned template don't persist."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    if hex_color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _fill_new_table(table, header, rows, fills=None):
    """Fill a freshly-created table (emit mode) — add_row is safe here."""
    for i, h in enumerate(header):
        _set_cell(table.rows[0].cells[i], h, bold=True)
    _shade_header(table.rows[0])
    for ri, r in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(r):
            _set_cell(cells[i], val)
            if fills is not None:
                _set_cell_fill(cells[i], fills[ri][i])


def _rebuild_body_by_cloning(table, rows, fills=None):
    """Replace a docx table's body rows, preserving the header and cell styling.

    Clones an existing body row's XML (so we inherit its cell formatting) instead
    of table.add_row(), which crashes on report tables that store fractional
    column widths in <w:gridCol w:w="1991.97"/>. Any shading on the clone is
    cleared and re-applied per `fills` so colours track the new data.
    """
    import copy
    body = table.rows[1:]
    if not body:
        raise ValueError("table has no body row to use as a clone template")
    template_tr = copy.deepcopy(body[0]._tr)
    for r in body:
        r._tr.getparent().remove(r._tr)
    for ri, values in enumerate(rows):
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)
        cells = table.rows[-1].cells
        for i, val in enumerate(values):
            _set_cell(cells[i], val)
            if fills is not None:
                _set_cell_fill(cells[i], fills[ri][i])


def _find_doc_table(doc, locate: Locate):
    if locate is None:
        return None
    if locate.index is not None:
        return doc.tables[locate.index]
    wanted = [w.lower() for w in locate.header_contains]
    banned = [w.lower() for w in locate.header_excludes]
    for t in doc.tables:
        hdr = " | ".join(c.text.strip().lower() for c in t.rows[0].cells)
        if not all(w in hdr for w in wanted):
            continue
        if any(w in hdr for w in banned):
            continue
        if locate.first_cell is not None:
            if len(t.rows) < 2:
                continue
            if locate.first_cell.lower() not in t.rows[1].cells[0].text.strip().lower():
                continue
        return t
    return None


# ── per-spec runners ──────────────────────────────────────────────────────────

def write_markdown(spec, header, rows):
    os.makedirs(MD_OUT_DIR, exist_ok=True)
    path = os.path.join(MD_OUT_DIR, f"{spec.name}.md")
    title = spec.title or spec.name
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"### {title}\n\n")
        f.write(to_markdown(header, rows))
        f.write("\n")
    return path


def apply_inplace(spec, doc, header, rows, fills=None):
    table = _find_doc_table(doc, spec.locate)
    if table is None:
        print(f"  [{spec.name}] WARNING: no matching table found in the report "
              f"(locate={spec.locate}); skipped in-place update.")
        return False
    # Keep the report's existing header wording; only refresh the body.
    _rebuild_body_by_cloning(table, rows, fills)
    print(f"  [{spec.name}] updated in-place ({len(rows)} rows).")
    return True


def emit_standalone(spec, header, rows, fills=None):
    """Write a standalone one-table .docx for a table not yet in the report."""
    import docx
    from docx.shared import Pt
    d = docx.Document()
    d.add_heading(spec.title or spec.name, level=2)
    t = d.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    _fill_new_table(t, header, rows, fills)
    path = os.path.join(MD_OUT_DIR, f"{spec.name}.docx")
    os.makedirs(MD_OUT_DIR, exist_ok=True)
    d.save(path)
    print(f"  [{spec.name}] emitted standalone table -> {path}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE REGISTRY — the dictionary mapping report tables to source CSVs.
#  Add a TableSpec here for each report table you want kept in sync.
# ══════════════════════════════════════════════════════════════════════════════

# Aggregation-method slugs -> report display names.
METHOD_DISPLAY = {
    "nashBargaining": "Nash Bargaining",
    "marketplace": "Marketplace",
    "MEC": "MEC",
    "lexicographicMaximin": "Lexicographic Maximin",
    "splitCycle": "Split Cycle",
    "borda": "Borda",
    "MET": "MET",
}


def _fmt_mult(v):
    """'2.5' -> '2.5', '10' / '10.0' -> '10' (for 'Cap at Nx budget' labels)."""
    f = float(v)
    return f"{f:.0f}" if f == int(f) else f"{f:g}"


def _fmt_pct(v):
    return f"{float(v):.0%}"


# Fund -> cause cluster (for the "out-of-cause" column).
FUND_CAUSE = {
    "givewell": "ghd", "leaf": "ghd",
    "ea_awf": "aw", "navigation_fund_cagefree": "aw", "navigation_fund_general": "aw",
    "sentinel_bio": "gcr", "longview_nuclear": "gcr", "longview_ai": "gcr",
}

# across-the-board: fund_varied slug -> short report label, and -> cause cluster.
VARIED_DISPLAY = {
    "all_gcr": "All GCR", "all_aw": "All AW", "all_ghd": "All GHD",
    "givewell": "GW", "leaf": "LEAF", "ea_awf": "EA AWF",
    "longview_ai": "AI", "longview_nuclear": "Nuclear", "sentinel_bio": "Bio",
    "navigation_fund_cagefree": "TNF - Cage-free",
    "navigation_fund_general": "TNF - General",
}
VARIED_CAUSE = {"all_gcr": "gcr", "all_aw": "aw", "all_ghd": "ghd"}

# risk aversion: profile slug -> display name (test = "<start>_to_<final>").
PROFILE_DISPLAY = {
    "neutral": "Neutral",
    "specialblend": "Baseline",
    "combined": "Upside skeptical + Downside critical",
    "wlu_5": "WLU (moderate)",
    "wlu_10": "WLU (high)",
    "downside_critical": "Downside critical",
    "cont_upside_skep": "Upside skeptical",
}


def _across_out_of_cause(ctx):
    """Largest allocation change among funds OUTSIDE the varied fund's cause."""
    fv = ctx["meta"]["fund_varied"]
    cause = VARIED_CAUSE.get(fv, FUND_CAUSE.get(fv))
    out = {f: d for f, d in ctx["metrics"].deltas.items()
           if FUND_CAUSE.get(f) != cause}
    if not out:
        return "NA"
    f = max(out, key=lambda k: abs(out[k]))
    if abs(out[f]) < 0.05:
        return "NA"
    return f"{FUND_DISPLAY[f]} ({fmt_signed(out[f])})"


def _risk_label(r):
    start, final = str(r["test"]).split("_to_", 1)
    return [PROFILE_DISPLAY.get(start, start), PROFILE_DISPLAY.get(final, final)]


def _worldview_label(r):
    wv = str(r["worldview"])
    name = f"{wv.split(' — ', 1)[0]} ({wv.split(' — ', 1)[1]})" if " — " in wv else wv
    return [name, f"{_fmt_pct(r['credence_base'])} to {_fmt_pct(r['credence_scenario'])}"]


TABLE_SPECS = [

    # ── GCR parameter sensitivity ────────────────────────────────────────────
    # Active analysis; not yet in the report, so emit a standalone table to paste
    # in. Ranked by total SI; between-cause SI = si_cluster.
    # Row count: max(min_rows, #scenarios with total > 10). GCR isn't in the
    # report yet so there's no "current count" to read — min_rows=8 keeps the
    # known real signals visible (NOTES.md: 8 scenarios clear of the 2.54 pp
    # fund-level noise floor). Change min_rows (or drop it) to taste; once this
    # table is pasted into the report and switched to mode="inplace", the floor
    # becomes the table's live row count automatically.
    TableSpec(
        name="gcr-params",
        mode="emit",
        title="GCR Parameter Sensitivity",
        fund_csv="gcr-params/outputs/fund/gcr_sensitivity_index.csv",
        key_cols=["scenario"],
        columns=[
            Col.label("Scenario"),
            Col.total(),
            Col.between(),
            Col.share(),
            Col.gain(),
            Col.loss(),
        ],
        select="auto_top",
        threshold=10.0,
        min_rows=8,
        exclude=["baseline", "noise_check"],
        label_map={
            "cause_fractions_equal": "Cause fractions equal",
            "cause_fractions_bio_nuclear_5x_higher": "Bio/nuclear fractions 5x higher",
            "cause_fractions_bio_nuclear_unequal": "Bio/nuclear fractions unequal",
            "near_pessimistic_outcomes": "Near-pessimistic outcomes",
            "r_inf_100x_up": "Background x-risk floor 100x up",
            "r_inf_100x_down": "Background x-risk floor 100x down",
            "rel_risk_10x_up": "Relative risk reduction 10x up",
            "rel_risk_10x_down": "Relative risk reduction 10x down",
            "rel_risk_100x_down": "Relative risk reduction 100x down",
            "no_cubic_growth": "No stellar expansion",
            "p_zero_75_pct": "P(zero impact) = 75%",
            "p_zero_5x_lower": "P(zero impact) 5x lower",
            "p_harm_25pct_higher": "P(harm) 25% higher",
            "s_10x_faster": "Settlement speed 10x faster",
            "s_current_speeds": "Settlement speed = current",
        },
        sort="total_desc",
    ),

    # ── Time discounting ─────────────────────────────────────────────────────
    # Already in the report (the table with the "Time discount multiplier" col).
    # Updated in place. Editorial: only the 1e-7 … 0 multipliers are shown.
    TableSpec(
        name="time-discounts",
        mode="inplace",
        title="Time Discounting",
        fund_csv="time-discounts/outputs/fund/discount_fund_si.csv",
        key_cols=["scenario_group", "multiplier"],
        locate=Locate(header_contains=["time discount multiplier", "largest gain"]),
        columns=[
            Col.label("Scenario"),
            Col.label("Time discount multiplier"),
            Col.total("Total portfolio shift"),
            Col.between("Between-cause portfolio shift"),
            Col.share("Share of SI is between clusters (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        rows=[
            Row(["Discount 500+", "1E-07"], [["discount 500+", "1e-7"]]),
            Row(["Discount 500+", "1E-08"], [["discount 500+", "1e-8"]]),
            Row(["Discount 500+", "1E-09"], [["discount 500+", "1e-9"]]),
            Row(["Discount 500+", "1E-10"], [["discount 500+", "1e-10"]]),
            Row(["Discount 500+", "0"],     [["discount 500+", "0"]]),
            Row(["Discount 100-500 and 500+", "1E-07"], [["discount 100-500 and 500+", "1e-7"]]),
            Row(["Discount 100-500 and 500+", "1E-08"], [["discount 100-500 and 500+", "1e-8"]]),
            Row(["Discount 100-500 and 500+", "1E-09"], [["discount 100-500 and 500+", "1e-9"]]),
            Row(["Discount 100-500 and 500+", "1E-10"], [["discount 100-500 and 500+", "1e-10"]]),
            Row(["Discount 100-500 and 500+", "0"],     [["discount 100-500 and 500+", "0"]]),
        ],
        sort="none",   # keep the editorial scenario/multiplier order
    ),

    # ── Aggregation methods (credence on each voting method) ─────────────────
    # In the report (the "Method | Scenario | …" table). Auto-ranked by total SI;
    # labels derived fully from the CSV (method name + "X% to Y%"), so no manual
    # mapping needed. Floor = the report table's current row count.
    TableSpec(
        name="aggregation-methods",
        mode="inplace",
        title="Aggregation Methods",
        fund_csv="aggregation-methods/outputs/fund/split_credences_index.csv",
        key_cols=["scenario"],
        locate=Locate(header_contains=["method", "largest gain"]),
        columns=[
            Col.label("Method"),
            Col.label("Scenario"),
            Col.total("Total portfolio shift"),
            Col.between("Between-cause shift"),
            Col.share("Share of shift is between causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        exclude=["baseline"],
        label_fn=lambda r: [
            METHOD_DISPLAY.get(r["method"], r["method"]),
            f"{_fmt_pct(r['credence_base'])} to {_fmt_pct(r['credence_scenario'])}",
        ],
    ),

    # ── Diminishing returns: max additional spend cap ────────────────────────
    # In the report (the "Cap at Nx budget" table). Auto-ranked; label derived
    # from the spend multiplier. between-cause SI = ca_sensitivity_index.
    TableSpec(
        name="diminishing-returns-maxspend",
        mode="inplace",
        title="Diminishing Returns — Max Additional Spend",
        fund_csv="diminishing-returns/outputs/fund/max_spend_sensitivity_by_fund.csv",
        key_cols=["scenario"],
        locate=Locate(header_contains=["scenario", "between-cause shift (pp)"]),
        columns=[
            Col.label("Scenario"),
            Col.total("Total portfolio shift (pp)"),
            Col.between("Between-cause shift (pp)"),
            Col.share("Share of shift between-causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        exclude=["baseline_5x"],
        label_fn=lambda r: [f"Cap at {_fmt_mult(r['max_addl_spend_multiplier'])}x budget"],
    ),

    # ── Moral weights ─────────────────────────────────────────────────────────
    # Labels confirmed: low_animals="Lowest weights", moderate_high="Baseline
    # weights", moderate="Moderate weights", sentience_only="Sentience-only
    # adjustment", and the numeric 0.1/0.2/0.5 are the "Nx multiplier" rows.
    TableSpec(
        name="moral-weights",
        mode="inplace",
        title="Moral Weights",
        fund_csv="moral-weights/outputs/fund/moral_weights_overall_si.csv",
        key_cols=["multiplier"],
        locate=Locate(header_contains=["between-causes portfolio shift (pp)"],
                      header_excludes=["starting risk profile"]),
        columns=[
            Col.label("Scenario"),
            Col.total("Total portfolio shift (pp)"),
            Col.between("Between-causes portfolio shift (pp)"),
            Col.share("Share of shift between-causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        exclude=["baseline"],
        label_map={
            "low_animals": "Lowest weights",
            "0.1": "0.1x multiplier",
            "sentience_only": "Sentience-only adjustment",
            "moderate_high": "Baseline weights",
            "0.2": "0.2x multiplier",
            "moderate": "Moderate weights",
            "0.5": "0.5x multiplier",
        },
    ),

    # ── Risk aversion: starting profile = Neutral ────────────────────────────
    # test = "<start>_to_<final>"; two report tables share an identical header,
    # disambiguated by the first body cell (Neutral vs Baseline).
    TableSpec(
        name="risk-aversion-neutral",
        mode="inplace",
        title="Risk Aversion (from Neutral)",
        fund_csv="risk-aversion/outputs/fund/risk_aversion_summary.csv",
        key_cols=["test"],
        locate=Locate(header_contains=["starting risk profile", "largest gain"],
                      first_cell="Neutral"),
        columns=[
            Col.label("Starting Risk Profile"),
            Col.label("Final Risk Profile"),
            Col.total("Total portfolio shift (pp)"),
            Col.between("Between-causes portfolio shift (pp)"),
            Col.share("Share of shift between-causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        row_filter=lambda r: str(r["test"]).startswith("neutral_to_"),
        label_fn=_risk_label,
    ),

    # ── Risk aversion: starting profile = Baseline (specialblend) ────────────
    TableSpec(
        name="risk-aversion-baseline",
        mode="inplace",
        title="Risk Aversion (from Baseline)",
        fund_csv="risk-aversion/outputs/fund/risk_aversion_summary.csv",
        key_cols=["test"],
        locate=Locate(header_contains=["starting risk profile", "largest gain"],
                      first_cell="Baseline"),
        columns=[
            Col.label("Starting Risk Profile"),
            Col.label("Final Risk Profile"),
            Col.total("Total portfolio shift (pp)"),
            Col.between("Between-causes portfolio shift (pp)"),
            Col.share("Share of shift between-causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        row_filter=lambda r: str(r["test"]).startswith("specialblend_to_"),
        label_fn=_risk_label,
    ),

    # ── Across-the-board CE multipliers ──────────────────────────────────────
    # One row per (fund_varied, multiplier) — no grouping, ranked high→low.
    # between-cause SI comes from the cause-area CSV; the 7th column is the
    # largest change among funds outside the varied fund's cause.
    TableSpec(
        name="across-the-board",
        mode="inplace",
        title="Across-the-Board CE Multipliers",
        fund_csv="across-the-board/outputs/fund/ce_multiplier_si.csv",
        cause_csv="across-the-board/outputs/cause/cause_area_si.csv",
        cause_si_col="sensitivity_index",
        key_cols=["fund_varied", "multiplier"],
        locate=Locate(header_contains=["fund varied"]),
        columns=[
            Col.label("Fund Varied"),
            Col.total("Total allocation shift (pp)"),
            Col.between("Between-cause allocation shift (pp)"),
            Col.share("Between-cause reallocation pct. (%)"),
            Col.gain("Largest Gain"),
            Col.loss("Largest Loss"),
            Col.custom("Largest out-of-cause change", _across_out_of_cause),
        ],
        select="auto_top",
        threshold=10.0,
        exclude=["baseline"],
        label_fn=lambda r: [
            f"{VARIED_DISPLAY.get(r['fund_varied'], r['fund_varied'])} "
            f"x{_fmt_mult(r['multiplier'])}"],
    ),

    # ── Worldview credences ──────────────────────────────────────────────────
    # Label = "Family (Variant)" from the worldview column + "X% to Y%".
    # Keep every shift with total SI >= 5; drop single-bound rows.
    TableSpec(
        name="worldview-credences",
        mode="inplace",
        title="Worldview Credences",
        fund_csv="worldview-sensitivity/outputs/fund/split_credences_index.csv",
        cause_csv="worldview-sensitivity/outputs/cause/cause_area_index.csv",
        cause_si_col="sensitivity_index",
        key_cols=["scenario"],
        locate=Locate(header_contains=["worldview", "credence shift"]),
        columns=[
            Col.label("Worldview"),
            Col.label("Credence shift"),
            Col.total("Total portfolio shift"),
            Col.between("Between-causes shift"),
            Col.share("Share of total shift is between causes (%)"),
            Col.gain("Largest Gain"),
            Col.loss("Largest Loss"),
        ],
        select="auto_top",
        min_total=5.0,
        row_filter=lambda r: str(r.get("bound", "")).strip() != "single",
        label_fn=_worldview_label,
        row_shade_fn=lambda meta: (NEUTRAL_GREY
                                   if str(meta.get("worldview", "")).strip()
                                   in RISK_NEUTRAL_WORLDVIEWS else None),
    ),

    # ── Diminishing returns: power (speed of impact) ─────────────────────────
    # Source = dr_sensitivity_by_fund.csv (combos vary the DR power for GCR/AW).
    # This is the report's "GCR and AW slow / fast" table.
    TableSpec(
        name="diminishing-returns-power",
        mode="inplace",
        title="Diminishing Returns — Power (speed of impact)",
        fund_csv="diminishing-returns/outputs/fund/dr_sensitivity_by_fund.csv",
        key_cols=["combo"],
        locate=Locate(header_contains=["share of shift that is between causes"]),
        columns=[
            Col.label(""),
            Col.total("Total portfolio shift (pp)"),
            Col.between("Between-cause portfolio shift (pp)"),
            Col.share("Share of shift that is between causes (%)"),
            Col.gain("Largest gain"),
            Col.loss("Largest loss"),
        ],
        select="auto_top",
        threshold=10.0,
        exclude=["baseline"],
        label_map={
            "gcr_and_aw_slow": "GCR and AW slow",
            "gcr_slow_aw_fast": "GCR slow, AW fast",
            "gcr_and_aw_fast": "GCR and AW fast",
            "aw_fast": "AW fast",
            "gcr_slow": "GCR slow",
            "aw_slow": "AW slow",
            "gcr_fast_aw_slow": "GCR fast, AW slow",
            "gcr_fast": "GCR fast",
        },
    ),

    # ── Baseline portfolio allocation (cluster subtotals + funds + total) ─────
    # Source = the baseline rows of the across-the-board allocation CSV
    # (weighted_allocation_pct); budget = sum of baseline.json stage budgets.
    AllocSpec(
        name="baseline-allocation",
        mode="inplace",
        title="Baseline Allocation",
        alloc_csv="across-the-board/outputs/fund/ce_multiplier_allocations.csv",
        fund_col="recipient_fund",
        pct_col="weighted_allocation_pct",
        baseline_filter={"fund_varied": "baseline"},
        locate=Locate(header_contains=["cluster/fund", "percent of funding"]),
        columns=["Cluster/Fund", "Allocation ($M)", "Percent of Funding"],
        clusters=[
            ("Animal Welfare",
             ["ea_awf", "navigation_fund_general", "navigation_fund_cagefree"]),
            ("Global Health and Development", ["givewell", "leaf"]),
            ("Global Catastrophic Risks",
             ["sentinel_bio", "longview_ai", "longview_nuclear"]),
        ],
        fund_display={
            "ea_awf": "EA Animal Welfare Fund",
            "navigation_fund_general": "TNF - General",
            "navigation_fund_cagefree": "TNF - Cage-free",
            "givewell": "GiveWell",
            "leaf": "LEAF",
            "sentinel_bio": "Sentinel Bio",
            "longview_ai": "Longview Frontier AI Fund",
            "longview_nuclear": "Longview Nuclear Weapons Policy Fund",
        },
    ),
]


SPECS_BY_NAME = {s.name: s for s in TABLE_SPECS}


# ── CLI ───────────────────────────────────────────────────────────────────────

def cmd_list_csv(name):
    spec = SPECS_BY_NAME.get(name)
    if not spec:
        print(f"Unknown spec '{name}'. Known: {sorted(SPECS_BY_NAME)}")
        return 1
    df = pd.read_csv(os.path.join(BASE, spec.fund_csv), nrows=5)
    print(f"{spec.fund_csv}\n  columns: {list(df.columns)}")
    print(f"  detected deltas: {detect_delta_cols(df.columns)}")
    print(f"  detected between-col: {detect_between_col(df.columns)}")
    return 0


def cmd_list_doc_tables(report_path):
    import docx
    d = docx.Document(report_path)
    for i, t in enumerate(d.tables):
        hdr = " | ".join(c.text.strip() for c in t.rows[0].cells)
        print(f"[{i}] ({len(t.rows)}x{len(t.columns)}) {hdr[:120]}")
    return 0


def main(argv):
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("specs", nargs="*", help="spec names to run (default: all)")
    ap.add_argument("--in", dest="report_in", default=DEFAULT_REPORT,
                    help="source report .docx (for inplace tables)")
    ap.add_argument("--out", dest="report_out", default=None,
                    help="output .docx (default: '<report> (auto).docx')")
    ap.add_argument("--no-docx", action="store_true",
                    help="markdown previews only; don't touch any .docx")
    ap.add_argument("--list-csv", metavar="NAME", help="show a spec's CSV columns")
    ap.add_argument("--list-doc-tables", action="store_true",
                    help="list every table in the report .docx with its index")
    args = ap.parse_args(argv[1:])

    if args.list_csv:
        return cmd_list_csv(args.list_csv)
    if args.list_doc_tables:
        return cmd_list_doc_tables(args.report_in)

    targets = args.specs or [s.name for s in TABLE_SPECS]
    unknown = [t for t in targets if t not in SPECS_BY_NAME]
    if unknown:
        print(f"Unknown spec(s): {unknown}. Known: {sorted(SPECS_BY_NAME)}")
        return 1
    specs = [SPECS_BY_NAME[t] for t in targets]

    # Load the report once (read-only here) so we can read each in-place table's
    # current row count — the auto_top floor ("however many are currently in
    # there"). The same in-memory doc is mutated and saved later; the original
    # file is never touched until we save to the output path.
    doc = None
    if any(s.mode == "inplace" for s in specs) and os.path.exists(args.report_in):
        import docx
        doc = docx.Document(args.report_in)

    def current_count_for(spec):
        if spec.mode != "inplace" or doc is None:
            return None
        t = _find_doc_table(doc, spec.locate)
        return (len(t.rows) - 1) if t is not None else None

    # Build every spec's rows first (so a CSV/spec error aborts before we write).
    built = []
    for spec in specs:
        cc = current_count_for(spec)
        builder = build_alloc_rows if isinstance(spec, AllocSpec) else build_rows
        header, rows, fills = builder(spec, current_count=cc)
        built.append((spec, header, rows, fills))
        note = f" (floor from current table: {cc})" if cc is not None and \
            spec.select == "auto_top" else ""
        print(f"  [{spec.name}] computed {len(rows)} rows.{note}")

    # Markdown previews (always).
    for spec, header, rows, fills in built:
        write_markdown(spec, header, rows)
    print(f"  Markdown previews -> {MD_OUT_DIR}")

    if args.no_docx:
        return 0

    inplace = [b for b in built if b[0].mode == "inplace"]
    emit = [b for b in built if b[0].mode == "emit"]

    for spec, header, rows, fills in emit:
        emit_standalone(spec, header, rows, fills)

    if inplace:
        if doc is None:
            import docx
            doc = docx.Document(args.report_in)
        any_applied = False
        for spec, header, rows, fills in inplace:
            any_applied |= apply_inplace(spec, doc, header, rows, fills)
        if any_applied:
            out = args.report_out
            if out is None:
                root, ext = os.path.splitext(args.report_in)
                out = f"{root} (auto){ext}"
            doc.save(out)
            print(f"  Updated report saved -> {out}")
            print(f"  (original untouched: {args.report_in})")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
