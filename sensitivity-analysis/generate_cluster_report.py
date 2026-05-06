"""
generate_cluster_report.py
Reads cause-area sensitivity CSVs and produces sensitivity_cluster_report.docx.

The key difference from generate_report.py: every SI in this report measures
movement *between* cause areas (GHD / GCR / AW), not between individual funds.
A test that shuffles money between funds within the same cause area will show
SI = 0 here but nonzero SI in the fund-level report.  Comparing the two reports
lets you distinguish cross-cause shifts from within-cause reshuffling.

Run: python sensitivity-analysis/generate_cluster_report.py
"""

import os
import io
import sys
from datetime import date

try:
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Run: pip install pandas matplotlib python-docx")
    sys.exit(1)

# ── paths ────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))

def p(*parts):
    return os.path.join(BASE, *parts)

# ── constants ─────────────────────────────────────────────────────────────────
LEAST_SENSITIVE    = 1
MODERATE_SENSITIVE = 5
HIGH_SENSITIVE     = 10
EXTREME_SENSITIVE  = 20

CAUSE_AREAS = ["ghd", "gcr", "aw"]

CAUSE_AREA_NAMES = {
    "ghd": "GHD (Global Health & Dev.)",
    "gcr": "GCR (Global Catastrophic Risks)",
    "aw":  "AW (Animal Welfare)",
}

# Funds that belong to each cause area (for computing baseline from fund-level CSV)
CAUSE_AREA_FUNDS = {
    "ghd": ["givewell", "leaf"],
    "gcr": ["longview_ai", "longview_nuclear", "sentinel_bio"],
    "aw":  ["ea_awf", "navigation_fund_cagefree", "navigation_fund_general"],
}

CAUSE_AREA_COLORS = {
    "ghd": "#E87B4C",
    "gcr": "#9B59B6",
    "aw":  "#4C9BE8",
}

DIM_COLORS = {
    "Worldview Credences":     "#1F77B4",
    "CE Multipliers":          "#2CA02C",
    "Dim. Returns (Power)":    "#FF7F0E",
    "Dim. Returns (Max Spend)":"#9467BD",
    "Aggregation Methods":     "#8C564B",
    "Risk Aversion":           "#E377C2",
    "Time Discounts":          "#17BECF",
    "Moral Weights":           "#BCBD22",
    "Single Worldview":        "#AEC7E8",
    "Single Agg Method":       "#C49C94",
}

# ── helpers ───────────────────────────────────────────────────────────────────

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def shade_row(row, hex_color="D9D9D9"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def fmt_si(v):
    return f"{v:.2f}"


def fmt_pp(v):
    if v > 0:
        return f"+{v:.2f} pp"
    return f"{v:.2f} pp"


def ca_label(k):
    return CAUSE_AREA_NAMES.get(k, k)


def add_picture(doc, stream, width=Inches(6)):
    doc.add_picture(stream, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── chart helpers ─────────────────────────────────────────────────────────────

def cause_area_change_chart(tests_data, title, top_n=5):
    """
    Horizontal grouped bar chart of cause-area allocation changes.
    tests_data: list of (label, {cause_area: delta_pp})
    """
    tests_data = tests_data[:top_n]
    n_tests = len(tests_data)
    n_ca = len(CAUSE_AREAS)
    fig_h = max(3, n_tests * 1.1)
    fig, ax = plt.subplots(figsize=(8, fig_h))

    y_base = np.arange(n_tests)
    bar_h = 0.22
    offsets = np.linspace(-(n_ca - 1) / 2, (n_ca - 1) / 2, n_ca) * bar_h

    for ci, ca in enumerate(CAUSE_AREAS):
        vals = [td[1].get(ca, 0) for td in tests_data]
        ax.barh(y_base + offsets[ci], vals, height=bar_h,
                color=CAUSE_AREA_COLORS[ca], label=ca_label(ca))

    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_yticks(y_base)
    ax.set_yticklabels([td[0] for td in tests_data], fontsize=9)
    ax.set_xlabel("Allocation change (percentage points)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left", borderaxespad=0)
    fig.tight_layout()
    return fig_to_stream(fig)


def ca_delta_chart(scenarios_data, title, top_n=15):
    """
    Horizontal grouped bar chart: cause-area allocation change vs baseline for each scenario.
    scenarios_data: list of (label, {ca: delta_pp})
    Sorted descending by Σ|Δ|; limited to top_n if given.
    """
    scenarios_data = sorted(
        scenarios_data, key=lambda x: sum(abs(v) for v in x[1].values()), reverse=True
    )
    if top_n:
        scenarios_data = scenarios_data[:top_n]
    n = len(scenarios_data)
    if n == 0:
        return None
    n_ca = len(CAUSE_AREAS)
    bar_h = 0.22
    fig_h = max(3.0, n * (n_ca * bar_h + 0.15) + 1.2)
    fig, ax = plt.subplots(figsize=(8, fig_h))
    y_base = np.arange(n)
    offsets = np.linspace(-(n_ca - 1) / 2, (n_ca - 1) / 2, n_ca) * bar_h
    for ci, ca in enumerate(CAUSE_AREAS):
        vals = [td[1].get(ca, 0) for td in scenarios_data]
        ax.barh(y_base + offsets[ci], vals, height=bar_h,
                color=CAUSE_AREA_COLORS[ca], label=ca_label(ca))
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_yticks(y_base)
    ax.set_yticklabels([td[0] for td in scenarios_data], fontsize=8)
    ax.set_xlabel("Allocation change vs baseline (pp)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left", borderaxespad=0)
    fig.tight_layout()
    return fig_to_stream(fig)


def tornado_chart(rows_si, title, xlabel="Cause-Area Sensitivity Index (pp)",
                  moderate=MODERATE_SENSITIVE, high=HIGH_SENSITIVE, extreme=EXTREME_SENSITIVE):
    rows_si = sorted(rows_si, key=lambda x: x[1])
    labels = [r[0] for r in rows_si]
    sis    = [r[1] for r in rows_si]

    fig, ax = plt.subplots(figsize=(8, max(3, len(rows_si) * 0.45)))
    colors = ["#8B0000" if s >= extreme else "#D62728" if s >= high else "#FF7F0E" if s >= moderate else "steelblue" for s in sis]
    ax.barh(range(len(labels)), sis, color=colors)
    ax.axvline(extreme,  color="darkred", linestyle="--", linewidth=0.8, label=f"Extreme (SI = {extreme})")
    ax.axvline(high,     color="red",     linestyle="--", linewidth=0.8, label=f"High (SI = {high})")
    ax.axvline(moderate, color="orange",  linestyle="--", linewidth=0.8, label=f"Moderate (SI = {moderate})")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel(xlabel, fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig_to_stream(fig)


def grouped_alloc_chart(scenarios, alloc_df, scenario_col, title):
    """
    Grouped bar chart: one group per cause area, one bar per scenario.
    alloc_df must have columns: scenario_col + ghd + gcr + aw.
    """
    scenario_colors = plt.cm.tab10(np.linspace(0, 0.7, len(scenarios)))
    n_ca   = len(CAUSE_AREAS)
    n_scen = len(scenarios)
    x = np.arange(n_ca)
    bar_w = 0.7 / n_scen

    fig, ax = plt.subplots(figsize=(7, 4))
    for si, scen in enumerate(scenarios):
        row = alloc_df[alloc_df[scenario_col] == scen]
        if row.empty:
            continue
        vals = [float(row[ca].iloc[0]) for ca in CAUSE_AREAS]
        ax.bar(x + si * bar_w - (n_scen - 1) * bar_w / 2, vals,
               width=bar_w, label=scen, color=scenario_colors[si])

    ax.set_xticks(x)
    ax.set_xticklabels([ca_label(ca) for ca in CAUSE_AREAS], fontsize=9)
    ax.set_ylabel("Allocation (%)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left")
    fig.tight_layout()
    return fig_to_stream(fig)


def combo_heatmap(combo_index):
    power_order = [
        "gcr_low_power", "gcr_high_power",
        "aw_low_power", "aw_high_power",
        "gcr_and_aw_low_power", "gcr_and_aw_high_power",
        "gcr_low_aw_high_power", "gcr_high_aw_low_power",
    ]
    spend_order = [2.5, 7.5, 10.0]

    pivot = combo_index.pivot(index="combo", columns="max_spend_multiplier",
                              values="sensitivity_index")
    pivot = pivot.reindex(index=power_order, columns=spend_order)

    fig, ax = plt.subplots(figsize=(6, 4.5))
    im = ax.imshow(pivot.values, cmap="RdYlGn_r", aspect="auto",
                   vmin=combo_index["sensitivity_index"].min(),
                   vmax=combo_index["sensitivity_index"].max())

    ax.set_xticks(range(len(spend_order)))
    ax.set_xticklabels([f"{s}×" for s in spend_order], fontsize=9)
    ax.set_yticks(range(len(power_order)))
    ax.set_yticklabels(power_order, fontsize=8)
    ax.set_xlabel("Max-spend cap multiplier", fontsize=9)
    ax.set_ylabel("DR power scenario", fontsize=9)
    ax.set_title("Cause-Area SI heatmap: DR power × max-spend cap", fontsize=11, fontweight="bold")

    for i in range(len(power_order)):
        for j in range(len(spend_order)):
            val = pivot.values[i, j]
            if not np.isnan(val):
                ax.text(j, i, f"{val:.1f}", ha="center", va="center",
                        fontsize=8, color="black")

    plt.colorbar(im, ax=ax, label="Cause-Area Sensitivity Index (pp)")
    fig.tight_layout()
    return fig_to_stream(fig)


def si_histogram(all_si_records):
    dims = list(DIM_COLORS.keys())
    max_val = max(r[0] for r in all_si_records)
    bins = np.arange(0, max_val + 3, 2)

    fig, ax = plt.subplots(figsize=(9, 4))
    bottom = np.zeros(len(bins) - 1)
    for dim in dims:
        vals = [r[0] for r in all_si_records if r[1] == dim]
        if not vals:
            continue
        counts, _ = np.histogram(vals, bins=bins)
        ax.bar(bins[:-1], counts, width=np.diff(bins), bottom=bottom,
               color=DIM_COLORS[dim], label=dim, align="edge", edgecolor="white")
        bottom += counts

    ax.axvline(LEAST_SENSITIVE,    color="gray",    linestyle="--", linewidth=1.2,
               label=f"Low threshold (SI={LEAST_SENSITIVE})")
    ax.axvline(MODERATE_SENSITIVE, color="orange",  linestyle="--", linewidth=1.2,
               label=f"Moderate threshold (SI={MODERATE_SENSITIVE})")
    ax.axvline(HIGH_SENSITIVE,     color="red",     linestyle="--", linewidth=1.2,
               label=f"High threshold (SI={HIGH_SENSITIVE})")
    ax.axvline(EXTREME_SENSITIVE,  color="darkred", linestyle="--", linewidth=1.2,
               label=f"Extreme threshold (SI={EXTREME_SENSITIVE})")
    ax.set_xlabel("Cause-Area Sensitivity Index (pp)", fontsize=10)
    ax.set_ylabel("Number of tests", fontsize=10)
    ax.set_title("Distribution of cause-area SI values across all tests",
                 fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left")
    fig.tight_layout()
    return fig_to_stream(fig)


def max_si_bar(max_si_per_dim):
    dims = list(max_si_per_dim.keys())
    vals = [max_si_per_dim[d] for d in dims]
    colors = [DIM_COLORS.get(d, "gray") for d in dims]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.barh(dims, vals, color=colors)
    ax.axvline(LEAST_SENSITIVE,    color="gray",    linestyle="--", linewidth=1,
               label=f"SI = {LEAST_SENSITIVE} (low)")
    ax.axvline(MODERATE_SENSITIVE, color="orange",  linestyle="--", linewidth=1,
               label=f"SI = {MODERATE_SENSITIVE} (moderate)")
    ax.axvline(HIGH_SENSITIVE,     color="red",     linestyle="--", linewidth=1,
               label=f"SI = {HIGH_SENSITIVE} (high)")
    ax.axvline(EXTREME_SENSITIVE,  color="darkred", linestyle="--", linewidth=1,
               label=f"SI = {EXTREME_SENSITIVE} (extreme)")
    ax.set_xlabel("Max Cause-Area Sensitivity Index", fontsize=9)
    ax.set_title("Maximum cause-area SI by dimension", fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig_to_stream(fig)


# ── load data ─────────────────────────────────────────────────────────────────

def load():
    def read(path):
        if not os.path.exists(path):
            raise FileNotFoundError(
                f"Missing output file: {path}\n"
                "Run all sensitivity analyses first (node sensitivity-analysis/run_all.js)"
            )
        return pd.read_csv(path)

    data = {}

    data["baseline"] = read(p("outputs", "baseline_staged.csv"))

    data["ra_cause"] = read(
        p("risk-aversion", "outputs", "risk_aversion_cause_area_summary.csv"))

    data["wv_ca_index"] = read(
        p("worldview-sensitivity", "outputs", "cause_area_index.csv"))
    data["wv_ca_alloc"] = read(
        p("worldview-sensitivity", "outputs", "split_credences_cause_areas.csv"))
    data["wv_single_ca"] = read(
        p("worldview-sensitivity", "outputs", "single_worldview_cause_areas.csv"))

    data["ce_ca_si"] = read(
        p("across-the-board", "outputs", "cause_area_si.csv"))
    data["ce_ca_alloc"] = read(
        p("across-the-board", "outputs", "cause_area_allocations.csv"))

    data["dr_ca_index"] = read(
        p("diminishing-returns", "outputs", "dr_sensitivity_cause_area_index.csv"))
    data["dr_ca_alloc"] = read(
        p("diminishing-returns", "outputs", "dr_sensitivity_cause_area_allocations.csv"))
    data["ms_ca_index"] = read(
        p("diminishing-returns", "outputs", "max_spend_cause_area_index.csv"))
    data["ms_ca_alloc"] = read(
        p("diminishing-returns", "outputs", "max_spend_cause_area_allocations.csv"))
    data["combo_ca_index"] = read(
        p("diminishing-returns", "outputs", "combo_max_spend_cause_area_index.csv"))
    data["combo_ca_alloc"] = read(
        p("diminishing-returns", "outputs", "combo_max_spend_cause_area_allocations.csv"))

    data["agg_ca_index"] = read(
        p("aggregation-methods", "outputs", "cause_area_index.csv"))
    data["agg_ca_alloc"] = read(
        p("aggregation-methods", "outputs", "split_credences_cause_areas.csv"))
    data["agg_single_ca"] = read(
        p("aggregation-methods", "outputs", "method_cause_areas.csv"))

    data["disc_ca_si"]   = read(
        p("time-discounts", "outputs", "discount_cause_area_si.csv"))
    data["disc_ca_alloc"]= read(
        p("time-discounts", "outputs", "discount_cause_area_allocations.csv"))
    data["mw_overall_si"]= read(
        p("moral-weights", "outputs", "moral_weights_overall_si.csv"))
    data["mw_ranked"]    = read(
        p("moral-weights", "outputs", "moral_weights_ranked_summary.csv"))

    return data


def baseline_cause_areas(baseline_df):
    """
    Derive baseline cause-area allocation (%) from fund-level baseline CSV.
    Expects a 'fund' column and a 'weighted_combined' column.
    Falls back to the first numeric column if 'weighted_combined' is absent.
    """
    val_col = "weighted_combined"
    if val_col not in baseline_df.columns:
        numeric_cols = baseline_df.select_dtypes(include="number").columns.tolist()
        if not numeric_cols:
            return {ca: 0.0 for ca in CAUSE_AREAS}
        val_col = numeric_cols[0]

    result = {}
    for ca, funds in CAUSE_AREA_FUNDS.items():
        rows = baseline_df[baseline_df["fund"].isin(funds)]
        result[ca] = float(rows[val_col].sum()) if not rows.empty else 0.0
    return result


# ── document builder ──────────────────────────────────────────────────────────

def build_doc(d):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ═══════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Donor Compass Sensitivity Analysis Report")
    run.bold = True
    run.font.size = Pt(20)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Cause-Area Cluster View")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    doc.add_paragraph(
        f"Generated {date.today().strftime('%B %d, %Y')}  ·  "
        "Rethink Priorities"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "Executive Summary", 1)

    doc.add_paragraph(
        "This report measures sensitivity at the cause-area level rather than the "
        "individual-fund level. Funds are grouped into three cause areas:"
    )
    for ca in CAUSE_AREAS:
        funds_str = " + ".join(CAUSE_AREA_FUNDS[ca])
        doc.add_paragraph(
            f"{ca_label(ca)}: {funds_str}",
            style="List Bullet"
        )

    doc.add_paragraph(
        "The cause-area sensitivity index (SI) measures how much the total allocation "
        "to each cause area changes when an assumption is shifted. "
        "SI = Σ|Δ cause area| / 2 (half the total flow between causes) for all analyses. "
        "A test that reshuffles money between funds within the same cause area — "
        "for example shifting dollars from GiveWell to LEAF — will show SI = 0 here "
        "but nonzero SI in the companion fund-level report. Comparing the two reports "
        "reveals whether a sensitivity driver is a cross-cause or within-cause effect."
    )

    add_heading(doc, "Baseline Cause-Area Allocation", 2)
    bl = d["baseline"]
    bl_ca = baseline_cause_areas(bl)
    bl_rows = [(ca_label(ca), f"{bl_ca[ca]:.1f}%") for ca in CAUSE_AREAS]
    add_table(doc,
              ["Cause Area", "Baseline Allocation (%)"],
              bl_rows,
              col_widths=[3.5, 1.5])
    doc.add_paragraph()

    # Split single vs bound rows for wv and agg
    wv_all   = d["wv_ca_index"].copy()
    wv_main  = wv_all[wv_all["bound"] != "single"].copy()

    agg_all  = d["agg_ca_index"].copy()
    agg_main = agg_all[agg_all["bound"] != "single"].copy()

    # Collect SI records for histogram and max-SI bar
    all_si = []

    wv_si_vals  = wv_main["sensitivity_index"].tolist()
    all_si += [(v, "Worldview Credences") for v in wv_si_vals]

    ce_nobase   = d["ce_ca_si"][d["ce_ca_si"]["fund_varied"] != "baseline"]
    ce_si_vals  = ce_nobase["sensitivity_index"].tolist()
    all_si += [(v, "CE Multipliers") for v in ce_si_vals]

    dr_si_vals  = d["dr_ca_index"]["sensitivity_index"].tolist()
    all_si += [(v, "Dim. Returns (Power)") for v in dr_si_vals]

    ms_si_vals  = d["ms_ca_index"]["sensitivity_index"].tolist()
    all_si += [(v, "Dim. Returns (Max Spend)") for v in ms_si_vals]

    agg_si_vals = agg_main["sensitivity_index"].tolist()
    all_si += [(v, "Aggregation Methods") for v in agg_si_vals]

    disc_nonbase = d["disc_ca_si"][d["disc_ca_si"]["scenario_group"] != "baseline"]
    disc_ca_si_vals = disc_nonbase[disc_nonbase["sensitivity_index"] > 0]["sensitivity_index"].tolist()
    all_si += [(v, "Time Discounts") for v in disc_ca_si_vals]

    mw_nonbase_ca = d["mw_overall_si"][d["mw_overall_si"]["multiplier"] != 1.0]
    mw_ca_si_vals = mw_nonbase_ca["ca_sensitivity_index"].tolist()
    all_si += [(v, "Moral Weights") for v in mw_ca_si_vals]

    max_si_per_dim = {
        "Worldview Credences":      max(wv_si_vals)      if wv_si_vals      else 0,
        "CE Multipliers":           max(ce_si_vals)      if ce_si_vals      else 0,
        "Dim. Returns (Power)":     max(dr_si_vals)      if dr_si_vals      else 0,
        "Dim. Returns (Max Spend)": max(ms_si_vals)      if ms_si_vals      else 0,
        "Aggregation Methods":      max(agg_si_vals)     if agg_si_vals     else 0,
        "Time Discounts":           max(disc_ca_si_vals) if disc_ca_si_vals else 0,
        "Moral Weights":            max(mw_ca_si_vals)   if mw_ca_si_vals   else 0,
    }

    add_heading(doc, "Maximum Cause-Area SI by Dimension", 2)
    add_picture(doc, max_si_bar(max_si_per_dim), width=Inches(5.5))
    doc.add_paragraph()

    add_heading(doc, "Distribution of Cause-Area SI Values — Main Analyses Only", 2)
    if all_si and max(r[0] for r in all_si) > 0:
        add_picture(doc, si_histogram(all_si), width=Inches(6.5))
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: risk-aversion scenario tests and analyses placing 100% credence in a single "
        "worldview or aggregation method are presented separately in Section 5 as illustrative "
        "scenarios; they are excluded from the chart above."
    )
    note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Distribution of Cause-Area SI Values — Including Illustrative Tests", 2)
    all_si_full = list(all_si)
    for _, r in d["ra_cause"].iterrows():
        all_si_full.append((r["sensitivity_index"], "Risk Aversion"))
    for _, r in d["wv_single_ca"].iterrows():
        si_val = sum(abs(float(r[ca]) - bl_ca[ca]) for ca in CAUSE_AREAS) / 2
        all_si_full.append((si_val, "Single Worldview"))
    for _, r in d["agg_single_ca"].iterrows():
        si_val = sum(abs(float(r[ca]) - bl_ca[ca]) for ca in CAUSE_AREAS) / 2
        all_si_full.append((si_val, "Single Agg Method"))
    if all_si_full and max(r[0] for r in all_si_full) > 0:
        add_picture(doc, si_histogram(all_si_full), width=Inches(6.5))
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 1 — WORLDVIEW CREDENCES
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "1. Worldview Credences", 1)

    doc.add_paragraph(
        "Each worldview's credence is shifted to its low-end or high-end uncertainty "
        "bound one at a time, with others renormalised. SI here measures the total flow "
        "between cause areas (GHD ↔ GCR ↔ AW). A test that moves money from GiveWell "
        "to LEAF (both GHD) will show SI = 0 in this report."
    )

    wv_sorted = wv_main.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    add_heading(doc, "Cause-Area Sensitivity — Ranked by SI (top 20)", 2)
    tbl_rows = []
    for _, r in wv_sorted.head(20).iterrows():
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        scaled = fmt_si(r["scaled_SI"]) if "scaled_SI" in r.index and pd.notna(r["scaled_SI"]) else "—"
        tbl_rows.append([
            r["worldview"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            scaled,
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Worldview", "Bound", "Credence Range", "SI",
               "Scaled SI*", "Most Affected Cause Area", "Max Δ"],
              tbl_rows,
              col_widths=[2.0, 0.5, 1.0, 0.55, 0.65, 1.55, 0.75])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ absolute credence change (pp)."
    )
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI (top 20)", 2)
    wv_scaled = wv_sorted[wv_sorted["scaled_SI"].notna()].copy()
    wv_scaled = wv_scaled.sort_values("scaled_SI", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(wv_scaled.head(20).iterrows()):
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        tbl_rows2.append([
            str(i + 1),
            r["worldview"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["scaled_SI"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Rank", "Worldview", "Bound", "Credence Range", "SI",
               "Scaled SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows2,
              col_widths=[0.4, 1.75, 0.5, 1.0, 0.55, 0.65, 1.35, 0.75])
    doc.add_paragraph()

    # Allocation-change chart for tests with SI ≥ moderate threshold
    top_wv = wv_sorted[wv_sorted["sensitivity_index"] >= MODERATE_SENSITIVE].head(8)
    if not top_wv.empty:
        add_heading(doc, f"Cause-Area Allocation Changes — Tests with SI ≥ {MODERATE_SENSITIVE} (moderate+)", 2)
        wv_ca_alloc = d["wv_ca_alloc"].copy()
        tests_data = []
        for _, r in top_wv.iterrows():
            scen = r["scenario"]
            row = wv_ca_alloc[wv_ca_alloc["scenario"] == scen]
            if row.empty:
                continue
            deltas = {ca: float(row[ca].iloc[0]) - bl_ca[ca] for ca in CAUSE_AREAS}
            label = f"{r['worldview'][:35]}… ({r['bound']})"
            tests_data.append((label, deltas))
        if tests_data:
            add_picture(doc, cause_area_change_chart(
                tests_data, "Worldview credences: cause-area allocation change (pp)"))

    add_heading(doc, "Cause-Area Allocation Changes — All Tests (Top 15 by Σ|Δ|)", 2)
    all_wv_deltas = [
        (f"{r['worldview'][:40]} ({r['bound']})",
         {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in d["wv_ca_alloc"].iterrows()
    ]
    stream = ca_delta_chart(all_wv_deltas,
                            "Worldview credences: cause-area change vs baseline (pp)", top_n=15)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 2 — COST-EFFECTIVENESS MULTIPLIERS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "2. Cost-Effectiveness Multipliers", 1)

    doc.add_paragraph(
        "Each fund's CE was multiplied by a constant factor, one fund or group at a time. "
        "SI = Σ|Δ cause area| / 2 (the same convention as all other analyses). "
        "Multiplying the CE of a fund affects its cause area's total allocation; funds in "
        "other cause areas gain or lose what that cause area gains or loses."
    )

    ce = d["ce_ca_si"].copy()
    ce_nobase = ce[ce["fund_varied"] != "baseline"].copy()
    ce_sorted = ce_nobase.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    add_heading(doc, "Cause-Area Sensitivity — Ranked by SI (all tests)", 2)
    tbl_rows = []
    for _, r in ce_sorted.iterrows():
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(r["si_scaled_pp_per_oom"]) if pd.notna(r["si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            r["fund_varied"],
            f"{r['multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            scaled,
            f"{ca_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{ca_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Fund Varied", "Multiplier", "SI", "Scaled SI*", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[1.2, 0.7, 0.55, 0.65, 2.1, 2.1])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier."
    )
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI", 2)
    ce_scaled = ce_sorted[ce_sorted["si_scaled_pp_per_oom"].notna()].copy()
    ce_scaled = ce_scaled.sort_values("si_scaled_pp_per_oom", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(ce_scaled.iterrows()):
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        tbl_rows2.append([
            str(i + 1),
            r["fund_varied"],
            f"{r['multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["si_scaled_pp_per_oom"]),
            f"{ca_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{ca_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Rank", "Fund Varied", "Multiplier", "SI", "Scaled SI",
               "Biggest Gainer", "Biggest Loser"],
              tbl_rows2,
              col_widths=[0.4, 1.15, 0.65, 0.55, 0.65, 2.0, 2.0])
    doc.add_paragraph()

    add_heading(doc, "Tornado Chart — All CE Multiplier Tests", 2)
    tornado_rows = []
    for _, r in ce_sorted.iterrows():
        label = f"{r['fund_varied']} ×{r['multiplier']}"
        tornado_rows.append((label, r["sensitivity_index"], "", 0, "", 0))
    add_picture(doc, tornado_chart(
        tornado_rows,
        "CE Multipliers: Cause-Area Sensitivity Index per test",
        xlabel="Cause-Area SI = Σ|Δ cause area| / 2 (pp)"),
        width=Inches(6.5))

    add_heading(doc, "Cause-Area Allocation Changes — All CE Tests (Top 15 by Σ|Δ|)", 2)
    ce_alloc_nobase = d["ce_ca_alloc"][d["ce_ca_alloc"]["fund_varied"] != "baseline"]
    ce_deltas = [
        (f"{r['fund_varied']} ×{r['multiplier']}",
         {ca: float(r[f"diff_{ca}"]) for ca in CAUSE_AREAS})
        for _, r in ce_alloc_nobase.iterrows()
    ]
    stream = ca_delta_chart(ce_deltas,
                            "CE multipliers: cause-area allocation change (pp)", top_n=15)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 3 — DIMINISHING RETURNS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "3. Diminishing Returns", 1)

    # 3a — Power Parameters
    add_heading(doc, "3a. Power Parameters", 2)
    doc.add_paragraph(
        "Eight combinations of DR power parameters across GCR and AW clusters. "
        "Because both clusters are separate cause areas, changes in their DR curves "
        "may shift money between GCR and AW (cross-cause) or between GCR/AW and GHD."
    )

    dr = d["dr_ca_index"].copy().sort_values("sensitivity_index", ascending=False)
    tbl_rows = [(r["combo"],
                 fmt_si(r["sensitivity_index"]),
                 ca_label(r["most_affected_cause"]),
                 fmt_pp(r["most_affected_delta"]))
                for _, r in dr.iterrows()]
    add_table(doc,
              ["Scenario", "Cause-Area SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows, col_widths=[2.5, 0.85, 2.2, 0.9])

    dr_alloc = d["dr_ca_alloc"].copy()
    scenarios_dr = ["baseline"] + dr["combo"].tolist()
    add_picture(doc, grouped_alloc_chart(
        scenarios_dr, dr_alloc, "combo",
        "DR Power: Cause-Area Allocations by Scenario"),
        width=Inches(6.5))

    add_heading(doc, "Cause-Area Allocation Changes by DR Power Scenario", 2)
    dr_base_row = dr_alloc[dr_alloc["combo"] == "baseline"]
    dr_base_ca = {ca: float(dr_base_row[ca].iloc[0]) for ca in CAUSE_AREAS} if not dr_base_row.empty else bl_ca
    dr_deltas = [
        (r["combo"], {ca: float(r[ca]) - dr_base_ca[ca] for ca in CAUSE_AREAS})
        for _, r in dr_alloc.iterrows() if r["combo"] != "baseline"
    ]
    stream = ca_delta_chart(dr_deltas,
                            "DR power: cause-area allocation change vs baseline (pp)", top_n=None)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # 3b — Max Spending Cap
    add_heading(doc, "3b. Maximum Spending Cap", 2)
    doc.add_paragraph(
        "Three cap multipliers (2.5×, 7.5×, 10×) were tested. The cap limits total "
        "spending on any fund; tighter caps redistribute within and between cause areas."
    )

    ms = d["ms_ca_index"].copy().sort_values("sensitivity_index", ascending=False)
    tbl_rows = []
    for _, r in ms.iterrows():
        tbl_rows.append([
            r["scenario"],
            f"{r['max_addl_spend_multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Scenario", "Cap Multiplier", "Cause-Area SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows, col_widths=[1.4, 0.85, 0.85, 2.1, 0.9])

    ms_alloc = d["ms_ca_alloc"].copy()
    baseline_ms_row = dr_alloc[dr_alloc["combo"] == "baseline"].copy()
    if not baseline_ms_row.empty:
        baseline_ms_row = baseline_ms_row.rename(columns={"combo": "scenario"})
        baseline_ms_row["scenario"] = "baseline"
        ms_alloc_full = pd.concat(
            [baseline_ms_row[["scenario"] + CAUSE_AREAS], ms_alloc[["scenario"] + CAUSE_AREAS]],
            ignore_index=True
        )
    else:
        ms_alloc_full = ms_alloc
    scenarios_ms = ["baseline", "max_spend_2_5x", "max_spend_7_5x", "max_spend_10x"]
    add_picture(doc, grouped_alloc_chart(
        scenarios_ms, ms_alloc_full, "scenario",
        "Max Spend Cap: Cause-Area Allocations by Scenario"),
        width=Inches(6.5))

    add_heading(doc, "Cause-Area Allocation Changes by Max Spend Cap", 2)
    ms_nobase = d["ms_ca_alloc"][~d["ms_ca_alloc"]["scenario"].str.startswith("baseline")]
    ms_deltas = [
        (r["scenario"], {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in ms_nobase.iterrows()
    ]
    stream = ca_delta_chart(ms_deltas,
                            "Max spend cap: cause-area allocation change vs baseline (pp)", top_n=None)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # 3c — Combined
    add_heading(doc, "3c. Combined: DR Power × Maximum Spending Cap", 2)
    doc.add_paragraph(
        "24 joint tests crossing 8 DR power scenarios with 3 cap multipliers."
    )

    combo_idx = d["combo_ca_index"].copy()

    add_heading(doc, "Cause-Area SI Heatmap: Power × Cap", 3)
    add_picture(doc, combo_heatmap(combo_idx), width=Inches(5.5))
    doc.add_paragraph()

    add_heading(doc, "All 24 Tests — Ranked by Cause-Area SI", 3)
    combo_sorted = combo_idx.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in combo_sorted.iterrows():
        tbl_rows.append([
            r["combo"],
            f"{r['max_spend_multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Power Scenario", "Cap Multiplier", "Cause-Area SI",
               "Most Affected Cause Area", "Max Δ"],
              tbl_rows, col_widths=[2.2, 0.9, 0.85, 2.0, 0.9])

    add_heading(doc, "Cause-Area Allocation Changes — Top 10 DR Combo Scenarios", 3)
    combo_nobase = d["combo_ca_alloc"][~d["combo_ca_alloc"]["scenario"].str.startswith("baseline")]
    combo_deltas = [
        (r["scenario"], {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in combo_nobase.iterrows()
    ]
    stream = ca_delta_chart(combo_deltas,
                            "DR power × max spend: cause-area allocation change (pp)", top_n=10)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 4 — AGGREGATION METHOD WEIGHTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "4. Aggregation Method Weights", 1)

    doc.add_paragraph(
        "Each aggregation method's credence is shifted to its low or high uncertainty "
        "bound, one at a time. SI measures cross-cause-area reallocation. Methods that "
        "primarily reshuffle within a cause area will appear low-sensitivity here."
    )

    agg_sorted = agg_main.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in agg_sorted.iterrows():
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        scaled = fmt_si(r["scaled_SI"]) if "scaled_SI" in r.index and pd.notna(r["scaled_SI"]) else "—"
        tbl_rows.append([
            r["scenario"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            scaled,
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Test", "Bound", "Credence Range", "SI",
               "Scaled SI*", "Most Affected Cause Area", "Max Δ"],
              tbl_rows, col_widths=[1.8, 0.55, 1.05, 0.55, 0.65, 1.65, 0.85])
    p_note = doc.add_paragraph("* Scaled SI = SI ÷ absolute credence change (pp).")
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI", 2)
    agg_scaled = agg_sorted[agg_sorted["scaled_SI"].notna()].copy()
    agg_scaled = agg_scaled.sort_values("scaled_SI", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(agg_scaled.iterrows()):
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        tbl_rows2.append([
            str(i + 1),
            r["scenario"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["scaled_SI"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Rank", "Test", "Bound", "Credence Range", "SI",
               "Scaled SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows2, col_widths=[0.4, 1.6, 0.55, 1.05, 0.55, 0.65, 1.6, 0.8])

    add_heading(doc, "Cause-Area Allocation Changes — All Aggregation Method Tests (Top 15)", 2)
    agg_deltas = [
        (f"{r['scenario']} ({r['bound']})", {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in d["agg_ca_alloc"].iterrows()
    ]
    stream = ca_delta_chart(agg_deltas,
                            "Aggregation method weight shifts: cause-area allocation change (pp)",
                            top_n=15)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # LEAST SENSITIVE TESTS (main analyses only)
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "Least Sensitive Tests at Cause-Area Level (SI < 1)", 1)

    doc.add_paragraph(
        "The following tests produced less than 1 pp of cross-cause reallocation. "
        "Note that some of these may appear in the fund-level report as sensitive "
        "(they move money within a cause area but not between cause areas)."
    )

    low_tests = []

    for _, r in wv_main.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Worldview Credences",
                               r["scenario"],
                               r["sensitivity_index"],
                               f"{r['worldview'][:50]} ({r['bound']} bound: "
                               f"{r['credence_base']:.0%}→{r['credence_scenario']:.0%})"))

    ce_low = ce_sorted[ce_sorted["sensitivity_index"] < LEAST_SENSITIVE]
    for _, r in ce_low.iterrows():
        low_tests.append(("CE Multipliers",
                           f"{r['fund_varied']} ×{r['multiplier']}",
                           r["sensitivity_index"],
                           f"Multiply {r['fund_varied']} CE by {r['multiplier']}×"))

    for _, r in d["dr_ca_index"].iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Dim. Returns (Power)", r["combo"],
                               r["sensitivity_index"],
                               f"DR power combo: {r['combo']}"))

    for _, r in d["ms_ca_index"].iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Dim. Returns (Max Spend)", r["scenario"],
                               r["sensitivity_index"],
                               f"Max spend cap ×{r['max_addl_spend_multiplier']}"))

    for _, r in agg_main.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Aggregation Methods", r["scenario"],
                               r["sensitivity_index"],
                               f"Method weight shift: {r['scenario']} "
                               f"({r['credence_base']:.0%}→{r['credence_scenario']:.0%})"))

    disc_all_ca = d["disc_ca_si"][d["disc_ca_si"]["scenario_group"] != "baseline"].copy()
    for _, r in disc_all_ca.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Time Discounts",
                               f"{r['scenario_group']} ×{r['multiplier']}",
                               r["sensitivity_index"],
                               f"Scale discount factor(s) by {r['multiplier']}× "
                               f"in group '{r['scenario_group']}'"))

    mw_all_ca = d["mw_overall_si"][d["mw_overall_si"]["multiplier"] != 1.0].copy()
    for _, r in mw_all_ca.iterrows():
        if r["ca_sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Moral Weights",
                               f"All weights ×{r['multiplier']}",
                               r["ca_sensitivity_index"],
                               f"Scale all animal moral weights by {r['multiplier']}×"))

    low_tests.sort(key=lambda x: (x[0], x[2]))

    current_dim = None
    for dim, _name, si, desc in low_tests:
        if dim != current_dim:
            current_dim = dim
            p_obj = doc.add_paragraph()
            run = p_obj.add_run(dim)
            run.bold = True
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.add_run(f"SI = {si:.2f}  ·  ").bold = False
        bullet.add_run(desc)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 5 — TIME DISCOUNTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "5. Time Discounts", 1)

    doc.add_paragraph(
        "This analysis tests how the allocation shifts when far-future timeframe weights "
        "are reduced. 'Discount 500+' modifies the 500+ year weight only; "
        "'Discount 100–500 and 500+' modifies both the 100–500 and 500+ year weights. "
        "Multipliers tested: 0.1×, 0.01×, 0.001×, and 0× (complete elimination)."
    )

    disc_ca = d["disc_ca_si"].copy()
    disc_ca_nonbase = disc_ca[disc_ca["scenario_group"] != "baseline"].copy()

    add_heading(doc, "Cause-Area Sensitivity — All Scenarios", 2)
    tbl_rows = []
    for _, r in disc_ca_nonbase.iterrows():
        diffs = {ca: r[f"diff_{ca}"] for ca in CAUSE_AREAS}
        if r["sensitivity_index"] > 0:
            gainer = max(diffs, key=lambda k: diffs[k])
            loser  = min(diffs, key=lambda k: diffs[k])
            g_str = f"{ca_label(gainer)} ({fmt_pp(diffs[gainer])})"
            l_str = f"{ca_label(loser)} ({fmt_pp(diffs[loser])})"
        else:
            g_str = l_str = "—"
        mult_str = "0×" if r["multiplier"] == 0 else f"{r['multiplier']}×"
        tbl_rows.append([
            r["scenario_group"],
            mult_str,
            fmt_si(r["sensitivity_index"]),
            g_str,
            l_str,
        ])
    add_table(doc,
              ["Scenario Group", "Multiplier", "Cause-Area SI", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[1.9, 0.65, 0.9, 2.0, 2.0])
    doc.add_paragraph()

    disc_ca_zero = disc_ca_nonbase[disc_ca_nonbase["multiplier"] == 0].copy()
    if not disc_ca_zero.empty:
        add_heading(doc, "Cause-Area Allocation Changes — 0× Scenarios", 2)
        tests_data = []
        for _, r in disc_ca_zero.iterrows():
            deltas = {ca: r[f"diff_{ca}"] for ca in CAUSE_AREAS}
            tests_data.append((r["scenario_group"], deltas))
        add_picture(doc, cause_area_change_chart(
            tests_data,
            "Time discounts: cause-area allocation change when discount factor eliminated (pp)",
            top_n=len(tests_data)))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  Partial reductions (0.1× through 0.001×) produce SI = 0 pp at the "
        "cause-area level — the allocation is completely insensitive. Complete elimination "
        "(0×) of the 500+ year weight shifts 15.38 pp out of GCR (−15.38 pp) into AW "
        "(+12.39 pp) and GHD (+2.99 pp). Eliminating both the 100–500 and 500+ year "
        "weights (SI = 20.87 pp) amplifies the GCR loss to −20.87 pp, with AW gaining "
        "+14.90 pp and GHD +5.97 pp. This is a cross-cause effect: far-future weighting "
        "directly benefits GCR (which has long-horizon expected value) at the expense of "
        "near-term AW and GHD."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 6 — ANIMAL MORAL WEIGHTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "6. Animal Moral Weights", 1)

    doc.add_paragraph(
        "Animal moral weights are scaled across two levels: (1) the overall blend — "
        "all worldviews' weights scaled simultaneously; and (2) per-worldview — each "
        "worldview run at 100% credence. Multipliers: 0.1×, 2×, 5×, 10×, 100×. "
        "Weights are capped at species-level upper bounds."
    )

    # 6a. Overall blend
    add_heading(doc, "6a. Overall Blend Sensitivity", 2)
    mw_osi = d["mw_overall_si"].copy()
    mw_osi_nonbase = mw_osi[mw_osi["multiplier"] != 1.0].copy()

    tbl_rows = []
    for _, r in mw_osi_nonbase.sort_values("ca_sensitivity_index", ascending=False).iterrows():
        diffs = {ca: r[f"diff_{ca}"] for ca in CAUSE_AREAS}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(r["ca_si_scaled_pp_per_oom"]) if pd.notna(r["ca_si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            f"{r['multiplier']}×",
            fmt_si(r["ca_sensitivity_index"]),
            scaled,
            f"{ca_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{ca_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Multiplier", "Cause-Area SI", "Scaled SI*", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[0.65, 0.9, 0.65, 2.2, 2.2])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Cause-Area Allocation Changes — All Multipliers", 2)
    tests_data = []
    for _, r in mw_osi_nonbase.sort_values("ca_sensitivity_index", ascending=False).iterrows():
        deltas = {ca: r[f"diff_{ca}"] for ca in CAUSE_AREAS}
        tests_data.append((f"{r['multiplier']}×", deltas))
    stream = ca_delta_chart(tests_data,
                            "Moral weights (overall blend): cause-area allocation change (pp)",
                            top_n=None)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # 6b. Per-worldview ranked summary
    add_heading(doc, "6b. Per-Worldview Sensitivity — Ranked Summary (Top 15)", 2)
    doc.add_paragraph(
        "Each worldview run in isolation with its animal weights scaled. "
        "Cause-area SI measured against that worldview's own unmodified baseline. "
        "Only non-zero scenarios shown."
    )

    mw_ranked_df = d["mw_ranked"].copy()

    tbl_rows = []
    for _, r in mw_ranked_df.head(15).iterrows():
        diffs = {ca: float(r[f"diff_{ca}"]) for ca in CAUSE_AREAS}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(float(r["ca_si_scaled_pp_per_oom"])) if pd.notna(r["ca_si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            str(int(r["worldview_idx"])),
            r["worldview_name"][:28],
            str(int(r["risk_profile"])),
            f"{r['multiplier']}×",
            fmt_si(float(r["ca_sensitivity_index"])),
            scaled,
            f"{ca_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{ca_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Idx", "Worldview", "Risk", "Mult.", "CA SI", "Scaled SI*",
               "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[0.35, 1.65, 0.45, 0.5, 0.6, 0.65, 1.85, 1.85])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Per-Worldview Cause-Area SI — Tornado Chart (Top 20 Scenarios)", 2)
    tornado_rows_mw = []
    for _, r in mw_ranked_df.head(20).iterrows():
        label = (f"{r['worldview_name'][:28]} "
                 f"(idx={int(r['worldview_idx'])}, {r['multiplier']}×)")
        tornado_rows_mw.append((label, float(r["ca_sensitivity_index"]), "", 0, "", 0))
    add_picture(doc, tornado_chart(
        tornado_rows_mw,
        "Animal moral weights (per worldview): cause-area SI vs own baseline"),
        width=Inches(6.5))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings (overall blend):  All moral weight changes — up or down — shift "
        "money between AW and GHD/GCR. Reducing weights (0.1×) causes the largest "
        "cross-cause shift (SI = 24.18 pp), with AW losing 24 pp primarily to GHD. "
        "Increasing weights drives the reverse: AW gains at the expense of GHD (100× "
        "gives SI = 14 pp at cause-area level). GCR is barely affected by moral weight "
        "changes at the overall blend level, as the GCR–AW tradeoff is driven "
        "predominantly by worldviews where one or the other is already dominant. "
        "Key findings (per worldview):  The most sensitive worldviews for cross-cause "
        "shifts are Contractualism idx 11 (CA SI = 75 pp at 100×, full GHD→AW reversal) "
        "and Non-Utilitarian Consequentialism idx 7 and 8 (CA SI ≈ 43–47 pp at 10×, "
        "GCR→AW shift). TU worldviews are sensitive downward only and shift GHD↔AW, "
        "while the Non-Util Consequentialist worldviews primarily shift GCR↔AW."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 7 — ILLUSTRATIVE SINGLE-PERSPECTIVE ANALYSES
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "7. Illustrative: Single-Perspective Analyses", 1)

    doc.add_paragraph(
        "These analyses place 100% credence in one perspective. Because credence shifts "
        "are extreme, cause-area SIs can be large. Sensitivity tiers: "
        f"moderate ≥ {MODERATE_SENSITIVE} pp, high ≥ {HIGH_SENSITIVE} pp, "
        f"extreme ≥ {EXTREME_SENSITIVE} pp. "
        "Large cause-area SIs here mean the perspective concentrates the portfolio in one "
        "cause area rather than just reshuffling within it."
    )

    # ── 5a. Risk Aversion ──────────────────────────────────────────────────────
    add_heading(doc, "5a. Risk Aversion", 2)

    ra = d["ra_cause"].copy()
    ra_sorted = ra.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    tbl_rows = []
    for _, r in ra_sorted.iterrows():
        tbl_rows.append([
            r["test"],
            fmt_si(r["sensitivity_index"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Test", "Cause-Area SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows,
              col_widths=[2.8, 0.85, 2.0, 0.9])
    doc.add_paragraph()

    top_ra = ra_sorted[ra_sorted["sensitivity_index"] >= EXTREME_SENSITIVE].head(5)
    tests_data = []
    for _, r in top_ra.iterrows():
        deltas = {ca: r[f"{ca}_delta"] for ca in CAUSE_AREAS if f"{ca}_delta" in r.index}
        tests_data.append((r["test"], deltas))
    if tests_data:
        add_picture(doc, cause_area_change_chart(
            tests_data,
            f"Risk Aversion: cause-area allocation change (pp) — SI ≥ {EXTREME_SENSITIVE} (extreme)",
            top_n=5))

    add_heading(doc, "Cause-Area Allocation Changes — All Risk Aversion Tests", 2)
    ra_deltas = [
        (r["test"], {ca: float(r[f"{ca}_delta"]) for ca in CAUSE_AREAS})
        for _, r in ra_sorted.iterrows()
    ]
    stream = ca_delta_chart(ra_deltas,
                            "Risk aversion: cause-area allocation change (baseline → new version, pp)",
                            top_n=15)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ── 5b. 100% Credence in One Worldview ────────────────────────────────────
    add_heading(doc, "5b. 100% Credence in One Worldview", 2)

    doc.add_paragraph(
        "Each worldview run in isolation. Cause-area SI is computed against the "
        "baseline special-blend allocation. A worldview that directs 100% of the "
        "budget to a single cause area will show SI ≈ 50 pp here (the maximum "
        "when one cause area gains 50 pp and the others lose 50 pp combined)."
    )

    wv_single = d["wv_single_ca"].copy()
    # Compute SI against baseline cause-area allocation
    wv_single_si = []
    for _, r in wv_single.iterrows():
        si_val = sum(abs(float(r[ca]) - bl_ca[ca]) for ca in CAUSE_AREAS) / 2
        most_aff = max(CAUSE_AREAS, key=lambda ca: abs(float(r[ca]) - bl_ca[ca]))
        delta = float(r[most_aff]) - bl_ca[most_aff]
        wv_single_si.append({
            "worldview": r["worldview"],
            "sensitivity_index": si_val,
            "most_affected_cause": most_aff,
            "most_affected_delta": delta,
        })
    wv_single_df = pd.DataFrame(wv_single_si).sort_values(
        "sensitivity_index", ascending=False).reset_index(drop=True)

    tbl_rows = []
    for _, r in wv_single_df.iterrows():
        tbl_rows.append([
            r["worldview"][:60],
            fmt_si(r["sensitivity_index"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Worldview", "Cause-Area SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows,
              col_widths=[3.0, 0.85, 2.0, 0.9])
    doc.add_paragraph()

    tornado_rows_wv = [
        (r["worldview"][:50], r["sensitivity_index"], "", 0, "", 0)
        for _, r in wv_single_df.iterrows()
    ]
    add_picture(doc, tornado_chart(
        tornado_rows_wv,
        "100% Credence in One Worldview: Cause-Area SI vs baseline"),
        width=Inches(6.5))

    add_heading(doc, "Cause-Area Allocation Changes — All Single Worldviews vs Baseline", 2)
    wv_single_deltas = [
        (r["worldview"][:50], {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in wv_single.iterrows()
    ]
    stream = ca_delta_chart(
        wv_single_deltas,
        "100% credence in one worldview: cause-area change vs special-blend baseline (pp)",
        top_n=None)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    # ── 5c. 100% Credence in One Aggregation Method ───────────────────────────
    add_heading(doc, "5c. 100% Credence in One Aggregation Method", 2)

    doc.add_paragraph(
        "Each method run on the full budget with all worldviews. Cause-area SI is "
        "computed against the baseline allocation."
    )

    agg_single = d["agg_single_ca"].copy()
    agg_single_si = []
    for _, r in agg_single.iterrows():
        si_val = sum(abs(float(r[ca]) - bl_ca[ca]) for ca in CAUSE_AREAS) / 2
        most_aff = max(CAUSE_AREAS, key=lambda ca: abs(float(r[ca]) - bl_ca[ca]))
        delta = float(r[most_aff]) - bl_ca[most_aff]
        agg_single_si.append({
            "method": r["method"],
            "sensitivity_index": si_val,
            "most_affected_cause": most_aff,
            "most_affected_delta": delta,
        })
    agg_single_df = pd.DataFrame(agg_single_si).sort_values(
        "sensitivity_index", ascending=False).reset_index(drop=True)

    tbl_rows = []
    for _, r in agg_single_df.iterrows():
        tbl_rows.append([
            r["method"],
            fmt_si(r["sensitivity_index"]),
            ca_label(r["most_affected_cause"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Method", "Cause-Area SI", "Most Affected Cause Area", "Max Δ"],
              tbl_rows,
              col_widths=[2.0, 0.85, 2.2, 0.9])
    doc.add_paragraph()

    tornado_rows_agg = [
        (r["method"], r["sensitivity_index"], "", 0, "", 0)
        for _, r in agg_single_df.iterrows()
    ]
    add_picture(doc, tornado_chart(
        tornado_rows_agg,
        "100% Credence in One Aggregation Method: Cause-Area SI vs baseline"),
        width=Inches(6.5))

    add_heading(doc, "Cause-Area Allocation Changes — All Single Methods vs Baseline", 2)
    agg_single_deltas = [
        (r["method"], {ca: float(r[ca]) - bl_ca[ca] for ca in CAUSE_AREAS})
        for _, r in agg_single.iterrows()
    ]
    stream = ca_delta_chart(
        agg_single_deltas,
        "100% credence in one aggregation method: cause-area change vs baseline (pp)",
        top_n=None)
    if stream:
        add_picture(doc, stream)
    doc.add_paragraph()

    return doc


# ── main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Loading data…")
    d = load()
    print("Building document…")
    doc = build_doc(d)
    out = p("sensitivity_cluster_report.docx")
    doc.save(out)
    print(f"Saved: {out}")
