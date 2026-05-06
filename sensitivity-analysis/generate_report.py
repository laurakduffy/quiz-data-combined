"""
generate_report.py
Reads sensitivity analysis CSVs and produces sensitivity_report.docx.
Run: python sensitivity-analysis/generate_report.py
"""

import os
import io
import sys
from datetime import date

try:
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Run: pip install pandas matplotlib python-docx")
    sys.exit(1)

# ── paths ────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))

def p(*parts):
    return os.path.join(BASE, *parts)

# ── constants ─────────────────────────────────────────────────────────────────
HIGHLY_SENSITIVE = 10
LEAST_SENSITIVE  = 2

# Thresholds for Section 5 illustrative single-perspective analyses
SINGLE_HIGHLY_SENSITIVE = 50
SINGLE_LEAST_SENSITIVE  = 15

FUND_NAMES = {
    "ea_awf":                    "EA Animal Welfare Fund",
    "givewell":                  "GiveWell",
    "leaf":                      "LEAF",
    "longview_ai":               "Longview AI Safety",
    "longview_nuclear":          "Longview Nuclear Security",
    "navigation_fund_cagefree":  "Navigation Fund (Cage-Free)",
    "navigation_fund_general":   "Navigation Fund (General)",
    "sentinel_bio":              "Sentinel Biosecurity",
}
FUNDS = list(FUND_NAMES.keys())

# One distinct color per fund (consistent across all charts)
FUND_COLORS = {
    "ea_awf":                   "#4C9BE8",
    "givewell":                 "#E87B4C",
    "leaf":                     "#5DB85D",
    "longview_ai":              "#9B59B6",
    "longview_nuclear":         "#E84C4C",
    "navigation_fund_cagefree": "#E8C84C",
    "navigation_fund_general":  "#4CCBE8",
    "sentinel_bio":             "#E84CA0",
}

# One color per dimension (for histogram — illustrative analyses excluded)
DIM_COLORS = {
    "Worldview Credences":    "#1F77B4",
    "CE Multipliers":         "#2CA02C",
    "Dim. Returns (Power)":   "#FF7F0E",
    "Dim. Returns (Max Spend)":"#9467BD",
    "Aggregation Methods":    "#8C564B",
    "Risk Aversion":          "#E377C2",
    "Time Discounts":         "#17BECF",
    "Moral Weights":          "#BCBD22",
    "Single Worldview":       "#AEC7E8",
    "Single Agg Method":      "#C49C94",
}

# ── helpers ───────────────────────────────────────────────────────────────────

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def shade_row(row, hex_color="D9D9D9"):
    """Apply gray shading to a table row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table. rows = list of lists."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def fmt_si(v):
    return f"{v:.2f}"


def fmt_pp(v):
    if v > 0:
        return f"+{v:.2f} pp"
    return f"{v:.2f} pp"


def fund_label(k):
    return FUND_NAMES.get(k, k)


def add_picture(doc, stream, width=Inches(6)):
    doc.add_picture(stream, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── chart helpers ─────────────────────────────────────────────────────────────

def allocation_change_chart(tests_data, title, top_n=5):
    """
    Horizontal grouped bar chart of fund allocation changes.
    tests_data: list of (label, {fund: delta_pp})
    top_n: max tests to show
    """
    tests_data = tests_data[:top_n]
    n_tests = len(tests_data)
    n_funds = len(FUNDS)
    fig_h = max(4, n_tests * 1.4)
    fig, ax = plt.subplots(figsize=(8, fig_h))

    y_base = np.arange(n_tests)
    bar_h = 0.09
    offsets = np.linspace(-(n_funds-1)/2, (n_funds-1)/2, n_funds) * bar_h

    for fi, fund in enumerate(FUNDS):
        vals = [td[1].get(fund, 0) for td in tests_data]
        ax.barh(y_base + offsets[fi], vals, height=bar_h,
                       color=FUND_COLORS[fund], label=fund_label(fund))

    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_yticks(y_base)
    ax.set_yticklabels([td[0] for td in tests_data], fontsize=9)
    ax.set_xlabel("Allocation change (percentage points)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=7, bbox_to_anchor=(1.01, 1), loc="upper left",
              borderaxespad=0)
    fig.tight_layout()
    return fig_to_stream(fig)


def tornado_chart(rows_si, title, high=HIGHLY_SENSITIVE, low=LEAST_SENSITIVE):
    """
    Tornado chart: one horizontal bar per test, length = SI.
    rows_si: list of (label, si, gaining_fund, gaining_delta, losing_fund, losing_delta)
    """
    rows_si = sorted(rows_si, key=lambda x: x[1])
    labels = [r[0] for r in rows_si]
    sis    = [r[1] for r in rows_si]

    fig, ax = plt.subplots(figsize=(8, max(3, len(rows_si)*0.45)))
    colors = ["#D62728" if s >= high else "#FF7F0E" for s in sis]
    ax.barh(range(len(labels)), sis, color=colors)
    ax.axvline(high, color="red",    linestyle="--", linewidth=0.8, label=f"SI = {high}")
    ax.axvline(low,  color="orange", linestyle="--", linewidth=0.8, label=f"SI = {low}")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Sensitivity Index (½ Σ|Δ pp|)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig_to_stream(fig)


def grouped_alloc_chart(scenarios, alloc_df, title):
    """
    Grouped bar chart: one group per fund, one bar per scenario.
    scenarios: list of scenario strings (first = baseline)
    alloc_df: DataFrame with 'scenario' + fund columns (pct values)
    """
    scenario_colors = plt.cm.tab10(np.linspace(0, 0.7, len(scenarios)))
    n_funds = len(FUNDS)
    n_scen  = len(scenarios)
    x = np.arange(n_funds)
    bar_w = 0.8 / n_scen

    fig, ax = plt.subplots(figsize=(10, 5))
    for si, scen in enumerate(scenarios):
        row = alloc_df[alloc_df["scenario"] == scen]
        if row.empty:
            continue
        vals = [float(row[f].iloc[0]) for f in FUNDS]
        ax.bar(x + si * bar_w - (n_scen-1)*bar_w/2, vals,
               width=bar_w, label=scen, color=scenario_colors[si])

    ax.set_xticks(x)
    ax.set_xticklabels([fund_label(f) for f in FUNDS], rotation=30,
                        ha="right", fontsize=8)
    ax.set_ylabel("Allocation (%)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left")
    fig.tight_layout()
    return fig_to_stream(fig)


def cause_area_chart(df, title):
    """
    Grouped bar chart: GHD / GCR / AW allocation changes per test.
    df must have columns: test, ghd_delta, gcr_delta, aw_delta.
    """
    ca_colors = {"GHD": "#E87B4C", "GCR": "#9B59B6", "AW": "#5DB85D"}
    tests = df["test"].tolist()
    n = len(tests)
    x = np.arange(n)
    bar_w = 0.25

    fig, ax = plt.subplots(figsize=(max(8, n * 0.9), 4.5))
    for i, (col_key, label) in enumerate([("ghd_delta", "GHD"),
                                           ("gcr_delta", "GCR"),
                                           ("aw_delta", "AW")]):
        vals = df[col_key].tolist()
        ax.bar(x + (i - 1) * bar_w, vals, width=bar_w,
               label=label, color=ca_colors[label])

    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(tests, rotation=30, ha="right", fontsize=8)
    ax.set_ylabel("Allocation change (pp)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.legend(fontsize=9)
    fig.tight_layout()
    return fig_to_stream(fig)


def alloc_delta_chart(alloc_df, scenario_col, baseline_label, scenario_labels, title):
    """
    Allocation change chart computed from an allocation table with a baseline row.
    alloc_df: DataFrame with scenario_col + fund columns (% values).
    Returns an allocation_change_chart stream.
    """
    base_row = alloc_df[alloc_df[scenario_col] == baseline_label]
    if base_row.empty:
        return None
    base_vals = {f: float(base_row[f].iloc[0]) for f in FUNDS if f in alloc_df.columns}

    tests_data = []
    for label in scenario_labels:
        row = alloc_df[alloc_df[scenario_col] == label]
        if row.empty:
            continue
        deltas = {f: float(row[f].iloc[0]) - base_vals.get(f, 0)
                  for f in FUNDS if f in alloc_df.columns}
        tests_data.append((label, deltas))
    return allocation_change_chart(tests_data, title, top_n=len(tests_data))


def combo_heatmap(combo_index):
    """
    Heatmap: rows = power combos, columns = max-spend multipliers, cells = SI.
    """
    power_order = [
        "gcr_low_power", "gcr_high_power",
        "aw_low_power", "aw_high_power",
        "gcr_and_aw_low_power", "gcr_and_aw_high_power",
        "gcr_low_aw_high_power", "gcr_high_aw_low_power",
    ]
    spend_order = [2.5, 7.5, 10.0]

    pivot = combo_index.pivot(index="combo", columns="max_spend_multiplier",
                              values="sensitivity_index")
    pivot = pivot.reindex(index=power_order, columns=spend_order)

    fig, ax = plt.subplots(figsize=(6, 4.5))
    im = ax.imshow(pivot.values, cmap="RdYlGn_r", aspect="auto",
                   vmin=combo_index["sensitivity_index"].min(),
                   vmax=combo_index["sensitivity_index"].max())

    ax.set_xticks(range(len(spend_order)))
    ax.set_xticklabels([f"{s}×" for s in spend_order], fontsize=9)
    ax.set_yticks(range(len(power_order)))
    ax.set_yticklabels(power_order, fontsize=8)
    ax.set_xlabel("Max-spend cap multiplier", fontsize=9)
    ax.set_ylabel("DR power scenario", fontsize=9)
    ax.set_title("SI heatmap: DR power × max-spend cap", fontsize=11, fontweight="bold")

    for i in range(len(power_order)):
        for j in range(len(spend_order)):
            val = pivot.values[i, j]
            if not np.isnan(val):
                ax.text(j, i, f"{val:.1f}", ha="center", va="center",
                        fontsize=8, color="black")

    plt.colorbar(im, ax=ax, label="Sensitivity Index (pp)")
    fig.tight_layout()
    return fig_to_stream(fig)


def si_histogram(all_si_records):
    """
    Stacked histogram of SI values, color-coded by dimension.
    all_si_records: list of (si_value, dimension_label)
    """
    dims = list(DIM_COLORS.keys())
    bins = np.arange(0, max(r[0] for r in all_si_records) + 5, 5)

    fig, ax = plt.subplots(figsize=(9, 4))
    bottom = np.zeros(len(bins)-1)
    for dim in dims:
        vals = [r[0] for r in all_si_records if r[1] == dim]
        if not vals:
            continue
        counts, _ = np.histogram(vals, bins=bins)
        ax.bar(bins[:-1], counts, width=np.diff(bins), bottom=bottom,
               color=DIM_COLORS[dim], label=dim, align="edge", edgecolor="white")
        bottom += counts

    ax.axvline(LEAST_SENSITIVE,  color="orange", linestyle="--", linewidth=1.2,
               label=f"Least sensitive threshold (SI={LEAST_SENSITIVE})")
    ax.axvline(HIGHLY_SENSITIVE, color="red",    linestyle="--", linewidth=1.2,
               label=f"Highly sensitive threshold (SI={HIGHLY_SENSITIVE})")
    ax.set_xlabel("Sensitivity Index (½ Σ|Δ pp|)", fontsize=10)
    ax.set_ylabel("Number of tests", fontsize=10)
    ax.set_title("Distribution of sensitivity indices across all tests", fontsize=11, fontweight="bold")
    ax.legend(fontsize=8, bbox_to_anchor=(1.01, 1), loc="upper left")
    fig.tight_layout()
    return fig_to_stream(fig)


def max_si_bar(max_si_per_dim):
    """Simple horizontal bar of max SI per dimension."""
    dims = list(max_si_per_dim.keys())
    vals = [max_si_per_dim[d] for d in dims]
    colors = [DIM_COLORS.get(d, "gray") for d in dims]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.barh(dims, vals, color=colors)
    ax.axvline(LEAST_SENSITIVE,  color="orange", linestyle="--", linewidth=1,
               label=f"SI = {LEAST_SENSITIVE}")
    ax.axvline(HIGHLY_SENSITIVE, color="red",    linestyle="--", linewidth=1,
               label=f"SI = {HIGHLY_SENSITIVE}")
    ax.set_xlabel("Max Sensitivity Index", fontsize=9)
    ax.set_title("Maximum SI by dimension", fontsize=11, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig_to_stream(fig)


# ── load data ─────────────────────────────────────────────────────────────────

def load():
    data = {}

    data["baseline"] = pd.read_csv(p("outputs", "baseline_staged.csv"))

    data["ra_summary"] = pd.read_csv(
        p("risk-aversion", "outputs", "risk_aversion_summary.csv"))
    data["ra_by_fund"] = pd.read_csv(
        p("risk-aversion", "outputs", "risk_aversion_by_fund.csv"))

    data["wv_index"] = pd.read_csv(
        p("worldview-sensitivity", "outputs", "split_credences_index.csv"))
    data["wv_by_fund"] = pd.read_csv(
        p("worldview-sensitivity", "outputs", "split_credences_by_fund.csv"))

    data["ce_si"] = pd.read_csv(
        p("across-the-board", "outputs", "ce_multiplier_si.csv"))
    data["ce_alloc"] = pd.read_csv(
        p("across-the-board", "outputs", "ce_multiplier_allocations.csv"))

    data["dr_index"] = pd.read_csv(
        p("diminishing-returns", "outputs", "dr_sensitivity_index.csv"))
    data["dr_alloc"] = pd.read_csv(
        p("diminishing-returns", "outputs", "dr_sensitivity_allocations.csv"))
    data["ms_index"] = pd.read_csv(
        p("diminishing-returns", "outputs", "max_spend_sensitivity_index.csv"))
    data["ms_alloc"] = pd.read_csv(
        p("diminishing-returns", "outputs", "max_spend_sensitivity_allocations.csv"))
    data["combo_index"] = pd.read_csv(
        p("diminishing-returns", "outputs", "combo_max_spend_index.csv"))
    data["combo_alloc"] = pd.read_csv(
        p("diminishing-returns", "outputs", "combo_max_spend_allocations.csv"))
    data["combo_by_fund"] = pd.read_csv(
        p("diminishing-returns", "outputs", "combo_max_spend_by_fund.csv"))

    data["agg_index"] = pd.read_csv(
        p("aggregation-methods", "outputs", "split_credences_index.csv"))

    data["ra_cause_area"] = pd.read_csv(
        p("risk-aversion", "outputs", "risk_aversion_cause_area_summary.csv"))
    data["agg_by_fund"] = pd.read_csv(
        p("aggregation-methods", "outputs", "split_credences_by_fund.csv"))

    data["disc_fund_si"]    = pd.read_csv(
        p("time-discounts", "outputs", "discount_fund_si.csv"))
    data["disc_fund_alloc"] = pd.read_csv(
        p("time-discounts", "outputs", "discount_fund_allocations.csv"))
    data["mw_overall_si"]   = pd.read_csv(
        p("moral-weights", "outputs", "moral_weights_overall_si.csv"))
    data["mw_overall_alloc"]= pd.read_csv(
        p("moral-weights", "outputs", "moral_weights_overall_allocations.csv"))
    data["mw_ranked"]       = pd.read_csv(
        p("moral-weights", "outputs", "moral_weights_ranked_summary.csv"))

    return data


# ── document builder ──────────────────────────────────────────────────────────

def build_doc(d):
    doc = Document()

    # ── narrow margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ═══════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Donor Compass Sensitivity Analysis Report")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph(
        f"Generated {date.today().strftime('%B %d, %Y')}  ·  "
        "Rethink Priorities"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "Executive Summary", 1)

    doc.add_paragraph(
        "The sensitivity index (SI) used throughout this report measures the total absolute "
        "change (in percentage points) summed across all funds, divided by two, when an "
        "assumption is shifted from its baseline to an alternative value. Because allocations "
        "sum to 100%, total gains always equal total losses; dividing by two gives the net "
        "amount redistributed. An SI of 10 means 10 pp was redistributed from losing funds "
        "to gaining funds."
    )

    # ── Key drivers summary ─────────────────────────────────────────────────────
    add_heading(doc, "Most Important Sensitivity Drivers", 2)

    doc.add_paragraph(
        "The table below ranks the most impactful parameter changes across all main "
        "analyses, sorted by SI. Risk-aversion scenarios (Section 5) are excluded "
        "because they represent extreme single-perspective cases rather than "
        "uncertainty about a shared parameter."
    )

    # Build ranked rows dynamically from data
    summary_rows = []

    # Worldview credences (main)
    wv_tmp = d["wv_index"].copy()
    wv_tmp = wv_tmp[~wv_tmp["bound"].isin(["single", "baseline"])]
    for _, r in wv_tmp.sort_values("sensitivity_index", ascending=False).head(3).iterrows():
        crange = f"{r['credence_base']:.0%}→{r['credence_scenario']:.0%}"
        summary_rows.append((
            r["sensitivity_index"],
            "Worldview Credences",
            f"{r['worldview'][:45]} ({r['bound']}, {crange})",
            fund_label(r["most_affected_fund"]),
            r["most_affected_delta"],
        ))

    # CE multipliers
    ce_tmp = d["ce_si"][d["ce_si"]["fund_varied"] != "baseline"].copy()
    for _, r in ce_tmp.sort_values("sensitivity_index", ascending=False).head(3).iterrows():
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        top_fund = max(diffs, key=lambda k: abs(diffs[k]))
        summary_rows.append((
            r["sensitivity_index"],
            "CE Multipliers",
            f"{r['fund_varied']} ×{r['multiplier']}",
            fund_label(top_fund),
            diffs[top_fund],
        ))

    # DR max spending cap
    for _, r in d["ms_index"].sort_values("sensitivity_index", ascending=False).iterrows():
        summary_rows.append((
            r["sensitivity_index"],
            "Dim. Returns (Max Spend)",
            f"Max spend cap ×{r['max_addl_spend_multiplier']}",
            fund_label(r["most_affected_fund"]),
            r["most_affected_delta"],
        ))

    # DR power
    for _, r in d["dr_index"].sort_values("sensitivity_index", ascending=False).head(2).iterrows():
        summary_rows.append((
            r["sensitivity_index"],
            "Dim. Returns (Power)",
            r["combo"],
            fund_label(r["most_affected_fund"]),
            r["most_affected_delta"],
        ))

    # Aggregation methods
    agg_tmp = d["agg_index"][~d["agg_index"]["bound"].isin(["single", "baseline"])].copy()
    for _, r in agg_tmp.sort_values("sensitivity_index", ascending=False).head(2).iterrows():
        summary_rows.append((
            r["sensitivity_index"],
            "Aggregation Methods",
            f"{r['scenario']} ({r['bound']} bound)",
            fund_label(r["most_affected_fund"]),
            r["most_affected_delta"],
        ))

    # Time discounts (only 0× scenarios produce SI > 0)
    disc_top = d["disc_fund_si"][
        (d["disc_fund_si"]["scenario_group"] != "baseline") &
        (d["disc_fund_si"]["sensitivity_index"] > 0)
    ].copy()
    for _, r in disc_top.sort_values("sensitivity_index", ascending=False).head(2).iterrows():
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        top_fund = max(diffs, key=lambda k: abs(diffs[k]))
        summary_rows.append((
            r["sensitivity_index"],
            "Time Discounts",
            f"{r['scenario_group']} ×{r['multiplier']}",
            fund_label(top_fund),
            diffs[top_fund],
        ))

    # Moral weights (overall blend, top 2 by SI)
    mw_top = d["mw_overall_si"][d["mw_overall_si"]["multiplier"] != 1.0].copy()
    fund_diff_cols = [c for c in mw_top.columns
                      if c.startswith("diff_") and c not in ("diff_ghd", "diff_gcr", "diff_aw")]
    for _, r in mw_top.sort_values("sensitivity_index", ascending=False).head(2).iterrows():
        diffs = {c.replace("diff_", ""): r[c] for c in fund_diff_cols}
        top_fund = max(diffs, key=lambda k: abs(diffs[k]))
        summary_rows.append((
            r["sensitivity_index"],
            "Moral Weights",
            f"All animal weights ×{r['multiplier']}",
            fund_label(top_fund),
            diffs[top_fund],
        ))

    summary_rows.sort(key=lambda x: x[0], reverse=True)
    tbl_rows = [
        [str(i + 1), fmt_si(si), dim, test[:50], fund, fmt_pp(delta)]
        for i, (si, dim, test, fund, delta) in enumerate(summary_rows)
    ]
    add_table(doc,
              ["Rank", "SI", "Dimension", "Test / Parameter", "Most Affected Fund", "Max Δ"],
              tbl_rows,
              col_widths=[0.35, 0.55, 1.45, 2.2, 1.45, 0.65])

    doc.add_paragraph(
        "Worldview credences and CE multipliers dominate: both reach SI > 27 pp, driven "
        "by parameters that directly control how much weight the model places on GCR "
        "spending. Increasing credence on either Total Utilitarianism — Default or "
        "Non-Utilitarian Consequentialism — Default (both GCR-favouring worldviews) "
        "shifts up to ~27 pp to Longview AI Safety. Boosting the GCR group's CE by "
        "100× similarly adds ~18 pp to Longview AI Safety. "
        "Within GHD, LEAF and GiveWell are highly sensitive to each other's CE: "
        "doubling LEAF's CE redistributes ~16 pp from GiveWell to LEAF. "
        "The max-spending cap is the most impactful DR parameter (SI up to 15 pp): "
        "a tight 2.5× cap prevents money from concentrating in Longview AI Safety, "
        "sending ~13 pp to GiveWell instead. "
        "Aggregation method weights and GHD timing are the least sensitive dimensions "
        "tested (SI ≤ 5 pp)."
    )
    doc.add_paragraph()

    # Baseline table
    add_heading(doc, "Baseline Allocation", 2)
    baseline = d["baseline"]

    # weighted_combined column is the headline
    if "weighted_combined" in baseline.columns:
        bl_rows = [(fund_label(r["fund"]), f"{r['weighted_combined']:.1f}%")
                   for _, r in baseline.iterrows()]
        add_table(doc,
                  ["Fund", "Baseline Allocation (%)"],
                  bl_rows,
                  col_widths=[3.5, 1.5])
    doc.add_paragraph()

    # ── Split single-perspective rows from main analyses ────────
    wv_all  = d["wv_index"].copy()
    wv_main = wv_all[~wv_all["bound"].isin(["single", "baseline"])].copy()
    wv_single = wv_all[wv_all["bound"] == "single"].copy()

    agg_all  = d["agg_index"].copy()
    agg_main = agg_all[~agg_all["bound"].isin(["single", "baseline"])].copy()
    agg_single = agg_all[agg_all["bound"] == "single"].copy()

    # ── Collect SI values for main histogram (illustratives excluded) ──
    all_si = []

    wv_si_vals = wv_main["sensitivity_index"].tolist()
    all_si += [(v, "Worldview Credences") for v in wv_si_vals]

    ce_si_vals = d["ce_si"][d["ce_si"]["fund_varied"] != "baseline"][
        "sensitivity_index"].tolist()
    all_si += [(v, "CE Multipliers") for v in ce_si_vals]

    dr_si_vals = d["dr_index"]["sensitivity_index"].tolist()
    all_si += [(v, "Dim. Returns (Power)") for v in dr_si_vals]

    ms_si_vals = d["ms_index"]["sensitivity_index"].tolist()
    all_si += [(v, "Dim. Returns (Max Spend)") for v in ms_si_vals]

    agg_si_vals = agg_main["sensitivity_index"].tolist()
    all_si += [(v, "Aggregation Methods") for v in agg_si_vals]

    disc_nonbase = d["disc_fund_si"][d["disc_fund_si"]["scenario_group"] != "baseline"]
    disc_si_vals = disc_nonbase[disc_nonbase["sensitivity_index"] > 0]["sensitivity_index"].tolist()
    all_si += [(v, "Time Discounts") for v in disc_si_vals]

    mw_nonbase = d["mw_overall_si"][d["mw_overall_si"]["multiplier"] != 1.0]
    mw_si_vals = mw_nonbase["sensitivity_index"].tolist()
    all_si += [(v, "Moral Weights") for v in mw_si_vals]

    max_si_per_dim = {
        "Worldview Credences":     max(wv_si_vals),
        "CE Multipliers":          max(ce_si_vals),
        "Dim. Returns (Power)":    max(dr_si_vals),
        "Dim. Returns (Max Spend)":max(ms_si_vals),
        "Aggregation Methods":     max(agg_si_vals),
        "Time Discounts":          max(disc_si_vals) if disc_si_vals else 0,
        "Moral Weights":           max(mw_si_vals)   if mw_si_vals   else 0,
    }

    add_heading(doc, "Maximum SI by Dimension", 2)
    add_picture(doc, max_si_bar(max_si_per_dim), width=Inches(5.5))
    doc.add_paragraph()

    add_heading(doc, "Distribution of SI Values — Main Analyses Only", 2)
    add_picture(doc, si_histogram(all_si), width=Inches(6.5))
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: risk-aversion scenario tests and analyses placing 100% credence in a single "
        "worldview or aggregation method are presented separately in Section 5 as illustrative "
        "scenarios; they are excluded from the chart above."
    )
    note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Distribution of SI Values — Including Illustrative Tests", 2)
    all_si_full = list(all_si)
    for _, r in d["ra_summary"].iterrows():
        all_si_full.append((float(r["sensitivity_index"]), "Risk Aversion"))
    for _, r in wv_single.iterrows():
        all_si_full.append((float(r["sensitivity_index"]), "Single Worldview"))
    for _, r in agg_single.iterrows():
        all_si_full.append((float(r["sensitivity_index"]), "Single Agg Method"))
    if all_si_full and max(r[0] for r in all_si_full) > 0:
        add_picture(doc, si_histogram(all_si_full), width=Inches(6.5))
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # SECTION 1 — WORLDVIEW CREDENCES
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "1. Worldview Credences", 1)

    doc.add_paragraph(
        "Each of the 14 worldviews carries a credence weight (best-guess estimate). "
        "Each worldview's credence is shifted to its low-end or high-end uncertainty "
        "bound, one worldview at a time, holding the rest at their best-guess values "
        "and re-normalising. For example, Total Utilitarianism — Default is shifted "
        "from 10% credence to 25% (high bound) or 5% (low bound). "
        "Single-worldview (100% credence) scenarios are presented in Section 5."
    )

    wv_sorted = wv_main.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    add_heading(doc, "Sensitivity by Test — Ranked by SI (top 20 of 28)", 2)
    tbl_rows = []
    for _, r in wv_sorted.head(20).iterrows():
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        scaled = fmt_si(r["scaled_SI"]) if "scaled_SI" in r.index and pd.notna(r["scaled_SI"]) else "—"
        tbl_rows.append([
            r["worldview"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            scaled,
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Worldview", "Bound", "Credence Range", "SI",
               "Scaled SI*", "Most Affected Fund", "Max Δ"],
              tbl_rows,
              col_widths=[2.1, 0.5, 1.0, 0.55, 0.65, 1.35, 0.8])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ absolute credence change (pp). "
        "Higher values indicate greater sensitivity per unit shift in credence weight."
    )
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI (top 20)", 2)
    wv_scaled = wv_sorted[wv_sorted["scaled_SI"].notna()].copy()
    wv_scaled = wv_scaled.sort_values("scaled_SI", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(wv_scaled.head(20).iterrows()):
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        tbl_rows2.append([
            str(i + 1),
            r["worldview"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["scaled_SI"]),
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Rank", "Worldview", "Bound", "Credence Range", "SI",
               "Scaled SI", "Most Affected Fund", "Max Δ"],
              tbl_rows2,
              col_widths=[0.4, 1.85, 0.5, 1.0, 0.55, 0.65, 1.3, 0.75])
    doc.add_paragraph()

    # Chart for SI ≥ 10
    top_wv = wv_sorted[wv_sorted["sensitivity_index"] >= HIGHLY_SENSITIVE].head(5)
    if not top_wv.empty:
        add_heading(doc, "Fund Allocation Changes — Tests with SI ≥ 10", 2)
        wv_bf = d["wv_by_fund"].copy()
        tests_data = []
        for _, r in top_wv.iterrows():
            scen = r["scenario"]
            rows = wv_bf[wv_bf["scenario"] == scen]
            deltas = {row["project_id"]: row["alloc_delta"]
                      for _, row in rows.iterrows()}
            label = f"{r['worldview'][:35]}… ({r['bound']})"
            tests_data.append((label, deltas))
        add_picture(doc, allocation_change_chart(
            tests_data, "Worldview credences: allocation change (pp)", top_n=5))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  Raising credence on Total Utilitarianism — Default or "
        "Non-Utilitarian Consequentialism — Default (both utilitarian-leaning worldviews "
        "that favour GCR spending) strongly increases Longview AI Safety's allocation "
        "(up to +27 pp). Raising credences on Contractualist or Kantian worldviews "
        "instead boosts GiveWell and Sentinel Biosecurity. Low-bound shifts are generally "
        "less impactful than high-bound shifts."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 2 — COST-EFFECTIVENESS MULTIPLIERS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "2. Cost-Effectiveness Multipliers", 1)

    doc.add_paragraph(
        "Each fund's cost-effectiveness (CE) was multiplied by a constant factor, "
        "one fund (or group) at a time. GCR individual funds (Longview AI, Longview Nuclear, "
        "Sentinel) were tested at 0.1× and 10×; GHD funds (GiveWell, LEAF) at 0.25×–2×; "
        "AW individual funds at 0.5×–2×. "
        "The GCR group (all three GCR funds jointly) was tested at 0.0001×–10,000× "
        "(∼8 orders of magnitude). The AW group (all three AW funds jointly) was tested "
        "at 0.25× and 4×."
    )

    ce = d["ce_si"].copy()
    ce_nobase = ce[ce["fund_varied"] != "baseline"].copy()
    ce_sorted = ce_nobase.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    add_heading(doc, "Sensitivity by Test — Ranked by SI (all 31 tests)", 2)
    tbl_rows = []
    for _, r in ce_sorted.iterrows():
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(r["si_scaled_pp_per_oom"]) if "si_scaled_pp_per_oom" in r.index and pd.notna(r["si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            r["fund_varied"],
            f"{r['multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            scaled,
            f"{fund_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{fund_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Fund Varied", "Multiplier", "SI", "Scaled SI*", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[1.2, 0.7, 0.55, 0.65, 1.95, 1.95])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier. "
        "Normalises for the size of the perturbation so tests at very different multiplier "
        "sizes can be compared on a like-for-like basis."
    )
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI", 2)
    ce_scaled = ce_sorted[ce_sorted["si_scaled_pp_per_oom"].notna()].copy()
    ce_scaled = ce_scaled.sort_values("si_scaled_pp_per_oom", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(ce_scaled.iterrows()):
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        tbl_rows2.append([
            str(i + 1),
            r["fund_varied"],
            f"{r['multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["si_scaled_pp_per_oom"]),
            f"{fund_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{fund_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Rank", "Fund Varied", "Multiplier", "SI", "Scaled SI",
               "Biggest Gainer", "Biggest Loser"],
              tbl_rows2,
              col_widths=[0.4, 1.15, 0.65, 0.55, 0.65, 1.85, 1.85])
    doc.add_paragraph()

    # Tornado chart for all tests, highlight SI ≥ 10
    add_heading(doc, "Tornado Chart — All CE Multiplier Tests", 2)
    tornado_rows = []
    for _, r in ce_sorted.iterrows():
        diff_cols = [c for c in r.index if c.startswith("diff_")]
        diffs = {c.replace("diff_", ""): r[c] for c in diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        label = f"{r['fund_varied']} ×{r['multiplier']}"
        tornado_rows.append((
            label, r["sensitivity_index"],
            fund_label(gainer), diffs[gainer],
            fund_label(loser),  diffs[loser],
        ))
    add_picture(doc, tornado_chart(tornado_rows,
        "CE Multipliers: Sensitivity Index per test (red = SI ≥ 10)"),
        width=Inches(6.5))
    doc.add_paragraph()

    # Fund-level allocation change chart for SI ≥ 10 tests
    add_heading(doc, "Fund Allocation Changes — Tests with SI ≥ 10", 2)
    ce_high = ce_sorted[ce_sorted["sensitivity_index"] >= HIGHLY_SENSITIVE].copy()
    diff_cols_all = [c for c in ce_high.columns if c.startswith("diff_")]
    ce_tests_data = []
    for _, r in ce_high.iterrows():
        deltas = {c.replace("diff_", ""): r[c] for c in diff_cols_all}
        label = f"{r['fund_varied']} ×{r['multiplier']}"
        ce_tests_data.append((label, deltas))
    if ce_tests_data:
        add_picture(doc, allocation_change_chart(
            ce_tests_data,
            "CE Multipliers: fund allocation changes for SI ≥ 10 tests (pp)",
            top_n=len(ce_tests_data)))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  The GCR group at 100× is the most impactful test, adding "
        "~18 pp to Longview AI Safety while reducing Navigation General (~11 pp), "
        "EA AWF (~7 pp), and GiveWell (~5 pp). Reducing Longview AI Safety's CE to "
        "10% (×0.1) redistributes ~13 pp, mainly to Sentinel Bio (+4.5 pp), "
        "Longview Nuclear (+3.9 pp), and GiveWell (+3.5 pp). GiveWell is also highly "
        "sensitive to LEAF's CE (×1.5 and ×2), as the two funds closely compete within "
        "GHD. Navigation Fund (Cage-Free) is the least sensitive individual fund — "
        "its allocation barely changes across any CE variation (max SI 2.75 pp)."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 3 — DIMINISHING RETURNS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "3. Diminishing Returns", 1)

    doc.add_paragraph(
        "Two sub-analyses explore how the diminishing-returns curves affect allocations."
    )

    # 4a — Power Parameters
    add_heading(doc, "3a. Power Parameters (SI range: 0.96 – 6.66)", 2)
    doc.add_paragraph(
        "The power parameter controls how steeply returns diminish as more money flows "
        "into a fund cluster. 'Low power' means gradual diminishing returns; 'high power' "
        "means steep. Eight combinations were tested across the GCR and AW clusters."
    )

    dr = d["dr_index"].copy().sort_values("sensitivity_index", ascending=False)
    tbl_rows = [(r["combo"],
                 fmt_si(r["sensitivity_index"]),
                 fund_label(r["most_affected_fund"]),
                 fmt_pp(r["most_affected_delta"]))
                for _, r in dr.iterrows()]
    add_table(doc,
              ["Scenario", "SI", "Most Affected Fund", "Max Δ"],
              tbl_rows, col_widths=[2.5, 0.6, 2.0, 0.9])
    add_heading(doc, "Fund Allocation Changes — All DR Power Scenarios", 2)
    dr_alloc_all = d["dr_alloc"].copy()
    dr_scenarios = [r["combo"] for _, r in dr.iterrows()]  # already sorted by SI desc
    dr_stream = alloc_delta_chart(
        dr_alloc_all.rename(columns={"combo": "scenario"}),
        "scenario", "baseline", dr_scenarios,
        "DR Power: allocation change from baseline (pp)")
    if dr_stream:
        add_picture(doc, dr_stream)
    doc.add_paragraph()

    doc.add_paragraph(
        "No tests exceed SI = 10. The most sensitive combination (gcr_high_aw_low_power, "
        "SI 6.66) reduces Navigation Fund (General) by 4.6 pp. GCR power parameters "
        "primarily affect Sentinel Biosecurity; AW power primarily affects Navigation General."
    )
    doc.add_paragraph()

    # 4b — Max Spending Cap
    add_heading(doc, "3b. Maximum Spending Cap (SI range: 9.74 – 15.13)", 2)
    doc.add_paragraph(
        "The maximum spending cap limits how much of the total budget can flow to any "
        "one fund cluster. Three cap multipliers were tested: 2.5×, 7.5×, and 10× the "
        "baseline per-fund spending ceiling. All three tests exceed the SI = 10 threshold."
    )

    ms = d["ms_index"].copy().sort_values("sensitivity_index", ascending=False)
    tbl_rows = []
    for _, r in ms.iterrows():
        tbl_rows.append([
            r["scenario"],
            f"{r['max_addl_spend_multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Scenario", "Cap Multiplier", "SI", "Most Affected Fund", "Max Δ"],
              tbl_rows, col_widths=[1.4, 1.0, 0.6, 2.0, 0.9])
    doc.add_paragraph()

    # Grouped bar chart: baseline + 3 max-spend scenarios
    ms_alloc = d["ms_alloc"].copy()
    # Add baseline row from dr_alloc
    dr_alloc = d["dr_alloc"].copy()
    baseline_row = dr_alloc[dr_alloc["combo"] == "baseline"].copy()
    if not baseline_row.empty:
        baseline_row = baseline_row.rename(columns={"combo": "scenario"})
        baseline_row["scenario"] = "baseline"
        ms_alloc_full = pd.concat([baseline_row[["scenario"] + FUNDS],
                                   ms_alloc.rename(columns={"scenario": "scenario"}
                                   )[["scenario"] + FUNDS]], ignore_index=True)
    else:
        ms_alloc_full = ms_alloc

    scenarios_ms = ["baseline", "max_spend_2_5x", "max_spend_7_5x", "max_spend_10x"]
    add_picture(doc, grouped_alloc_chart(
        scenarios_ms, ms_alloc_full,
        "Diminishing Returns — Max Spending Cap: Fund Allocations by Scenario"),
        width=Inches(7.0))
    doc.add_paragraph()

    add_heading(doc, "Fund Allocation Changes — Max-Spend Scenarios vs Baseline", 2)
    ms_delta_stream = alloc_delta_chart(
        ms_alloc_full, "scenario", "baseline",
        ["max_spend_2_5x", "max_spend_7_5x", "max_spend_10x"],
        "DR Max Spend: allocation change from baseline (pp)")
    if ms_delta_stream:
        add_picture(doc, ms_delta_stream)
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  At 2.5× cap, GiveWell gains ~13 pp because the tighter cap "
        "limits how much goes to Longview AI Safety, redistributing to the next-best GHD "
        "option. At 10× cap, Navigation General loses ~9 pp as funds spread to higher-CE "
        "options that are now unconstrained. The 7.5× cap produces intermediate effects."
    )

    # 4c — Combined: Power × Max Spend
    add_heading(doc, "3c. Combined: DR Power × Maximum Spending Cap", 2)
    doc.add_paragraph(
        "This sub-analysis crosses each of the 8 DR power scenarios with each of the "
        "3 max-spend cap multipliers, producing 24 joint tests. It shows whether the "
        "effects of the two parameters reinforce or offset each other."
    )

    combo_idx = d["combo_index"].copy()
    combo_alloc = d["combo_alloc"].copy()
    combo_bf = d["combo_by_fund"].copy()

    # Heatmap
    add_heading(doc, "SI Heatmap: Power Scenario × Cap Multiplier", 3)
    add_picture(doc, combo_heatmap(combo_idx), width=Inches(5.5))
    doc.add_paragraph()

    # Full SI table sorted by SI
    add_heading(doc, "All 24 Tests — Ranked by SI", 3)
    combo_sorted = combo_idx.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in combo_sorted.iterrows():
        tbl_rows.append([
            r["combo"],
            f"{r['max_spend_multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Power Scenario", "Cap Multiplier", "SI",
               "Most Affected Fund", "Max Δ"],
              tbl_rows, col_widths=[2.2, 0.9, 0.6, 1.9, 0.9])
    doc.add_paragraph()

    # Chart: allocation changes for top 4 most sensitive combos
    top_combo = combo_sorted.head(4)
    add_heading(doc, "Fund Allocation Changes — Top 4 Most Sensitive Combos", 3)
    tests_data = []
    for _, r in top_combo.iterrows():
        scen = r["scenario"]
        rows = combo_bf[combo_bf["scenario"] == scen]
        deltas = {row["project_id"]: row["alloc_delta"] for _, row in rows.iterrows()}
        label = f"{r['combo']} ×{r['max_spend_multiplier']}"
        tests_data.append((label, deltas))
    add_picture(doc, allocation_change_chart(
        tests_data,
        "Combined DR power × cap: allocation change from baseline (pp)", top_n=4))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  The highest SI across all 24 combinations is "
        f"gcr_and_aw_high_power × 10× (SI {combo_sorted.iloc[0]['sensitivity_index']:.2f}), "
        "driven by Navigation Fund (Cage-Free) gaining ~13 pp. "
        "The 2.5× cap dominates at most power settings — it consistently produces the "
        "largest GiveWell gains (10–14 pp) by constraining the preferred GCR/AW cluster. "
        "At 7.5× and 10× caps, high AW power scenarios redirect money toward Navigation "
        "Fund (Cage-Free), while low AW power scenarios redirect it toward EA Animal Welfare "
        "Fund. The GCR power parameter has little additional impact once the cap multiplier "
        "is fixed, suggesting the two parameters interact primarily through the AW cluster."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 4 — AGGREGATION METHOD WEIGHTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "4. Aggregation Method Weights", 1)

    doc.add_paragraph(
        "Seven voting/aggregation methods are used to combine worldview scores into "
        "an allocation. Each method carries a credence weight. This dimension shifts "
        "each method's credence from its best-guess value to its low or high "
        "uncertainty bound, one method at a time. "
        "NashBargaining: 35% (range 15–50%); Marketplace: 30% (20–50%); "
        "MEC: 15% (10–20%); MET: 10% (5–15%); SplitCycle: 8% (2–10%); "
        "Borda: 2% (0–5%); LexMaximin: 0% (0–5%). "
        "Single-method (100% credence) scenarios are presented in Section 5."
    )

    agg_sorted = agg_main.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in agg_sorted.iterrows():
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        scaled = fmt_si(r["scaled_SI"]) if "scaled_SI" in r.index and pd.notna(r["scaled_SI"]) else "—"
        tbl_rows.append([
            r["scenario"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            scaled,
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Test", "Bound", "Credence Range", "SI",
               "Scaled SI*", "Most Affected Fund", "Max Δ"],
              tbl_rows, col_widths=[1.8, 0.55, 1.05, 0.55, 0.65, 1.5, 0.85])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ absolute credence change (pp)."
    )
    p_note.runs[0].italic = True

    doc.add_paragraph()
    add_heading(doc, "Same Tests — Ranked by Scaled SI", 2)
    agg_scaled = agg_sorted[agg_sorted["scaled_SI"].notna()].copy()
    agg_scaled = agg_scaled.sort_values("scaled_SI", ascending=False).reset_index(drop=True)
    tbl_rows2 = []
    for i, (_, r) in enumerate(agg_scaled.iterrows()):
        crange = f"{r['credence_base']:.0%} → {r['credence_scenario']:.0%}"
        tbl_rows2.append([
            str(i + 1),
            r["scenario"],
            r["bound"],
            crange,
            fmt_si(r["sensitivity_index"]),
            fmt_si(r["scaled_SI"]),
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Rank", "Test", "Bound", "Credence Range", "SI",
               "Scaled SI", "Most Affected Fund", "Max Δ"],
              tbl_rows2, col_widths=[0.4, 1.6, 0.55, 1.05, 0.55, 0.65, 1.45, 0.8])
    doc.add_paragraph(
        "No tests exceed SI = 10. The most sensitive test by raw SI is NashBargaining at its "
        "low bound (credence 35% → 15%, SI 4.72), which reduces Longview AI Safety by "
        "3.6 pp. The allocation is relatively stable to changes in method weights within "
        "their uncertainty ranges."
    )

    add_heading(doc, "Fund Allocation Changes — All Credence-Variation Tests", 2)
    agg_bf = d["agg_by_fund"].copy()
    agg_bf_main = agg_bf[~agg_bf["bound"].isin(["single", "baseline"])].copy()
    agg_tests_data = []
    for scen in agg_sorted["scenario"].tolist():
        rows = agg_bf_main[agg_bf_main["scenario"] == scen]
        if rows.empty:
            continue
        deltas = {row["project_id"]: row["alloc_delta"] for _, row in rows.iterrows()}
        agg_tests_data.append((scen, deltas))
    if agg_tests_data:
        add_picture(doc, allocation_change_chart(
            agg_tests_data,
            "Aggregation method weights: allocation change (pp)",
            top_n=len(agg_tests_data)))
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════
    # LEAST SENSITIVE TESTS (main analyses only)
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "Least Sensitive Tests (SI < 2)", 1)

    doc.add_paragraph(
        "The following main-analysis tests produced a maximum fund allocation change of "
        "less than 2 percentage points. These assumptions can be considered robust within "
        "their tested range. Illustrative single-perspective scenarios are excluded here "
        "and appear in Section 5."
    )

    low_tests = []

    for _, r in wv_main.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Worldview Credences",
                               r["scenario"],
                               r["sensitivity_index"],
                               f"{r['worldview'][:50]} ({r['bound']} bound: "
                               f"{r['credence_base']:.0%}→{r['credence_scenario']:.0%})"))

    ce_low = ce_sorted[ce_sorted["sensitivity_index"] < LEAST_SENSITIVE]
    for _, r in ce_low.iterrows():
        low_tests.append(("CE Multipliers",
                           f"{r['fund_varied']} ×{r['multiplier']}",
                           r["sensitivity_index"],
                           f"Multiply {r['fund_varied']} CE by {r['multiplier']}×"))

    for _, r in d["dr_index"].iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Dim. Returns (Power)", r["combo"],
                               r["sensitivity_index"],
                               f"DR power combo: {r['combo']}"))

    for _, r in d["ms_index"].iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Dim. Returns (Max Spend)", r["scenario"],
                               r["sensitivity_index"],
                               f"Max spend cap ×{r['max_addl_spend_multiplier']}"))

    for _, r in agg_main.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Aggregation Methods", r["scenario"],
                               r["sensitivity_index"],
                               f"Method weight shift: {r['scenario']} "
                               f"({r['credence_base']:.0%}→{r['credence_scenario']:.0%})"))

    disc_all = d["disc_fund_si"][d["disc_fund_si"]["scenario_group"] != "baseline"].copy()
    for _, r in disc_all.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Time Discounts",
                               f"{r['scenario_group']} ×{r['multiplier']}",
                               r["sensitivity_index"],
                               f"Scale discount factor(s) by {r['multiplier']}× "
                               f"in group '{r['scenario_group']}'"))

    mw_all = d["mw_overall_si"][d["mw_overall_si"]["multiplier"] != 1.0].copy()
    for _, r in mw_all.iterrows():
        if r["sensitivity_index"] < LEAST_SENSITIVE:
            low_tests.append(("Moral Weights",
                               f"All weights ×{r['multiplier']}",
                               r["sensitivity_index"],
                               f"Scale all animal moral weights by {r['multiplier']}×"))

    low_tests.sort(key=lambda x: (x[0], x[2]))

    current_dim = None
    for dim, _name, si, desc in low_tests:
        if dim != current_dim:
            current_dim = dim
            p_obj = doc.add_paragraph()
            run = p_obj.add_run(dim)
            run.bold = True
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.add_run(f"SI = {si:.2f}  ·  ").bold = False
        bullet.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: GCR-params sub-directory contains only analysis scripts with no output "
        "CSV files and is not included in this report."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 5 — TIME DISCOUNTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "5. Time Discounts", 1)

    doc.add_paragraph(
        "This analysis tests how the allocation shifts when the weights assigned to "
        "far-future effects are reduced. Two scenario groups were tested: "
        "'Discount 500+' modifies only the 500-year-plus timeframe weight (index 5 of the "
        "discount_factors array); 'Discount 100–500 and 500+' modifies both the 100–500 and "
        "500+ year weights (indices 4 and 5). Within each group, multipliers of 0.1×, 0.01×, "
        "0.001×, and 0× were applied to the target discount factor(s)."
    )

    disc_si = d["disc_fund_si"].copy()
    disc_nonbase_all = disc_si[disc_si["scenario_group"] != "baseline"].copy()
    disc_fund_diff_cols = [c for c in disc_si.columns if c.startswith("diff_")]

    add_heading(doc, "Sensitivity by Test — All Scenarios", 2)
    tbl_rows = []
    for _, r in disc_nonbase_all.iterrows():
        diffs = {c.replace("diff_", ""): r[c] for c in disc_fund_diff_cols}
        if r["sensitivity_index"] > 0:
            gainer = max(diffs, key=lambda k: diffs[k])
            loser  = min(diffs, key=lambda k: diffs[k])
            g_str = f"{fund_label(gainer)} ({fmt_pp(diffs[gainer])})"
            l_str = f"{fund_label(loser)} ({fmt_pp(diffs[loser])})"
        else:
            g_str = l_str = "—"
        mult_str = "0×" if r["multiplier"] == 0 else f"{r['multiplier']}×"
        tbl_rows.append([
            r["scenario_group"],
            mult_str,
            fmt_si(r["sensitivity_index"]),
            g_str,
            l_str,
        ])
    add_table(doc,
              ["Scenario Group", "Multiplier", "SI", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[1.9, 0.65, 0.55, 2.1, 2.1])
    doc.add_paragraph()

    disc_zero = disc_nonbase_all[disc_nonbase_all["multiplier"] == 0].copy()
    if not disc_zero.empty:
        add_heading(doc, "Fund Allocation Changes — 0× Scenarios (Complete Elimination)", 2)
        tests_data = []
        for _, r in disc_zero.iterrows():
            deltas = {c.replace("diff_", ""): r[c] for c in disc_fund_diff_cols}
            tests_data.append((r["scenario_group"], deltas))
        add_picture(doc, allocation_change_chart(
            tests_data,
            "Time discounts: fund allocation change when discount factor eliminated (pp)",
            top_n=len(tests_data)))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  The allocation is completely insensitive to any partial reduction "
        "in far-future discount factors — multipliers of 0.1×, 0.01×, and 0.001× all "
        "produce SI = 0 pp. Only complete elimination (0×) of the 500+ year weight shifts "
        "the allocation significantly (SI = 15.38 pp), primarily redistributing from GCR "
        "(Longview AI Safety −6.39 pp, Sentinel Bio −6.63 pp) to Animal Welfare (Navigation "
        "General +8.36 pp, EA AWF +3.84 pp). Including the 100–500 year weight in the "
        "elimination amplifies the effect (SI = 20.87 pp, with Navigation General gaining "
        "+10.74 pp). The insensitivity to fractional reductions reflects that these timeframe "
        "factors are already small in most worldviews, so partial scaling barely changes "
        "effective scores."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 6 — ANIMAL MORAL WEIGHTS
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "6. Animal Moral Weights", 1)

    doc.add_paragraph(
        "This analysis scales the moral weights assigned to animal welfare effects "
        "across two levels: (1) the overall blend — all worldviews' animal weights "
        "scaled simultaneously, measuring aggregate portfolio sensitivity; and "
        "(2) per-worldview — each worldview run at 100% credence independently, "
        "measuring how sensitive each worldview's own recommended allocation is to "
        "its animal weight assumptions. Multipliers tested: 0.1×, 2×, 5×, 10×, 100×. "
        "Weights are capped at species-level upper bounds: chickens 0.75, fish 0.60, "
        "shrimp 0.25, invertebrates 0.20, mammals 0.90."
    )

    # 6a. Overall blend
    add_heading(doc, "6a. Overall Blend Sensitivity", 2)
    mw_si_df = d["mw_overall_si"].copy()
    mw_nonbase_df = mw_si_df[mw_si_df["multiplier"] != 1.0].copy()
    mw_fund_diff_cols = [c for c in mw_si_df.columns
                         if c.startswith("diff_") and c not in ("diff_ghd", "diff_gcr", "diff_aw")]

    tbl_rows = []
    for _, r in mw_nonbase_df.sort_values("sensitivity_index", ascending=False).iterrows():
        diffs = {c.replace("diff_", ""): r[c] for c in mw_fund_diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(r["si_scaled_pp_per_oom"]) if pd.notna(r["si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            f"{r['multiplier']}×",
            fmt_si(r["sensitivity_index"]),
            scaled,
            f"{fund_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{fund_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Multiplier", "SI", "Scaled SI*", "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[0.65, 0.55, 0.65, 2.3, 2.3])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Fund Allocation Changes — All Multipliers", 2)
    tests_data = []
    for _, r in mw_nonbase_df.sort_values("sensitivity_index", ascending=False).iterrows():
        deltas = {c.replace("diff_", ""): r[c] for c in mw_fund_diff_cols}
        tests_data.append((f"{r['multiplier']}×", deltas))
    if tests_data:
        add_picture(doc, allocation_change_chart(
            tests_data,
            "Moral weights (overall blend): fund allocation change (pp)",
            top_n=len(tests_data)))
    doc.add_paragraph()

    # 6b. Per-worldview ranked summary
    add_heading(doc, "6b. Per-Worldview Sensitivity — Ranked Summary (Top 15)", 2)
    doc.add_paragraph(
        "Each worldview is run in isolation (100% credence) with its animal weights "
        "scaled. SI is measured against that worldview's own unmodified baseline, not "
        "the special-blend baseline. Only non-zero scenarios are shown."
    )

    mw_ranked_df = d["mw_ranked"].copy()
    mw_ranked_fund_diff_cols = [c for c in mw_ranked_df.columns
                                 if c.startswith("diff_") and c not in ("diff_ghd", "diff_gcr", "diff_aw")]

    tbl_rows = []
    for _, r in mw_ranked_df.head(15).iterrows():
        diffs = {c.replace("diff_", ""): float(r[c]) for c in mw_ranked_fund_diff_cols}
        gainer = max(diffs, key=lambda k: diffs[k])
        loser  = min(diffs, key=lambda k: diffs[k])
        scaled = fmt_si(float(r["si_scaled_pp_per_oom"])) if pd.notna(r["si_scaled_pp_per_oom"]) else "—"
        tbl_rows.append([
            str(int(r["worldview_idx"])),
            r["worldview_name"][:28],
            str(int(r["risk_profile"])),
            f"{r['multiplier']}×",
            fmt_si(float(r["sensitivity_index"])),
            scaled,
            f"{fund_label(gainer)} ({fmt_pp(diffs[gainer])})",
            f"{fund_label(loser)} ({fmt_pp(diffs[loser])})",
        ])
    add_table(doc,
              ["Idx", "Worldview", "Risk", "Mult.", "SI", "Scaled SI*",
               "Biggest Gainer", "Biggest Loser"],
              tbl_rows,
              col_widths=[0.35, 1.65, 0.45, 0.5, 0.55, 0.65, 1.85, 1.85])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ orders-of-magnitude change in multiplier."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    add_heading(doc, "Per-Worldview SI — Tornado Chart (Top 20 Scenarios)", 2)
    tornado_rows_mw = []
    for _, r in mw_ranked_df.head(20).iterrows():
        label = (f"{r['worldview_name'][:28]} "
                 f"(idx={int(r['worldview_idx'])}, {r['multiplier']}×)")
        tornado_rows_mw.append((label, float(r["sensitivity_index"]), "", 0, "", 0))
    add_picture(doc, tornado_chart(
        tornado_rows_mw,
        "Animal moral weights (per worldview): SI vs own baseline",
        high=HIGHLY_SENSITIVE, low=LEAST_SENSITIVE),
        width=Inches(6.5))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings (overall blend):  Reducing all animal weights to 10% (0.1×) is "
        "the most impactful scenario (SI = 24.18 pp), shifting funds from AW to GHD as "
        "lower animal weights depress AW fund scores. Increasing weights (2×–100×) is "
        "less impactful (SI 1.84–14.01 pp) because many TU worldviews are already at "
        "or near species caps and cannot move further upward. "
        "Key findings (per worldview):  Contractualism idx 11 (risk-averse) is the most "
        "sensitive single worldview at 100× (SI = 75 pp, full GHD→AW reversal) — its "
        "baseline animal weights are so low that no species caps are hit. "
        "Non-Utilitarian Consequentialist worldviews (idx 7, 8) show a hard plateau: SI "
        "is essentially unchanged between 10× and 100× (~43–47 pp), because AW funds hit "
        "their diminishing-returns ceiling at approximately 10×. Their shift is GCR→AW "
        "rather than GHD→AW. TU worldviews are almost exclusively sensitive in the "
        "downward direction — they are already near the species caps at baseline so "
        "upward multipliers produce SI = 0."
    )

    # ═══════════════════════════════════════════════════════════
    # SECTION 7 — ILLUSTRATIVE SINGLE-PERSPECTIVE ANALYSES
    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "7. Illustrative: Single-Perspective Analyses", 1)

    doc.add_paragraph(
        "The analyses in this section place 100% credence in one ethical perspective — "
        "a single risk profile applied to all worldviews, a single worldview, or a single "
        "aggregation method. We do not endorse any of these positions; they are included to "
        "show the range of outcomes implied by individual views taken alone. Because the "
        "credence shifts here are far larger than in the main sensitivity analyses, higher "
        f"thresholds apply: SI ≥ {SINGLE_HIGHLY_SENSITIVE} pp is considered highly "
        f"sensitive (red); SI < {SINGLE_LEAST_SENSITIVE} pp is considered low-sensitivity "
        "(orange threshold line)."
    )

    # ── 5a. Risk Aversion ────────────────────────────────────────
    add_heading(doc, "5a. Risk Aversion", 2)

    doc.add_paragraph(
        "Each test assigns one risk-attitude profile uniformly to all worldviews. "
        "Two starting baselines are compared: \"Special blend\" (the default mixed profiles) "
        "and \"Neutral\" (all worldviews risk-neutral). Alternative profiles: Neutral, "
        "Downside Critical, Continuous Upside Skeptic, Bilateral Skeptic, WLU-5, WLU-10, "
        "and Combined. Tests are named <baseline>_to_<new_profile>."
    )

    ra = d["ra_summary"].copy()
    ra_sorted = ra.sort_values("sensitivity_index", ascending=False).reset_index(drop=True)

    tbl_rows = []
    for _, r in ra_sorted.iterrows():
        tbl_rows.append([
            r["test"],
            fmt_si(r["sensitivity_index"]),
            fund_label(r["most_affected_fund"]),
            fmt_pp(r["most_affected_delta"]),
        ])
    add_table(doc,
              ["Test", "SI", "Most Affected Fund", "Max Δ"],
              tbl_rows,
              col_widths=[2.8, 0.6, 1.8, 0.9])
    doc.add_paragraph()

    # Fund allocation changes — all 11 tests
    add_heading(doc, "Fund Allocation Changes — All Tests", 2)
    fund_delta_cols = {f: f"{f}_delta" for f in FUNDS}
    all_ra_tests_data = []
    for _, r in ra_sorted.iterrows():
        deltas = {f: r[dc] for f, dc in fund_delta_cols.items() if dc in r.index}
        all_ra_tests_data.append((r["test"], deltas))
    if all_ra_tests_data:
        add_picture(doc, allocation_change_chart(
            all_ra_tests_data,
            "Risk Aversion: fund allocation change from baseline (pp) — all tests",
            top_n=len(all_ra_tests_data)))
    doc.add_paragraph()

    # Cause-area chart
    add_heading(doc, "Cause-Area Allocation Changes (GHD / GCR / AW)", 2)
    ra_ca = d["ra_cause_area"].copy()
    ra_ca = ra_ca.set_index("test").reindex(ra_sorted["test"]).reset_index()
    add_picture(doc, cause_area_chart(
        ra_ca,
        "Risk Aversion: cause-area allocation change from baseline (pp)"))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  Moving from the special-blend baseline to risk-neutral or "
        "bilateral-skeptic profiles causes Longview AI Safety to dominate (≤89% of budget) "
        "because risk-neutral worldviews heavily favour the highest-EV option. "
        "Switching to WLU, Combined, or Downside Critical profiles redistributes sharply "
        "toward GiveWell, Navigation General, and EA Animal Welfare Fund. The "
        "neutral → combined test shows the largest SI (90 pp), driven by Longview AI "
        "collapsing from 89% to 0%."
    )

    # ── 5b. 100% Credence in One Worldview ───────────────────────
    add_heading(doc, "5b. 100% Credence in One Worldview", 2)

    doc.add_paragraph(
        "Each worldview is run in isolation — as if it received 100% credence — "
        "to show the allocation it would produce on its own. SI is measured against the "
        "baseline special-blend allocation. SIs are large (30–84 pp) because the credence "
        "shift (from a small best-guess weight to 100%) is extreme."
    )

    wv_single_sorted = wv_single.sort_values(
        "sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in wv_single_sorted.iterrows():
        scaled_val = r.get("scaled_SI")
        tbl_rows.append([
            r["worldview"][:60],
            f"{float(r['credence_base']):.0%}",
            fmt_si(r["sensitivity_index"]),
            fmt_si(scaled_val) if pd.notna(scaled_val) else "—",
            fund_label(r["most_affected_fund"]),
            fmt_pp(float(r["most_affected_delta"])),
        ])
    add_table(doc,
              ["Worldview", "Base Credence", "SI", "Scaled SI*",
               "Most Affected Fund", "Max Δ"],
              tbl_rows,
              col_widths=[2.5, 0.85, 0.6, 0.7, 1.5, 0.85])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ absolute credence change (pp). "
        "Normalises for the fact that smaller-credence worldviews have a larger shift to 100%."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    tornado_rows_wv = [
        (r["worldview"][:50], float(r["sensitivity_index"]), "", 0, "", 0)
        for _, r in wv_single_sorted.iterrows()
    ]
    add_picture(doc, tornado_chart(
        tornado_rows_wv,
        "100% Credence in One Worldview: SI vs baseline",
        high=SINGLE_HIGHLY_SENSITIVE, low=SINGLE_LEAST_SENSITIVE),
        width=Inches(6.5))
    doc.add_paragraph()

    # ── 5c. 100% Credence in One Aggregation Method ──────────────
    add_heading(doc, "5c. 100% Credence in One Aggregation Method", 2)

    doc.add_paragraph(
        "Each aggregation method is run on the full budget with all worldviews, "
        "as if it were the only method used. SI is measured against the baseline "
        "weighted-average allocation."
    )

    agg_single_sorted = agg_single.sort_values(
        "sensitivity_index", ascending=False).reset_index(drop=True)
    tbl_rows = []
    for _, r in agg_single_sorted.iterrows():
        scaled_val = r.get("scaled_SI")
        tbl_rows.append([
            r["method"],
            f"{float(r['credence_base']):.0%}",
            fmt_si(r["sensitivity_index"]),
            fmt_si(scaled_val) if pd.notna(scaled_val) else "—",
            fund_label(r["most_affected_fund"]),
            fmt_pp(float(r["most_affected_delta"])),
        ])
    add_table(doc,
              ["Method", "Base Credence", "SI", "Scaled SI*",
               "Most Affected Fund", "Max Δ"],
              tbl_rows,
              col_widths=[1.7, 0.85, 0.6, 0.7, 1.7, 0.85])
    p_note = doc.add_paragraph(
        "* Scaled SI = SI ÷ absolute credence change (pp)."
    )
    p_note.runs[0].italic = True
    doc.add_paragraph()

    tornado_rows_agg = [
        (r["method"], float(r["sensitivity_index"]), "", 0, "", 0)
        for _, r in agg_single_sorted.iterrows()
    ]
    add_picture(doc, tornado_chart(
        tornado_rows_agg,
        "100% Credence in One Aggregation Method: SI vs baseline",
        high=SINGLE_HIGHLY_SENSITIVE, low=SINGLE_LEAST_SENSITIVE),
        width=Inches(6.5))
    doc.add_paragraph()

    doc.add_paragraph(
        "Key findings:  LexicographicMaximin produces the most extreme single-method "
        "allocation (SI 83 pp), placing nearly 96% of the budget into GiveWell. MEC "
        "and the cycle/Borda methods produce high SIs (35–40 pp) by concentrating on "
        "Navigation General. NashBargaining and Marketplace, which dominate the baseline "
        "blend, produce the smallest single-method SIs (≈15 pp each) because the "
        "baseline already reflects their preferences heavily."
    )

    return doc


# ── main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Loading data…")
    d = load()
    print("Building document…")
    doc = build_doc(d)
    out = p("sensitivity_report.docx")
    doc.save(out)
    print(f"Saved: {out}")
