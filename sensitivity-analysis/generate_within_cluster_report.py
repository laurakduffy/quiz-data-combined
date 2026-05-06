"""
generate_within_cluster_report.py
Identifies within-cluster reshuffling: tests with low cluster-level SI but
high fund-level SI, where money moves between funds inside the same cause area
without changing that area's total allocation much.

Run: python sensitivity-analysis/generate_within_cluster_report.py
"""

import os
import io
import sys
from datetime import date

try:
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Run: pip install pandas matplotlib python-docx")
    sys.exit(1)

BASE = os.path.dirname(os.path.abspath(__file__))


def p(*parts):
    return os.path.join(BASE, *parts)


# ── constants ─────────────────────────────────────────────────────────────────
GAP_THRESHOLD = 3.0   # fund_SI - cluster_SI flags within-cluster reshuffling
CA_HIGHLY     = 5
CA_LEAST      = 1

CAUSE_AREAS = ["ghd", "gcr", "aw"]
CAUSE_AREA_NAMES = {
    "ghd": "GHD (Global Health & Dev.)",
    "gcr": "GCR (Global Catastrophic Risks)",
    "aw":  "AW (Animal Welfare)",
}
CAUSE_AREA_FUNDS = {
    "ghd": ["givewell", "leaf"],
    "gcr": ["longview_ai", "longview_nuclear", "sentinel_bio"],
    "aw":  ["ea_awf", "navigation_fund_cagefree", "navigation_fund_general"],
}
FUNDS = [
    "ea_awf", "givewell", "leaf", "longview_ai",
    "longview_nuclear", "navigation_fund_cagefree",
    "navigation_fund_general", "sentinel_bio",
]
FUND_NAMES = {
    "ea_awf":                    "EA AWF",
    "givewell":                  "GiveWell",
    "leaf":                      "LEAF",
    "longview_ai":               "Longview AI",
    "longview_nuclear":          "Longview Nuclear",
    "navigation_fund_cagefree":  "Nav. Cage-Free",
    "navigation_fund_general":   "Nav. General",
    "sentinel_bio":              "Sentinel Bio",
}
FUND_TO_CA = {f: ca for ca, flist in CAUSE_AREA_FUNDS.items() for f in flist}

CAUSE_AREA_COLORS = {"ghd": "#E87B4C", "gcr": "#9B59B6", "aw": "#4C9BE8"}
DIM_COLORS = {
    "Worldview Credences":      "#1F77B4",
    "CE Multipliers":           "#2CA02C",
    "Dim. Returns (Power)":     "#FF7F0E",
    "Dim. Returns (Max Spend)": "#9467BD",
    "DR Combined":              "#D62728",
    "Aggregation Methods":      "#8C564B",
}

# ── generic helpers ───────────────────────────────────────────────────────────

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def shade_row(row, hex_color="D9D9D9"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    shade_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def add_picture(doc, stream, width=Inches(6)):
    if stream is None:
        return
    doc.add_picture(stream, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def fmt_si(v):
    return f"{v:.2f}"


def fmt_pp(v):
    return f"+{v:.2f} pp" if v > 0 else f"{v:.2f} pp"


def ca_label(k):
    return CAUSE_AREA_NAMES.get(k, k)


def fund_label(k):
    return FUND_NAMES.get(k, k)


# ── data helpers ──────────────────────────────────────────────────────────────

def baseline_fund_allocs(baseline_df):
    val_col = "weighted_combined"
    if val_col not in baseline_df.columns:
        numeric = baseline_df.select_dtypes(include="number").columns.tolist()
        val_col = numeric[0] if numeric else None
    if val_col is None:
        return {}
    return {r["fund"]: float(r[val_col]) for _, r in baseline_df.iterrows()}


def wide_fund_deltas(alloc_df, key_col, test_key, bl_key):
    bl = alloc_df[alloc_df[key_col] == bl_key]
    tr = alloc_df[alloc_df[key_col] == test_key]
    if bl.empty or tr.empty:
        return {}
    bl_row = bl.iloc[0]
    tr_row = tr.iloc[0]
    return {f: float(tr_row[f]) - float(bl_row[f])
            for f in FUNDS if f in alloc_df.columns}


def wide_fund_deltas_from_dict(alloc_df, key_col, test_key, bl_dict):
    tr = alloc_df[alloc_df[key_col] == test_key]
    if tr.empty:
        return {}
    tr_row = tr.iloc[0]
    return {f: float(tr_row[f]) - bl_dict.get(f, 0)
            for f in FUNDS if f in alloc_df.columns}


def long_ce_deltas(ce_alloc, fund_varied, multiplier):
    rows = ce_alloc[
        (ce_alloc["fund_varied"] == fund_varied) &
        (ce_alloc["multiplier"].astype(float) == float(multiplier))
    ]
    return {r["recipient_fund"]: float(r["allocation_diff_pp"])
            for _, r in rows.iterrows() if r["recipient_fund"] in FUNDS}


def long_wv_deltas(wv_alloc, scenario):
    rows = wv_alloc[wv_alloc["scenario"] == scenario]
    return {r["project_id"]: float(r["alloc_delta"])
            for _, r in rows.iterrows() if r["project_id"] in FUNDS}


def ra_fund_deltas(ra_fund_row):
    return {f: float(ra_fund_row[f"{f}_delta"])
            for f in FUNDS if f"{f}_delta" in ra_fund_row.index}


def most_within_cluster_ca(fund_deltas):
    """Return the cause area with the largest within-cluster flow."""
    best_ca, best_flow = None, -1
    for ca, funds in CAUSE_AREA_FUNDS.items():
        abs_sum = sum(abs(fund_deltas.get(f, 0)) for f in funds)
        net = abs(sum(fund_deltas.get(f, 0) for f in funds))
        flow = abs_sum - net
        if flow > best_flow:
            best_flow, best_ca = flow, ca
    return best_ca, best_flow


# ── chart helpers ─────────────────────────────────────────────────────────────

def scatter_plot(records, title):
    """records: list of (ca_si, fund_si, dim, label)"""
    if not records:
        return None
    fig, ax = plt.subplots(figsize=(8, 6))
    all_vals = [r[0] for r in records] + [r[1] for r in records]
    max_v = max(all_vals) * 1.08 + 1
    ax.plot([0, max_v], [0, max_v], "k--", lw=0.8, alpha=0.4, label="y = x (no within-cluster gap)")
    xs = np.linspace(0, max_v, 100)
    ax.fill_between(xs, xs, max_v, alpha=0.04, color="orange")

    for dim, col in DIM_COLORS.items():
        pts = [(r[0], r[1], r[3]) for r in records if r[2] == dim]
        if not pts:
            continue
        ax.scatter([t[0] for t in pts], [t[1] for t in pts],
                   color=col, label=dim, alpha=0.75, s=45, zorder=3)
        for cx, cy, lbl in pts:
            if cy - cx >= GAP_THRESHOLD:
                ax.annotate(lbl[:28], (cx, cy), fontsize=5.5, alpha=0.85,
                            xytext=(4, 2), textcoords="offset points")

    ax.set_xlim(0, max_v)
    ax.set_ylim(0, max_v)
    ax.set_xlabel("Cluster-Level SI (pp)", fontsize=10)
    ax.set_ylabel("Fund-Level SI (pp)", fontsize=10)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.text(0.02, 0.97,
            f"Orange region: fund SI > cluster SI\n(within-cluster reshuffling, gap ≥ {GAP_THRESHOLD} pp = flagged)",
            transform=ax.transAxes, fontsize=7, va="top",
            bbox=dict(boxstyle="round,pad=0.3", fc="wheat", alpha=0.7))
    ax.legend(fontsize=7, bbox_to_anchor=(1.01, 1), loc="upper left", borderaxespad=0)
    fig.tight_layout()
    return fig_to_stream(fig)


def gap_bar(records, title, top_n=25):
    """Horizontal bar chart: gap = fund_SI - cluster_SI, sorted descending."""
    recs = sorted(records, key=lambda x: x[1] - x[0], reverse=True)[:top_n]
    labels = [r[3][:45] for r in recs]
    gaps   = [r[1] - r[0] for r in recs]
    cols   = [DIM_COLORS.get(r[2], "gray") for r in recs]

    fig, ax = plt.subplots(figsize=(8, max(3.5, len(recs) * 0.38)))
    ax.barh(range(len(labels)), gaps, color=cols, edgecolor="white")
    ax.axvline(GAP_THRESHOLD, color="red", ls="--", lw=0.9,
               label=f"Threshold = {GAP_THRESHOLD} pp")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Gap: Fund SI − Cluster SI (pp)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    patches = [mpatches.Patch(color=DIM_COLORS.get(d, "gray"), label=d)
               for d in DIM_COLORS if any(r[2] == d for r in recs)]
    patches.append(plt.Line2D([0], [0], color="red", ls="--", lw=0.9,
                               label=f"Threshold = {GAP_THRESHOLD} pp"))
    ax.legend(handles=patches, fontsize=7, loc="lower right")
    fig.tight_layout()
    return fig_to_stream(fig)


def fund_delta_chart(fund_deltas, title, flag_ca=None):
    """Fund-level bar chart grouped by cause area, with cluster background bands."""
    fund_order = []
    for ca in CAUSE_AREAS:
        fund_order.extend(CAUSE_AREA_FUNDS[ca])
    labels = [fund_label(f) for f in fund_order]
    values = [fund_deltas.get(f, 0) for f in fund_order]
    colors = [CAUSE_AREA_COLORS[FUND_TO_CA[f]] for f in fund_order]

    fig, ax = plt.subplots(figsize=(7, max(3.5, len(fund_order) * 0.55)))
    y_pos = 0
    for ca in CAUSE_AREAS:
        nf = len(CAUSE_AREA_FUNDS[ca])
        alpha = 0.18 if flag_ca == ca else 0.06
        ax.axhspan(y_pos - 0.5, y_pos + nf - 0.5,
                   color=CAUSE_AREA_COLORS[ca], alpha=alpha, zorder=0)
        if flag_ca == ca:
            ax.text(ax.get_xlim()[0] if ax.get_xlim()[0] != 0 else -0.5,
                    y_pos + nf / 2 - 0.5, "◄",
                    fontsize=10, color=CAUSE_AREA_COLORS[ca], va="center", ha="right")
        y_pos += nf

    ax.barh(range(len(labels)), values, color=colors, edgecolor="white", height=0.7)
    ax.axvline(0, color="black", lw=0.8)
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel("Allocation change vs baseline (pp)", fontsize=9)
    ax.set_title(title, fontsize=10, fontweight="bold")
    patches = [mpatches.Patch(color=CAUSE_AREA_COLORS[ca],
                               label=ca_label(ca) + (" ◄ within-cluster shift" if ca == flag_ca else ""))
               for ca in CAUSE_AREAS]
    ax.legend(handles=patches, fontsize=7, loc="lower right")
    fig.tight_layout()
    return fig_to_stream(fig)


def within_cluster_detail(doc, label, ca_si, fund_si, gap, fund_deltas, flag_ca):
    """Add a detail block for one flagged within-cluster test."""
    doc.add_paragraph(
        f"Cluster SI = {ca_si:.2f} pp  |  Fund SI = {fund_si:.2f} pp  |  Gap = {gap:.2f} pp"
    ).runs[0].bold = False

    if flag_ca:
        funds_in_ca = CAUSE_AREA_FUNDS[flag_ca]
        movers = [(f, fund_deltas.get(f, 0)) for f in funds_in_ca if abs(fund_deltas.get(f, 0)) > 0.3]
        movers.sort(key=lambda x: abs(x[1]), reverse=True)
        if movers:
            mover_str = ", ".join(f"{fund_label(f)} {fmt_pp(v)}" for f, v in movers)
            doc.add_paragraph(
                f"Within-cluster movement in {ca_label(flag_ca)}: {mover_str}."
            ).runs[0].italic = True

    add_picture(doc, fund_delta_chart(fund_deltas, label[:65] + ": fund allocation change (pp)",
                                      flag_ca=flag_ca), width=Inches(5.5))


# ── load ──────────────────────────────────────────────────────────────────────

def load():
    def read(path):
        if not os.path.exists(path):
            raise FileNotFoundError(
                f"Missing: {path}\n"
                "Run all sensitivity analysis scripts first."
            )
        return pd.read_csv(path)

    d = {}
    d["baseline"]   = read(p("outputs", "baseline_staged.csv"))

    d["wv_fund"]    = read(p("worldview-sensitivity",  "outputs", "split_credences_index.csv"))
    d["wv_ca"]      = read(p("worldview-sensitivity",  "outputs", "cause_area_index.csv"))
    d["wv_alloc"]   = read(p("worldview-sensitivity",  "outputs", "split_credences_by_fund.csv"))

    d["ce_fund"]    = read(p("across-the-board",       "outputs", "ce_multiplier_si.csv"))
    d["ce_ca"]      = read(p("across-the-board",       "outputs", "cause_area_si.csv"))
    d["ce_alloc"]   = read(p("across-the-board",       "outputs", "ce_multiplier_allocations.csv"))

    d["dr_fund"]    = read(p("diminishing-returns",    "outputs", "dr_sensitivity_index.csv"))
    d["dr_ca"]      = read(p("diminishing-returns",    "outputs", "dr_sensitivity_cause_area_index.csv"))
    d["dr_alloc"]   = read(p("diminishing-returns",    "outputs", "dr_sensitivity_allocations.csv"))

    d["ms_fund"]    = read(p("diminishing-returns",    "outputs", "max_spend_sensitivity_index.csv"))
    d["ms_ca"]      = read(p("diminishing-returns",    "outputs", "max_spend_cause_area_index.csv"))
    d["ms_alloc"]   = read(p("diminishing-returns",    "outputs", "max_spend_sensitivity_allocations.csv"))

    d["combo_fund"] = read(p("diminishing-returns",    "outputs", "combo_max_spend_index.csv"))
    d["combo_ca"]   = read(p("diminishing-returns",    "outputs", "combo_max_spend_cause_area_index.csv"))
    d["combo_alloc"]= read(p("diminishing-returns",    "outputs", "combo_max_spend_allocations.csv"))

    d["agg_fund"]   = read(p("aggregation-methods",    "outputs", "split_credences_index.csv"))
    d["agg_ca"]     = read(p("aggregation-methods",    "outputs", "cause_area_index.csv"))
    d["agg_alloc"]  = read(p("aggregation-methods",    "outputs", "split_credences_allocations.csv"))

    d["ra_fund"]    = read(p("risk-aversion",          "outputs", "risk_aversion_summary.csv"))
    d["ra_ca"]      = read(p("risk-aversion",          "outputs", "risk_aversion_cause_area_summary.csv"))
    d["ra_alloc"]   = read(p("risk-aversion",          "outputs", "risk_aversion_by_fund.csv"))

    return d


# ── section helpers ───────────────────────────────────────────────────────────

def _comparison_table_rows(merged_df, label_col, extra_cols=None):
    """
    Build rows for the merged SI comparison table.
    merged_df must have: label_col, ca_si, fund_si, gap columns.
    extra_cols: list of (header, colname) pairs inserted before cluster SI.
    Returns (headers, rows, col_widths).
    """
    extra_cols = extra_cols or []
    headers = ([label_col] + [h for h, _ in extra_cols] +
               ["Cluster SI", "Fund SI", "Gap", "Flag"])
    col_widths = ([2.0] + [0.9] * len(extra_cols) +
                  [0.7, 0.7, 0.6, 0.35])
    # Squeeze label if many extra cols
    if len(extra_cols) >= 2:
        col_widths[0] = 1.6

    rows = []
    for _, r in merged_df.iterrows():
        row = [str(r[label_col])[:50]]
        for _, c in extra_cols:
            row.append(str(r[c]) if c in r.index else "—")
        row += [fmt_si(r["ca_si"]), fmt_si(r["fund_si"]), fmt_si(r["gap"]),
                "⚠" if r["gap"] >= GAP_THRESHOLD else ""]
        rows.append(row)
    return headers, rows, col_widths


def _flagged_detail_blocks(doc, flagged_df, label_col, get_deltas_fn, top_n=5):
    """Emit detail blocks (narrative + chart) for each flagged test."""
    if flagged_df.empty:
        return
    add_heading(doc, f"Within-Cluster Shifts — Tests with Gap ≥ {GAP_THRESHOLD} pp", 2)
    for _, r in flagged_df.head(top_n).iterrows():
        fund_deltas = get_deltas_fn(r)
        if not fund_deltas:
            continue
        flag_ca, _ = most_within_cluster_ca(fund_deltas)
        lbl = str(r[label_col])[:65]
        add_heading(doc, lbl, 3)
        within_cluster_detail(doc, lbl, r["ca_si"], r["fund_si"], r["gap"],
                               fund_deltas, flag_ca)


# ── document builder ──────────────────────────────────────────────────────────

def build_doc(d):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)

    bl_fund = baseline_fund_allocs(d["baseline"])
    bl_ca   = {ca: sum(bl_fund.get(f, 0) for f in funds)
               for ca, funds in CAUSE_AREA_FUNDS.items()}

    # ── Title ─────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("Donor Compass Sensitivity Analysis Report")
    r.bold = True; r.font.size = Pt(20)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Within-Cluster Reshuffling Analysis")
    sr.bold = True; sr.font.size = Pt(14)
    doc.add_paragraph(
        f"Generated {date.today().strftime('%B %d, %Y')}  ·  Rethink Priorities"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "This report compares cluster-level and fund-level sensitivity indices across "
        "all sensitivity dimensions. A test is flagged as a 'within-cluster shift' when "
        "the fund-level SI significantly exceeds the cluster-level SI — meaning money "
        "is moving between funds inside the same cause area, producing little net change "
        "at the cause-area level but large changes at the individual-fund level."
    )
    doc.add_paragraph(
        f"Metric used: Gap = Fund SI − Cluster SI. "
        f"Tests with Gap ≥ {GAP_THRESHOLD} pp are flagged and detailed in each section."
    )

    # Baseline table
    add_heading(doc, "Baseline Allocations", 2)
    fund_order = [f for ca in CAUSE_AREAS for f in CAUSE_AREA_FUNDS[ca]]
    bl_rows = [(fund_label(f), ca_label(FUND_TO_CA[f]), f"{bl_fund.get(f, 0):.1f}%")
               for f in fund_order]
    add_table(doc, ["Fund", "Cause Area", "Baseline (%)"], bl_rows,
              col_widths=[2.0, 2.5, 1.2])
    doc.add_paragraph()

    # ── Pre-compute merged DataFrames ─────────────────────────────────────────

    # Worldview credences (main only: bound != "single")
    wv_fund_m = d["wv_fund"][~d["wv_fund"]["bound"].isin(["single"])].copy()
    wv_ca_m   = d["wv_ca"][~d["wv_ca"]["bound"].isin(["single"])].copy()
    wv_mg = pd.merge(
        wv_fund_m[["scenario", "worldview", "bound", "credence_base",
                   "credence_scenario", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        wv_ca_m[["scenario", "sensitivity_index", "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on="scenario"
    )
    wv_mg["gap"] = wv_mg["fund_si"] - wv_mg["ca_si"]
    wv_mg = wv_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # CE multipliers (exclude baseline)
    ce_fund_nb = d["ce_fund"][d["ce_fund"]["fund_varied"] != "baseline"].copy()
    ce_ca_nb   = d["ce_ca"][d["ce_ca"]["fund_varied"] != "baseline"].copy()
    # CE cluster SI has diff_ghd/gcr/aw columns instead of most_affected_cause
    ce_ca_nb["most_affected_cause"] = ce_ca_nb[["diff_ghd", "diff_gcr", "diff_aw"]].abs().idxmax(axis=1).str.replace("diff_", "")
    ce_mg = pd.merge(
        ce_fund_nb[["fund_varied", "multiplier", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        ce_ca_nb[["fund_varied", "multiplier", "sensitivity_index",
                  "most_affected_cause"]].rename(columns={"sensitivity_index": "ca_si"}),
        on=["fund_varied", "multiplier"]
    )
    ce_mg["gap"] = ce_mg["fund_si"] - ce_mg["ca_si"]
    ce_mg["label"] = ce_mg["fund_varied"] + " ×" + ce_mg["multiplier"].astype(str)
    ce_mg = ce_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # DR power
    dr_mg = pd.merge(
        d["dr_fund"][["combo", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        d["dr_ca"][["combo", "sensitivity_index", "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on="combo"
    )
    dr_mg["gap"] = dr_mg["fund_si"] - dr_mg["ca_si"]
    dr_mg = dr_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # DR max spend
    ms_mg = pd.merge(
        d["ms_fund"][["scenario", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        d["ms_ca"][["scenario", "sensitivity_index", "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on="scenario"
    )
    ms_mg["gap"] = ms_mg["fund_si"] - ms_mg["ca_si"]
    ms_mg = ms_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # DR combo (join on combo + max_spend_multiplier)
    combo_mg = pd.merge(
        d["combo_fund"][["combo", "max_spend_multiplier", "scenario",
                         "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        d["combo_ca"][["combo", "max_spend_multiplier", "sensitivity_index",
                       "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on=["combo", "max_spend_multiplier"]
    )
    combo_mg["gap"]   = combo_mg["fund_si"] - combo_mg["ca_si"]
    combo_mg["label"] = (combo_mg["combo"] + " ×" +
                         combo_mg["max_spend_multiplier"].astype(str))
    combo_mg = combo_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # Aggregation methods (main only)
    agg_fund_m = d["agg_fund"][~d["agg_fund"]["bound"].isin(["single"])].copy()
    agg_ca_m   = d["agg_ca"][~d["agg_ca"]["bound"].isin(["single"])].copy()
    agg_mg = pd.merge(
        agg_fund_m[["scenario", "method", "bound", "credence_base",
                    "credence_scenario", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        agg_ca_m[["scenario", "sensitivity_index", "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on="scenario"
    )
    agg_mg["gap"] = agg_mg["fund_si"] - agg_mg["ca_si"]
    agg_mg = agg_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # Risk aversion (illustrative)
    ra_mg = pd.merge(
        d["ra_fund"][["test", "sensitivity_index"]].rename(
            columns={"sensitivity_index": "fund_si"}),
        d["ra_ca"][["test", "sensitivity_index", "most_affected_cause"]].rename(
            columns={"sensitivity_index": "ca_si"}),
        on="test"
    )
    ra_mg["gap"] = ra_mg["fund_si"] - ra_mg["ca_si"]
    ra_mg = ra_mg.sort_values("gap", ascending=False).reset_index(drop=True)

    # ── Build global records list for scatter + gap bar ───────────────────────
    all_records = []  # (ca_si, fund_si, dim, label)

    for _, r in wv_mg.iterrows():
        lbl = f"{r['worldview'][:35]} ({r['bound']})"
        all_records.append((r["ca_si"], r["fund_si"], "Worldview Credences", lbl))

    for _, r in ce_mg.iterrows():
        all_records.append((r["ca_si"], r["fund_si"], "CE Multipliers", r["label"]))

    for _, r in dr_mg.iterrows():
        all_records.append((r["ca_si"], r["fund_si"], "Dim. Returns (Power)", r["combo"]))

    for _, r in ms_mg.iterrows():
        all_records.append((r["ca_si"], r["fund_si"], "Dim. Returns (Max Spend)", r["scenario"]))

    for _, r in combo_mg.iterrows():
        all_records.append((r["ca_si"], r["fund_si"], "DR Combined", r["label"]))

    for _, r in agg_mg.iterrows():
        lbl = f"{r['method']} ({r['bound']})"
        all_records.append((r["ca_si"], r["fund_si"], "Aggregation Methods", lbl))

    # ── Executive summary charts ──────────────────────────────────────────────
    add_heading(doc, "Cluster SI vs Fund SI — All Main-Analysis Tests", 2)
    doc.add_paragraph(
        "Each point represents one test. Points above the diagonal (y = x) have "
        "higher fund-level SI than cluster-level SI, indicating within-cluster "
        "reshuffling. Labeled points have a gap ≥ " + str(GAP_THRESHOLD) + " pp."
    )
    add_picture(doc,
                scatter_plot(all_records, "Fund-Level SI vs Cluster-Level SI: all tests"),
                width=Inches(6.5))
    doc.add_paragraph()

    # Gap bar chart
    add_heading(doc, f"Top Tests by Gap (Fund SI − Cluster SI)", 2)
    add_picture(doc, gap_bar(all_records,
                             "Gap = Fund SI − Cluster SI, top 25 tests",
                             top_n=25),
                width=Inches(6.5))
    doc.add_paragraph()

    # Summary table: top within-cluster shifts
    flagged_all = [(r[0], r[1], r[2], r[3]) for r in all_records if r[1] - r[0] >= GAP_THRESHOLD]
    flagged_sorted = sorted(flagged_all, key=lambda x: x[1] - x[0], reverse=True)
    n_flagged = len(flagged_all)
    n_total   = len(all_records)

    add_heading(doc, f"Summary: {n_flagged} Flagged Tests (Gap ≥ {GAP_THRESHOLD} pp)", 2)
    if flagged_sorted:
        tbl_rows = [
            [lbl[:55], dim[:25], fmt_si(ca_si), fmt_si(fund_si), fmt_si(fund_si - ca_si)]
            for ca_si, fund_si, dim, lbl in flagged_sorted[:30]
        ]
        add_table(doc,
                  ["Test", "Dimension", "Cluster SI", "Fund SI", "Gap"],
                  tbl_rows,
                  col_widths=[2.5, 1.5, 0.75, 0.75, 0.65])
        doc.add_paragraph(
            f"{n_flagged} of {n_total} main-analysis tests ({100 * n_flagged / n_total:.0f}%) "
            f"show a within-cluster gap ≥ {GAP_THRESHOLD} pp."
        )
    else:
        doc.add_paragraph(f"No tests found with gap ≥ {GAP_THRESHOLD} pp.")

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — WORLDVIEW CREDENCES
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Worldview Credences", 1)
    doc.add_paragraph(
        "Each worldview's credence is shifted one at a time to its low or high "
        "uncertainty bound. A within-cluster gap here means the credence shift "
        "primarily reallocates money between funds within the same cause area "
        "(e.g., GCR funds rearranging among Longview AI, Nuclear, Sentinel) "
        "rather than shifting money between GHD, GCR, and AW."
    )

    # Comparison table
    add_heading(doc, "Cluster SI vs Fund SI — All Worldview Credence Tests", 2)
    wv_tbl = wv_mg.copy()
    wv_tbl["credence_range"] = (wv_tbl["credence_base"].map(lambda x: f"{x:.0%}") +
                                 "→" +
                                 wv_tbl["credence_scenario"].map(lambda x: f"{x:.0%}"))
    tbl_rows = []
    for _, r in wv_tbl.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["worldview"][:42],
            r["bound"],
            r["credence_range"],
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Worldview", "Bound", "Credence Range", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[2.05, 0.5, 0.9, 0.7, 0.7, 0.6, 0.3])
    doc.add_paragraph(
        "Note: A '⚠' in the last column indicates Gap ≥ " + str(GAP_THRESHOLD) + " pp "
        "(within-cluster shift). Cluster SI and Fund SI can differ even without "
        "within-cluster movement if a cause area's net flow is concentrated in one fund."
    )

    # Flagged detail blocks
    flagged_wv = wv_mg[wv_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_wv.empty:
        def get_wv_deltas(row):
            return long_wv_deltas(d["wv_alloc"], row["scenario"])
        _flagged_detail_blocks(doc, flagged_wv, "worldview", get_wv_deltas, top_n=5)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — COST-EFFECTIVENESS MULTIPLIERS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Cost-Effectiveness Multipliers", 1)
    doc.add_paragraph(
        "Each fund's cost-effectiveness score is multiplied by a constant factor. "
        "The key within-cluster effect: scaling a fund within GHD (e.g. LEAF) "
        "changes which GHD fund 'wins' but leaves GHD's total allocation largely "
        "unchanged. Cluster-level SI stays low while fund-level SI can be high."
    )

    # Comparison table
    add_heading(doc, "Cluster SI vs Fund SI — All CE Multiplier Tests", 2)
    tbl_rows = []
    for _, r in ce_mg.iterrows():
        tbl_rows.append([
            r["fund_varied"],
            f"{r['multiplier']}×",
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Fund Varied", "Multiplier", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[1.5, 0.65, 0.7, 0.7, 0.6, 0.3])

    # Flagged detail blocks
    flagged_ce = ce_mg[ce_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_ce.empty:
        def get_ce_deltas(row):
            return long_ce_deltas(d["ce_alloc"], row["fund_varied"], row["multiplier"])
        _flagged_detail_blocks(doc, flagged_ce, "label", get_ce_deltas, top_n=6)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — DIMINISHING RETURNS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Diminishing Returns", 1)

    # 3a — DR Power Parameters
    add_heading(doc, "3a. DR Power Parameters", 2)
    doc.add_paragraph(
        "DR power curves control how steeply returns diminish as spending increases. "
        "Altering curves for GCR or AW can reshape fund shares within those clusters "
        "without necessarily moving money between clusters."
    )

    tbl_rows = []
    for _, r in dr_mg.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["combo"],
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["DR Power Scenario", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[2.3, 0.75, 0.75, 0.65, 0.35])

    flagged_dr = dr_mg[dr_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_dr.empty:
        def get_dr_deltas(row):
            return wide_fund_deltas(d["dr_alloc"], "combo", row["combo"], "baseline")
        _flagged_detail_blocks(doc, flagged_dr, "combo", get_dr_deltas, top_n=4)

    doc.add_paragraph()

    # 3b — Max Spending Cap
    add_heading(doc, "3b. Maximum Spending Cap", 2)
    doc.add_paragraph(
        "Three cap multipliers (2.5×, 7.5×, 10×) limit maximum fund-level spending. "
        "A tighter cap can cause large shifts between funds in the same cause area "
        "if one fund was the main beneficiary of uncapped spending."
    )

    # baseline for ms_alloc is "baseline_5x"
    ms_bl_key = "baseline_5x"
    tbl_rows = []
    for _, r in ms_mg.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["scenario"],
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Scenario", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[1.8, 0.75, 0.75, 0.65, 0.35])

    flagged_ms = ms_mg[ms_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_ms.empty:
        def get_ms_deltas(row):
            return wide_fund_deltas(d["ms_alloc"], "scenario", row["scenario"], ms_bl_key)
        _flagged_detail_blocks(doc, flagged_ms, "scenario", get_ms_deltas, top_n=4)

    doc.add_paragraph()

    # 3c — Combined DR Power × Max Spend
    add_heading(doc, "3c. Combined: DR Power × Maximum Spending Cap", 2)
    doc.add_paragraph(
        "24 joint tests. The interaction of power curves and spending caps can "
        "amplify within-cluster shifts beyond what either dimension alone produces."
    )

    # baseline row in combo_alloc is where scenario starts with "baseline"
    combo_bl_rows = d["combo_alloc"][d["combo_alloc"]["scenario"].str.startswith("baseline")]
    combo_bl_key = combo_bl_rows["scenario"].iloc[0] if not combo_bl_rows.empty else None

    tbl_rows = []
    for _, r in combo_mg.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["combo"],
            f"{r['max_spend_multiplier']}×",
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Power Scenario", "Cap Mult.", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[2.0, 0.65, 0.7, 0.7, 0.6, 0.3])

    if combo_bl_key:
        flagged_combo = combo_mg[combo_mg["gap"] >= GAP_THRESHOLD].copy()
        if not flagged_combo.empty:
            def get_combo_deltas(row):
                return wide_fund_deltas(d["combo_alloc"], "scenario",
                                        row["scenario"], combo_bl_key)
            _flagged_detail_blocks(doc, flagged_combo, "label", get_combo_deltas, top_n=4)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — AGGREGATION METHOD WEIGHTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Aggregation Method Weights", 1)
    doc.add_paragraph(
        "Each aggregation method's credence is shifted to its low or high bound. "
        "Some methods favour particular funds within a cause area without substantially "
        "changing the cause area's total budget share."
    )

    agg_tbl = agg_mg.copy()
    agg_tbl["credence_range"] = (
        agg_tbl["credence_base"].map(lambda x: f"{x:.0%}") + "→" +
        agg_tbl["credence_scenario"].map(lambda x: f"{x:.0%}")
    )
    tbl_rows = []
    for _, r in agg_tbl.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["method"],
            r["bound"],
            r["credence_range"],
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Method", "Bound", "Credence Range", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[1.4, 0.5, 0.9, 0.7, 0.7, 0.6, 0.3])

    flagged_agg = agg_mg[agg_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_agg.empty:
        def get_agg_deltas(row):
            return wide_fund_deltas_from_dict(
                d["agg_alloc"], "scenario", row["scenario"], bl_fund)
        _flagged_detail_blocks(doc, flagged_agg, "scenario", get_agg_deltas, top_n=4)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — ILLUSTRATIVE: RISK AVERSION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Illustrative: Risk Aversion", 1)
    doc.add_paragraph(
        "Risk-aversion scenarios place extreme credence in a single risk profile. "
        "These are presented separately as illustrative analyses. Because the shifts "
        "are often total reallocation toward one cause area, the gap here reflects "
        "whether individual funds within a favoured cause area also redistribute."
    )

    tbl_rows = []
    for _, r in ra_mg.sort_values("gap", ascending=False).iterrows():
        tbl_rows.append([
            r["test"],
            fmt_si(r["ca_si"]),
            fmt_si(r["fund_si"]),
            fmt_si(r["gap"]),
            "⚠" if r["gap"] >= GAP_THRESHOLD else "",
        ])
    add_table(doc,
              ["Test", "Cluster SI", "Fund SI", "Gap", ""],
              tbl_rows,
              col_widths=[2.8, 0.75, 0.75, 0.65, 0.35])

    flagged_ra = ra_mg[ra_mg["gap"] >= GAP_THRESHOLD].copy()
    if not flagged_ra.empty:
        add_heading(doc, f"Within-Cluster Shifts — Tests with Gap ≥ {GAP_THRESHOLD} pp", 2)
        for _, r in flagged_ra.head(4).iterrows():
            ra_fund_row = d["ra_fund"][d["ra_fund"]["test"] == r["test"]]
            if ra_fund_row.empty:
                continue
            fund_deltas = ra_fund_deltas(ra_fund_row.iloc[0])
            flag_ca, _ = most_within_cluster_ca(fund_deltas)
            add_heading(doc, r["test"], 3)
            within_cluster_detail(doc, r["test"], r["ca_si"], r["fund_si"],
                                   r["gap"], fund_deltas, flag_ca)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # CLOSING: ROBUSTLY INSENSITIVE TESTS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Robustly Insensitive Tests (Cluster SI < 1 and Fund SI < 2)", 1)
    doc.add_paragraph(
        "The following tests show low sensitivity at both the cluster and fund level — "
        "the allocation is genuinely robust to these parameter changes."
    )

    robust = [(r[0], r[1], r[2], r[3]) for r in all_records
              if r[0] < CA_LEAST and r[1] < 2.0]
    if robust:
        by_dim: dict = {}
        for ca_si, fund_si, dim, lbl in robust:
            by_dim.setdefault(dim, []).append((ca_si, fund_si, lbl))

        for dim, tests in sorted(by_dim.items()):
            p_obj = doc.add_paragraph()
            p_obj.add_run(dim).bold = True
            for ca_si, fund_si, lbl in sorted(tests, key=lambda x: x[1]):
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(f"Cluster SI = {ca_si:.2f}  |  Fund SI = {fund_si:.2f}  ·  {lbl[:70]}")
    else:
        doc.add_paragraph("No tests found meeting both criteria.")

    doc.add_paragraph()

    # Within-cluster shifters that also have low cluster SI
    add_heading(doc, "Within-Cluster Shifters with Low Cluster SI (SI < 2)", 1)
    doc.add_paragraph(
        "These tests appear nearly stable in a cluster-only view (cluster SI < 2) "
        "but produce meaningful fund-level reallocation (fund SI ≥ 2). "
        "Donors targeting specific funds should treat these as sensitive."
    )

    hidden = [(r[0], r[1], r[2], r[3]) for r in all_records
              if r[0] < 2.0 and r[1] >= 2.0]
    if hidden:
        by_dim2: dict = {}
        for ca_si, fund_si, dim, lbl in hidden:
            by_dim2.setdefault(dim, []).append((ca_si, fund_si, lbl))

        for dim, tests in sorted(by_dim2.items()):
            p_obj = doc.add_paragraph()
            p_obj.add_run(dim).bold = True
            for ca_si, fund_si, lbl in sorted(tests, key=lambda x: -(x[1] - x[0])):
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(
                    f"Cluster SI = {ca_si:.2f}  |  Fund SI = {fund_si:.2f}  "
                    f"|  Gap = {fund_si - ca_si:.2f}  ·  {lbl[:65]}"
                )
    else:
        doc.add_paragraph("No tests found meeting these criteria.")

    return doc


# ── main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Loading data…")
    d = load()
    print("Building document…")
    doc = build_doc(d)
    out = p("sensitivity_within_cluster_report.docx")
    doc.save(out)
    print(f"Saved: {out}")
