"""Generate a Word document describing the sensitivity analysis methodology.

Produces sensitivity-analysis/sensitivity_methodology.docx containing:
  Part I  — Fixed one-pager on the three output metrics and why we use them.
  Part II — Dynamic summary of every scenario: parameter modified, magnitude
            of change, and rationale.

NOTE: Scenario metadata is defined directly in this file (SCENARIO_METADATA
below) rather than imported from run_sensitivity.py, to avoid pulling in the
full model import chain. Keep SCENARIO_METADATA in sync with the SCENARIOS
dict in run_sensitivity.py when adding or removing scenarios.

Requires python-docx:
    pip install python-docx

Usage:
    cd sensitivity-analysis
    python generate_methodology_doc.py
"""

from datetime import date
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("python-docx is required.  Install it with:  pip install python-docx")
    sys.exit(1)

SCRIPT_DIR = Path(__file__).parent

# ---------------------------------------------------------------------------
# Scenario metadata
# Mirrors SCENARIOS in run_sensitivity.py — keep in sync when editing that file.
#
# Each entry: (scenario_id, parameter_label, base_distribution, perturbation, rationale)
# ---------------------------------------------------------------------------

# Each entry: (scenario_id, parameter_label, base_distribution, perturbation,
#              perturbation_ratio, rationale)
# perturbation_ratio: numeric factor (e.g. 100) for scalar shifts;
#                     "N/A (categorical)" for binary/structural changes.
SCENARIO_METADATA = [
    (
        "r_inf_100x_up",
        "Long-run x-risk floor  (r_inf)",
        "loguniform CI [1×10⁻⁸, 5×10⁻⁵]",
        "100× higher → CI [1×10⁻⁶, 5×10⁻³]",
        "100",
        "Tests how results change if the long-run background annual extinction risk is "
        "100× higher than our central estimate. A higher floor risk means extinction "
        "remains likely even in the very long run, which reduces the value of pushing "
        "risk far into the future and compresses the relative advantage of long-term "
        "x-risk funds over near-term ones.",
    ),
    (
        "r_inf_100x_down",
        "Long-run x-risk floor  (r_inf)",
        "loguniform CI [1×10⁻⁸, 5×10⁻⁵]",
        "100× lower → CI [1×10⁻¹⁰, 5×10⁻⁷]",
        "100",
        "Symmetric downward test. With a near-zero floor, the long-run future is much "
        "more recoverable, increasing the leverage of any near-term risk reduction and "
        "raising the relative value of x-risk funds compared to the base case.",
    ),
    (
        "no_cubic_growth",
        "Probability of stellar expansion  (p_cubic_growth)",
        "beta CI [0.01, 0.15], mean ≈ 0.065",
        "Degenerate: p_cubic_growth = 0 (stellar expansion never occurs)",
        "N/A (categorical)",
        "A 'terrestrial only' scenario: humanity never expands beyond Earth. This is "
        "the single most important structural assumption in the model — removing cubic "
        "growth eliminates the dominant source of long-run value and sharply reduces "
        "the expected benefit of x-risk prevention relative to near-term welfare "
        "interventions.",
    ),
    (
        "s_10x_faster",
        "Stellar settlement speed  (s, ly/yr)",
        "loguniform CI [4×10⁻⁵, 1×10⁻²]",
        "10× faster → CI [4×10⁻⁴, 0.1]",
        "10",
        "Faster settlement means the long-run value of the future accumulates sooner, "
        "raising the present value of any intervention that preserves access to that "
        "future.",
    ),
    (
        "s_10x_slower",
        "Stellar settlement speed  (s, ly/yr)",
        "loguniform CI [4×10⁻⁵, 1×10⁻²]",
        "10× slower → CI [4×10⁻⁶, 1×10⁻³]",
        "10",
        "The bulk of cosmic value is delayed further into the future, reducing the "
        "present value of risk-reduction interventions that act on near-term timescales.",
    ),
    (
        "cause_fractions_bio_nuclear_higher",
        "Cause fractions of x-risk",
        "Dirichlet means [bio=0.03, nuclear=0.03, AI=0.90, other=0.04]",
        "bio → 0.30, nuclear → 0.30, AI → 0.36  (~10× shift for bio and nuclear)",
        "10",
        "Tests the allocation sensitivity to the assumed share of total x-risk "
        "attributable to each cause. Since each fund's expected impact is proportional "
        "to its cause fraction, this directly tests whether the AI fund's dominance is "
        "driven by the AI cause-fraction assumption rather than its cost-effectiveness.",
    ),
    (
        "rel_risk_10x_up",
        "Relative cost-effectiveness  (rel_risk_reduction, all GCR funds)",
        "loguniform CI scaled by fund budget",
        "10× higher for all three GCR funds",
        "10",
        "Tests whether the GCR funds' absolute cost-effectiveness is high enough to "
        "dominate the portfolio when assumed to be 10× more effective. If even a 10× "
        "improvement does not substantially shift the allocation, GCR funds are near a "
        "score ceiling relative to diminishing returns. If it does shift dramatically, "
        "cost-effectiveness assumptions are the key bottleneck.",
    ),
    (
        "rel_risk_10x_down",
        "Relative cost-effectiveness  (rel_risk_reduction, all GCR funds)",
        "loguniform CI scaled by fund budget",
        "10× lower for all three GCR funds",
        "10",
        "Identifies the threshold below which GCR funds lose out to near-term "
        "alternatives, and how rapidly the portfolio rebalances toward GiveWell, "
        "LEAF, and animal welfare funds.",
    ),
    (
        "rel_risk_100x_down",
        "Relative cost-effectiveness  (rel_risk_reduction, all GCR funds)",
        "loguniform CI scaled by fund budget",
        "100× lower for all three GCR funds",
        "100",
        "Extends rel_risk_10x_down by a further order of magnitude. If GCR funds "
        "still hold appreciable allocation under a 100× cost-effectiveness reduction, "
        "it signals that the longtermist component of the worldview mixture is "
        "sufficiently dominant to sustain GCR investment even at very low effectiveness.",
    ),
    (
        "p_zero_5x_lower",
        "Outcome probabilities  (p_harm / p_zero / p_positive)",
        "Dirichlet means — Sentinel/Nuclear: [harm=0.05, zero=0.50, pos=0.45]; "
        "AI: [harm=0.15, zero=0.50, pos=0.35]",
        "p_zero ÷5 (0.50→0.10); surplus redistributed proportionally to p_harm and p_positive",
        "5",
        "Tests the scenario in which interventions are much less likely to have zero "
        "impact, with released probability mass redistributed proportionally between "
        "harm and positive outcomes. This represents greater certainty that the "
        "intervention does something — but in proportion to prior beliefs about "
        "whether that something is good or bad.",
    ),
    (
        "near_pessimistic_outcomes",
        "Outcome probabilities  (p_harm / p_zero / p_positive)",
        "Dirichlet means — Sentinel/Nuclear: [harm=0.05, zero=0.50, pos=0.45]; "
        "AI: [harm=0.15, zero=0.50, pos=0.35]",
        "All funds: p_harm=0.225, p_zero=0.50, p_positive=0.275 → gap = 0.05 pp",
        "N/A (categorical)",
        "Stress-tests sign uncertainty. In the base case the net positive signal "
        "(p_positive − p_harm) is ~0.40 for Sentinel/Nuclear and ~0.20 for AI. "
        "This scenario narrows the gap to 0.05 for all funds, making outcomes nearly "
        "as likely to be harmful as beneficial, while holding p_zero constant.",
    ),
]


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _shade_row(row, hex_color="E8EAF6"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _set_col_widths(table, widths_inches):
    for i, width in enumerate(widths_inches):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)


def _add_metric_block(doc, title, formula, description, rationale):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Formula:  {formula}")
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph(description)

    p3 = doc.add_paragraph()
    p3.add_run("Why this metric: ").bold = True
    p3.add_run(rationale)
    p3.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("GCR Sensitivity Analysis: Methodology & Scenarios", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Rethink Priorities  ·  Generated {date.today().strftime('%B %d, %Y')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PART I — Output Metrics
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("Part I — Output Metrics", 1)

    doc.add_paragraph(
        "For every scenario we report five metrics. The first two summarise the whole "
        "portfolio in one or two numbers; the remaining three break down what happened "
        "fund-by-fund and perspective-by-perspective. Together they answer: how much "
        "did the recommended donation split change, why did it change, and did anything "
        "cross a meaningful threshold?"
    )
    doc.add_paragraph()

    # Metric 1
    _add_metric_block(
        doc,
        title="1.  Portfolio shift  (headline — one number per scenario)",
        formula="portfolio shift  =  ( Σ |new allocation% − old allocation%| )  ÷  2",
        description=(
            "For each fund, we take the absolute difference between its allocation "
            "percentage before and after the parameter change, sum those differences "
            "across all funds, then divide by two. The result is in percentage points "
            "(pp).\n\n"
            "Example: if the scenario moves 15% of the budget from GiveWell to "
            "Sentinel Bio and nothing else changes, the portfolio shift = 15 pp."
        ),
        rationale=(
            "Because the allocations across all funds must always add up to 100%, "
            "every percentage point gained by one fund is lost by others. The raw sum "
            "of all changes therefore double-counts every transfer; dividing by two "
            "gives the actual amount of money that moved. This produces a single, "
            "intuitive number: 'X% of the total donation budget was reallocated.'\n\n"
            "Why this is the primary metric: The size of each scenario's parameter "
            "change is not arbitrary — it is set to represent a realistic range of "
            "expert uncertainty. So the portfolio shift directly answers the decision- "
            "relevant question: 'If our assumptions about this parameter are wrong by "
            "this much, how different would a donor's recommended split be?' That is "
            "what a donor needs to know.\n\n"
            "Why the per-order-of-magnitude version (Metric 1b) is also reported: "
            "Some scenarios test a 100× change in a parameter; others test only a 10× "
            "change. All else equal, larger perturbations produce larger portfolio "
            "shifts, making it hard to compare parameter importance across scenarios. "
            "Dividing by the size of the perturbation (on a log scale) puts scenarios "
            "on a common footing — expressing sensitivity as 'how many percentage "
            "points shift per factor-of-10 change in this parameter.' Both versions "
            "are reported; neither should be used alone. The per-OOM version is left "
            "blank for scenarios that cannot be expressed as a factor change (e.g. "
            "'assume stellar expansion never happens')."
        ),
    )

    # Metric 1b
    _add_metric_block(
        doc,
        title="1b.  Portfolio shift per order of magnitude  (secondary)",
        formula=(
            "shift per OOM  =  portfolio shift  ÷  log₁₀(perturbation factor)\n"
            "e.g. for a 100× change:  shift per OOM  =  portfolio shift  ÷  2"
        ),
        description=(
            "The portfolio shift (Metric 1) divided by the base-10 logarithm of the "
            "perturbation factor. Units: percentage points per order of magnitude "
            "(pp/OOM). A value of 8 pp/OOM means that each factor-of-10 change in "
            "the parameter is associated with 8 pp of budget reallocation.\n\n"
            "Left blank for scenarios that represent a categorical change rather than "
            "a scalar one."
        ),
        rationale=(
            "See discussion under Metric 1. Reported alongside the raw portfolio shift "
            "— not instead of it — because the actual uncertainty ranges matter for "
            "real decisions."
        ),
    )

    # Metric 2
    _add_metric_block(
        doc,
        title="2.  Score magnitude change, per perspective  (explains the mechanism)",
        formula=(
            "magnitude change  =  log₁₀(|score after|) − log₁₀(|score before|)\n"
            "equivalently:      =  log₁₀( |score after| ÷ |score before| )"
        ),
        description=(
            "For each fund and each ethical perspective, we measure how many orders "
            "of magnitude the fund's cost-effectiveness score changed. A value of 1 "
            "means the score became 10× larger in magnitude; 0.3 ≈ 2×; −1 means it "
            "shrank to one-tenth its original size. The sign of the score (whether "
            "the fund looks beneficial or harmful) is reported separately in Metric 2b.\n\n"
            "This is computed three times per fund — once under each ethical "
            "perspective — so a reader can see whether a scenario affects all "
            "perspectives equally or only some."
        ),
        rationale=(
            "Fund scores in this model range across many orders of magnitude. Under "
            "perspectives that place high weight on the very long-run future, a fund "
            "that prevents extinction may score trillions of times higher than one "
            "that improves animal welfare. A simple percentage change would be "
            "meaningless at these scales. Log-scale differences are the natural "
            "unit: a change of 1 always means a factor of 10, regardless of whether "
            "the baseline score is 100 or 10¹⁸.\n\n"
            "We use the absolute value of scores before taking the logarithm because "
            "some ethical perspectives assign negative scores to GCR funds (reflecting "
            "risk aversion about outcomes that could also be harmful). Without this, "
            "the calculation would be undefined whenever a score is negative or zero, "
            "and the sign change would be hidden inside the magnitude metric. Keeping "
            "them separate makes both signals unambiguous.\n\n"
            "We report one value per ethical perspective rather than a single weighted "
            "average, because collapsing across perspectives is itself a value judgment "
            "we do not want to embed in the analysis. A parameter change might shift "
            "scores dramatically for someone who cares about the long-run future but "
            "be completely irrelevant to someone focused on near-term animal welfare — "
            "and a weighted average would hide that distinction."
        ),
    )

    # Metric 2b
    _add_metric_block(
        doc,
        title="2b.  Score direction, per perspective  (captures sign flips)",
        formula='direction  =  "<before sign> / <after sign>"   where sign is + or −',
        description=(
            "For each fund and each ethical perspective, we record whether the "
            "fund's score was positive or negative before the scenario, and positive "
            "or negative after. The four possible values are:\n"
            "  +/+  score was positive and stayed positive\n"
            "  +/−  score flipped from positive to negative\n"
            "  −/+  score flipped from negative to positive\n"
            "  −/−  score was negative and stayed negative\n\n"
            "A score of exactly zero is marked 0 (rare in practice)."
        ),
        rationale=(
            "A fund whose score changes sign has crossed a fundamental threshold: "
            "it has gone from looking beneficial to looking harmful (or vice versa) "
            "under that ethical perspective. This is a qualitatively different kind "
            "of change from a shift in magnitude — and it cannot be detected from "
            "the magnitude metric alone, which works on absolute values. Reporting "
            "direction separately ensures that both 'how much did the score change?' "
            "and 'did it cross from good to bad?' are always clearly visible."
        ),
    )

    # Metric 3
    _add_metric_block(
        doc,
        title="3.  Ranking change, per perspective  (detects threshold crossings)",
        formula=(
            "ranking change  =  rank after scenario − rank before scenario\n"
            "(rank 1 = highest-scoring fund;  positive value = fund fell in ranking)"
        ),
        description=(
            "For each ethical perspective, funds are ranked by their cost-effectiveness "
            "score from highest (rank 1) to lowest. We report how each fund's rank "
            "changed between the base case and the scenario. A positive number means "
            "the fund fell; a negative number means it rose."
        ),
        rationale=(
            "The allocation model tends to concentrate funding: the highest-ranked "
            "fund receives money first, then the next-highest, and so on, with each "
            "fund receiving less as more money flows to it. A fund that slips from "
            "second place to third — even if its score barely changed — may drop from "
            "'receives 20% of the budget' to 'receives nothing,' because the fund "
            "that overtook it now absorbs all the marginal dollars.\n\n"
            "Rank change separates two distinct reasons a fund's allocation might "
            "shift: (a) its score changed in magnitude, or (b) the scores of other "
            "funds changed and it was overtaken. The magnitude metric (Metric 2) "
            "captures (a); rank change captures (b). Together they explain whether "
            "a portfolio shift was driven by a fund improving in absolute terms, or "
            "by the competitive landscape changing around it.\n\n"
            "Ranks are based on scores alone — before the allocation process applies "
            "diminishing returns — so rank reflects pure cost-effectiveness ordering, "
            "not the final funding outcome."
        ),
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PART II — Parameter Scenarios
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("Part II — Parameter Scenarios", 1)

    doc.add_paragraph(
        f"The following {len(SCENARIO_METADATA)} scenarios each test what happens "
        "when one group of assumptions is changed while everything else is held fixed "
        "at its central estimate. The size of each change is chosen to be large "
        "enough to produce a clear signal — typically a factor-of-10 or factor-of-100 "
        "shift for quantities that vary continuously, or a complete redistribution of "
        "probability for quantities that sum to one. The 'Ratio' column records the "
        "factor by which the parameter was shifted, used to compute the per-OOM "
        "sensitivity metric."
    )
    doc.add_paragraph()

    # Table
    col_headers = ["Scenario ID", "Parameter", "Base distribution",
                   "Perturbation", "Ratio", "Rationale"]
    col_widths  = [1.2, 1.2, 1.4, 1.3, 0.65, 2.25]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    _shade_row(hdr_row, "1A237E")
    for i, hdr in enumerate(col_headers):
        cell = hdr_row.cells[i]
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    for row_idx, (sc_id, param, base, perturb, ratio, rationale) in enumerate(SCENARIO_METADATA):
        data_row = table.add_row()
        _shade_row(data_row, "F3F4FE" if row_idx % 2 == 0 else "FFFFFF")
        for i, val in enumerate([sc_id, param, base, perturb, ratio, rationale]):
            cell = data_row.cells[i]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    if i == 0:
                        run.font.name = "Courier New"

    _set_col_widths(table, col_widths)
    doc.add_paragraph()

    # ── Ethical perspectives used for comparison ─────────────────────────
    doc.add_heading("Ethical Perspectives Used in the Comparison", 2)
    doc.add_paragraph(
        "Rather than assuming a single set of values, the analysis runs under three "
        "distinct ethical perspectives simultaneously. The budget is divided across "
        "perspectives in proportion to their credence weights. At each $2M step, each "
        "perspective independently directs its share of the money to whichever fund "
        "looks most cost-effective from that perspective's point of view. The final "
        "allocation reflects all three perspectives in proportion to how much weight "
        "is placed on each."
    )
    doc.add_paragraph()

    wv_table = doc.add_table(rows=1, cols=4)
    wv_table.style = "Table Grid"
    wv_hdr = wv_table.rows[0]
    _shade_row(wv_hdr, "1A237E")
    for i, hdr in enumerate(["Perspective", "Weight", "Attitude to risk", "Extinction discount"]):
        wv_hdr.cells[i].text = hdr
        run = wv_hdr.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    wv_rows = [
        ("Human-focused, near-term",
         "35%",
         "Risk-neutral (takes expected values at face value)",
         "10% — relatively low concern that civilisation will not survive to benefit "
         "from long-run investments"),
        ("Animal welfare, near-term",
         "40%",
         "Mildly loss-averse (bad outcomes weighted slightly more than good ones)",
         "5% — very low extinction concern; near-term welfare dominates"),
        ("Longtermist, high extinction concern",
         "25%",
         "Moderately loss-averse",
         "50% — high concern that future generations may not exist; strongly discounts "
         "benefits that only materialise if civilisation survives"),
    ]
    for row_idx, (name, cred, rp, pext) in enumerate(wv_rows):
        data_row = wv_table.add_row()
        _shade_row(data_row, "F3F4FE" if row_idx % 2 == 0 else "FFFFFF")
        for i, val in enumerate([name, cred, rp, pext]):
            data_row.cells[i].text = val
            for para in data_row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    _set_col_widths(wv_table, [1.8, 0.7, 2.3, 3.2])
    doc.add_paragraph()

    doc.add_paragraph(
        "Each perspective also assigns different relative values to human and animal "
        "welfare (the human-focused perspective places near-zero weight on animal "
        "wellbeing; the animal welfare perspective weights chickens, fish, and mammals "
        "substantially) and discounts the future at different rates (the longtermist "
        "perspective treats outcomes 500 years from now as nearly as important as "
        "outcomes today; the human-focused perspective discounts sharply beyond 20 "
        "years). These differences mean the three perspectives often disagree sharply "
        "on which fund deserves the most money."
    )
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph(
        "Generated automatically by generate_methodology_doc.py. "
        "Part II reflects SCENARIO_METADATA in that file; keep it in sync with "
        "SCENARIOS in run_sensitivity.py when adding or removing scenarios."
    )
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    out_path = SCRIPT_DIR / "sensitivity_methodology.docx"
    print("Building methodology document...")
    doc = build_document()
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    print(f"  Part I:  3 output metrics")
    print(f"  Part II: {len(SCENARIO_METADATA)} parameter scenarios")


if __name__ == "__main__":
    main()
