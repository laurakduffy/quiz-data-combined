"""Generate a plain-language Word (.docx) overview of the Animal Welfare model.

Intended for an intelligent non-technical audience. Parameter values are drawn
from the YAML fund configs and intervention estimates so the document stays in sync.
"""

import os
import sys
import yaml

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_HERE, "..")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(_HERE, "aw_model_lay_overview.docx")

_DATA_DIR = os.path.join(_ROOT, "aw-models", "data", "inputs")

# ---------------------------------------------------------------------------
# Load parameters from YAML files
# ---------------------------------------------------------------------------

def _load_yaml(path):
    with open(path) as f:
        return yaml.safe_load(f)

_estimates = _load_yaml(os.path.join(_DATA_DIR, "aw_model_intervention_estimates.yaml"))
_combined  = _load_yaml(os.path.join(_DATA_DIR, "funds", "aw_combined.yaml"))["fund"]
_ea_awf    = _load_yaml(os.path.join(_DATA_DIR, "funds", "ea_awf.yaml"))["fund"]
_nav_cf    = _load_yaml(os.path.join(_DATA_DIR, "funds", "navigation_fund_cagefree.yaml"))["fund"]
_nav_gen   = _load_yaml(os.path.join(_DATA_DIR, "funds", "navigation_fund_general.yaml"))["fund"]

_interventions = _estimates["interventions"]

# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_callout(doc, text, color_fill="EAF4FB", text_color=(0x1A, 0x5C, 0x82)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    return p

def add_note(doc, text):
    return add_callout(doc, "Note: " + text, color_fill="FFFADC", text_color=(0x50, 0x3C, 0x00))

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D2DCF0")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    for ri, row in enumerate(rows):
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = t.rows[ri + 1].cells[i]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How We Estimate the Value of Animal Welfare Grants")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A plain-language guide to the Animal Welfare cost-effectiveness model")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. What this model does
# ---------------------------------------------------------------------------
add_heading(doc, "1.  What this model does", 1)
add_body(doc,
    "Animals — especially those raised in industrial farming — vastly outnumber humans and "
    "may experience substantial suffering. If their welfare matters morally, even a small "
    "improvement in how billions of chickens, fish, or shrimp are treated could dwarf the "
    "impact of many human-focused interventions."
)
add_body(doc,
    "Rethink Priorities evaluates animal welfare funds on behalf of Anthropic staff "
    "considering where to direct their personal philanthropy. This model estimates how many "
    "animal suffering-years are averted per $1 million donated to animal welfare funds, "
    "using analyst-derived cost-effectiveness distributions for each intervention type."
)
add_callout(doc,
    "The numbers in this model are in 'animal suffering-years averted' — not yet converted "
    "to human-equivalent welfare units. That final conversion requires a view on how much "
    "animal welfare matters relative to human welfare, which varies widely across ethical "
    "frameworks and is applied separately downstream."
)

# ---------------------------------------------------------------------------
# 2. The funds modelled
# ---------------------------------------------------------------------------
add_heading(doc, "2.  The funds modelled", 1)
add_body(doc,
    "Three funds are currently modelled:"
)
add_table(doc,
    ["Fund", "Annual budget", "Description"],
    [
        [_ea_awf["display_name"],
         f"${_ea_awf['annual_budget_M']}M/yr",
         "EA Animal Welfare Fund; a re-granting fund that distributes to effective "
         "animal welfare organisations worldwide"],
        [_nav_cf["display_name"],
         f"${_nav_cf['annual_budget_M']}M/yr",
         "Navigation Fund cage-free sub-portfolio; focused on corporate and political "
         "strategies to end cages for layer hens"],
        [_nav_gen["display_name"],
         f"${_nav_gen['annual_budget_M']}M/yr",
         "Navigation Fund general farm animal grantmaking across species and regions"],
    ],
)

# ---------------------------------------------------------------------------
# 3. What interventions are funded
# ---------------------------------------------------------------------------
add_heading(doc, "3.  What interventions are funded", 1)
add_body(doc,
    "Animal welfare funds spread money across several types of intervention. Each type "
    "targets a different species or advocacy pathway:"
)

# Build intervention table from combined fund splits + CCM data
_split_display = {
    "chicken_corporate_campaigns": ("Chicken corporate campaigns",
        "Lobbying companies to adopt cage-free, higher-welfare standards for egg-laying hens"),
    "movement_building":           ("Movement building",
        "Growing the effective animal advocacy community through media, education, and outreach"),
    "policy_advocacy": ("Policy advocacy",
        "Legislative and regulatory campaigns for multi-species welfare protections"),
    "fish_welfare":                ("Fish welfare",
        "Improving conditions for farmed fish — stunning, stocking density, water quality"),
    "shrimp_welfare":              ("Shrimp welfare",
        "Improving conditions for farmed shrimp, including slaughter practices"),
    "wild_animal_welfare":         ("Wild animal welfare",
        "Research and advocacy for reducing suffering in wild animal populations"),
    "invertebrate_welfare":        ("Invertebrate welfare",
        "Research into whether and how invertebrates (insects, crustaceans) experience suffering"),
}

combined_splits = _combined["splits"]
table_rows = []
for key, split in sorted(combined_splits.items(), key=lambda x: -x[1]):
    label, desc = _split_display.get(key, (key, ""))
    ccm = _interventions.get(key, {})
    p50 = ccm.get("percentiles_per_1000", {}).get("p50")
    mean = ccm.get("percentiles_per_1000", {}).get("mean")
    persist = ccm.get("persistence_years", "?")
    if p50 and mean:
        eff = f"Median {p50 * 1000 * split:,.0f} / mean {mean * 1000 * split:,.0f} aDALYs per $1M"
    else:
        eff = "—"
    table_rows.append([label, f"{split:.0%}", eff, f"{persist} years"])

add_table(doc,
    ["Intervention", "Share of combined fund", "Est. impact (weighted by fund share)", "Effect duration"],
    table_rows,
)
add_body(doc,
    "The 'Est. impact' column shows the suffering-years averted per $1M at the fund level "
    "(i.e. already scaled by the intervention's share of the fund budget). The wide gap "
    "between median and mean for most interventions reflects high uncertainty — in most "
    "optimistic scenarios, interventions perform dramatically better than the median."
)
add_note(doc,
    "Chicken corporate campaigns receive more than half the combined fund budget because "
    "the Longview Philanthropy Navigation Fund (which dominates by budget) allocates heavily "
    "to cage-free campaigns and movement building. As more fund data becomes available, "
    "these shares will be refined."
)

# ---------------------------------------------------------------------------
# 4. Where the cost-effectiveness estimates come from: the CCM
# ---------------------------------------------------------------------------
add_heading(doc, "4.  Where the estimates come from", 1)
add_body(doc,
    "The effectiveness numbers come from analyst-derived distributions built by Rethink "
    "Priorities. For each intervention, the model synthesises the best available evidence — "
    "field studies, expert surveys, and meta-analyses — and produces a full distribution of "
    "plausible cost-effectiveness estimates, not just a single point."
)
add_body(doc,
    "For each intervention, 100,000 random samples of 'animal suffering-years averted per "
    "$1,000 spent' are drawn. These samples capture everything we know and don't know about "
    "how effective each intervention is. The wide spreads reflect genuine uncertainty — we "
    "simply don't have precise data on, for example, exactly how many chickens are affected "
    "by a successful cage-free campaign."
)
add_callout(doc,
    "These are animal suffering-years, not human-equivalent DALYs. This model uses "
    "pre-moral-weight values; the cross-cause moral weighting step (applied downstream) "
    "converts them to a comparable basis with human welfare."
)

# Illustrative uncertainty table
add_body(doc, "To give a sense of the uncertainty involved, here are estimates for the two largest interventions:")
for key in ["chicken_corporate_campaigns", "movement_building"]:
    est = _interventions.get(key, {})
    pcts = est.get("percentiles_per_1000", {})
    label, _ = _split_display.get(key, (key, ""))
    split = combined_splits.get(key, 1.0)
    if pcts:
        add_table(doc,
            ["Percentile", "Per $1,000 spent", "Per $1M (at fund share)"],
            [
                ["10th (pessimistic)", f"{pcts.get('p10', '?'):,.0f}", f"{pcts.get('p10', 0) * 1000 * split:,.0f}"],
                ["50th (median)",      f"{pcts.get('p50', '?'):,.0f}", f"{pcts.get('p50', 0) * 1000 * split:,.0f}"],
                ["Mean",               f"{pcts.get('mean', '?'):,.0f}", f"{pcts.get('mean', 0) * 1000 * split:,.0f}"],
                ["90th (optimistic)",  f"{pcts.get('p90', '?'):,.0f}", f"{pcts.get('p90', 0) * 1000 * split:,.0f}"],
            ],
        )
        add_body(doc, f"Table: {label} (uncertainty at fund share {split:.0%})")
        doc.add_paragraph()

# ---------------------------------------------------------------------------
# 5. How time periods work for animal welfare
# ---------------------------------------------------------------------------
add_heading(doc, "5.  When do benefits happen?", 1)
add_body(doc,
    "Unlike global health interventions (where a vaccine prevents a death this year), animal "
    "welfare interventions work through corporate commitments, policy changes, and cultural "
    "shifts. Their effects take time to materialise and persist for a limited window."
)
add_body(doc,
    "Each intervention specifies two timing parameters:"
)
add_bullet(doc, "Start year — when the effect first kicks in (typically year 1 for most interventions).", bold_prefix="Effect start. ")
add_bullet(doc, "Persistence — how many years the effect continues (10–15 years for most AW interventions).", bold_prefix="Duration. ")
add_body(doc, "The effect is then spread proportionally across whichever time windows overlap with the persistence period:")
add_table(doc,
    ["Intervention", "Effect starts", "Lasts", "Which time windows receive weight"],
    [
        [_split_display.get(k, (k,))[0],
         f"Year {_interventions[k].get('effect_start_year', 1)}",
         f"{_interventions[k].get('persistence_years', '?')} years",
         "0–5 yr, 5–10 yr, 10–20 yr (proportionally)"]
        for k in combined_splits if k in _interventions
    ],
)
add_note(doc,
    "The AW model uses four time windows: 0–5, 5–10, 10–20, and 20–100 years. All current "
    "AW interventions have their effects within this 100-year horizon — the model does not "
    "include a 100–500 year or 500+ year window for animal welfare. This contrasts sharply "
    "with GCR models, where the entire value of survival extends across all of human (and "
    "post-human) history."
)

# ---------------------------------------------------------------------------
# 6. Diminishing returns
# ---------------------------------------------------------------------------
add_heading(doc, "6.  What happens if we give more?", 1)
add_body(doc,
    "Animal welfare organisations have limited absorptive capacity — there are only so many "
    "high-quality cage-free campaigns to run, so many lobbyists to hire, and so much movement "
    "infrastructure to build. The model now includes a formal diminishing returns curve for the "
    "combined AW fund, based on analyst estimates of how marginal cost-effectiveness changes "
    "as total funding increases."
)

# Build diminishing returns table from aw_combined anchors
_dr_anchors = _combined.get("diminishing_anchors", [])
_rfmf = _combined.get("room_for_more_M", "?")
add_body(doc,
    f"The curve is parameterised by anchor points that specify CE as a fraction of baseline "
    f"at different cumulative spend levels. The combined AW fund has an estimated room for more "
    f"funding (RFMF) of ${_rfmf}M — this is the total amount the market can absorb while "
    f"maintaining meaningful cost-effectiveness."
)

if _dr_anchors:
    add_table(doc,
        ["Cumulative spend", "Marginal CE (relative to baseline)"],
        [[f"${a[0]:.0f}M", f"{a[1]:.0%}"] for a in _dr_anchors],
    )

add_body(doc,
    "For example, if you are the only funder up to $10M, you get the full (baseline) CE. "
    "As total giving scales toward the RFMF and beyond, the marginal impact of each additional "
    "dollar declines. Beyond the last anchor, the curve continues to fall following a 1/x "
    "hyperbolic trajectory — cost-effectiveness keeps declining but never reaches zero."
)
add_note(doc,
    "These anchor points are rough analyst estimates based on conversations with fund managers "
    "about absorptive capacity. They should be treated as indicative rather than precise."
)

# ---------------------------------------------------------------------------
# 7. Risk profiles
# ---------------------------------------------------------------------------
add_heading(doc, "7.  Nine views of the same uncertainty", 1)
add_body(doc,
    "Because the model provides 100,000 samples per intervention (capturing the full "
    "distribution from pessimistic to optimistic), we can ask: how should we summarise "
    "this distribution? Different ethical frameworks disagree."
)
add_body(doc,
    "Animal welfare interventions tend to have particularly heavy right tails — the most "
    "optimistic scenarios are dramatically better than the median. This makes the choice "
    "of risk profile very consequential:"
)
add_table(doc,
    ["Profile", "Relative to neutral (mean)", "What drives the difference for AW"],
    [
        ["Neutral (mean)",   "Baseline",        "Includes full weight of extremely optimistic scenarios"],
        ["Upside",           "Somewhat lower",  "Clips the most extreme upper-tail scenarios"],
        ["Downside",         "Lower",           "Penalises variance; AW has high variance"],
        ["Combined",         "Much lower",      "Both tail clipping and variance penalty"],
        ["DMREU",            "Lower",           "Formal risk aversion; grows with tail magnitude"],
        ["WLU (low/mod/high)","Lower–much lower","Concave weighting; very sensitive to AW's heavy tail"],
        ["Ambiguity",        "Much lower",      "Discounts deep uncertainty; AW estimates are especially uncertain"],
    ],
)
add_callout(doc,
    "For animal welfare, the gap between the 'neutral' and 'combined' or 'ambiguity' profiles "
    "is often larger than for other cause areas. This is because AW cost-effectiveness estimates "
    "have very wide distributions — the difference between a campaign that succeeds and one that "
    "fails is enormous, and we genuinely don't know which will happen."
)

# ---------------------------------------------------------------------------
# 8. Key limitations
# ---------------------------------------------------------------------------
add_heading(doc, "8.  Key limitations and honest caveats", 1)
add_bullet(doc,
    "The numbers are in pre-moral-weight animal suffering-years. To compare with human welfare "
    "or GCR work, you need to apply a cross-cause moral weight — which is deeply contested. "
    "Some people think animal welfare is negligible; others think it may be the most important "
    "cause area. The model cannot resolve that debate.",
    bold_prefix="Moral weights not yet applied. ")
add_bullet(doc,
    "Estimates are derived from limited empirical evidence. Corporate campaigns in particular "
    "rely heavily on extrapolation from a small number of documented successes. The wide "
    "confidence intervals reflect genuine ignorance, not well-characterised uncertainty.",
    bold_prefix="Evidence base is thin. ")
add_bullet(doc,
    "The model covers seven intervention categories. Important types of work — including "
    "alternative proteins, regulatory advocacy, and farmed mammal welfare — are not yet "
    "included or are represented by placeholder values.",
    bold_prefix="Incomplete coverage. ")
add_bullet(doc,
    "The fund splits (how much of each fund's budget goes to which intervention type) are "
    "estimated from grant reports and qualitative descriptions, not directly confirmed by "
    "fund managers. Particularly for newer funds, these estimates carry significant uncertainty.",
    bold_prefix="Fund split uncertainty. ")
add_bullet(doc,
    "All effects are modelled as occurring within 20 years. If some welfare improvements "
    "(e.g. cultural or regulatory shifts) have effects that compound over much longer periods, "
    "the model will underestimate their value.",
    bold_prefix="No long-run effects. ")

# ---------------------------------------------------------------------------
# 9. Summary at a glance
# ---------------------------------------------------------------------------
add_heading(doc, "9.  Summary at a glance", 1)

add_table(doc,
    ["Item", "EA AWF", "Nav Fund — Cage-Free", "Nav Fund — General"],
    [
        ["Annual budget",
         f"${_ea_awf['annual_budget_M']}M",
         f"${_nav_cf['annual_budget_M']}M",
         f"${_nav_gen['annual_budget_M']}M"],
        ["Number of interventions modelled",
         str(sum(1 for v in _ea_awf["splits"].values() if v and v > 0)),
         str(sum(1 for v in _nav_cf["splits"].values() if v and v > 0)),
         str(sum(1 for v in _nav_gen["splits"].values() if v and v > 0))],
        ["Largest intervention (by share)",
         max(_ea_awf["splits"], key=lambda k: _ea_awf["splits"][k]).replace("_", " "),
         max(_nav_cf["splits"], key=lambda k: _nav_cf["splits"][k]).replace("_", " "),
         max(_nav_gen["splits"], key=lambda k: _nav_gen["splits"][k]).replace("_", " ")],
        ["Samples per intervention",
         "100,000 (full) or 10,000 (fallback)",
         "100,000 (full) or 10,000 (fallback)",
         "100,000 (full) or 10,000 (fallback)"],
        ["Time horizon of effects", "0–20 years", "0–20 years", "0–20 years"],
        ["Units",
         "Animal suffering-years / $1M (pre-moral-weight)",
         "Animal suffering-years / $1M (pre-moral-weight)",
         "Animal suffering-years / $1M (pre-moral-weight)"],
    ],
)

# Save
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
