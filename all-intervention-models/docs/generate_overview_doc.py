"""Generate a Word (.docx) overview document describing what all three models
compute, how effects are broken out by time and type, and how risk modeling works."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "model_overview.docx")

# ---------------------------------------------------------------------------
# Helpers (same style as generate_model_doc.py)
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_body(doc, text, monospace=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if monospace:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Note: " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x3C, 0x00)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFADC")
    pPr.append(shd)
    return p

def add_table(doc, headers, rows):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "D2DCF0")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
        if hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_bg(cells[i], bg)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Intervention Valuation Models: Overview")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("What is compared, how effects are decomposed, and how risk modeling works")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Covers: GiveWell (GW)  |  Global Catastrophic Risk (GCR)  |  Animal Welfare (AW)")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x78, 0x78, 0x78)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. What Is Being Compared
# ---------------------------------------------------------------------------
add_heading(doc, "1.  What Is Being Compared", 1)

add_body(doc,
    "The three models evaluate the cost-effectiveness of philanthropic spending across three "
    "distinct cause areas: global health (GiveWell), existential and global catastrophic risk "
    "reduction (GCR), and farmed and wild animal welfare (AW). Each model asks the same "
    "underlying question: how much expected value does an additional $1 million of spending "
    "generate, broken down by time horizon and adjusted for risk preferences?"
)
add_body(doc,
    "The outputs of all three models share a common structure: a matrix of values indexed by "
    "6 time periods and 9 risk profiles. This allows a downstream analysis tool to compare "
    "across cause areas on a level playing field, regardless of the domain-specific units "
    "used internally by each model."
)

add_heading(doc, "1.1  The three cause areas", 2)

add_table(doc,
    ["Model", "What is funded", "Funds / interventions modelled", "Internal unit"],
    [
        ["GiveWell (GW)",
         "Global health charities (malaria, vaccines, nutrition, water, livelihoods, family planning)",
         "Single GiveWell portfolio (8 cause areas combined by funding fraction)",
         "Life-years saved, YLDs averted, income doublings"],
        ["GCR",
         "Existential and catastrophic risk reduction organisations",
         "Sentinel Bio, Longview Nuclear, Longview AI",
         "Expected survival-weighted civilizational value (welfare units)"],
        ["Animal Welfare (AW)",
         "Farmed and wild animal welfare interventions",
         "Chicken corporate campaigns, shrimp welfare, fish welfare, invertebrate welfare, "
         "policy advocacy, movement building, wild animal welfare",
         "Suffering-years averted (pre-moral-weight)"],
    ],
)

add_heading(doc, "1.2  What is calculated", 2)
add_body(doc,
    "Each model calculates the marginal expected value per $1 million of additional spending. "
    "This is computed as a distribution across Monte Carlo samples, not a single point "
    "estimate, so that the 9 risk profiles can be applied to the resulting sample array."
)
add_body(doc, "The key steps common to all three models are:")
add_bullet(doc, "Draw a large number of samples from the uncertainty distribution for that cause area (10,000–100,000 samples depending on the model).")
add_bullet(doc, "For each sample, compute the expected effect of $1M of spending in welfare-relevant units.")
add_bullet(doc, "Allocate the total effect across the 6 standard time periods.")
add_bullet(doc, "Apply 9 risk-adjustment transformations to the sample array to produce 9 risk-profile values.")
add_bullet(doc, "Store the resulting 6 × 9 = 54 values as the model output for that fund or intervention.")

add_note(doc,
    "The models do not use a shared moral weight across cause areas. GW outputs are in "
    "GiveWell's own welfare-weighted units; GCR outputs are in civilizational value units; "
    "AW outputs are in suffering-years. Comparison across cause areas requires the downstream "
    "analysis tool to apply its own cross-cause moral weights."
)

# ---------------------------------------------------------------------------
# 2. Time Decomposition
# ---------------------------------------------------------------------------
add_heading(doc, "2.  How Effects Are Broken Out by Time", 1)

add_body(doc,
    "All three models report effects across the same 6 time periods. These periods are "
    "intentionally unequal in length: the early periods are short to capture near-term "
    "effects precisely, while the later periods are long because effects there are heavily "
    "discounted by survival probability and moral uncertainty."
)

add_table(doc,
    ["Period label", "Years covered", "Typical focus"],
    [
        ["t0", "0 – 5 years",     "Immediate near-term effects"],
        ["t1", "5 – 10 years",    "Short-to-medium term"],
        ["t2", "10 – 20 years",   "Medium term"],
        ["t3", "20 – 100 years",  "Long run within a century"],
        ["t4", "100 – 500 years", "Multi-generational"],
        ["t5", "500+ years",      "Very long run / longtermist tail"],
    ],
)

add_heading(doc, "2.1  GiveWell: fixed temporal breakdowns by effect type", 2)
add_body(doc,
    "GiveWell effects are heavily front-loaded. Direct health effects (mortality and morbidity) "
    "occur almost entirely within the first 10 years because the interventions produce "
    "near-term outcomes for present-day beneficiaries. Economic effects are spread further "
    "into the future because income gains compound across generations."
)
add_table(doc,
    ["Effect type", "t0 (0–5y)", "t1 (5–10y)", "t2 (10–20y)", "t3 (20–100y)", "t4 (100–500y)", "t5 (500y+)"],
    [
        ["YLDs averted",     "90.0%", "7.0%", "1.5%", "1.0%", "0.5%", "0.0%"],
        ["Life-years saved", "90.0%", "7.0%", "1.5%", "1.0%", "0.5%", "0.0%"],
        ["Income doublings", "18.0%", "1.4%", "12.5%", "68.1%", "0.0%", "0.0%"],
    ],
)
add_note(doc,
    "Income doubling effects are concentrated in t3 (20–100 years) because they represent "
    "the long-run compounding of economic development, not just the immediate income gain "
    "to current recipients."
)

add_heading(doc, "2.2  GCR: survival-probability-weighted integration", 2)
add_body(doc,
    "The GCR model does not use fixed temporal breakdowns. Instead, it computes the "
    "expected value contribution of each time period by integrating the product of "
    "(a) the change in survival probability caused by the intervention and (b) the value "
    "of the world conditional on survival, over each time window."
)
add_body(doc,
    "Because civilization must survive to year T for t3–t5 to matter, those periods "
    "receive weight only in proportion to the cumulative survival probability. The model "
    "includes scenarios where stellar settlement begins at T_c (500 years in the modal "
    "case), which concentrates long-run value in t4 and t5 under the 10% of samples "
    "where cubic growth occurs."
)
add_body(doc,
    "In practice, the GCR model's time decomposition is determined by the convolution of "
    "three factors: the Gaussian risk peak (when the intervention's risk reduction is "
    "most powerful), the timing parameters year_effect_starts and persistence_effect "
    "(which window the intervention is active), and the value trajectory (when value "
    "accumulates most rapidly)."
)

add_heading(doc, "2.3  Animal Welfare: persistence windows", 2)
add_body(doc,
    "AW effects are allocated to time periods based on two intervention-specific parameters: "
    "the year the effect starts (start_year) and how many years it persists (persistence_years). "
    "The fraction of the total effect assigned to each period equals the overlap between "
    "the persistence window and that period, divided by the total persistence."
)
add_body(doc,
    "Example for chicken corporate campaigns (start year 1, persistence 15 years, so active "
    "years 1–16):"
)
add_table(doc,
    ["Period", "Years covered", "Overlap with [1,16)", "Share of effect"],
    [
        ["t0", "0–5",   "4 years (1–5)",   "4/15 = 26.7%"],
        ["t1", "5–10",  "5 years (5–10)",  "5/15 = 33.3%"],
        ["t2", "10–20", "6 years (10–16)", "6/15 = 40.0%"],
        ["t3", "20–100","0 years",         "0%"],
        ["t4", "100–500","0 years",        "0%"],
        ["t5", "500+",  "0 years",         "0%"],
    ],
)
add_note(doc,
    "Unlike GW and GCR, AW effects do not extend beyond the persistence window. All seven "
    "AW interventions have persistence windows entirely within the first 20 years, so t3–t5 "
    "receive zero weight unless a specific intervention parameter is changed."
)

# ---------------------------------------------------------------------------
# 3. Effect Type Decomposition
# ---------------------------------------------------------------------------
add_heading(doc, "3.  How Effects Are Broken Out by Effect Type", 1)

add_body(doc,
    "Effect type decomposition differs substantially across the three models, reflecting "
    "the different nature of the interventions."
)

add_heading(doc, "3.1  GiveWell: three effect types", 2)
add_body(doc,
    "GiveWell is the only model with multiple effect types. Each cause area within the "
    "GiveWell portfolio is characterised by the fraction of its total cost-effectiveness "
    "attributable to each of three types:"
)
add_bullet(doc, "YLDs averted — ", bold_prefix="YLDs averted: ")
add_bullet(doc, "reduction in years lived with disability (morbidity)")
add_bullet(doc, "Life-years saved — ", bold_prefix="Life-years saved: ")
add_bullet(doc, "mortality reduction, converted to life-years using cause-specific life expectancy")
add_bullet(doc, "Income doublings — ", bold_prefix="Income doublings: ")
add_bullet(doc, "economic development effects on beneficiary households")

# Rebuild as proper bullets with bold prefix
# Reset and redo these as a table instead
doc.paragraphs[-5].clear()
doc.paragraphs[-4].clear()
doc.paragraphs[-3].clear()
doc.paragraphs[-2].clear()
doc.paragraphs[-1].clear()

add_table(doc,
    ["Effect type", "What it measures", "GW moral weight (multiplier)"],
    [
        ["YLDs averted",     "Morbidity reduction (disability-free life-years gained)",             "2.3"],
        ["Life-years saved", "Mortality reduction (life-years gained from deaths averted)",         "115.6"],
        ["Income doublings", "Economic development (household income gains, compounded over time)", "1.0"],
    ],
)
add_body(doc,
    "The model draws 10,000 samples from each cause area's cost-effectiveness distribution "
    "and attributes the total to the three effect types using cause-specific fractions. "
    "For example, malaria spending is approximately 14% YLDs, 58% lives saved, and 27% "
    "income doublings. The output stores one separate row per effect type."
)

add_heading(doc, "3.2  GCR: single aggregate effect type", 2)
add_body(doc,
    "The GCR model does not decompose by effect type. It computes a single quantity: "
    "the expected increase in survival-weighted civilizational value caused by the "
    "intervention. This combines near-term and long-term value, extinction-pathway effects, "
    "and (for Sentinel Bio only) sub-extinction tier effects into one number per sample."
)
add_body(doc,
    "Sub-extinction tiers for Sentinel Bio are modelled separately using a simple expected "
    "value formula (annual probability × expected deaths × relative risk reduction × "
    "persistence × counterfactual factor) and then added to the extinction-pathway EV."
)

add_heading(doc, "3.3  Animal Welfare: single aggregate effect type", 2)
add_body(doc,
    "The AW model outputs a single effect type per intervention: suffering-years averted "
    "per $1M. This is drawn directly from Rethink Priorities' Cross-Cause Cost-Effectiveness "
    "Model (CCM) empirical sample distribution, which already embeds moral weights for "
    "different animal species and welfare states. The 100,000 samples preserve the full "
    "shape of the CCM distribution, including its heavy right tail."
)

# ---------------------------------------------------------------------------
# 4. Risk Modeling
# ---------------------------------------------------------------------------
add_heading(doc, "4.  How Risk Modeling Works", 1)

add_body(doc,
    "All three models share the same risk-adjustment framework. The framework is defined "
    "in risk_profiles.py and applied identically across all cause areas. It transforms a "
    "raw array of Monte Carlo samples into 9 risk-adjusted expected values, each "
    "corresponding to a different attitude toward uncertainty and tail outcomes."
)
add_body(doc,
    "The 9 profiles are organised into three groups: informal adjustments, percentile-weighted "
    "adjustments, and formal decision-theoretic models. They range from risk-neutral (simple "
    "mean) to highly risk-averse (strong downweighting of large positive outcomes)."
)

add_heading(doc, "4.1  The 9 risk profiles", 2)

add_table(doc,
    ["Profile", "Group", "Description"],
    [
        ["neutral",       "Informal",            "Simple mean of all samples. Risk-neutral expected value."],
        ["upside",        "Informal",            "Clips samples at the 99th percentile before taking the mean. "
                                                  "Reflects skepticism that very large outcomes are real."],
        ["downside",      "Informal",            "Loss-averse utility around the median. Losses below the "
                                                  "median are amplified by a factor of 5.0 relative to equal-sized gains."],
        ["combined",      "Percentile-weighted", "Applies exponential decay to weights for samples above the "
                                                  "97.5th percentile, zeroing out samples above the 99.9th percentile, "
                                                  "then adds loss aversion around the median."],
        ["ambiguity",     "Percentile-weighted", "Same exponential decay weighting as 'combined' but without "
                                                  "the loss-aversion component. Reflects pure ambiguity aversion."],
        ["dmreu",         "Decision-theoretic",  "Difference-Making Risk-Weighted Expected Utility (Duffy 2023). "
                                                  "Applies probability weighting m(P) = P^0.86 to tail outcomes, "
                                                  "producing moderate risk aversion."],
        ["wlu - low",     "Decision-theoretic",  "Weighted Linear Utility with concavity c = 0.01. "
                                                  "Magnitude-sensitive: larger outcomes receive lower weight."],
        ["wlu - moderate","Decision-theoretic",  "Weighted Linear Utility with concavity c = 0.05."],
        ["wlu - high",    "Decision-theoretic",  "Weighted Linear Utility with concavity c = 0.10. "
                                                  "The most risk-averse of the WLU profiles."],
    ],
)

add_heading(doc, "4.2  How the profiles are computed", 2)

add_heading(doc, "Neutral", 3)
add_body(doc, "    E[X]  =  mean(samples)", monospace=True)

add_heading(doc, "Upside", 3)
add_body(doc, "    E[X | X <= p99]  =  mean(clip(samples, max=p99))", monospace=True)
add_body(doc,
    "Samples above the 99th percentile are replaced by the 99th percentile value before "
    "taking the mean. This effectively removes the influence of extreme upper-tail draws."
)

add_heading(doc, "Downside (loss aversion)", 3)
add_body(doc, "    ref  =  median(samples)", monospace=True)
add_body(doc, "    gain_i  =  sample_i - ref", monospace=True)
add_body(doc, "    adjusted_gain_i  =  gain_i  if gain_i >= 0  else  5.0 * gain_i", monospace=True)
add_body(doc, "    E_downside  =  ref + mean(adjusted_gains)", monospace=True)
add_body(doc,
    "A loss below the reference point (the median) counts 5.0 times as much as an equal "
    "gain above it (LOSS_AVERSION_LAMBDA = 5.0)."
)

add_heading(doc, "Combined and Ambiguity (percentile weighting)", 3)
add_body(doc,
    "Both profiles apply a weight function to each sample based on its percentile rank:"
)
add_body(doc, "    w(p)  =  1.0                                   if p <= 97.5th percentile", monospace=True)
add_body(doc, "    w(p)  =  exp(-ln(100)/1.5 * (p - 97.5))       if 97.5 < p <= 99.9th percentile", monospace=True)
add_body(doc, "    w(p)  =  0.0                                   if p > 99.9th percentile", monospace=True)
add_body(doc,
    "This creates an exponential fade-out: samples between the 97.5th and 99.9th percentiles "
    "receive decreasing weight, and samples above the 99.9th percentile are ignored entirely. "
    "'Combined' then additionally applies loss aversion as in the downside profile; "
    "'ambiguity' uses only the percentile weighting."
)

add_heading(doc, "DMREU (Difference-Making Risk-Weighted Expected Utility)", 3)
add_body(doc,
    "Based on Duffy (2023). Applies a probability-weighting function to the ranked sample "
    "distribution:"
)
add_body(doc, "    a  =  -2 / log10(0.05)  ≈  0.862", monospace=True)
add_body(doc, "    m(P)  =  P^a", monospace=True)
add_body(doc,
    "The weight of each sample reflects the difference between the probability of doing "
    "at least as well as that sample under the intervention versus under the alternative. "
    "This produces moderate risk aversion that grows with the magnitude of tail outcomes."
)

add_heading(doc, "WLU (Weighted Linear Utility)", 3)
add_body(doc,
    "Also from Duffy (2023). Applies a concave weighting function that reduces the weight "
    "of larger outcomes:"
)
add_body(doc, "    w(x)  =  1 / (1 + c * |x|)", monospace=True)
add_body(doc,
    "where c is the concavity parameter (0.01, 0.05, or 0.10 for the low, moderate, and "
    "high variants). Higher c means more weight on smaller outcomes and less on larger ones, "
    "producing stronger risk aversion. The weighted mean is:"
)
add_body(doc, "    E_WLU  =  sum(w(x_i) * x_i) / sum(w(x_i))", monospace=True)

add_heading(doc, "4.3  How risk profiles interact with time decomposition", 2)
add_body(doc,
    "The risk profiles are applied to the full distribution of total intervention value "
    "per $1M (summed across all time periods), not to each time period independently. "
    "The time decomposition then apportions the risk-adjusted total across periods "
    "using fixed fractions (GW, AW) or the survival-weighted integrals by period (GCR)."
)
add_body(doc,
    "This means that a risk profile which reduces the overall expected value (e.g., 'downside') "
    "scales all time periods proportionally. The relative share of value in each time period "
    "is the same under all 9 risk profiles for GW and AW; for GCR, the time decomposition "
    "can vary slightly across samples depending on which risk scenarios drive large outcomes."
)

add_heading(doc, "4.4  How risk profiles interact with the three models", 2)
add_table(doc,
    ["Model", "Sample size", "What the samples represent", "Risk profile application"],
    [
        ["GW",  "10,000",  "Portfolio cost-effectiveness draws from mixture distributions over 8 cause areas",
         "Applied separately to each effect type (YLDs, life-years, income doublings); 3 × 9 output rows"],
        ["GCR", "100,000", "Joint draws over risk trajectory, value growth, and intervention effectiveness parameters",
         "Applied to total EV per sample (one fund at a time); 1 × 9 output rows per fund"],
        ["AW",  "100,000", "Empirical CCM samples per intervention, preserving the original distribution shape",
         "Applied to the raw CCM sample array; 1 × 9 output rows per intervention"],
    ],
)
add_note(doc,
    "The GCR model uses stratified sampling to ensure low-probability but high-importance "
    "strata (e.g., cubic growth scenarios at 10% probability) are represented proportionally. "
    "The AW model uses the full empirical distribution from the CCM without any distribution "
    "fitting, which preserves the heavy right tail that drives large differences between "
    "risk profiles."
)

# ---------------------------------------------------------------------------
# 5. How the Models Relate and Are Combined
# ---------------------------------------------------------------------------
add_heading(doc, "5.  How the Three Models Relate", 1)

add_body(doc,
    "Despite their different domains, units, and modelling approaches, all three models "
    "produce outputs in an identical format: a matrix of values with one row per fund "
    "or intervention and 54 columns (6 time periods × 9 risk profiles). A downstream "
    "combine_data.py script reads all three CSV outputs and merges them into a single "
    "JSON structure for comparison."
)

add_table(doc,
    ["Dimension", "GW", "GCR", "AW"],
    [
        ["Domain",               "Global health",                "Existential & catastrophic risk",    "Animal welfare"],
        ["Primary time horizon", "0–100 years (income to 500y)", "0 to 10^14 years",                   "0–20 years (persistence-limited)"],
        ["Sampling method",      "Monte Carlo mixture (10k)",    "Stratified Monte Carlo (100k)",       "Empirical CCM distribution (100k)"],
        ["Effect types",         "3 (YLDs, lives, income)",      "1 (survival-weighted value)",         "1 (suffering-years averted)"],
        ["Uncertainty source",   "CE distribution per cause",    "Risk trajectory + value parameters",  "CCM empirical spread"],
        ["Diminishing returns",  "Fixed by fund",                "Piecewise linear + 1/x tail",         "Piecewise linear + 1/x tail"],
        ["Output rows",          "3 (one per effect type)",      "3 (one per fund)",                    "7 (one per intervention)"],
    ],
)

add_body(doc,
    "The key difference between the models is what drives their sample distributions. "
    "GW uncertainty comes from not knowing the exact cost-effectiveness of each cause "
    "area within the GiveWell portfolio. GCR uncertainty comes from deep uncertainty about "
    "the shape of the risk trajectory, the value of the long-run future, and how much the "
    "intervention actually reduces risk. AW uncertainty comes from empirical uncertainty "
    "about how effective welfare interventions are, as quantified by the CCM."
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
