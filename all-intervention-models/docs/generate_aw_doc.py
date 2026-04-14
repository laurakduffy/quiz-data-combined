"""Generate a Word document for the aw_intervention_models.py technical documentation."""

import os
from docx import Document

_DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(16)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(13)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(11)
style_h3.font.bold = True
style_h3.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold_prefix=None):
    """Add a paragraph. If bold_prefix is given, that part is bolded."""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def bold_para(label, text):
    """Add paragraph with a bold label followed by normal text."""
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)
    return p

def bullet(text, bold_prefix=None, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def hr():
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Technical Documentation: aw_intervention_models.py')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

# ── Overview ─────────────────────────────────────────────────────────────────

h1('Overview')

para(
    'The script aw_intervention_models.py is a probabilistic cost-effectiveness model for '
    'animal welfare interventions. Its purpose is to estimate, in units of suffering-years '
    'averted per $1,000 spent, how effective various interventions are across different animal '
    'populations. Rather than producing single point estimates, the script uses Monte Carlo '
    'simulation — drawing 100,000 samples from uncertainty distributions for each input '
    'parameter — to characterize the full range of plausible outcomes for each intervention. '
    'These distributions are then summarized as percentiles for human readability and '
    'downsampled to 10,000 draws for use in downstream risk analysis pipelines.'
)

para(
    'The script partially replicates and extends the distribution parameters used in Rethink '
    'Priorities\' Cross-Cause Cost-Effectiveness Model (CCM), supplementing those with '
    'analyst-constructed estimates. The output unit — suffering-years averted per $1,000 — '
    'is a pre-moral-weight quantity. Moral weight conversion (i.e., how much weight to assign '
    'to animal suffering relative to human suffering) is handled in a separate downstream step.'
)

# ── Sampling Methodology ──────────────────────────────────────────────────────

h1('Sampling Methodology')

h2('Random Seed and Sample Size')

para(
    'The script fixes a random seed at the start (seed 42), ensuring that all results are '
    'exactly reproducible across runs. All distributions are sampled at a size of N = 100,000 '
    'draws. This large sample size ensures that the tails of distributions are well-characterized '
    'and that compounded uncertainty across multiple parameters (which involves elementwise '
    'multiplication of arrays) produces stable results.'
)

h2('Lognormal Distributions')

para(
    'Many parameters in this model are inherently positive-valued and right-skewed — for '
    'example, the number of animals affected per dollar, or the cost of a campaign. For these, '
    'the script uses lognormal distributions. A lognormal distribution is parameterized here by '
    'specifying a confidence interval: the analyst provides a low value and a high value, and a '
    'credibility level (defaulting to 90%, meaning those bounds correspond to the 5th and 95th '
    'percentiles). The script back-calculates the underlying mean and standard deviation of the '
    'log-transformed variable, then draws samples. Clipping parameters (lclip and rclip) are '
    'available to truncate the distribution at hard physical or logical bounds — for example, '
    'hours of suffering cannot be negative, and some rates cannot meaningfully exceed known '
    'physical maxima.'
)

h2('Normal Distributions')

para(
    'For parameters that may be more symmetrically distributed, the script uses normal '
    'distributions, again specified by a confidence interval. Clipping is again available to '
    'enforce bounds.'
)

h2('Beta Distributions')

para(
    'For parameters that represent proportions or probabilities (values bounded between 0 and 1), '
    'the script uses beta distributions, parameterized by shape parameters a and b. A beta(a, b) '
    'distribution has a mean of a / (a + b). For example, beta(7, 3) has a mean of 0.70, '
    'representing a prior belief that the intervention averts roughly 70% of suffering. Beta '
    'distributions are naturally bounded to the unit interval, making them well-suited to '
    'probability and proportion parameters.'
)

h2('The Output Unit')

para(
    'Every intervention\'s final output is expressed in suffering-years (SY) averted per $1,000 '
    'spent. A suffering-year is conceptually equivalent to a DALY (Disability-Adjusted Life Year) '
    'for an animal — one year of life lived in full suffering. This is computed by converting '
    'hours of suffering into years (dividing by HOURS_PER_YEAR = 24 × 365.25 = 8,766 hours), '
    'multiplying by the number of animals affected, and scaling by cost.'
)

h2('Downsampling')

para(
    'The script includes a downsampling function that reduces the 100,000-sample arrays to '
    '10,000 samples using quantile spacing — that is, evenly spaced quantiles are drawn from '
    'the full distribution rather than a simple random subsample. This preserves the shape of '
    'the distribution, particularly in the tails, while reducing file size and computational '
    'load for downstream consumers.'
)

# ── Interventions ─────────────────────────────────────────────────────────────

h1('Interventions')

# ── Chickens ──────────────────────────────────────────────────────────────────

h2('Chickens: Corporate Campaign Advocacy')

bold_para('Purpose and source.  ',
    'This intervention models the cost-effectiveness of corporate campaigns advocating for '
    'cage-free egg production, primarily targeting laying hens. The estimates are based on a '
    'blend of two independent assessments: one by The Humane League (THL) using THL\'s own '
    'data, and one by Animal Charity Evaluators (ACE), with 80% weight assigned to the THL '
    'estimate and 20% to ACE. The blended estimate was then verified against a Causal model '
    'built from Welfare Footprint Initiative data.'
)

bold_para('Input parameters.  ',
    'The primary input is the number of hen-DALYs averted per $1,000 spent, drawn from a '
    'lognormal distribution with a 90% credibility interval of 177 to 3,600 hen-DALYs per '
    '$1,000, clipped to a minimum of 50 and a maximum of 10,000. This represents the combined '
    'effect of: the number of hens spared per dollar (weighted average of THL: 0.2–8 hens/$ '
    'and ACE: 2–44 hens/$), the duration of a laying hen\'s lifespan, and the degree to which '
    'welfare is improved. The prior RP estimate of 9–120 chickens per dollar was judged '
    'approximately 8× too high based on updated source data and a Causal model yielding a mean '
    'of ~1,200 hen-DALYs/$1,000.'
)

bold_para('Calculation.  ',
    'Because the per-$1,000 hen-DALY estimate directly encodes all relevant factors, the '
    'calculation for chickens is a single sampling step. The lognormal draw is used directly '
    'as the suffering-years per $1,000 output.'
)

# ── Shrimp ────────────────────────────────────────────────────────────────────

h2('Shrimp: Weighted Average of Intervention Sub-Types')

para(
    'The shrimp model is the most structurally complex in the script. Shrimp welfare '
    'interventions are divided into two categories — humane slaughter and combined sludge '
    'removal/stocking density — and a final weighted average is computed based on the estimated '
    'share of funding going to each.'
)

h3('Humane Slaughter Intervention (HSI)')

bold_para('Input parameters.  ',
    'Four parameters drive this sub-model:'
)

bullet('Shrimp affected per dollar per year: lognormal, 90% CI of 800 to 2,200 shrimp/$/year '
       '(from Shrimp Welfare Project estimates), clipped to 100–100,000.')
bullet('Persistence of the intervention\'s effect: lognormal, 90% CI of 6 to 15 years (mean '
       '~10 years). Multiplying reach-per-year by persistence gives lifetime shrimp affected per dollar.')
bullet('Hours of suffering per shrimp from conventional slaughter: lognormal, 90% CI of 0.28 '
       'to 6.4 hours (from McKay 2023), clipped to 0.1–24 hours, with a mean of ~2 hours. '
       'Divided by HOURS_PER_YEAR to convert to DALYs per shrimp.')
bullet('Proportion of suffering reduced by HSI: beta(7, 3), mean ~70%, 90% CI ~46%–90%.')

bold_para('Calculation.  ',
    'Per-dollar lifetime shrimp reach × per-shrimp suffering (in DALYs) × proportion of '
    'suffering reduced = DALYs reduced per dollar.'
)

h3('Sludge Removal and Stocking Density')

para(
    'These two interventions share a common reach estimate derived from ACE data: a campaign '
    'costing a fixed $71,342 is estimated to affect 50M–90M shrimp (lognormal, 90% CI), '
    'clipped to 10M–200M. Dividing affected shrimp by cost gives shrimp per dollar.'
)

bold_para('Sludge removal.  ',
    'Suffering from poor water quality: lognormal, 90% CI of 32–180 hours per shrimp (from '
    'McKay 2023 water quality data). Proportion of that suffering reduced: beta(4, 4), mean '
    '50%, reflecting symmetric uncertainty.'
)

bold_para('Stocking density.  ',
    'Suffering from high stocking density: lognormal, 90% CI of 50–150 hours per shrimp. '
    'Proportion reduced: beta(3, 12), mean ~20%, reflecting a more modest expected welfare gain.'
)

para(
    'The DALYs reduced per dollar from sludge removal and stocking density are summed to '
    'produce the combined sludge-and-density estimate.'
)

h3('Weighted Average')

para(
    'The proportion of shrimp welfare spending going to HSI is modeled as beta(18, 2), mean '
    '~90%, indicating HSI dominates spending. The final shrimp estimate is: (HSI share × HSI '
    'DALYs/dollar) + ((1 − HSI share) × sludge/density DALYs/dollar), then scaled to per $1,000.'
)

# ── Fish ──────────────────────────────────────────────────────────────────────

h2('Fish: Carp as a Proxy for Farmed Fish')

bold_para('Purpose and source.  ',
    'This model covers farmed fish welfare interventions, using carp as a representative '
    'species. The marginal cost-effectiveness is estimated using data from Fish Welfare '
    'Initiative (FWI), which reports approximately 1–7 fish affected per dollar depending '
    'on program type.'
)

bold_para('Input parameters.  ', 'Three parameters drive this model:')

bullet('Fish affected per dollar: lognormal, 90% CI of 0.5–15 fish/$, clipped to 0.2–50. '
       'Mean ~4.6 fish/dollar, reflecting the range across FWI program types.')
bullet('Hours of suffering per fish: normal distribution with bounds derived from a carp '
       'culture cycle of 383 days. Lower bound = 383 × 24 / 6 hours (⅙ of a year of '
       'suffering); upper bound = 383 × 24 / 3 hours (⅓ of a year), reflecting the '
       'analyst\'s estimate that a typical fish suffers ⅙ to ⅓ DALYs per year.')
bullet('Proportion of suffering reduced: beta(3, 17), mean ~15%, representing the assumption '
       'that interventions (e.g., improved slaughter or handling) avert a modest fraction of '
       'total farmed fish suffering (90% CI: 5%–30%).')

bold_para('Calculation.  ',
    'Fish per dollar × DALY fraction per fish × proportion reduced = DALYs per dollar, '
    'scaled to per $1,000.'
)

# ── Invertebrates ─────────────────────────────────────────────────────────────

h2('Invertebrates: Black Soldier Fly (BSF) as a Proxy')

bold_para('Purpose and source.  ',
    'This model covers farmed invertebrate welfare interventions, using black soldier fly '
    '(BSF) larvae as a representative species. BSF is industrially farmed at very large '
    'scales, primarily for animal feed production. Parameters are based on CCM defaults.'
)

bold_para('Input parameters.  ', 'Seven parameters drive this model:')

bullet('Number of BSF born per year: normal, bounds 200B–300B, clipped to 20B–1T.')
bullet('Proportion affected by the intervention: lognormal, 90% CI 0.005%–0.1% '
       '(5×10⁻⁵ to 10⁻³). Reflects the small fraction of total production a campaign '
       'can realistically reach.')
bullet('Hours of suffering per BSF: normal distribution derived from a 22–24 day larval '
       'stage, assuming suffering corresponds to ⅕ to 1/20 of a DALY over that period.')
bullet('Proportion of suffering reduced: beta(9, 4), mean ~69%.')
bullet('Probability of campaign success: beta(4, 16), mean 20%. The low probability '
       'reflects the speculative and early-stage nature of invertebrate welfare advocacy.')
bullet('Campaign cost: lognormal, 90% CI $150,000–$1,000,000.')
bullet('Persistence: lognormal, 90% CI 5–20 years.')

bold_para('Calculation.  ',
    'A binary success indicator is drawn by comparing the success probability draw against '
    'a uniform random draw — a Bernoulli trial performed element-wise across all 100,000 '
    'samples. In samples where the intervention fails, all welfare gains are zeroed out. '
    'Annual suffering averted = (BSF born × proportion affected × DALY fraction per BSF '
    '× proportion reduced) × success indicator. This annual figure is multiplied by '
    'persistence and divided by cost, then scaled to per $1,000.'
)

# ── Wild Mammals ──────────────────────────────────────────────────────────────

h2('Wild Mammals: Rat Contraception as a BOTEC')

bold_para('Purpose and source.  ',
    'This model estimates the cost-effectiveness of wild mammal welfare interventions '
    'using rodent contraception as a concrete back-of-the-envelope calculation (BOTEC). '
    'It models a campaign to reduce rat poisoning in urban environments, using pain '
    'duration data from the Welfare Footprint Initiative\'s GPT Pain Track.'
)

bold_para('Input parameters.  ', 'Six parameters drive this model:')

bullet('Target population: lognormal, 90% CI of 4,100–56,000 rats. This corresponds to '
       '0.1% of urban rats in a city of ~300,000 people, derived from the assumption of '
       '0.05–0.25 rats per person in a typical urban area, applied to the US urban population.')
bullet('Suffering hours per rat from conventional poisoning: lognormal, 90% CI 60–330 '
       'hours, mean ~170 hours. The Welfare Footprint Initiative BOTEC breaks this down as: '
       '7.2–28.8 hours "annoying," 19.2–64.8 hours "hurtful," and 14.4–45.6 hours '
       '"disabling" pain, with 0 hours excruciating.')
bullet('Probability of campaign success: beta(4, 16), mean 20%. The campaign requires '
       'lobbying a municipality to implement rat birth control.')
bullet('Proportion of poisoning deaths averted if successful: beta(2, 2), symmetric mean '
       '50%, reflecting uncertainty about whether the campaign substantially displaces poisoning.')
bullet('Impact period: lognormal, 90% CI 5–20 years.')
bullet('Campaign cost: lognormal, 90% CI $100,000–$10,000,000, clipped to '
       '$10,000–$15,000,000.')

bold_para('Calculation.  ',
    'As with BSF, a binary success draw is applied element-wise. Suffering-years per dollar '
    '= (target population × suffering hours / HOURS_PER_YEAR × proportion of deaths averted '
    '× impact period × success indicator) / campaign cost, scaled to per $1,000.'
)

# ── Wild Invertebrates ────────────────────────────────────────────────────────

h2('Wild Invertebrates')

bold_para('Purpose and source.  ',
    'This model parallels the BSF farmed invertebrate model but represents wild '
    'invertebrate welfare interventions, considered even more speculative.'
)

bold_para('Input parameters.  ',
    'Most parameters are shared with the BSF model — the same distributions for population '
    'size, proportion affected, hours of suffering, proportion of suffering reduced, cost, '
    'and persistence. The key difference is the probability of success: beta(1, 9), mean '
    '10% — half the BSF probability — reflecting the substantially greater difficulty of '
    'achieving welfare improvements in wild contexts.'
)

bold_para('Calculation.  ',
    'Identical structure to BSF: Bernoulli success draw, annual averted suffering from '
    'population and per-individual suffering, multiplied by persistence, divided by cost, '
    'scaled to per $1,000.'
)

# ── Wild Animal Welfare Mixture ───────────────────────────────────────────────

h2('Wild Animal Welfare: Mixture Distribution')

para(
    'Rather than committing to a single representative intervention for wild animal welfare, '
    'the script constructs a mixture of the wild mammal and wild invertebrate estimates. '
    'This reflects uncertainty about whether future funding will flow primarily toward '
    'mammal-focused or invertebrate-focused work.'
)

para(
    'A funding share for wild mammals is drawn from a beta(1, 1) distribution, which is '
    'uniform over [0, 1] — meaning every split from 100% mammals to 100% invertebrates is '
    'considered equally plausible. The overall estimate is: mammal share × wild mammal '
    'SY/$1,000 + (1 − mammal share) × wild invertebrate SY/$1,000. Because both the shares '
    'and the underlying estimates are random draws, the result is a proper mixture that '
    'integrates over uncertainty in funding allocation.'
)

# ── Policy Advocacy ───────────────────────────────────────────────────────────

h2('Policy Advocacy')

para(
    'Policy advocacy is modeled as a derived intervention — its effectiveness is pegged as '
    'a discounted version of direct farmed animal interventions. The reasoning is that policy '
    'work influences the same underlying systems (cage-free standards, slaughter practices) '
    'but operates more indirectly.'
)

bold_para('Calculation.  ',
    'A 60% chicken / 40% shrimp weighted average of those interventions\' per-$1,000 '
    'distributions is computed, then a 50% discount is applied. This reflects the analyst\'s '
    'view that the indirect pathway through policy roughly halves expected impact per dollar '
    'relative to direct corporate campaigns.'
)

# ── Movement Building ─────────────────────────────────────────────────────────

h2('Movement Building')

para(
    'Movement building is similarly a derived, indirect intervention — one that supports '
    'the growth and capacity of the broader animal advocacy ecosystem.'
)

bold_para('Calculation.  ',
    'The same 60% chicken / 40% shrimp blend is used, but a 75% discount is applied '
    '(leaving 25% of the blend value). This reflects a more conservative view of the '
    'causal pathway from movement-building activities to on-the-ground animal welfare impact.'
)

# ── Outputs ───────────────────────────────────────────────────────────────────

h1('Outputs')

para(
    'The script produces four types of outputs for each intervention, written at the end '
    'of the script after all distributions have been computed.'
)

h2('YAML File')

para(
    'A primary YAML file contains metadata and results for all interventions. For each, '
    'the file records the percentile summary (p1, p5, p10, p50, p90, p95, p99, and mean) '
    'and a downsampled array of 10,000 empirical samples drawn from the full 100,000-sample '
    'distribution using quantile spacing. Including actual samples — rather than only '
    'percentiles or fitted distribution parameters — eliminates the need for downstream '
    'distribution fitting and avoids information loss and fitting errors.'
)

h2('.npz File')

para(
    'A NumPy compressed archive (.npz) contains the full 100,000-sample arrays for all '
    'interventions. This is intended for direct programmatic use in Python-based downstream '
    'pipelines where the full distributional fidelity is needed — for example, in '
    'risk-adjustment calculations or cross-cause comparisons.'
)

h2('CSV Statistics File')

para(
    'A CSV file contains extended percentile summaries for each intervention, including '
    'additional quantiles (p0.15, p2.5, p97.5, p99.85) beyond the standard seven. This '
    'file is intended for quality control review, making it easy to spot anomalies such as '
    'unexpectedly heavy tails or implausible minima.'
)

h2('Histograms')

para(
    'For each intervention, a PNG histogram of the distribution is saved on a log scale. '
    'Log-scale plotting is used because the distributions are highly right-skewed. The '
    'histograms mark the median, mean, and 10th–90th percentile range. Samples that are '
    'exactly zero (from failed Bernoulli draws) are excluded from the plot but their '
    'fraction is noted in the title. These images are intended for visual quality control: '
    'an analyst can quickly inspect whether a distribution has the expected shape, whether '
    'clipping is producing artifacts at the boundaries, and whether any intervention shows '
    'an implausibly extreme tail.'
)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = os.path.join(_DOCS_DIR, 'aw_intervention_models_documentation.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
