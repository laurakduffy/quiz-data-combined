"""Generate a Word doc walkthrough of every file and function in the project."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)

def h4(doc, text):
    doc.add_heading(text, level=4)

def body(doc, text):
    doc.add_paragraph(text)

def fn(doc, name, description):
    """Add a function entry: bold name + plain description."""
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("  —  " + description)

def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(text)

def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

add_title(doc, "All-Intervention-Models: Complete Code Walkthrough")
body(doc, "Plain-English description of every file and function.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "Contents")
for item in [
    "1. Root-Level Files",
    "   risk_profiles.py",
    "   combine_data.py",
    "   run_all.py",
    "   validate_output.py",
    "2. GCR Models (gcr-models/)",
    "   gcr_model.py",
    "   fund_profiles.py",
    "   export_rp_csv.py",
    "3. GiveWell Models (gw-models/)",
    "   gw_cea_modeling.py",
    "4. LEAF Models (leaf-models/)",
    "   leaf_cea_model.py",
    "5. Animal Welfare Models (aw-models/)",
    "   src/models/effects.py",
    "   src/models/allocate_to_periods.py",
    "   src/pipeline/build_dataset.py",
    "   src/pipeline/export.py",
    "6. Things Worth Vetting",
]:
    body(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — ROOT-LEVEL FILES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "1. Root-Level Files")

# ── risk_profiles.py ─────────────────────────────────────────────────────────

h2(doc, "risk_profiles.py")
body(doc,
    "This is the shared math library used by all three models. It takes an array of "
    "simulated outcome values (e.g., life-years per $1M) and computes 9 differently "
    "risk-weighted summaries of that distribution. Every fund model calls "
    "compute_risk_profiles() and gets back the same 9 named floats."
)

fn(doc, "compute_dmreu(samples, p=0.05)",
    "Implements the Difference-Making Risk-Weighted Expected Utility formula from Duffy (2023). "
    "It sorts the samples from smallest to largest, then assigns each sample a weight based on "
    "a probability-distortion function: weight = P^a, where a = -2 / log10(0.05) ≈ 0.431 and P "
    "is how often you'd see an outcome at least that large. This downweights very high-outcome "
    "tail samples relative to a plain mean, producing moderate risk aversion. The p=0.05 parameter "
    "was chosen by Duffy as 'the probability of a rare but clearly significant event' — it "
    "controls how much aversion you get."
)

fn(doc, "compute_wlu(samples, c)",
    "Implements Weighted Linear Utility from Duffy (2023). For each sample value, computes a "
    "weight of 1 / (1 + |value|^c). Because this fraction shrinks as value grows, smaller "
    "outcomes get higher weight than larger ones — the utility function is concave. Negative "
    "outcomes get a flipped weight (2 - w), making them weighted more heavily still. The "
    "weights are renormalized to sum to 1 before computing the weighted mean. Higher c = more "
    "risk aversion. Called three times with c=0.01, 0.05, 0.10 to produce the three WLU variants "
    "(low, moderate, high)."
)

fn(doc, "compute_ambiguity(samples)",
    "Applies a pure upside-skepticism weighting with no loss aversion component. Samples below "
    "the 97.5th percentile keep full weight (1.0). Samples between the 97.5th and 99.9th percentile "
    "are exponentially down-weighted — the decay rate is chosen so that weight falls by a factor "
    "of 100 over that 2.4-percentile band. Samples above the 99.9th percentile get weight 0. "
    "The interpretation: 'I don't trust outcomes that look implausibly good, so I discount them "
    "progressively as they approach and then exceed the 99.9th percentile.'"
)

fn(doc, "compute_risk_profiles(samples)",
    "The main entry point. Calls all of the above helpers and also computes four informal "
    "adjustments inline: "
    "(1) neutral — plain mean. "
    "(2) upside — clips everything above the 99th percentile to the 99th percentile value, then "
    "takes the mean (eliminates the very top tail). "
    "(3) downside — computes a loss-averse utility around the median: gains above the median are "
    "counted at face value, losses below the median are multiplied by 5.0 (lambda), then the "
    "mean of those adjusted gains/losses is added back to the median. "
    "(4) combined — same exponential weight decay as ambiguity, but also applies the 5.0× loss "
    "aversion for below-median outcomes before averaging. Returns a dict mapping each of the 9 "
    "profile names to a float."
)

# ── combine_data.py ───────────────────────────────────────────────────────────

h2(doc, "combine_data.py")
body(doc,
    "The master assembly script. It reads the output CSVs from all three models, normalizes "
    "column names and IDs to a shared standard, and produces output_data_{scenario}.json "
    "(the full nested structure consumed by the front-end) and all_risk_adjusted.csv "
    "(a flat version of the same data). This file runs as a script — "
    "the top-level code executes directly rather than inside a main() function."
)
body(doc,
    "Key constants at the top: BUDGET_M=897 (total portfolio budget in $M), "
    "INCREMENT_SIZE=10 ($M steps for the front-end slider). "
    "GCR_DMR_SCENARIO is set via the --gcr-dmr-scenario argument (default: 'median'); "
    "it selects which GCR diminishing returns CSV to load from gcr-models/diminishing_returns/."
)

fn(doc, "parse_diminishing_returns(df)",
    "Takes the combined diminishing-returns dataframe (all three models concatenated) and turns "
    "it into a dict keyed by standardized project ID, where each value is a list of DR "
    "multipliers at $10M, $20M, ..., $900M. Handles GiveWell's percentage-string format "
    "(e.g., '75%') by converting to 0.75. Translates raw fund keys to canonical names using "
    "FUND_NAME_MAP (e.g., 'sentinel' → 'sentinel_bio')."
)

fn(doc, "parse_effects(df)",
    "Takes the combined risk-scores dataframe and builds the nested JSON structure. For each row: "
    "(1) maps the project ID and effect ID to canonical names; "
    "(2) maps recipient_type strings, with special overrides for certain AW effect IDs "
    "(e.g., movement_building is always chickens_birds regardless of the generic type field); "
    "(3) builds the 6×8 values matrix (6 time periods × 8 risk profiles) by trying both the "
    "standard {rp}_{t0} column naming and the AW model's {rp}_{0_to_5} naming — whichever "
    "exists in the dataframe; for AW funds with only 4 periods (0–100 years), the t4 and t5 "
    "entries are set to 0.0 by the absence of those columns; "
    "(4) applies NEAR_TERM_XRISK_OVERRIDES to force Sentinel Bio and Longview Nuclear to False "
    "regardless of what's in the CSV. After the function, the top-level code merges "
    "sub-extinction tier sub-projects (e.g., sentinel_bio_100m_1b) back into their parent "
    "projects and deletes the sub-fund entries. "
    "Note: combine_data.py reads AW diminishing returns per-fund (ea_awf_diminishing_returns.csv "
    "and navigation_fund_diminishing_returns.csv) and the LEAF fund "
    "(leaf_diminishing_returns.csv) in addition to GW and GCR files."
)

# ── validate_output.py ────────────────────────────────────────────────────────

h2(doc, "validate_output.py")
body(doc,
    "A standalone consistency checker. Run it after combine_data.py to confirm nothing was "
    "lost or mis-mapped in the assembly step. It cross-checks the JSON against both the "
    "all_risk_adjusted.csv and the source diminishing-returns CSVs."
)

fn(doc, "rel_close(a, b, tol=1e-6)",
    "Returns True if two floats are within 1-in-a-million relative tolerance. Handles NaN "
    "(both must be NaN to match), infinities (must be exactly equal), and near-zero "
    "denominators (falls back to absolute tolerance)."
)

fn(doc, "parse_pct(val)",
    "Strips the % sign from GiveWell percentage strings and divides by 100. Returns a plain float."
)

fn(doc, "validate_risk_scores(projects, csv_path)",
    "Cross-checks the JSON projects dict against all_risk_adjusted.csv. For every row in the "
    "CSV it verifies: (1) the project exists in JSON, (2) the effect exists under that project, "
    "(3) recipient_type matches, (4) near_term_xrisk matches, (5) all 48 values in the 6×8 "
    "matrix match within floating-point tolerance. Also checks the reverse — that every JSON "
    "effect appears in the CSV. Reports all mismatches found."
)

fn(doc, "validate_diminishing_returns(projects)",
    "Rebuilds the expected diminishing-returns dict from the three source CSVs (mirroring the "
    "logic in parse_diminishing_returns) and compares it element-by-element against the JSON. "
    "Skips sub-extinction tier IDs since those don't have DR curves."
)

fn(doc, "main()",
    "Loads output_data_{scenario}.json (where scenario is the --gcr-dmr-scenario argument, "
    "default: median), calls both validators, prints a pass/fail summary."
)

# ── run_all.py ────────────────────────────────────────────────────────────────

h2(doc, "run_all.py")
body(doc,
    "Top-level convenience script that runs every model pipeline in sequence with a single command. "
    "Calls subprocess.run() on each script in order: "
    "(1) aw-models/data/inputs/aw_intervention_models.py — regenerates CCM intervention samples; "
    "(2) aw-models/run.py — runs the AW fund pipeline for all three funds; "
    "(3) gw-models/gw_cea_modeling.py — runs the GiveWell model; "
    "(4) leaf-models/leaf_cea_model.py — runs the LEAF model; "
    "(5) gcr-models/export_rp_csv.py — runs all three GCR funds; "
    "(6) combine_data.py — assembles everything into output_data_{scenario}.json and all_risk_adjusted.csv; accepts --gcr-dmr-scenario (default: median). "
    "Each script is run with check=True, so a failure in any step aborts the rest. "
    "This is the recommended way to refresh all data after changing any parameters."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — GCR MODELS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "2. GCR Models (gcr-models/)")

# ── gcr_model.py ─────────────────────────────────────────────────────────────

h2(doc, "gcr_model.py")
body(doc,
    "The core physics and math engine for all three GCR funds. Implements the Tarsney (2020) "
    "framework for computing the expected value of reducing existential risk, including both "
    "the near-term 'time of perils' phase and (optionally) very long-run stellar expansion value."
)

h3(doc, "GCRParams (dataclass)")
body(doc,
    "A plain container for all input parameters. Every array-valued parameter has shape "
    "(n_sims,) — one value per scenario — so the model can run multiple scenarios simultaneously "
    "in a vectorized fashion. Parameters are grouped into: risk trajectory (Gaussian peak shape, "
    "background risk level), intervention effect (when it starts, how long it lasts, how much it "
    "reduces the Gaussian risk peak), Earth value growth (logistic curve parameters), and stellar "
    "expansion (cubic growth switch, colonization start year, settlement speed). There is no "
    "computation in this class — it just holds numbers."
)

h3(doc, "Module-level function: _solve_r_max(...)")
fn(doc, "_solve_r_max(cumulative_risk, year_max_risk, year_risk_1pct_max, r_inf)",
    "Solves backwards for the peak annual risk r_max of the Gaussian 'time of perils' curve that "
    "would produce a given cumulative 100-year extinction probability. For example, if you want "
    "the curve to imply a 10% chance of extinction over 100 years, this finds the r_max that "
    "achieves exactly that. Uses vectorized bisection (60 iterations, giving ~1e-18 precision). "
    "The Gaussian sigma is set as year_risk_1pct_max / 3, so that annual risk falls to 1% of its "
    "peak at 'year_risk_1pct_max' years away from the peak year. Called once on model initialization "
    "for every scenario in the batch."
)

h3(doc, "GCRModel methods")

fn(doc, "__init__ and _derive()",
    "On construction, calls _solve_r_max to get r_max for each scenario, then computes the "
    "intervention's relative risk reduction (rel_rr_from_int) in one of three ways depending "
    "on which parameter was specified: (1) rel_risk_reduction × cause_fraction if that was set "
    "directly; (2) abs_risk_reduction / r_max to express an absolute annual reduction as a "
    "fraction of the peak; or (3) bp_reduction_per_bn × budget / 1B for the legacy "
    "basis-points parameterization. Also pre-computes stellar expansion geometry constants "
    "(b2 = stellar density difference × galaxy radius; T_s = time to reach galaxy edge)."
)

fn(doc, "get_annual_risk_level(t, I)",
    "Returns the annual probability of catastrophe at year t. If I=0 (no intervention): "
    "r_inf + r_max × gaussian(t). If I=1 (with intervention): the Gaussian amplitude is "
    "reduced by rel_rr_from_int, but only during the years when the intervention is active "
    "(from year_effect_starts to year_effect_starts + persistence_effect). Outside that window, "
    "even with the intervention active, risk reverts to the baseline Gaussian."
)

fn(doc, "get_p_survival_vec(risk_2d)",
    "Converts an array of annual risks into cumulative survival probabilities using "
    "cumprod(1 - risk). Entry t gives the probability of surviving all the way from year 0 "
    "through year t."
)

fn(doc, "get_year_of_const_risk(I)",
    "Scans forward in time from the Gaussian peak year until annual risk has decayed to within "
    "1% of background rate r_inf. This marks the year the 'crisis period' is effectively over "
    "and no further year-by-year simulation is needed for the risk component."
)

fn(doc, "value_level_logistic(t)",
    "Returns Earth's welfare-weighted 'value at stake' at year t using a logistic growth curve "
    "parametrized by initial_value, carrying_capacity, and rate_growth. Think of this as the "
    "size of the stake that gets extinguished if catastrophe occurs."
)

fn(doc, "get_year_logistic_ends()",
    "Scans forward until the logistic value reaches 99% of carrying capacity. After this point "
    "Earth's value is held constant, so the explicit year-by-year simulation can stop."
)

fn(doc, "get_year_const_value_and_risk(I)",
    "Returns max(year_risk_converges, year_value_converges) — the first year after which both "
    "the risk trajectory and the value trajectory are effectively constant. This is the cutoff "
    "year for the explicit simulation; everything after is handled analytically."
)

fn(doc, "get_earth_value(t, y_const_value)",
    "Returns the logistic value up to the convergence year, then holds it at 99% of carrying "
    "capacity. Just a wrapper that applies the plateau."
)

fn(doc, "get_value_stars_settled(t)",
    "Models the value from interstellar colonization if cubic_growth=True and t >= T_c "
    "(colonization start year). Within the Milky Way, value grows as a1 × (t-T_c)^3. Beyond "
    "the Milky Way edge (reached at T_s = galaxy_radius / settlement_speed), it transitions "
    "to a2 × (t-T_c)^3 + b2, where a2 and b2 reflect the lower density of stars in the "
    "broader Virgo supercluster. The cubic exponent comes from volume = (4/3)π r^3 and r = s(t-T_c)."
)

fn(doc, "get_total_value_level(t, y_const_value)",
    "Simply adds Earth value and stellar value together. This is the total value at stake at year t."
)

fn(doc, "get_conditional_future_value_on_earth(n_years)",
    "Computes the analytical integral of expected Earth value from the convergence year to "
    "T_h (10^14 years), discounted by ongoing background extinction risk r_inf. The formula "
    "is v_const / r_inf × (1 - exp(-r_inf × (T_h - n_years))). Because r_inf is tiny, this "
    "is approximately v_const × (T_h - n_years) — the vast majority of long-run value."
)

fn(doc, "get_conditional_future_value_stars_to_Ts2() / Ts3() / Th3()",
    "Three variants that compute analytical integrals of expected stellar value, discounted "
    "by r_inf, for different segments of the future timeline. The Ts2/Ts3 distinction handles "
    "whether the colonization start year T_c is before or after the convergence year n_years "
    "(Ts3 = T_c is already in the past relative to n_years; Ts2 = T_c is still in the future). "
    "Th3 handles the segment from the galaxy edge T_s to the time horizon T_h. These use "
    "the standard integral of t^3 × exp(-rt), which has a closed-form solution involving "
    "polynomials in t times exp(-rt)."
)

fn(doc, "run(verbose=False)",
    "The main computation loop. Steps: "
    "(1) Determine how many years to simulate (convergence year for each scenario). "
    "(2) Compute stellar expansion coefficients a1, a2 from Earth's value at year T_c. "
    "(3) Build year-by-year risk arrays for both with-intervention (I=1) and "
    "without-intervention (I=0) worlds. "
    "(4) Compute cumulative survival arrays and take the difference: "
    "diff_in_survival[t] = P(survive to t WITH intervention) - P(survive to t WITHOUT). "
    "This is the probability that year t's value is realized *because of* the intervention. "
    "(5) For each time period (0-5, 5-10, ..., 100-500 years), sum value[t] × diff_in_survival[t] "
    "over all years in that window. This gives the expected additional value the intervention "
    "creates in that period. "
    "(6) After the convergence year, switch to the analytical long-term integrals. "
    "(7) Assemble everything into a per-period dict and scale to per-$1M."
)

fn(doc, "run_monte_carlo(sweep_params, fixed_params, n_samples, p_harm, p_zero, harm_multiplier, seed)",
    "Wraps run() to handle thousands of parameter samples with stratified sampling. "
    "Key design choices: "
    "(1) Stratified sampling — identifies cubic_growth, T_c, and r_inf as the parameters that "
    "most affect the shape of the time-period distribution. Divides n_samples proportionally "
    "across all combinations of those parameters, respecting user-specified probability weights "
    "(e.g., cubic_growth=True has probability 10%, so exactly 10% of samples get it). This "
    "prevents rare-but-important scenarios from being undersampled. "
    "(2) Three outcome types — within each stratum, assigns samples as 'positive' (helps), "
    "'zero' (no effect), or 'harm' (backfires), using the fund-profile probabilities p_positive "
    "= 1 - p_zero - p_harm. Uses probabilistic rounding to ensure exact integer counts with "
    "no systematic bias. Harm samples have their per-period EV values negated and multiplied "
    "by harm_multiplier. "
    "(3) All n_samples run through GCRModel at once as a single vectorized batch. "
    "Returns the full array of total EV values plus per-period arrays, summary percentiles, "
    "and a list of per-sample parameter dicts."
)

fn(doc, "make_original_notebook_params()",
    "Factory that returns GCRParams matching the original notebook's hardcoded defaults. "
    "Useful as a regression test — running these through GCRModel should reproduce the "
    "known outputs from the original implementation."
)

# ── fund_profiles.py ──────────────────────────────────────────────────────────

h2(doc, "fund_profiles.py")
body(doc,
    "All domain-knowledge assumptions about each GCR fund. No computation here — just "
    "configuration. The FUND_PROFILES dict holds one entry per fund with all parameters "
    "needed to run the Monte Carlo."
)
body(doc,
    "Key global constants: "
    "Cause fractions: AI ~90%, Nuclear ~3%, Bio ~3% of total 100-year x-risk. "
    "Total x-risk scenarios: [5%, 10%, 65%] for conservative/central/optimistic. "
    "All three funds use the same relative-risk-reduction ladder per $10M: "
    "[0.002/50, 0.002/10, 0.002], giving a 50× range from conservative to optimistic."
)

fn(doc, "_r_max_from_cumulative_risk(cumulative_risk_100_yrs, ...)",
    "A thin wrapper around _solve_r_max that accepts scalars or arrays and returns the peak "
    "annual Gaussian risk for a given cumulative 100-year extinction probability. Used in the "
    "__main__ block to generate calibration CSVs, not during normal model runs."
)

fn(doc, "get_fund_profile(fund_key)",
    "Looks up a fund by key (e.g., 'sentinel', 'longview_nuclear', 'longview_ai'), deep-copies "
    "its profile dict so downstream changes can't corrupt the master, adds a fund_key field, "
    "and sets adjustment_factor = counterfactual_factor. The deep copy is important because "
    "the sweep_params dicts are mutable."
)

fn(doc, "list_fund_profiles()",
    "Returns a sorted list of valid fund key strings."
)

fn(doc, "make_earth_only_profile(profile)",
    "Returns a variant of a fund profile with stellar expansion disabled: cubic_growth forced "
    "to False, T_c and s removed from sweep_params and fixed at default values. Used when you "
    "want to evaluate the fund without any long-termist stellar value."
)

fn(doc, "__main__ block",
    "When fund_profiles.py is run directly (not imported), prints a formatted calibration table "
    "showing, for each fund and scenario, the relative risk reduction per $10M, the total "
    "relative risk reduction given the fund budget, and what percentage of the Gaussian r_max "
    "that represents under each of the three cumulative-risk scenarios. Then writes two CSVs: "
    "calibration_abs_risk_reduction_detail.csv (all 9 scenario combinations per fund) and "
    "calibration_abs_risk_reduction_summary.csv (summary stats across scenarios)."
)

# ── export_rp_csv.py ──────────────────────────────────────────────────────────

h2(doc, "export_rp_csv.py")
body(doc,
    "The top-level runner for GCR models. Calls run_monte_carlo for each of the three funds, "
    "then exports effects CSVs, a separate diminishing-returns CSV, histograms, and summary "
    "statistics. Run with: python export_rp_csv.py [--n-samples N] [-o output.csv] [--quiet]"
)

fn(doc, "_eval_diminishing_raw(budget_m, anchors, spend_m)",
    "Evaluates a piecewise-linear diminishing-returns curve at a given cumulative spend level. "
    "Below the first anchor: flat at the first CE multiplier. Between anchors: linear "
    "interpolation. Beyond the last anchor: falls off as last_ce × (last_multiple / "
    "current_multiple) — a 1/x hyperbolic tail, meaning CE keeps declining but never hits zero."
)

fn(doc, "compute_diminishing_row(budget_m, anchors)",
    "Calls _eval_diminishing_raw at every $10M step from $10M to $900M (90 points total), "
    "then normalizes so the first point ($10M) = 1.0. Returns a list of 90 values representing "
    "the marginal CE at each cumulative spend level relative to the CE at $10M."
)

fn(doc, "_years_in_period(persistence, start, end)",
    "Given a total persistence duration (e.g., 15 years), computes how many of those years "
    "fall within a specific time window (e.g., years 5–10). Simple overlap: "
    "min(persistence, end) - max(0, start), floored at 0."
)

fn(doc, "_compute_sub_extinction_rows(profile, n_samples, verbose)",
    "Runs a simplified Monte Carlo for the sub-extinction tiers (100M-1B deaths, 10M-100M "
    "deaths). Instead of the full Tarsney model, uses: "
    "annual_EV = P(event/year) × expected_deaths × rel_risk_reduction × counterfactual_factor × discount. "
    "Stratifies samples across the sweep_rel_rr × sweep_persistence grid (same approach as the "
    "main MC). Applies the same harm mask (p_harm fraction of samples get their EV negated). "
    "Multiplies annual_EV by years-in-each-period to allocate across the 6 time windows, "
    "then calls compute_risk_profiles on the per-$1M array for each period."
)

fn(doc, "run_fund_and_extract(fund_key, n_samples, verbose)",
    "The main function that runs one fund end-to-end. "
    "Calls run_monte_carlo, then: "
    "(1) scales each period's EV array from total-EV to per-$1M by multiplying by "
    "counterfactual_factor / budget × 1e6; "
    "(2) calls compute_risk_profiles on each period's 100k-sample array to get the 9 "
    "risk-adjusted summaries; "
    "(3) computes the 'after 500+' period as a residual: total_per_1M minus the sum of "
    "all five explicit periods (so the six periods always sum to the total); "
    "(4) calls _compute_sub_extinction_rows for the two sub-extinction tiers. "
    "Returns everything needed for CSV export and histogram creation."
)

fn(doc, "write_rp_csv(fund_results, output_path, verbose)",
    "Writes gcr_output.csv in the standard RP format. Columns: project_id, near_term_xrisk, "
    "effect_id, recipient_type, then 54 value columns named {risk_profile}_t{0-5} for each "
    "of 9 risk profiles × 6 time periods. One row per fund (plus sub-extinction tier rows)."
)

fn(doc, "write_diminishing_returns_csv(fund_results, output_path, verbose)",
    "Writes {scenario}_diminishing_returns_gcr.csv (where scenario is --gcr-dmr-scenario) with one row per fund. Columns: project_id "
    "followed by 90 DR multiplier values at $10M, $20M, ..., $900M, normalized to the $10M value."
)

fn(doc, "validate_output(output_path, n_funds, n_effect_rows, verbose)",
    "Light structural validation: checks that the CSV starts with the correct header row and "
    "that all effect rows have a non-empty project_id. Does not validate numeric values — "
    "that is handled by the root-level validate_output.py."
)

fn(doc, "create_and_save_histograms(fund_results, output_dir, verbose)",
    "For each fund (and each sub-extinction tier), saves a PNG with two side-by-side panels: "
    "a linear-scale histogram of total QALYs/$1M across 100k samples, and a log-scale "
    "histogram showing only positive values (with the percentage of zero/negative samples "
    "noted in the title). Both panels show the mean and median as vertical dashed lines."
)

fn(doc, "write_summary_statistics(fund_results, output_path, verbose)",
    "Writes gcr_output_summary_stats.csv. One row per fund and per sub-extinction tier, "
    "with columns: fund ID, display name, tier name, n_samples, mean, and percentiles "
    "p1/p5/p10/p50/p90/p95/p99 of total QALYs/$1M."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — GIVEWELL MODELS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "3. GiveWell Models (gw-models/)")


h2(doc, "gw_cea_modeling.py")
body(doc,
    "Generates 10,000 simulated estimates of GiveWell's portfolio cost-effectiveness broken "
    "down by effect type (YLDs averted, life-years saved, income doublings) and time period, "
    "then applies all 9 risk profiles to produce the standard RP output format."
)
body(doc,
    "Key constants at the top: UNITS_VALUE_PER_M_PER_X_CASH = 3280 (GW units of moral value "
    "per $1M per 'x cash' multiple, calculated offline). LIFE_YEARS_PER_LIFE = 60 (assumed "
    "life-years per life saved). N_SAMPLES = 10,000."
)
body(doc,
    "Portfolio CE distribution (module-level): Three squigglepy distributions representing "
    "GW's portfolio split — 5% of spending below 8x cash effectiveness (lognormal 2-8), "
    "68% between 8x and 16x (normal 8-16), 27% above 16x (lognormal 16-44). These are "
    "sampling distributions for each tier, not point estimates."
)

fn(doc, "sample_units_value_per_m()",
    "Draws 10,000 samples from GiveWell's portfolio CE distribution. For each sample, "
    "independently draws from each of the three tier distributions, then computes the "
    "portfolio-weighted sum (5% × below_8x + 68% × 8to16x + 27% × above_16x). Multiplies "
    "by UNITS_VALUE_PER_M_PER_X_CASH to convert from 'cash multiples' to absolute units "
    "of GW moral value per $1M."
)

fn(doc, "get_weighted_average_percent_effect_by_type(percent_effect_by_type_dict, percent_funding_by_dist_dict)",
    "Computes the portfolio-weighted breakdown of what fraction of GW's total moral value "
    "is YLDs averted vs. lives saved vs. income doublings. Takes the 8-cause-area effect "
    "breakdowns, weights each by its share of GW funding, sums, and renormalizes. The "
    "result is three fractions (e.g., ~13% YLDs, ~52% lives, ~35% income) that reflect "
    "GW's actual portfolio mix."
)

fn(doc, "get_sample_units_value_by_type(sample_units_value_per_M, weighted_average_percent_effect_by_type)",
    "Multiplies the 10k portfolio CE samples by each effect-type fraction to produce three "
    "10k arrays: units of GW moral value attributable to YLDs averted, lives saved, and "
    "income doublings respectively, per $1M."
)

fn(doc, "get_distribution_effect_per_M(sample_effect_by_type)",
    "Divides each effect-type array by GW's moral weights (YLDs=2.3, lives=115.6, income=1). "
    "This converts from 'units of GW moral value' into natural physical units: "
    "YLD-equivalents per $1M, life-equivalents per $1M, income-doubling-equivalents per $1M."
)

fn(doc, "get_effect_per_M_by_time(distribution_effect_by_type, temporal_breakdown_by_type_dict)",
    "Multiplies each effect-type's 10k samples by its temporal allocation fractions to produce "
    "a nested dict {effect_type: {time_window: 10k array}}. "
    "Health effects (YLDs, lives) are heavily front-loaded: 90% in years 0-5, 7% in 5-10, "
    "2.5% in 10-20, 0.5% in 20-100, 0% thereafter — reflecting that health interventions "
    "save lives now. "
    "Income doublings are more spread out: 18% in 0-5, 1.4% in 5-10, 12.5% in 10-20, 68% "
    "in 20-100 — reflecting that economic gains compound across generations."
)

fn(doc, "convert_lives_saved_to_life_years_saved(effect_per_M_by_time)",
    "Multiplies the 'lives saved' arrays by 60 (LIFE_YEARS_PER_LIFE) and renames the key "
    "from lives_saved to life_years_saved. Modifies the dict in-place to match the standard "
    "effect_id naming used across all models."
)

fn(doc, "create_summary_statistics(effect_per_M_by_time)",
    "Computes mean, p5, and p95 for every effect-type × time-window combination and writes "
    "them to summary_statistics.csv."
)

fn(doc, "create_and_save_histograms(distribution_effect_by_type)",
    "Saves one histogram per effect type to the histograms/ directory, showing the "
    "distribution of that effect type's CE across 10k samples."
)

fn(doc, "apply_risk_adjustments_to_simulations(effect_per_M_by_time)",
    "The final step that produces the standard RP output. For each of the 3 effect types, "
    "for each of the 6 time windows, calls compute_risk_profiles on the 10k samples and "
    "stores all 9 risk-adjusted values as columns named {rp}_t{0-5}. Returns a 3-row "
    "DataFrame (one row per effect type) in the RP format, which is saved as gw_risk_adjusted.csv."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — LEAF MODELS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "4. LEAF Models (leaf-models/)")

h2(doc, "leaf_cea_model.py")
body(doc,
    "Generates 10,000 simulated estimates of LEAF's cost-effectiveness broken down by effect "
    "type (YLDs averted, life-years saved, income doublings) and time period, then applies all "
    "9 risk profiles. The structure mirrors gw_cea_modeling.py closely."
)
body(doc,
    "Key constants: N_SAMPLES = 10,000. Three GEV (Generalized Extreme Value) distributions — "
    "one per effect type — parameterized by shape, location, and scale. These were fit to "
    "analyst estimates of LEAF's impact distribution."
)

fn(doc, "sample_impacts_per_m()",
    "Draws 10,000 samples from each GEV distribution (YLDs averted, life-years saved, income "
    "doublings). Returns a dict of effect_type → sample array. The GEV distribution is used "
    "because LEAF's impact distribution is expected to have a heavier right tail than a normal "
    "or lognormal distribution."
)
body(doc,
    "Temporal breakdown: unlike GiveWell (heavily front-loaded), LEAF's effects are more "
    "back-loaded. For example, life-years saved allocates ~80% to the 20–100 year window. "
    "This reflects that LEAF's longevity research is expected to compound over decades. "
    "The model applies the same 6-period structure as GiveWell and GCR (0-5, 5-10, 10-20, "
    "20-100, 100-500, 500+ years), with 100-500 and 500+ always 0."
)
body(doc,
    "Output: leaf_risk_adjusted.csv (standard RP format), leaf_diminishing_returns.csv, "
    "histograms in leaf-models/histograms/, and leaf-models/summary_statistics.csv."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ANIMAL WELFARE MODELS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "5. Animal Welfare Models (aw-models/)")

h2(doc, "src/models/effects.py")
body(doc,
    "Loads CCM (Charity Cost-effectiveness Model) intervention estimates and fund allocation "
    "splits, then produces effect dicts for each intervention with the right units for the "
    "rest of the pipeline."
)

fn(doc, "load_intervention_estimates(path)",
    "Reads aw_model_intervention_estimates.yaml, which contains analyst-derived cost-effectiveness "
    "distributions for each AW intervention (percentile summaries and 10k downsampled samples). "
    "The full 100k samples are stored separately in the .npz file."
)

fn(doc, "load_full_samples(path)",
    "Loads aw_model_intervention_samples_100k.npz — a NumPy binary archive containing 100,000 "
    "pre-drawn samples per intervention from the CCM. Returns a dict mapping "
    "intervention_key to a numpy array, or None if the file doesn't exist. The 100k samples "
    "are preferred over the 10k YAML samples for higher accuracy in the tails."
)

fn(doc, "load_fund(fund_id, path)",
    "Reads the per-fund YAML file (e.g., aw_combined.yaml), which specifies the fund's "
    "annual budget, room-for-more-funding, and percentage allocation across interventions."
)

fn(doc, "compute_all_effects(fund_key, verbose, use_full_samples)",
    "Main function. For each intervention in the fund's allocation split: "
    "(1) Tries to load 100k samples from the .npz file first, then falls back to 10k samples "
    "from YAML, then to percentiles-only. "
    "(2) Converts from 'suffering-years averted per $1,000 spent on the intervention' to "
    "'suffering-years averted per $1M spent on the fund' by multiplying by 1,000 × the "
    "split fraction. For example, if an intervention gets 30% of fund dollars, multiply by 300. "
    "(3) Returns an effect dict for each intervention carrying the samples, percentile summaries "
    "(for reporting), species type, recipient type, effect start year, and persistence years."
)

h2(doc, "src/models/allocate_to_periods.py")
body(doc,
    "Standalone module that allocates an effect across the four AW time windows based on "
    "when the effect starts and how long it persists. Previously this logic lived in "
    "diminishing_returns.py (which has been removed); it now has its own dedicated file."
)

fn(doc, "years_in_period(persistence, start, end)",
    "Computes how many years of an effect's active lifetime (0 to persistence) overlap "
    "with a specific time window [start, end). Simple min/max overlap calculation. "
    "If end is None, returns all remaining years after start."
)

fn(doc, "allocate_to_periods(effect_start_year, persistence_years)",
    "Given when an effect starts (e.g., year 4) and how long it lasts (e.g., 15 years), "
    "computes what fraction of the total effect falls in each of the 4 AW time windows: "
    "0-5, 5-10, 10-20, and 20-100 years. Period keys are '0_to_5', '5_to_10', '10_to_20', "
    "'20_to_100'. Returns a dict mapping period_key → fraction. "
    "Note: the AW model only uses 4 periods (to year 100), unlike GCR and GW which go "
    "to 500+ years (6 periods). combine_data.py handles the missing t4 and t5 by treating "
    "the absence of those columns as 0.0, which is intentional — animal welfare interventions "
    "are assumed not to have impacts on a civilizational timescale."
)
body(doc,
    "PERIOD_BOUNDS and PERIOD_KEYS are module-level constants: "
    "PERIOD_BOUNDS = [(0,5), (5,10), (10,20), (20,100)] and "
    "PERIOD_KEYS = ['0_to_5', '5_to_10', '10_to_20', '20_to_100']. "
    "These drive the column names in the output CSV (e.g., neutral_0_to_5, upside_10_to_20)."
)

h2(doc, "src/pipeline/build_dataset.py")
body(doc,
    "The orchestrator for the AW pipeline. Calls compute_all_effects, then applies risk "
    "profiles directly to the empirical samples and allocates effects across time periods. "
    "No distribution fitting is performed — the 100k empirical samples from the .npz file "
    "are used directly."
)

fn(doc, "build_all_effects(fund_key, verbose)",
    "Main pipeline function. For each intervention effect returned by compute_all_effects: "
    "(1) Takes the empirical sample array (100k from .npz or 10k fallback from YAML) and "
    "converts it to a numpy array of floats. "
    "(2) Calls compute_risk_profiles(draws) to get the 9 risk-adjusted values. "
    "(3) Calls allocate_to_periods (from allocate_to_periods.py) to get the fraction of "
    "effect in each of the 4 time windows. "
    "(4) Assembles a flat row dict with all metadata plus {rp}_{period_key} columns for "
    "every combination of 9 risk profiles × 4 time windows (= 36 value columns), plus "
    "total_{rp} columns for the undistributed total. "
    "Returns {fund_config, rows (one per intervention), metadata}."
)

h2(doc, "src/pipeline/export.py")
body(doc,
    "Handles all file output for the AW pipeline. Called by aw-models/run.py after "
    "build_all_effects completes. Produces per-fund files (e.g., ea_awf_dataset.csv, "
    "navigation_fund_cagefree_assumptions.md) rather than a single combined file."
)

fn(doc, "export_dataset(dataset, output_path, verbose)",
    "Writes the per-fund dataset CSV (e.g., ea_awf_dataset.csv) in the standard RP format. "
    "Columns: metadata fields (project_id, effect_id, species, recipient_type, fund_split_pct, "
    "effect_start_year, persistence_years, data_source), percentile summary columns "
    "(animal_dalys_per_M_p10 etc.), total risk-adjusted columns (total_neutral, total_upside, ...), "
    "and the 9 risk profiles × 4 time periods = 36 period-allocated columns named {rp}_{period_key} "
    "(e.g., neutral_0_to_5, ambiguity_20_to_100)."
)

fn(doc, "export_assumptions(dataset, output_path, verbose)",
    "Writes a markdown assumptions register (e.g., ea_awf_assumptions.md) listing the fund "
    "configuration, CE source metadata, a summary table per intervention (species, split, "
    "persistence, neutral aDALYs/$1M), key data sources, and caveats. If diminishing returns "
    "data is present in the dataset dict, also includes a diminishing returns section."
)

fn(doc, "export_diminishing(dataset, output_path, verbose)",
    "Writes the fund-level diminishing returns curve to CSV if diminishing returns data is "
    "present in the dataset. If the dataset has no 'diminishing' key (the current default when "
    "DR is not computed), this function exits silently. The CSV has one row per spend point "
    "with columns: spend_M, marginal_ce_multiplier (normalized so the first point = 1.0)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — THINGS WORTH VETTING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "6. Things Worth Vetting")
body(doc,
    "Based on reading the actual code, here are specific things that may warrant a closer look:"
)

h3(doc, "1. validate_output.py: match --gcr-dmr-scenario to combine_data.py")
body(doc,
    "validate_output.py accepts --gcr-dmr-scenario (default: median) and loads the matching "
    "output_data_{scenario}.json and GCR DR CSV. Always pass the same --gcr-dmr-scenario "
    "value to both combine_data.py and validate_output.py, or the validator will compare "
    "against the wrong source file."
)

h3(doc, "2. validate_output.py is missing the dmreu risk profile")
body(doc,
    "The RISK_PROFILES list in validate_output.py (lines 12-21) has 8 entries and is missing "
    "'dmreu'. The list in the actual model outputs has all 9 profiles including dmreu. This "
    "means dmreu values are never validated, and the column-order mapping when indexing into "
    "the values matrix may be off for all profiles listed after where dmreu would appear."
)

h3(doc, "3. AW model only covers 4 time periods (to year 100)")
body(doc,
    "The AW allocate_to_periods.py PERIOD_BOUNDS stops at year 100, so t4 "
    "(100-500 years) and t5 (500+ years) are always 0.0 for animal welfare effects. This is "
    "intentional — AW interventions aren't assumed to have century-scale impacts — but it is "
    "worth confirming this is the intended design rather than an oversight. "
    "combine_data.py handles this gracefully by treating the absence of those columns as 0.0."
)

h3(doc, "4. GiveWell income-doublings temporal fractions sum to 99.9%")
body(doc,
    "In gw_cea_modeling.py (lines 144-146), the temporal fractions for income doublings are: "
    "0.180 + 0.014 + 0.125 + 0.681 + 0 + 0 = 1.000. This is fine numerically. "
    "However, note that 68.1% of income effect is allocated to years 20-100, reflecting "
    "an assumption about multi-generational economic compounding. Worth confirming this "
    "temporal assumption is intentional and matches the model documentation."
)

h3(doc, "5. counterfactual_factor applied outside the Monte Carlo, not inside")
body(doc,
    "In run_fund_and_extract (export_rp_csv.py lines 270-272), the counterfactual_factor "
    "is applied to the per-period EV after the Monte Carlo completes: per_1m = raw * adj / budget. "
    "This means all 100k samples are multiplied by the same scalar adjustment factor. If you "
    "believed there was uncertainty in the counterfactual factor itself, you'd want to sample "
    "it inside the Monte Carlo instead. Currently it's treated as a known constant."
)

h3(doc, "6. Sub-extinction tier harm uses a shared mask across tiers")
body(doc,
    "In _compute_sub_extinction_rows (export_rp_csv.py lines 158-175), a single "
    "shared_causes_harm boolean array is built once and applied to both the 100M-1B tier "
    "and the 10M-100M tier. This means if a sample is 'harmful' in one tier, it is also "
    "harmful in the other. This is a deliberate modeling choice (the same intervention "
    "would backfire on both scales), but it is worth confirming this correlation structure "
    "is intended rather than accidental."
)

h3(doc, "7. AW sensitivity analysis has been removed")
body(doc,
    "The original AW pipeline included export_sensitivity() which ran a one-way sensitivity "
    "analysis varying fund_split ±50% and persistence_years ±50% per intervention. This "
    "function no longer exists in export.py. If you want to understand which interventions "
    "drive the AW results, you would need to re-implement this or vary the fund YAML splits "
    "manually."
)

h3(doc, "8. AW diminishing returns are defined but not exported by run.py")
body(doc,
    "aw_combined.yaml now contains diminishing_anchors (piecewise-linear anchors at $10M, "
    "$26.5M, $55M, $100M, $160M defining the CE curve). The export_diminishing() function "
    "in export.py can write these to CSV, but run.py does not call it — it only calls "
    "export_dataset() and export_assumptions(). The per-fund diminishing returns CSVs that "
    "combine_data.py reads (ea_awf_diminishing_returns.csv, "
    "navigation_fund_diminishing_returns.csv) must be generated separately."
)

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = os.path.join(_DOCS_DIR, "code_walkthrough.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
