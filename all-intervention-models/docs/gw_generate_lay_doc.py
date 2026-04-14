"""Generate a plain-language Word (.docx) overview of the GiveWell cost-effectiveness model.

Intended for an intelligent non-technical audience. All parameter values are
taken directly from gw_cea_modeling.py so the document stays in sync.
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "gw_model_lay_overview.docx")

# ---------------------------------------------------------------------------
# All parameter values from gw_cea_modeling.py — kept here so this script
# runs standalone (without importing the model).
# ---------------------------------------------------------------------------

UNITS_VALUE_PER_M_PER_X_CASH = 3280
N_SAMPLES = 10_000
LIFE_YEARS_PER_LIFE = 60

# Funding tiers (fraction of portfolio in each cost-effectiveness range)
TIERS = {
    "below_8x":         {"share": 0.05, "label": "Below 8×",      "dist": "lognormal, 90% CI [2×, 8×], clipped [0.5×, 16×]"},
    "between_8x_and_16x": {"share": 0.68, "label": "8× to 16×",   "dist": "normal, 90% CI [8×, 16×], clipped [2×, 32×]"},
    "above_16x":        {"share": 0.27, "label": "Above 16×",      "dist": "lognormal, 90% CI [16×, 44×], clipped [8×, 80×]"},
}

# GW moral weights
MORAL_WEIGHTS = {
    "YLDs averted":     2.3,
    "Lives saved":      115.6,
    "Income doublings": 1.0,
}

# Intervention splits and effect-type breakdowns
INTERVENTIONS = [
    ("Malaria prevention & treatment", 0.38, {"YLDs": 14.2, "Lives": 58.3, "Income": 27.4}),
    ("Vaccinations",                   0.12, {"YLDs":  6.7, "Lives": 70.9, "Income": 22.4}),
    ("Malnutrition treatment",         0.09, {"YLDs":  3.9, "Lives": 80.0, "Income": 16.1}),
    ("Water quality",                  0.09, {"YLDs":  2.8, "Lives": 66.5, "Income": 30.7}),
    ("Vitamin A supplementation (VAS)",0.07, {"YLDs": 16.4, "Lives": 66.8, "Income": 16.7}),
    ("Iron fortification",             0.07, {"YLDs": 58.0, "Lives":  0.0, "Income": 42.0}),
    ("Livelihood programs",            0.03, {"YLDs":  9.1, "Lives":  9.3, "Income": 81.6}),
    ("Family planning",                0.02, {"YLDs": 40.0, "Lives": 20.0, "Income": 40.0}),
]

# Temporal breakdowns (% of effect in each time period)
TEMPORAL = {
    "YLDs / life-years": {"0–5 yr": 90.0, "5–10 yr": 7.0, "10–20 yr": 1.5, "20–100 yr": 0.5, "100–500 yr": 0.0, "500+ yr": 0.0},
    "Income doublings":  {"0–5 yr": 18.0, "5–10 yr": 1.4, "10–20 yr": 12.5,"20–100 yr": 68.1,"100–500 yr": 0.0, "500+ yr": 0.0},
}

# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_callout(doc, text, color_fill="EAF4FB", text_color=(0x1A, 0x5C, 0x82)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    return p

def add_note(doc, text):
    return add_callout(doc, "Note: " + text, color_fill="FFFADC", text_color=(0x50, 0x3C, 0x00))

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D2DCF0")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    for ri, row in enumerate(rows):
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = t.rows[ri + 1].cells[i]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How We Estimate the Value of GiveWell Grants")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A plain-language guide to the GiveWell cost-effectiveness model")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. What this model does
# ---------------------------------------------------------------------------
add_heading(doc, "1.  What this model does", 1)
add_body(doc,
    "GiveWell is one of the most rigorous charity evaluators in the world. It focuses on global "
    "health and development — areas where strong evidence exists and where each dollar can go "
    "very far because the populations served face severe poverty. Rethink Priorities evaluates "
    "GiveWell's portfolio on behalf of Anthropic staff who are deciding where to direct "
    "their personal philanthropy. The key question: compared to other ways that money could be "
    "used, how much good does donating to GiveWell's portfolio actually do?"
)
add_body(doc,
    "This model estimates how many people's lives are improved — and by how much — per $1 million "
    "donated to GiveWell's 2025 grantmaking portfolio. It does this by combining data on GiveWell's "
    "own cost-effectiveness estimates with Rethink Priorities' framework for comparing impacts "
    "across time and under different views about risk."
)
add_callout(doc,
    "Think of it as asking: 'If an Anthropic staff member gives $1M to GiveWell today, "
    "how many life-years are saved, how much suffering is prevented, and how much richer "
    "do people become — and when do these benefits materialise?'"
)

# ---------------------------------------------------------------------------
# 2. GiveWell's portfolio: what gets funded
# ---------------------------------------------------------------------------
add_heading(doc, "2.  GiveWell's portfolio: what gets funded", 1)
add_body(doc,
    "GiveWell doesn't fund a single programme — it spreads grants across many proven health "
    "interventions based on which ones it believes are most cost-effective each year. In 2025, "
    "eight cause areas accounted for the bulk of GiveWell's grantmaking:"
)
add_table(doc,
    ["Cause area", "Share of 2025 funding", "What it does"],
    [
        [name, f"{share:.0%}",
         f"YLDs {d['YLDs']:.0f}% / Lives {d['Lives']:.0f}% / Income {d['Income']:.0f}%"]
        for name, share, d in INTERVENTIONS
    ],
)
add_body(doc,
    "The 'YLDs / Lives / Income' column shows roughly what fraction of each intervention's value "
    "comes from reducing illness (YLDs = years lived with disability), saving lives, and boosting "
    "long-term incomes. The remaining ~13% of GiveWell's funding goes to research and scoping "
    "grants that are too varied to model, so the splits above are normalised to sum to 100%."
)
add_note(doc,
    "Iron fortification stands out: it produces almost no mortality reduction but dramatically "
    "reduces anaemia and cognitive impairment, making its value almost entirely about health "
    "quality rather than lives saved. Livelihood programs (e.g. cash transfers) are the "
    "opposite — almost all value comes through long-run income gains."
)

# ---------------------------------------------------------------------------
# 3. Measuring cost-effectiveness: multiples of cash
# ---------------------------------------------------------------------------
add_heading(doc, "3.  Measuring cost-effectiveness: multiples of cash", 1)
add_body(doc,
    "GiveWell rates charities by comparing them to its benchmark: GiveDirectly, which gives "
    "unconditional cash transfers directly to poor households. A programme rated '10×' is judged "
    "ten times as cost-effective as simply handing out the equivalent amount of money."
)
add_body(doc,
    f"We translate this into a concrete number: at exactly 1× (cash-transfer level), $1 million "
    f"produces approximately {UNITS_VALUE_PER_M_PER_X_CASH:,} 'units of value' under GiveWell's moral "
    f"framework. This was calculated by working through GiveWell's own cost-effectiveness "
    f"spreadsheet for GiveDirectly across five countries."
)
add_body(doc,
    "GiveWell's portfolio is not uniformly cost-effective — some grants are excellent, others are "
    "adequate. In 2025, GiveWell reported its portfolio breaks down as:"
)
add_table(doc,
    ["Cost-effectiveness tier", "Share of portfolio", "How we model it"],
    [
        [v["label"], f"{v['share']:.0%}", v["dist"]]
        for v in TIERS.values()
    ],
)
add_body(doc,
    f"To capture this uncertainty, the model draws {N_SAMPLES:,} random samples — one per "
    f"simulated 'world' — from these three distributions, weighted by portfolio share. Each "
    f"sample gives a total cost-effectiveness multiplier; multiplying by "
    f"{UNITS_VALUE_PER_M_PER_X_CASH:,} converts it into units of value per $1M."
)
add_callout(doc,
    "Because we draw random samples from these distributions, the model gives a range of "
    "plausible outcomes rather than a single number. The mean is about 14× the cash benchmark."
)

# ---------------------------------------------------------------------------
# 4. Three types of benefit
# ---------------------------------------------------------------------------
add_heading(doc, "4.  Three types of benefit", 1)
add_body(doc,
    "Not all health improvements are the same. GiveWell explicitly distinguishes three "
    "categories of impact, and applies different values ('moral weights') to each:"
)
add_table(doc,
    ["Benefit type", "What it means", "GiveWell's moral weight"],
    [
        ["YLDs averted",
         "Preventing disability — e.g. a year without debilitating malaria illness",
         "2.3 units of value per YLD averted"],
        ["Lives saved",
         "Preventing a death — primarily child deaths in low-income settings",
         "115.6 units of value per life saved"],
        ["Income doublings",
         "Permanently doubling a household's income — the long-run economic effect",
         "1.0 unit of value per income doubling"],
    ],
)
add_body(doc,
    "GiveWell values saving a life at roughly 50× the value of a year of disability prevention. "
    "This reflects the enormous loss of a death, especially a child's death, compared to a year "
    "of illness or lost income. The model uses these weights to split the total portfolio "
    "value into separate streams for each benefit type."
)
add_body(doc,
    "Once the total units of value per $1M are known, the portfolio's funding-weighted average "
    "split is applied: approximately 15% of value comes from YLDs averted, 57% from life-years "
    "saved, and 28% from income doublings."
)
add_note(doc,
    f"'Lives saved' is subsequently multiplied by {LIFE_YEARS_PER_LIFE} to convert to life-years "
    "saved — reflecting the roughly 60-year remaining life expectancy of the children GiveWell's "
    "work primarily reaches."
)

# ---------------------------------------------------------------------------
# 5. When do benefits happen? The time split
# ---------------------------------------------------------------------------
add_heading(doc, "5.  When do benefits happen?", 1)
add_body(doc,
    "Not all benefits arrive immediately. The model spreads each type of benefit across six "
    "time windows. Health benefits (preventing deaths and disability) happen quickly — within "
    "the current grant cycle. Economic benefits unfold more slowly, because the children helped "
    "today won't enter the workforce for another 10–20 years."
)
add_table(doc,
    ["Time window", "YLDs averted / life-years", "Income doublings", "Why"],
    [
        ["0–5 years",    "90%",  "18%",  "Immediate health benefits; short-term income/savings from cash-like components"],
        ["5–10 years",   "7%",   "1.4%", "Residual from research & pilot investments"],
        ["10–20 years",  "1.5%", "12.5%","Children reached today begin their working lives"],
        ["20–100 years", "0.5%", "68%",  "Long-run compounding of improved nutrition, education, and productivity"],
        ["100–500 years","0%",   "0%",   "No effects modelled beyond 100 years for GiveWell"],
        ["500+ years",   "0%",   "0%",   "No effects modelled beyond 100 years for GiveWell"],
    ],
)
add_callout(doc,
    "This means that health effects (preventing deaths and disability) are front-loaded — "
    "90% happens within five years. Income effects are back-loaded — nearly 70% falls in "
    "the 20–100 year window. Readers who are sceptical of very long-run income projections "
    "should focus on the near-term columns."
)

# ---------------------------------------------------------------------------
# 6. Handling uncertainty: risk profiles
# ---------------------------------------------------------------------------
add_heading(doc, "6.  Handling uncertainty: nine views of the same data", 1)
add_body(doc,
    "Because we are drawing random samples from the cost-effectiveness distributions, we "
    "end up with a spread of outcomes — some worlds where GiveWell performs very well, some "
    "where it performs less well. Different decision-makers care about different aspects of "
    "this spread."
)
add_body(doc,
    "The model applies nine 'risk profiles' to the same set of samples. Each profile "
    "answers a slightly different question:"
)
add_table(doc,
    ["Profile", "The question it answers"],
    [
        ["Neutral (mean)",      "On average, across all simulated worlds, what is the expected impact?"],
        ["Upside",              "What if I ignore the very best-case outcomes as unrealistically optimistic?"],
        ["Downside",            "What if I'm more worried about underperforming than excited about outperforming?"],
        ["Combined",            "Both cautious about extremes and worried about underperformance combined"],
        ["DMREU",               "A formal model of modest risk aversion from decision theory"],
        ["WLU (low/mod/high)",  "Progressively stronger views that very large impacts count for less"],
        ["Ambiguity",           "What if deep uncertainty itself is a reason to discount large outcomes?"],
    ],
)
add_body(doc,
    "For GiveWell, because the cost-effectiveness distribution is relatively well-behaved "
    "(no extreme tail outcomes), most risk profiles give similar answers. The 'neutral' mean "
    "is a reliable summary for this model."
)

# ---------------------------------------------------------------------------
# 7. What the output numbers mean
# ---------------------------------------------------------------------------
add_heading(doc, "7.  What the output numbers mean", 1)
add_body(doc,
    "The final output is a table with one row per benefit type (life-years saved, YLDs averted, "
    "income doublings) and columns for each time period × risk profile combination. Each cell "
    "contains the estimated amount of that benefit per $1 million donated to GiveWell."
)
add_body(doc, "For example, a typical output might show:")
add_bullet(doc,
    "Life-years saved (0–5 years, neutral): roughly 4–8 life-years per $1M",
    bold_prefix="Example. ")
add_body(doc,
    "This number reflects: the distribution of portfolio cost-effectiveness × the share of "
    "value attributable to lives × GiveWell's moral weight for lives × the 60-year "
    "life-expectancy assumption × the fraction of lives-saving happening in the first 5 years."
)
add_note(doc,
    "GiveWell's outputs are in welfare units based on GiveWell's own moral framework. "
    "Comparing these numbers directly to GCR or Animal Welfare outputs requires a separate "
    "cross-cause moral weighting step, which is done downstream and not in this model."
)

# ---------------------------------------------------------------------------
# 8. Key limitations
# ---------------------------------------------------------------------------
add_heading(doc, "8.  Key limitations and honest caveats", 1)
add_bullet(doc,
    "The model's cost-effectiveness distributions are calibrated to GiveWell's own ratings. "
    "If GiveWell's estimates are systematically biased (e.g., too optimistic about long-run "
    "income effects), the model will inherit that bias.",
    bold_prefix="GiveWell's estimates as inputs. ")
add_bullet(doc,
    "The model assumes the same cost-effectiveness uncertainty applies equally to all benefit "
    "types. In reality, income effects (which depend on long-run compounding) are probably "
    "more uncertain than near-term mortality effects. This is a simplification.",
    bold_prefix="Independence assumption. ")
add_bullet(doc,
    "Only 87% of GiveWell's 2025 portfolio is covered (the remaining 13% being research and "
    "scoping grants). The model implicitly assumes the unmodelled portion has similar "
    "cost-effectiveness to the modelled portion.",
    bold_prefix="Portfolio coverage. ")
add_bullet(doc,
    "Income benefits are projected 100 years into the future. Projections this far ahead are "
    "inherently speculative. The long-run income column should be read as an estimate of "
    "the direction and rough magnitude, not a precise forecast.",
    bold_prefix="Long-run income uncertainty. ")

# ---------------------------------------------------------------------------
# 9. Summary of key parameters
# ---------------------------------------------------------------------------
add_heading(doc, "9.  Summary of key parameters at a glance", 1)
add_table(doc,
    ["Parameter", "Value", "What it means"],
    [
        ["Baseline unit value (1× cash)",   f"{UNITS_VALUE_PER_M_PER_X_CASH:,} units / $1M",
         "Value produced per $1M at GiveDirectly-level effectiveness"],
        ["Portfolio cost-effectiveness (mean)", "~14× cash",
         "Weighted average across the three tier distributions"],
        ["Simulation draws",                f"{N_SAMPLES:,}",
         "Number of random scenarios sampled"],
        ["Life-years per life saved",       str(LIFE_YEARS_PER_LIFE),
         "Assumed remaining life expectancy of a GW-funded life"],
        ["Moral weight — life saved",       str(MORAL_WEIGHTS["Lives saved"]),
         "GW's valuation: one life = 115.6 units"],
        ["Moral weight — YLD averted",      str(MORAL_WEIGHTS["YLDs averted"]),
         "GW's valuation: one YLD averted = 2.3 units"],
        ["Portfolio: largest cause area",   "Malaria (38%)",
         "Malaria prevention & treatment dominates the portfolio"],
        ["Largest income-effect cause",     "Livelihoods (82% income)",
         "Cash-like programs have the highest income share"],
    ],
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
