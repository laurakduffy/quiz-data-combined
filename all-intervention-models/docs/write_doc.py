import os
from docx import Document
from docx.shared import Pt

_DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# Title
doc.add_heading('GiveWell Cost-Effectiveness Estimation: Methodology Notes', 0)

doc.add_paragraph(
    'Note on framing: This model was built by Rethink Priorities to evaluate the cost-'
    'effectiveness of donations to GiveWell\'s portfolio on behalf of Anthropic staff '
    'who are deciding where to direct their personal philanthropy. Rethink Priorities '
    'is the evaluator, not the funder. All references to "$1M spent" or "donations" '
    'refer to hypothetical Anthropic staff giving, not RP grantmaking.'
)

# ─────────────────────────────────────────────────────────────
# 0. OVERALL PROCESS
# ─────────────────────────────────────────────────────────────
doc.add_heading('0. Overall Process', level=1)
doc.add_paragraph(
    'The goal is to estimate, for an average dollar donated by an Anthropic staff member '
    'to GiveWell (GW), how many life-years saved, YLDs averted, and income doublings are '
    'produced per $1M spent, disaggregated across six time horizons (0-5, 5-10, 10-20, '
    '20-100, 100-500, and 500+ years). The pipeline has five main steps:'
)
steps = [
    'Estimate the number of units of value produced per $1M when GW\'s spending is 1x as cost-effective as its benchmark (i.e., equivalent to cash transfers). This is UNITS_VALUE_PER_M_PER_X_CASH.',
    'Sample GW\'s cost-effectiveness in multiples of cash from distributions calibrated to GW\'s 2025 grantmaking data, weighted by the share of funding in each cost-effectiveness range, to get a distribution of total units of value per $1M.',
    'Decompose total value into three benefit streams - YLDs averted, lives saved, and income doublings - using a weighted average of each intervention\'s benefit-type breakdown (weighted by share of 2025 funding).',
    'Convert units of value per benefit stream into actual outcomes (lives, YLDs, income doublings) by dividing by GW\'s moral weights for each benefit type.',
    'Apply temporal weights to spread each outcome across the six time horizons, then convert lives saved to life-years saved by multiplying by 60.',
]
for s in steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

# ─────────────────────────────────────────────────────────────
# 1. UNITS_VALUE_PER_M_PER_X_CASH
# ─────────────────────────────────────────────────────────────
doc.add_heading('1. How UNITS_VALUE_PER_M_PER_X_CASH Was Estimated', level=1)
doc.add_paragraph(
    'The constant UNITS_VALUE_PER_M_PER_X_CASH = 3,280 represents the number of units of value '
    '(under GW\'s moral weights) produced per $1M of spending that is exactly 1x as cost-effective '
    'as GW\'s cash-transfer benchmark. It was calculated offline using the "Cash Transfers Value" '
    'sheet in the accompanying spreadsheet.'
)
doc.add_paragraph('The steps were:')
cash_steps = [
    'Start with GiveWell\'s own cost-effectiveness analysis (CEA) for GiveDirectly, across five program countries: Kenya, Malawi, Mozambique, Rwanda, and Uganda.',
    'For each country, sum four components of value per $1M: (a) consumption benefits to recipients, (b) spillover benefits to non-recipients, (c) mortality benefits, and (d) additional benefits and downsides.',
    'Apply GW\'s standard adjustment factors: +10% reduced morbidity, +5% developmental effects, +2% leverage & funging, -2% risk of wastage/fraud, -5% within-org fungibility, -2% psychological spillovers. Net non-morbidity multiplier = 0.98. Morbidity benefits are computed separately as 10% x pre-adjustment value x 0.98.',
    'Take a weighted average across the five countries using estimated marginal allocation weights (Kenya 16%, Malawi 33%, Mozambique 2%, Rwanda 30%, Uganda 18%) to get total adjusted units of value per $1M for GiveDirectly.',
    'GiveWell reports that GiveDirectly is 3.3x their benchmark as of 2024. Dividing total adjusted units/$1M by 3.3 gives the units of value per $1M per 1x benchmark: approximately 3,280.',
]
for s in cash_steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

# ─────────────────────────────────────────────────────────────
# 2. DISTRIBUTIONS FOR GW COST-EFFECTIVENESS
# ─────────────────────────────────────────────────────────────
doc.add_heading('2. Distributions for GW\'s Cost-Effectiveness and Combining with % Funding', level=1)
doc.add_paragraph(
    'GW publicly reports that, in 2025, approximately 5% of grantmaking was below 8x the cash '
    'benchmark, 68% was between 8x and 16x, and 27% was above 16x. '
    'A separate distribution was chosen for each cost-effectiveness tier:'
)

tiers = [
    ('<8x tier (5% of funding)',
     'A lognormal distribution with a 90% CI of [2x, 8x], left-clipped at 0.5x and right-clipped '
     'at 16x. The lower bound of 2x reflects that some GW grants appear as low as ~2x (still better '
     'than half as effective as GiveDirectly at 3.3x). The distribution is symmetric on a log scale, '
     'implying a geometric mean of ~4x and an arithmetic mean of ~6x. There is implicitly a 5% chance '
     'a grant classified as <8x is actually >8x, consistent with classification uncertainty.'),
    ('8-16x tier (68% of funding)',
     'A normal distribution with a 90% CI of [8x, 16x], left-clipped at 2x and right-clipped at 32x. '
     'The symmetric distribution implies a mean of 12x, and allows ~10% probability of '
     'misclassification across the 8x and 16x thresholds.'),
    ('>16x tier (27% of funding)',
     'A lognormal distribution with a 90% CI of [16x, 44x], left-clipped at 8x and right-clipped '
     'at 80x. The upper bound of 44x is based on GiveWell\'s estimate of the maximum cost-effectiveness '
     'for New Incentives (a high-performing vaccination incentive program). A lognormal 90% CI of '
     '[16, 44] implies a mean of ~28x. GW reports its 2025 portfolio averaged ~16x, which the '
     'weighted sum across tiers approximately reproduces.'),
]
for label, text in tiers:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

doc.add_paragraph(
    'To simulate the overall cost-effectiveness distribution, 10,000 samples are drawn from each '
    'of the three tier distributions. The weighted sum is computed as:'
)
doc.add_paragraph(
    '    sample_multiples = 0.05 x below_8x_sample + 0.68 x between_8x_16x_sample + 0.27 x above_16x_sample'
)
doc.add_paragraph(
    'This is then multiplied by UNITS_VALUE_PER_M_PER_X_CASH (3,280) to convert multiples of the '
    'benchmark into total units of value per $1M:'
)
doc.add_paragraph(
    '    sample_units_value_per_M = sample_multiples x 3,280'
)

# ─────────────────────────────────────────────────────────────
# 3. PERCENT EFFECT BY TYPE
# ─────────────────────────────────────────────────────────────
doc.add_heading('3. Estimating the Percent of Value from YLDs, Lives Saved, and Income Doublings', level=1)
doc.add_paragraph(
    'For each intervention in GW\'s portfolio, the share of total value (under GW moral weights) '
    'attributable to YLDs averted, lives saved, and income doublings was estimated from GW\'s own '
    'CEA spreadsheets. The interventions and their benefit-type splits are shown in the '
    '"Percent Distribution Across Int" sheet. The sources and logic for each intervention are:'
)

interventions = [
    ('Malaria Prevention and Treatment (38% of funding)',
     'Based on AMF as representative. Malaria morbidity contributes 9% of value; '
     'short-term anemia effects add another 9% (of which 58% is quality-of-life, i.e., YLD). '
     'In Uganda, 59% of benefits come from reduced child mortality and 9% from adult mortality; '
     '32% from long-term income increases. These are renormalized to account for morbidity. '
     'Result: YLD 14.2%, lives 58.3%, income 27.4%.'),
    ('Vaccinations (12% of funding)',
     '6.7% morbidity (direct + indirect vaccine morbidity effects). 66% of value from under-5 '
     'lives saved, 10% from over-5 lives saved; 22% from long-term consumption increases and 2% '
     'immediate (before renormalization). Result: YLD 6.7%, lives 70.9%, income 22.4%.'),
    ('Malnutrition Treatment (9% of funding)',
     'Based on a Taimaka grant. GW estimates 3.9% of value from morbidity reduction (5.9% - 2.0%). '
     'Of 52,937 total units of value, 44,059 from mortality reductions and 8,878 from income. '
     'Result: YLD 3.9%, lives 80.0%, income 16.1%.'),
    ('Water Quality (9% of funding)',
     'Based on ILC interventions in Kenya as representative. Per 100k people covered: '
     '128 units from morbidity, 2,119 + 966 units from under-5 and over-5 lives saved, '
     '1,426 units from income. Result: YLD 2.8%, lives 66.5%, income 30.7%.'),
    ('VAS (7% of funding)',
     '6% from reduced infectious disease morbidity + 9% from reduced anemia (58% is morbidity) '
     '+ 9% from reduced blindness (58% I assume from morbidity). 80% of benefits from reduced child mortality; '
     '20% from long-term income increases. Result: YLD 16.4%, lives 66.8%, income 16.7%.'),
    ('Iron Fortification (7% of funding)',
     'GW estimates 58% of value from reduced anemia morbidity; no life-saving effects modeled. '
     '34% from cognitive benefits in adults, 8% from cognitive benefits in children (both income). '
     'Result: YLD 58.0%, lives 0%, income 42.0%.'),
    ('Livelihood Programs (3% of funding)',
     'Assumed to have a similar distribution as cash transfers via GiveDirectly (from the Cash '
     'Transfers Value sheet). Result: YLD 9.1%, lives 9.3%, income 81.6%.'),
    ('Family Planning (2% of funding)',
     'GW estimates each year of contraception produces 0.7 units of value. Morbidity: 0.044 from '
     'maternal morbidity, 0.006 from child morbidity, 0.23 from women\'s mental health. Mortality: '
     '0.039 maternal + 0.010 child + 0.091 other. Income: 0.051 women\'s earnings + 0.191 child resources '
     '+ 0.038 reduced medical expenses. Renormalized to sum to 100%. '
     'Result: YLD ~40%, lives ~20%, income ~40%.'),
]
for label, text in interventions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

doc.add_paragraph(
    'The remaining 13% of GW\'s 2025 funding went to "Other Research, Scoping, and Planning Grants" '
    '(4%) and "Other Health Programs" (8%), which were excluded as too vague to assess. '
    'The Python code normalizes the funding weights so they sum to 100% by dividing each '
    'intervention\'s weight by the sum of assessed funding (0.87). The funding-weighted average '
    'across the eight interventions yields approximately: YLDs ~15.0%, Lives ~56.6%, Income ~28.4%. '
    'The spreadsheet notes that, for this preliminary version, the distribution of value across '
    'benefit types is assumed to be independent of an intervention\'s cost-effectiveness.'
)

# ─────────────────────────────────────────────────────────────
# 4. CONVERTING UNITS TO OUTCOMES
# ─────────────────────────────────────────────────────────────
doc.add_heading('4. Converting Units of Value to Outcomes Using GW Moral Weights', level=1)
doc.add_paragraph(
    'Once total units of value per $1M (sample_units_value_per_M) and the weighted-average '
    'percent by benefit type are known, the pipeline splits the simulated total value into '
    'three streams:'
)
doc.add_paragraph(
    '    units_YLD/$1M    = sample_units_value_per_M x 0.150\n'
    '    units_lives/$1M  = sample_units_value_per_M x 0.566\n'
    '    units_income/$1M = sample_units_value_per_M x 0.284'
)
doc.add_paragraph(
    "GiveWell's moral weights define how many units of value one unit of each outcome is worth:"
)
mw_items = [
    'YLDs averted: 2.3 units of value per YLD averted',
    'Lives saved: 115.6 units of value per life saved',
    'Income doublings: 1.0 unit of value per income doubling',
]
for item in mw_items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph(
    'To convert from units of value back to actual outcomes, divide each stream by its moral weight:'
)
doc.add_paragraph(
    '    YLDs averted/$1M     = units_YLD/$1M    / 2.3\n'
    '    Lives saved/$1M      = units_lives/$1M  / 115.6\n'
    '    Income doublings/$1M = units_income/$1M / 1.0'
)
doc.add_paragraph(
    'The spreadsheet notes that this approach implicitly assumes the same level of uncertainty '
    'around each benefit stream. In reality, income effects (which occur far in the future) are '
    'likely higher-variance than mortality/morbidity effects; however, this is left as a '
    'simplification for this preliminary estimate.'
)

# ─────────────────────────────────────────────────────────────
# 5. TEMPORAL BREAKDOWN
# ─────────────────────────────────────────────────────────────
doc.add_heading('5. Temporal Breakdown of Effects', level=1)
doc.add_paragraph(
    'Effects are split across six time horizons: 0-5, 5-10, 10-20, 20-100, 100-500, and 500+ years. '
    'Separate breakdowns were chosen for lives/YLDs and income doublings.'
)

doc.add_heading('Lives Saved and YLDs Averted', level=2)
doc.add_paragraph(
    'These are assumed to have the same temporal distribution. The assumed split is:'
)
lives_temporal = [
    ('0-5 years: 90%', 'The vast majority of mortality and morbidity effects happen within the '
     'current grant cycle. GW primarily funds proven, near-term interventions. GW spends some '
     'percentage of its budget on research and pilot projects that should result in longer-term '
     'lives saved, but this is a small share.'),
    ('5-10 years: 7%', 'A small share comes from research and pilot projects that pay off slightly later.'),
    ('10-20 years: 2.5%', 'A residual share from longer-horizon investments.'),
    ('20-100 years: 0.5%', 'A very small residual.'),
    ('100-500 years and 500+ years: 0%', 'No lives or YLDs are expected to materialize beyond 100 years.'),
]
for label, text in lives_temporal:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

doc.add_heading('Income Doublings', level=2)
doc.add_paragraph(
    "Income effects are more complex because GW's programs primarily target under-5 children, "
    "who won't enter the workforce for 10-20 years. The reasoning is as follows:"
)
doc.add_paragraph(
    '80% of income effects are long-term (workforce entry and adult productivity). Of this 80%:'
)
doc.add_paragraph('15% occurs in the 10-20 year window (early working years)', style='List Bullet')
doc.add_paragraph('85% occurs in the 20-100 year window (full working life)', style='List Bullet')
doc.add_paragraph(
    'The remaining 20% of income effects are short-term (healthcare savings, immediate consumption '
    'increases from cash-like components of programs), distributed with the same 90/7/2.5/0.5 split '
    'as lives/YLDs. The resulting temporal breakdown for income is:'
)
income_temporal = [
    ('0-5 years: 18.0%', '20% x 90%'),
    ('5-10 years: 1.4%', '20% x 7%'),
    ('10-20 years: 12.5%', '20% x 2.5% + 15% x 80%'),
    ('20-100 years: 68.1%', '20% x 0.5% + 85% x 80%'),
    ('100-500 years and 500+ years: 0%', 'No income doublings beyond 100 years.'),
]
for label, text in income_temporal:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

doc.add_paragraph(
    'After applying temporal weights, lives saved/$1M per time horizon are multiplied by 60 '
    '(assumed life-years per life saved for the average GW-funded life) to convert to '
    'life-years saved/$1M. The 60-year assumption reflects an approximate remaining life expectancy '
    'for the population GW primarily serves (predominantly children under 5).'
)

# ─────────────────────────────────────────────────────────────
# 6. SUMMARY OF THE PYTHON FILE
# ─────────────────────────────────────────────────────────────
doc.add_heading('6. Summary of the Python File and What Happens to the Output', level=1)
doc.add_paragraph(
    'The file gw_cea_modeling.py implements the full pipeline described above and produces '
    'three outputs:'
)
outputs = [
    ('summary_statistics.csv',
     'For each of 18 (effect type x time horizon) combinations, stores the mean, 5th percentile, '
     'and 95th percentile across the 10,000 simulated samples. Effect types are YLDs averted, '
     'life-years saved, and income doublings per $1M.'),
    ('histograms/ directory',
     'A PNG histogram for each of the 18 (effect type x time horizon) combinations, showing the '
     'simulated distribution of outcomes per $1M.'),
    ('gw_risk_adjusted.csv',
     'The main downstream output. For each effect type (life_years_saved, YLDs_averted, '
     'income_doublings), the 10,000 simulated samples for each time horizon are passed through '
     'compute_risk_profiles() from the risk_profiles module. This produces, for each of nine '
     'risk preference profiles (neutral, upside, downside, combined, DMREU, WLU-low, WLU-moderate, '
     'WLU-high, ambiguity), a single summary statistic per time horizon. The resulting DataFrame '
     'has 3 rows (one per effect type) and columns for each (risk profile x time horizon) '
     'combination (9 profiles x 6 time horizons = 54 data columns), plus 4 metadata columns '
     '(project_id="givewell", near_term_xrisk="FALSE", effect_id, recipient_type).'),
]
for label, text in outputs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

doc.add_paragraph(
    'The gw_risk_adjusted.csv file is intended as a standardized input to a broader risk-adjusted '
    'impact assessment framework used by Rethink Priorities. It represents GiveWell\'s expected '
    'impact per $1M under different ethical frameworks that vary in how they weight uncertain, '
    'downside, or ambiguous outcomes.'
)

out_path = os.path.join(_DOCS_DIR, 'GW_CEA_Methodology_Notes.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
