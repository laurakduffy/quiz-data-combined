"""Generate a Word (.docx) technical documentation of the intervention model methodology.

Documents the methodology in aw_intervention_models.py: how intervention cost-effectiveness
distributions are constructed, what data sources are used, and how the output
YAML / NPZ files are structured for use in the AW model pipeline.

Pulls parameter values from aw_model_intervention_estimates.yaml to stay in sync.
"""

import os
import sys
import yaml

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.join(_HERE, "..")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(_HERE, "AW_MODEL_DOCUMENTATION.docx")

# ---------------------------------------------------------------------------
# Load estimates YAML to pull live values
# ---------------------------------------------------------------------------

_ESTIMATES_PATH = os.path.join(_ROOT, "aw-models", "data", "inputs", "aw_model_intervention_estimates.yaml")

with open(_ESTIMATES_PATH) as f:
    _ESTIMATES = yaml.safe_load(f)

_INTERVENTIONS = _ESTIMATES["interventions"]
_META = _ESTIMATES["metadata"]

# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_body(doc, text, monospace=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if monospace:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Note: " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x50, 0x3C, 0x00)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFADC")
    pPr.append(shd)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D2DCF0")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    for ri, row in enumerate(rows):
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = t.rows[ri + 1].cells[i]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AW Intervention Model: Methodology & Parameters")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How animal welfare intervention cost-effectiveness distributions are constructed")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Generated from aw_intervention_models.py and aw_model_intervention_estimates.yaml")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x78, 0x78, 0x78)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. Purpose and framing
# ---------------------------------------------------------------------------
add_heading(doc, "1.  Purpose and framing", 1)
add_body(doc,
    "This document describes the methodology behind aw_intervention_models.py, which produces "
    "the cost-effectiveness distribution inputs for Rethink Priorities' animal welfare "
    "intervention model."
)
add_body(doc,
    "Context: Rethink Priorities built this model to evaluate the cost-effectiveness of "
    "animal welfare donations on behalf of Anthropic staff who are deciding where to "
    "direct their personal philanthropy. Rethink Priorities is the evaluator, not the "
    "funder. The organisations and funds assessed here are candidates for Anthropic staff "
    "giving, not recipients of RP grants."
)
add_body(doc,
    "All distributions are based on analyst estimates. Full derivations are documented "
    "in the linked Google Doc (see aw_intervention_models.py header)."
)
add_note(doc,
    f"Unit: all outputs are in suffering-years averted per $1,000 spent on the intervention "
    f"(pre-moral-weight). Moral weight adjustments — reflecting how much weight to give "
    f"animal welfare relative to human welfare — are applied separately downstream. "
    f"This pipeline uses pre-weight values directly, labelled 'animal DALYs', pending "
    f"confirmation on which weights to apply."
)

# ---------------------------------------------------------------------------
# 2. Output files
# ---------------------------------------------------------------------------
add_heading(doc, "2.  Output files", 1)
add_body(doc, "aw_intervention_models.py produces three outputs:")
add_bullet(doc,
    "  aw_model_intervention_estimates.yaml — percentile summaries (p1/p5/p10/p50/p90/p95/p99/mean) "
    "and 10,000 downsampled empirical samples per intervention.",
    bold_prefix="YAML.")
add_bullet(doc,
    "  samples/aw_model_intervention_samples_100k.npz — full 100,000 samples per intervention, "
    "compressed. Used by the AW model pipeline for maximum accuracy.",
    bold_prefix="NPZ.")
add_bullet(doc,
    "  outputs/aw_model_extended_statistics.csv — extended percentile summary (0.15th to 99.85th) "
    "and histograms for QC purposes.",
    bold_prefix="CSV + histograms.")

add_table(doc,
    ["File", "Samples per intervention", "Primary use"],
    [
        ["aw_model_intervention_estimates.yaml", "10,000 (downsampled)", "Human-readable percentiles; fallback for model pipeline"],
        ["aw_model_intervention_samples_100k.npz", "100,000 (full)", "Primary input to AW model pipeline"],
        ["aw_model_extended_statistics.csv", "Summary stats only", "QC and documentation"],
    ],
)

# ---------------------------------------------------------------------------
# 3. Sampling methodology
# ---------------------------------------------------------------------------
add_heading(doc, "3.  Sampling methodology", 1)
add_body(doc,
    f"All distributions are sampled at N = {_META.get('n_samples', 100000):,} draws with "
    f"random seed {_META.get('seed', 42)} for reproducibility. Two helper functions are used:"
)
add_bullet(doc,
    "  sample_lognorm_ci(lo, hi, credibility=90) — fits a lognormal to a confidence "
    "interval [lo, hi] at the specified credibility level (default 90%, i.e. lo=p5, hi=p95). "
    "sigma = (log(hi) - log(lo)) / (2 * z_0.95).",
    bold_prefix="Lognormal from CI.")
add_bullet(doc,
    "  sample_norm_ci(lo, hi, credibility=90) — fits a normal distribution to a "
    "confidence interval. Used where the quantity can be negative or is symmetric on "
    "a linear scale.",
    bold_prefix="Normal from CI.")
add_bullet(doc,
    "  sample_beta(a, b) — draws from a Beta distribution. Used for proportions "
    "(e.g. fraction of suffering reduced, probability of campaign success).",
    bold_prefix="Beta.")
add_body(doc,
    "Downsampling to 10,000 uses quantile spacing (np.quantile at evenly-spaced quantiles) "
    "to preserve the full shape of the distribution with minimal information loss."
)

# ---------------------------------------------------------------------------
# 4. Per-intervention methodology
# ---------------------------------------------------------------------------
add_heading(doc, "4.  Per-intervention methodology", 1)
add_body(doc,
    "Seven interventions are currently modelled. For each, the table below shows the "
    "data source, the key distributional assumptions, and the current central estimate "
    "(median and mean per $1,000 spent on the intervention)."
)

# Build summary table from YAML
summary_rows = []
for key, info in _INTERVENTIONS.items():
    pct = info.get("percentiles_per_1000", {})
    p50  = f"{pct.get('p50', 0):,.0f}" if pct else "—"
    mean = f"{pct.get('mean', 0):,.0f}" if pct else "—"
    start = info.get("effect_start_year", "?")
    persist = info.get("persistence_years", "?")
    summary_rows.append([
        key.replace("_", " ").title(),
        info.get("method", "—"),
        info.get("species", "—").title(),
        p50,
        mean,
        f"yr {start}, {persist} yrs",
    ])

add_table(doc,
    ["Intervention", "Data source", "Species", "Median / $1k", "Mean / $1k", "Timing"],
    summary_rows,
)

# Detailed sub-sections
add_heading(doc, "4.1  Chicken corporate campaigns", 2)
add_body(doc,
    "Source: RP analyst estimates (Laura Duffy, updated 3 Feb 2026). "
    "Custom lognormal calibrated to recent field evidence on cage-free campaigns."
)
_ck = _INTERVENTIONS.get("chicken_corporate_campaigns", {})
_ck_p = _ck.get("percentiles_per_1000", {})
add_table(doc,
    ["Parameter", "Value", "Interpretation"],
    [
        ["Distribution",   "Lognormal 90% CI [177, 3600], clipped [50, 10000]",
         "Suffering-years averted per $1,000 (highly skewed right)"],
        ["p10 / p50 / mean / p90",
         f"{_ck_p.get('p10',0):,.0f} / {_ck_p.get('p50',0):,.0f} / {_ck_p.get('mean',0):,.0f} / {_ck_p.get('p90',0):,.0f}",
         "Wide range reflects genuine uncertainty about campaign success rates"],
        ["Effect start",   str(_ck.get("effect_start_year","?")), "Year 1 — commitments begin quickly"],
        ["Persistence",    f"{_ck.get('persistence_years','?')} years", "Corporate commitments take 10–15 years to fully implement"],
    ],
)

add_heading(doc, "4.2  Shrimp welfare", 2)
add_body(doc,
    "Source: RP analyst estimates combining two intervention pathways:"
)
add_bullet(doc,
    "  Humane slaughter (~90% of budget): "
    "Animals affected per dollar drawn from a lognormal on slaughter volume; "
    "hours of suffering per shrimp at conventional slaughter drawn from a lognormal "
    "(90% CI [0.28, 6.4] hours); pain reduction of a humane slaughter intervention "
    "drawn from a Beta(7,3) ~ 70% reduction. The resulting DALYs per dollar are "
    "multiplied by a lognormal persistence (90% CI [6, 15] years).",
    bold_prefix="Pathway 1.")
add_bullet(doc,
    "  Sludge removal and stocking density (~10% of budget): "
    "Animals affected modelled directly as shrimp affected per dollar given cost of the "
    "intervention; hours of suffering per shrimp drawn from lognormals; "
    "proportion suffering reduced from Beta distributions.",
    bold_prefix="Pathway 2.")
_sh = _INTERVENTIONS.get("shrimp_welfare", {})
_sh_p = _sh.get("percentiles_per_1000", {})
add_body(doc,
    f"Current percentiles per $1,000: p10 = {_sh_p.get('p10',0):,.0f}, "
    f"p50 = {_sh_p.get('p50',0):,.0f}, mean = {_sh_p.get('mean',0):,.0f}, "
    f"p90 = {_sh_p.get('p90',0):,.0f}."
)

add_heading(doc, "4.3  Fish welfare (carp proxy)", 2)
add_body(doc,
    "Source: RP analyst estimates using carp as a proxy for farmed fish, based on "
    "Fish Welfare Initiative data. Three components are multiplied together: "
    "(1) fish affected per dollar [lognormal 90% CI 0.5–15], "
    "(2) hours of suffering during a typical culture cycle [383-day cycle, normal CI "
    "on 1/6 to 1/3 of total hours], "
    "(3) proportion of suffering reduced by the intervention [Beta(3,17) ~ 15% reduction]. "
    "The product gives DALYs reduced per dollar."
)
_fi = _INTERVENTIONS.get("fish_welfare", {})
_fi_p = _fi.get("percentiles_per_1000", {})
add_body(doc,
    f"Current percentiles per $1,000: p10 = {_fi_p.get('p10',0):,.0f}, "
    f"p50 = {_fi_p.get('p50',0):,.0f}, mean = {_fi_p.get('mean',0):,.0f}, "
    f"p90 = {_fi_p.get('p90',0):,.0f}."
)

add_heading(doc, "4.4  Invertebrate welfare (Black Soldier Fly proxy)", 2)
add_body(doc,
    "Source: Default CCM BSF parameters. This is a bottom-up estimate with six components: "
    "(1) BSF born per batch [normal CI 200–300 billion], "
    "(2) proportion affected [lognormal CI 0.005–0.1%], "
    "(3) hours of suffering during larval stage [normal CI on 1/20 to 1/5 of 23-day larval cycle], "
    "(4) proportion of suffering reduced [Beta(9,4) ~ 69%], "
    "(5) probability of intervention success [Beta(4,16) ~ 20%], "
    "(6) cost [lognormal CI $150k–$1M]. "
    "A Bernoulli draw from (5) produces zero-inflated outcomes — many scenarios "
    "contribute nothing because the campaign fails."
)
_bsf = _INTERVENTIONS.get("invertebrate_welfare", {})
_bsf_p = _bsf.get("percentiles_per_1000", {})
add_body(doc,
    f"Effect starts year {_bsf.get('effect_start_year','?')} (reflects time to develop "
    f"regulatory or industry traction). Current percentiles per $1,000: "
    f"p10 = {_bsf_p.get('p10',0):,.0f}, p50 = {_bsf_p.get('p50',0):,.0f}, "
    f"mean = {_bsf_p.get('mean',0):,.0f}, p90 = {_bsf_p.get('p90',0):,.0f}."
)

add_heading(doc, "4.5  Policy advocacy (multi-species)", 2)
_pol = _INTERVENTIONS.get("policy_advocacy", {})
_pol_p = _pol.get("percentiles_per_1000", {})
add_body(doc,
    "Source: Analyst estimate. Policy advocacy is modelled as the chicken corporate "
    "campaigns distribution discounted by 50% to reflect the additional causal steps "
    "required (policy must pass, be enforced, and change actual conditions). "
    "This is a placeholder pending a dedicated model."
)
add_body(doc,
    "    policy_blend = 0.5 * chicken", )
add_body(doc,
    f"Effect starts year {_pol.get('effect_start_year','?')} (policy campaigns take time to "
    f"produce results). Current percentiles per $1,000: "
    f"p10 = {_pol_p.get('p10',0):,.0f}, p50 = {_pol_p.get('p50',0):,.0f}, "
    f"mean = {_pol_p.get('mean',0):,.0f}, p90 = {_pol_p.get('p90',0):,.0f}."
)

add_heading(doc, "4.6  Movement building", 2)
_mv = _INTERVENTIONS.get("movement_building", {})
_mv_p = _mv.get("percentiles_per_1000", {})
add_body(doc,
    "Source: Analyst estimate. Movement building (outreach, education, infrastructure) "
    "is modelled as an indirect multiplier on chicken corporate campaigns, set at 25% "
    "of the chicken value. This reflects the view that movement building has a similar "
    "direction of impact but is harder to attribute causally and works through longer "
    "causal chains."
)
add_body(doc,
    "    movement = 0.25 * chicken", )
add_body(doc,
    f"Effect starts year {_mv.get('effect_start_year','?')} (movement-building investments "
    f"take time to affect the field). Current percentiles per $1,000: "
    f"p10 = {_mv_p.get('p10',0):,.0f}, p50 = {_mv_p.get('p50',0):,.0f}, "
    f"mean = {_mv_p.get('mean',0):,.0f}, p90 = {_mv_p.get('p90',0):,.0f}."
)

add_heading(doc, "4.7  Wild animal welfare", 2)
_wd = _INTERVENTIONS.get("wild_animal_welfare", {})
_wd_p = _wd.get("percentiles_per_1000", {})
add_body(doc,
    "Source: Analyst estimate. Modelled as a mixture of two sub-models:"
)
add_bullet(doc,
    "  Wild mammal intervention (rodent as proxy): target population [lognormal CI 4,100–56,000], "
    "hours of suffering per mammal [lognormal CI 60–330 hours], probability of success "
    "[Beta(4,16) ~ 20%], proportion of deaths averted if successful [Beta(2,2) ~ 50%], "
    "years of impact [lognormal CI 5–20], cost [lognormal CI $100k–$10M]. "
    "A Bernoulli success draw zero-inflates the distribution.",
    bold_prefix="Mammals.")
add_bullet(doc,
    "  Wild invertebrate intervention: same parameters as BSF but with lower probability "
    "of success [Beta(1,9) ~ 10%] to reflect the even earlier stage of this field.",
    bold_prefix="Invertebrates.")
add_body(doc,
    "    wild_sy = share_mammals * wild_mammal_sy + (1 - share_mammals) * wild_invert_sy", )
add_body(doc,
    f"The mammal share is drawn from Beta(1,1) ~ uniform [0,1], reflecting high uncertainty "
    f"about portfolio composition. Effect starts year {_wd.get('effect_start_year','?')}. "
    f"Current percentiles per $1,000: "
    f"p10 = {_wd_p.get('p10',0):,.0f}, p50 = {_wd_p.get('p50',0):,.0f}, "
    f"mean = {_wd_p.get('mean',0):,.0f}, p90 = {_wd_p.get('p90',0):,.0f}."
)

# ---------------------------------------------------------------------------
# 5. How outputs are used
# ---------------------------------------------------------------------------
add_heading(doc, "5.  How outputs feed into the AW model", 1)
add_body(doc,
    "The AW model pipeline (effects.py) loads the intervention estimates and applies fund-specific "
    "budget split weights to convert per-intervention values to per-fund marginal impact:"
)
add_body(doc,
    "    animal_dalys_per_M = samples_per_1000 * 1000 * fund_split_pct", monospace=True)
add_body(doc,
    "where fund_split_pct is the fraction of the fund's budget allocated to that intervention. "
    "The resulting sample array (100,000 values per intervention) is then passed to the "
    "risk-profile framework, which computes 9 risk-adjusted summaries, and the time-period "
    "allocation function, which distributes the effect across four time windows (0–100 years)."
)
add_table(doc,
    ["Step", "Operation", "Output"],
    [
        ["1. Load samples",
         "Load 100k samples from .npz (or 10k from YAML fallback)",
         "Raw samples in suffering-years / $1,000 on the intervention"],
        ["2. Scale to fund level",
         "Multiply by 1,000 ($/1000 → $/1M) and by fund_split_pct",
         "Samples in suffering-years / $1M at fund level"],
        ["3. Time allocation",
         "Distribute across 4 periods using effect_start_year and persistence_years "
         "(periods: 0-5, 5-10, 10-20, 20-100 years)",
         "Fraction per period (period keys: 0_to_5, 5_to_10, 10_to_20, 20_to_100)"],
        ["4. Risk profiles",
         "Apply 9 risk-adjustment functions to the sample array",
         "9 risk-adjusted scalars per time period"],
        ["5. Output",
         "Write 4 × 9 = 36 period-allocated columns plus 9 total columns per intervention "
         "to per-fund CSVs (e.g., ea_awf_dataset.csv, navigation_fund_cagefree_dataset.csv)",
         "Columns named {rp}_{period_key} (e.g., neutral_0_to_5, ambiguity_20_to_100) "
         "and total_{rp} (e.g., total_neutral). Standard RP format for downstream combination."],
    ],
)

# ---------------------------------------------------------------------------
# 6. Limitations
# ---------------------------------------------------------------------------
add_heading(doc, "6.  Limitations and caveats", 1)
add_bullet(doc,
    "Policy advocacy and movement building are estimated as fixed multiples of the chicken/shrimp "
    "blend. These are rough approximations that should be replaced by dedicated models as "
    "better evidence becomes available.",
    bold_prefix="Derived interventions. ")
add_bullet(doc,
    "The wild animal welfare model is a rough placeholder "
    "based on rodent and invertebrate analogues. It should be treated as highly speculative.",
    bold_prefix="Wild animal welfare. ")
add_bullet(doc,
    "All distributions are in pre-moral-weight animal suffering-years. The gap between "
    "the pre-weight values here and human-comparable values depends on contested empirical "
    "and philosophical questions about animal sentience.",
    bold_prefix="Moral weights not applied. ")
add_bullet(doc,
    "The chicken estimate is based on RP analyst estimates and should be reviewed periodically as new field evidence becomes available.",
    bold_prefix="Chicken override. ")
add_bullet(doc,
    "Invertebrate and wild animal welfare estimates have zero-inflated distributions "
    "(many zero outcomes) because they depend on a binary 'does the intervention succeed' "
    "draw. The mean is therefore pulled above the median by rare successes.",
    bold_prefix="Zero-inflation. ")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
