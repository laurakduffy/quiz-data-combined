"""Generate a plain-language Word (.docx) overview of all three RP intervention models.

Intended for an intelligent non-technical audience. Explains why we have three models,
what they share, how they differ, and how to read the outputs.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "model_overview_lay.docx")

# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_callout(doc, text, color_fill="EAF4FB", text_color=(0x1A, 0x5C, 0x82)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    return p

def add_note(doc, text):
    return add_callout(doc, "Note: " + text, color_fill="FFFADC", text_color=(0x50, 0x3C, 0x00))

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "D2DCF0")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    for ri, row in enumerate(rows):
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = t.rows[ri + 1].cells[i]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Comparing Grants Across Cause Areas")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A plain-language guide to how Rethink Priorities evaluates "
                "GiveWell, GCR, and Animal Welfare grants side by side")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. Why compare at all?
# ---------------------------------------------------------------------------
add_heading(doc, "1.  Why compare grants across such different causes?", 1)
add_body(doc,
    "Rethink Priorities evaluates charitable giving opportunities on behalf of Anthropic "
    "staff who are deciding where to direct their personal philanthropy. The cause areas "
    "under consideration fall into three fundamentally different categories:"
)
add_bullet(doc,
    " Global health and development — saving lives and reducing poverty today, mostly "
    "in low-income countries, through proven interventions like malaria nets and vaccines.",
    bold_prefix="GiveWell (GW).")
add_bullet(doc,
    " Global catastrophic risk (GCR) — reducing the small but non-zero chance that "
    "civilisation is destroyed or permanently crippled this century, by AI, pandemics, or nuclear war.",
    bold_prefix="GCR.")
add_bullet(doc,
    " Animal welfare (AW) — reducing the suffering of billions of farmed and wild animals "
    "through corporate campaigns, policy advocacy, and research.",
    bold_prefix="Animal Welfare.")
add_body(doc,
    "These cause areas operate on entirely different time horizons, affect different groups, "
    "and involve different types of evidence. Yet an Anthropic staff member with a giving "
    "budget must choose — or at least rank — them. This project attempts to put all three "
    "on a common footing, so that comparisons are at least principled even if they remain "
    "deeply uncertain."
)
add_callout(doc,
    "This is genuinely hard. We are comparing preventing a child's death today against "
    "reducing the probability of extinction over a century against preventing suffering for "
    "billions of chickens. The models don't resolve the fundamental values question — "
    "they make the tradeoffs explicit and quantified so decision-makers can reason about them."
)

# ---------------------------------------------------------------------------
# 2. What each model produces
# ---------------------------------------------------------------------------
add_heading(doc, "2.  What each model produces", 1)
add_body(doc,
    "Each model calculates the expected impact of $1 million donated to a specific fund or "
    "portfolio. 'Expected' means averaged across all the different ways the world might unfold — "
    "the model runs thousands of simulated scenarios and combines them."
)
add_body(doc, "The three models differ substantially in their domain and approach:")
add_table(doc,
    ["Model", "What it funds", "How impact is measured", "Main source of uncertainty"],
    [
        ["GiveWell",
         "Global health charities (malaria, vaccines, nutrition, water, livelihoods, family planning)",
         "Life-years saved, years of disability prevented, income doublings",
         "GiveWell's exact cost-effectiveness (ranging from 8× to 44× a cash benchmark)"],
        ["GCR (3 funds)",
         "Existential risk reduction: Sentinel Bio (engineered pandemics), Longview Nuclear, Longview AI",
         "Expected survival-weighted value of all future humanity (and potentially civilisation)",
         "How dangerous this century is, how effective the intervention is, "
         "and whether humanity eventually reaches the stars"],
        ["Animal Welfare (7 interventions)",
         "Corporate campaigns, policy advocacy, fish/shrimp/invertebrate welfare, movement building",
         "Animal suffering-years averted (before moral weight conversion)",
         "How effective welfare campaigns actually are in changing animal lives"],
    ],
)
add_note(doc,
    "These models use different units internally. Comparing them directly requires applying "
    "moral weights — judgments about how much human life-years compare to animal suffering-years "
    "compare to future civilisational value. Those weights are applied downstream, not within "
    "the models described here."
)

# ---------------------------------------------------------------------------
# 3. The common output format
# ---------------------------------------------------------------------------
add_heading(doc, "3.  The common output format: a 6 × 9 grid", 1)
add_body(doc,
    "Despite their differences, all three models produce outputs in an identical format: "
    "a grid of values where each cell represents the expected impact per $1M under a "
    "specific combination of time window and ethical risk attitude."
)
add_body(doc, "The six time windows:")
add_table(doc,
    ["Window", "Years covered", "What it captures"],
    [
        ["t0", "0–5 years",     "Immediate, near-term effects"],
        ["t1", "5–10 years",    "Effects visible within a decade"],
        ["t2", "10–20 years",   "Medium-term effects (children growing up, policies taking hold)"],
        ["t3", "20–100 years",  "Long run within a century"],
        ["t4", "100–500 years", "Multi-generational effects"],
        ["t5", "500+ years",    "Civilisational / very long-run tail"],
    ],
)
add_body(doc,
    "Why split by time? Because different ethical frameworks weight the future differently. "
    "Someone who thinks only about living people today cares only about t0–t1. "
    "Someone who gives equal weight to all future generations cares equally about t5. "
    "The time split lets each perspective read the number most relevant to it."
)
add_callout(doc,
    "GiveWell effects are almost entirely in t0–t3 (no effects beyond 100 years). "
    "Animal welfare effects are entirely in t0–t2 (all within 20 years). "
    "GCR effects span all periods, with t4–t5 potentially dominant if humanity spreads "
    "beyond Earth."
)

# ---------------------------------------------------------------------------
# 4. The nine risk profiles
# ---------------------------------------------------------------------------
add_heading(doc, "4.  Nine ways to handle uncertainty: the risk profiles", 1)
add_body(doc,
    "Each model runs thousands of simulations and produces a spread of possible outcomes. "
    "How you summarise that spread depends on your values and risk tolerance. "
    "We apply nine different summaries — called 'risk profiles' — to the same underlying data."
)
add_body(doc,
    "Think of them as nine different people looking at the same probability distribution "
    "and asking different questions:"
)
add_table(doc,
    ["Profile", "The question it answers", "Key feature"],
    [
        ["Neutral (mean)",
         "What is the average outcome if we could repeat this grant many times?",
         "No risk aversion; counts all scenarios equally"],
        ["Upside",
         "What if the very best-case scenarios are unrealistically optimistic?",
         "Removes the top 1% of outcomes before averaging"],
        ["Downside",
         "What if I care much more about underperforming than outperforming?",
         "Losses below the median count 5.0× more than equivalent gains"],
        ["Combined",
         "Both cautious about extreme outcomes and worried about downside risk",
         "Tail clipping + loss aversion combined"],
        ["DMREU",
         "What does a formal model of moderate risk aversion say?",
         "Based on Duffy (2023); probability weights decline for rare large outcomes"],
        ["WLU – low",
         "What if very large impacts count for a little less per unit?",
         "Concavity c=0.01; mild diminishing returns to scale of impact"],
        ["WLU – moderate",
         "Moderate concavity (c=0.05)",
         "More diminishing returns to large-scale impact"],
        ["WLU – high",
         "Strong concavity (c=0.10)",
         "Strongest downweighting of large outcomes in this framework"],
        ["Ambiguity",
         "What if I'm especially sceptical of outcomes based on deep uncertainty?",
         "Exponential fade-out of top 2.5% of scenarios; zero weight above 99.9th percentile"],
    ],
)
add_body(doc,
    "For GiveWell, all nine profiles give similar answers because the cost-effectiveness "
    "distribution is well-behaved and doesn't have extreme outliers. "
    "For GCR, the profiles diverge sharply — the 'neutral' mean is dominated by rare but "
    "astronomically valuable futures, while 'combined' or 'ambiguity' is much more conservative. "
    "For Animal Welfare, the profiles also diverge significantly because CCM estimates have "
    "very heavy right tails."
)

# ---------------------------------------------------------------------------
# 5. How uncertainty differs across the three models
# ---------------------------------------------------------------------------
add_heading(doc, "5.  What drives uncertainty in each model", 1)
add_body(doc,
    "The three models are uncertain for very different reasons. Understanding the source "
    "of uncertainty helps interpret the results:"
)
add_table(doc,
    ["Model", "Main driver of uncertainty", "Effect on output spread"],
    [
        ["GiveWell",
         "How cost-effective is each cause area this year? GiveWell's portfolio ranges "
         "from 8× to 44× the cash benchmark.",
         "Moderate spread; well-characterised by GiveWell's own analysis"],
        ["GCR",
         "Three deep uncertainties compound: (1) Is this century actually dangerous? "
         "(2) Does this grant actually reduce risk? "
         "(3) Does humanity eventually colonise the stars?",
         "Enormous spread — top outcomes involve astronomical value from stellar expansion; "
         "bottom outcomes involve grants that have no effect"],
        ["Animal Welfare",
         "How effective is corporate advocacy actually? A successful cage-free campaign "
         "might affect billions of animals; a failed one affects none.",
         "Very heavy right tail; median is much lower than mean"],
    ],
)
add_callout(doc,
    "The GCR model's huge spread is not a modelling flaw — it reflects genuine uncertainty "
    "about catastrophic risk and the long-run future. Risk-averse profiles are particularly "
    "important for GCR because the mean is dominated by the 10% of scenarios where humanity "
    "reaches the stars."
)

# ---------------------------------------------------------------------------
# 6. How the time windows look different across models
# ---------------------------------------------------------------------------
add_heading(doc, "6.  When do effects arrive? A comparison across models", 1)
add_body(doc,
    "The three models allocate their effects to time windows very differently, "
    "which reflects the fundamentally different nature of each intervention:"
)
add_table(doc,
    ["Model / fund", "t0 (0–5 yr)", "t1–t2 (5–20 yr)", "t3 (20–100 yr)", "t4–t5 (100+ yr)"],
    [
        ["GiveWell — health effects",  "~90%",  "~9%",   "~1%",   "0%"],
        ["GiveWell — income effects",  "~18%",  "~14%",  "~68%",  "0%"],
        ["GCR (Sentinel Bio etc.)",    "Varies","Varies","Varies","Potentially dominant (stellar)"],
        ["Animal Welfare",             "~27%",  "~73%",  "0%",    "0%"],
    ],
)
add_body(doc,
    "GiveWell's health effects are highly front-loaded — a malaria net saves a life in the "
    "next five years, not a century from now. GiveWell's income effects are back-loaded — "
    "the children kept healthy today won't join the workforce for 10–20 years."
)
add_body(doc,
    "Animal welfare effects fall almost entirely in t0–t2, because corporate commitments and "
    "policy changes last for a limited number of years (10–15 typically) before needing renewal."
)
add_body(doc,
    "GCR effects are unique: they extend across all time windows, because reducing the risk "
    "of extinction this century preserves value that accumulates over trillions of years. "
    "The 10% probability of stellar expansion means that t4 and t5 often dominate the "
    "GCR expected value under risk-neutral analysis."
)

# ---------------------------------------------------------------------------
# 7. How the models are combined
# ---------------------------------------------------------------------------
add_heading(doc, "7.  How the models are combined", 1)
add_body(doc,
    "A separate combining script reads the outputs of all three models and merges them "
    "into a single data structure for downstream comparison. Each fund or intervention "
    "contributes one row to the combined dataset, with:"
)
add_bullet(doc, "A unique identifier (e.g. 'sentinel_bio', 'givewell', 'chicken_corporate_campaigns')")
add_bullet(doc, "54 impact values (6 time periods × 9 risk profiles)")
add_bullet(doc, "A diminishing returns curve showing how marginal cost-effectiveness changes with scale")
add_body(doc,
    "The combined dataset does not apply any cross-cause moral weights. It presents "
    "each model's output in its own units. The downstream analysis tool (or the analyst "
    "reading the data) must then decide: how do life-years compare to animal suffering-years "
    "compare to civilisational survival probability?"
)
add_note(doc,
    "The combined dataset currently includes: 3 GiveWell rows (life-years, YLDs, income), "
    "3 GCR rows (one per fund, plus Sentinel Bio sub-extinction tiers), and 7 Animal Welfare "
    "rows (one per intervention in the combined AW fund)."
)

# ---------------------------------------------------------------------------
# 8. What the models can and can't tell you
# ---------------------------------------------------------------------------
add_heading(doc, "8.  What these models can and can't tell you", 1)
add_body(doc, "These models are useful for:")
add_bullet(doc,
    "Comparing relative cost-effectiveness within a cause area (e.g. which GCR fund is most "
    "promising per dollar?)",
    bold_prefix="Within-cause comparisons. ")
add_bullet(doc,
    "Understanding how sensitive conclusions are to ethical assumptions (e.g. how does the "
    "ranking change if you use 'combined' instead of 'neutral'?)",
    bold_prefix="Sensitivity to risk attitudes. ")
add_bullet(doc,
    "Seeing when benefits arrive — which is crucial for anyone who discounts the future or "
    "who prioritises near-term versus long-term impact.",
    bold_prefix="Time decomposition. ")
add_bullet(doc,
    "Having a principled, quantitative basis for grant recommendations rather than purely "
    "qualitative judgment.",
    bold_prefix="Structured decision support. ")

add_body(doc, "These models cannot:")
add_bullet(doc,
    "Resolve the cross-cause comparison without additional moral weights. A number of "
    "life-years cannot be mechanically compared to a number of animal suffering-years without "
    "a view on how those two things compare in value.",
    bold_prefix="Make the cross-cause comparison for you. ")
add_bullet(doc,
    "Capture strategic, systemic, or portfolio effects — e.g. whether funding GCR reduces "
    "the incentive for others to fund it, or whether concentrating on one cause area is wiser "
    "than diversification.",
    bold_prefix="Account for portfolio strategy. ")
add_bullet(doc,
    "Substitute for qualitative judgment about whether an organisation is well-run, whether "
    "its theory of change is plausible, or whether its leadership is trustworthy.",
    bold_prefix="Replace qualitative assessment. ")
add_bullet(doc,
    "Be taken as precise. Every model contains significant simplifying assumptions. The "
    "output numbers are best interpreted as order-of-magnitude estimates with wide error bars.",
    bold_prefix="Give precise answers. ")

# ---------------------------------------------------------------------------
# 9. A quick guide for non-technical readers
# ---------------------------------------------------------------------------
add_heading(doc, "9.  Quick guide for non-technical readers", 1)
add_body(doc, "If you're looking at model output tables, here's how to orient yourself:")
add_table(doc,
    ["If you want to know...", "Look at..."],
    [
        ["How much direct, near-term impact a GiveWell grant has",
         "GW life-years row, t0 column (0–5 years), neutral or combined profile"],
        ["Whether a GCR grant looks good under a cautious view",
         "GCR fund row, all time periods, combined or ambiguity profile"],
        ["How much animal welfare spending depends on optimistic assumptions",
         "AW intervention row: compare neutral vs. combined — a large gap means "
         "the conclusion is sensitive to tail outcomes"],
        ["Whether GCR's long-run value survives risk adjustment",
         "GCR rows: if t4+t5 still dominate under 'combined', the long-run "
         "case is robust; if they collapse, it's tail-driven"],
        ["How quickly to expect results from a grant",
         "Look at where the row's value is concentrated: early columns = near-term; "
         "later columns = long-term"],
    ],
)
add_callout(doc,
    "The most important thing to remember: these models make uncertainty explicit rather "
    "than hiding it. A wide range between the 'neutral' and 'combined' numbers is not a "
    "sign that the model is unreliable — it's an honest acknowledgement that the world is "
    "genuinely uncertain, and that your conclusions should be held with appropriate humility."
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
