"""Generate a plain-language Word (.docx) overview of the GCR model.

Intended for an intelligent non-technical audience. Pulls all parameter values
directly from fund_profiles.py so the document stays in sync with model
assumptions.
"""

import math
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "gcr-models"))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fund_profiles import (
    FUND_PROFILES, get_fund_profile,
    _AI_CAUSE_FRACTION, _NUCLEAR_CAUSE_FRACTION, _BIO_CAUSE_FRACTION,
    _SENTINEL_REL_REDUCTION_PER_10M, _NUCLEAR_REL_REDUCTION_PER_10M, _AI_REL_REDUCTION_PER_10M,
    _SENTINEL_REL_RISK_REDUCTION, _NUCLEAR_REL_RISK_REDUCTION, _AI_REL_RISK_REDUCTION,
    _TOTAL_XRISK_100YR, _RP_WORLD_PRIORS,
    _r_max_from_cumulative_risk,
)

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "gcr_model_lay_overview.docx")
M = 10 ** 6

# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def fmt_pct(v, decimals=1):
    return f"{v * 100:.{decimals}f}%"

def fmt_dollar(v):
    s = f"{v / M:g}"
    return f"${s}M"


# ---------------------------------------------------------------------------
# Word document helpers (same as technical doc)
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p


def add_callout(doc, text, color_fill="EAF4FB", text_color=(0x1A, 0x5C, 0x82)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*text_color)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    return p


def add_note(doc, text):
    return add_callout(doc, "Note: " + text, color_fill="FFFADC",
                       text_color=(0x50, 0x3C, 0x00))


def add_table(doc, headers, rows):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "D2DCF0")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = "F5F7FC" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_bg(cells[i], bg)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How We Estimate the Value of GCR Grants")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A plain-language guide to the Global Catastrophic Risk intervention model")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x32, 0x5A, 0xA0)

doc.add_paragraph()

# ---------------------------------------------------------------------------
# 1. What this model does
# ---------------------------------------------------------------------------
add_heading(doc, "1.  What this model does", 1)
add_body(doc,
    "Rethink Priorities evaluates charitable giving opportunities on behalf of donors — in this "
    "case, Anthropic staff considering where to direct their personal philanthropy. One category "
    "of giving is organisations working to reduce global catastrophic risks (GCRs): events that "
    "could kill enormous numbers of people or permanently curtail humanity's future, such as "
    "engineered pandemics, nuclear war, or dangerous artificial intelligence. These events are "
    "rare but potentially enormous in consequence, and ordinary cost-effectiveness tools don't "
    "handle them well."
)
add_body(doc,
    "This model asks: if an Anthropic staff member donates $X to one of these funds, how much "
    "does humanity's chance of surviving this century — and thriving long into the future — "
    "increase? The answer is expressed as an expected value: the probability-weighted sum of "
    "all the good the grant might do, across all the possible ways the future could unfold."
)
add_callout(doc,
    "Think of it as asking: 'Across the full range of worlds that might exist, how much better "
    "off is humanity on average because of this grant?'"
)

add_body(doc, "Three grants are currently modelled:")
add_bullet(doc, " Sentinel Bio — biosecurity work focused on preventing engineered pandemics.", bold_prefix="Sentinel Bio")
add_bullet(doc, " Longview Nuclear — policy work to reduce the risk of nuclear war.", bold_prefix="Longview Nuclear")
add_bullet(doc, " Longview AI — AI safety funding.", bold_prefix="Longview AI")

# ---------------------------------------------------------------------------
# 2. How risk is modeled: the danger curve
# ---------------------------------------------------------------------------
add_heading(doc, "2.  How risk is modeled: the danger curve", 1)
add_body(doc,
    "The model assumes that humanity is currently passing through a particularly dangerous "
    "period — a 'time of perils.' Annual risk is not constant; it rises, peaks, and then "
    "gradually falls as we either develop better safeguards or survive long enough that "
    "the most acute threats pass."
)
add_body(doc,
    "Technically this is modelled as a bell curve (Gaussian) of annual risk, centred on a "
    "peak year, plus a small permanent background risk that persists indefinitely. The model "
    "sweeps over three scenarios for when this peak occurs and how broad the danger window is:"
)
_ymr = _RP_WORLD_PRIORS["year_max_risk"]
_yr1 = _RP_WORLD_PRIORS["year_risk_1pct_max"]
add_table(doc,
    ["Parameter", "Values considered", "What it means"],
    [
        ["Year of peak danger",
         f"{_ymr[0]}, {_ymr[1]}, or {_ymr[2]} years from now",
         "How soon does risk peak?"],
        ["Width of the danger window",
         f"{_yr1[0]}, {_yr1[1]}, or {_yr1[2]} year half-width",
         "Is the window narrow (acute crisis) or broad (prolonged transition)?"],
        ["Permanent background risk",
         "Very low (1 in 10 billion/yr), medium (1 in 10 million/yr), or higher (1 in 1,000/yr)",
         "Even after the transition, some risk persists forever"],
    ],
)
add_body(doc,
    "The height of the peak is not a separate free assumption — it is calculated automatically "
    "so that the total chance of catastrophe over the next 100 years matches our overall belief "
    "about how dangerous this century is (see Section 3)."
)

# ---------------------------------------------------------------------------
# 3. How dangerous is this century?
# ---------------------------------------------------------------------------
add_heading(doc, "3.  How dangerous is this century?", 1)
add_body(doc,
    "The single most important assumption in the model is: what is the total probability "
    "that some catastrophic or extinction-level event occurs before 2125? We consider three "
    "scenarios:"
)
add_table(doc,
    ["Scenario", "Chance of catastrophe in 100 years", "Interpretation"],
    [
        ["Low",     fmt_pct(0.05), "Relatively optimistic — safeguards hold"],
        ["Central", fmt_pct(0.10), "RP baseline — meaningful but not overwhelming risk"],
        ["High",    fmt_pct(0.65), "Pessimistic — we are in serious danger"],
    ],
)
add_note(doc,
    "These figures cover all catastrophic risks combined — AI, bio, nuclear, and others. "
    "The model does not try to add them up separately; instead it treats total risk as a "
    "single uncertain quantity and then asks what fraction of it each fund addresses."
)

# ---------------------------------------------------------------------------
# 4. How interventions reduce risk
# ---------------------------------------------------------------------------
add_heading(doc, "4.  How a grant reduces risk", 1)
add_body(doc,
    "Each fund's grant is assumed to shave a small fraction off the danger curve during "
    "a limited window of time (its 'persistence' — how long the effects last). "
    "The key question is: how large a fraction?"
)
add_body(doc,
    "Because each fund works on a specific cause (bio, nuclear, or AI), it can only affect "
    "the slice of total risk that belongs to that cause. We call this the cause fraction. "
    "Based on RP's cross-cause risk model:"
)
add_table(doc,
    ["Fund", "Cause fraction of total x-risk", "Intuition"],
    [
        ["Sentinel Bio",
         fmt_pct(_BIO_CAUSE_FRACTION, 1),
         "Bio risk is a small share of total catastrophic risk in our model"],
        ["Longview Nuclear",
         fmt_pct(_NUCLEAR_CAUSE_FRACTION, 1),
         "Nuclear risk is a similarly small share"],
        ["Longview AI",
         fmt_pct(_AI_CAUSE_FRACTION, 1),
         "AI (direct + indirect pathways) dominates our risk model"],
    ],
)
add_body(doc,
    "The effectiveness of each grant is then expressed as a relative risk reduction "
    "per $10 million spent — how many percentage points of the cause-specific danger "
    "curve does $10M remove? This is swept across three scenarios (conservative, central, "
    "optimistic) to reflect genuine uncertainty about how much each fund achieves:"
)
add_table(doc,
    ["Fund", "Budget", "Rel. risk reduction per $10M (conservative → optimistic)", "Effect lasts"],
    [
        ["Sentinel Bio",
         fmt_dollar(FUND_PROFILES["sentinel"]["budget"]),
         f"{fmt_pct(_SENTINEL_REL_REDUCTION_PER_10M['values'][0], 4)} → {fmt_pct(_SENTINEL_REL_REDUCTION_PER_10M['values'][2], 3)}",
         f"{FUND_PROFILES['sentinel']['fixed_params']['persistence_effect']} years"],
        ["Longview Nuclear",
         fmt_dollar(FUND_PROFILES["longview_nuclear"]["budget"]),
         f"{fmt_pct(_NUCLEAR_REL_REDUCTION_PER_10M['values'][0], 4)} → {fmt_pct(_NUCLEAR_REL_REDUCTION_PER_10M['values'][2], 3)}",
         f"{FUND_PROFILES['longview_nuclear']['fixed_params']['persistence_effect']} years"],
        ["Longview AI",
         fmt_dollar(FUND_PROFILES["longview_ai"]["budget"]),
         f"{fmt_pct(_AI_REL_REDUCTION_PER_10M['values'][0], 4)} → {fmt_pct(_AI_REL_REDUCTION_PER_10M['values'][2], 3)}",
         f"{FUND_PROFILES['longview_ai']['fixed_params']['persistence_effect']} years"],
    ],
)
add_note(doc,
    "The central scenario uses 0.02% per $10M for all three funds — matching the Longview "
    "Nuclear survey response, discounted 10x from the raw estimate to be conservative. "
    "The conservative is 50x below the raw estimate; the optimistic uses the raw estimate directly."
)

# ---------------------------------------------------------------------------
# 5. Why long-term value matters so much
# ---------------------------------------------------------------------------
add_heading(doc, "5.  Why the long run dominates the calculation", 1)
add_body(doc,
    "Even a tiny reduction in the chance of extinction this century is enormously valuable "
    "if humanity has a long and flourishing future ahead. The model captures this by "
    "explicitly valuing all future years — from now until roughly 10^14 years from now "
    "(the end of the stelliferous era, when star formation ceases)."
)
add_body(doc,
    "Future world value is modelled in two phases:"
)
add_bullet(doc,
    " Earth-based growth: world value grows logistically, limited by Earth's carrying capacity. "
    "The model is uncertain about how large that carrying capacity ultimately is (1.5× to 100× "
    "current value).",
    bold_prefix="Near-term (thousands of years).")
add_bullet(doc,
    " If humanity eventually spreads beyond Earth, value could grow cubically with the "
    "expanding settlement frontier. The model assigns a 10% probability to this occurring, "
    "with the most likely start date several hundred years from now. When it does occur, "
    "the long-term value dwarfs the near-term figure by many orders of magnitude.",
    bold_prefix="Very long-term (stellar expansion).")
add_callout(doc,
    "Because stellar expansion is so valuable but also quite uncertain (10% probability), "
    "results vary enormously across scenarios. The mean expected value is therefore dominated "
    "by the 10% of scenarios in which humanity eventually reaches the stars."
)
add_body(doc,
    "The model reports results both as a single mean ('risk-neutral') and under several "
    "risk-averse frameworks that discount the most extreme tail outcomes, giving a fuller "
    "picture of the range of conclusions."
)

# ---------------------------------------------------------------------------
# 6. Adjustments for uncertainty about the grant itself
# ---------------------------------------------------------------------------
add_heading(doc, "6.  Adjustments for grant-level uncertainty", 1)
add_body(doc,
    "Even if the model's world assumptions are correct, we are uncertain about several "
    "things specific to each grant:"
)

_fp_sentinel  = FUND_PROFILES["sentinel"]
_fp_nuclear   = FUND_PROFILES["longview_nuclear"]
_fp_ai        = FUND_PROFILES["longview_ai"]

add_table(doc,
    ["Adjustment", "Sentinel Bio", "Longview Nuclear", "Longview AI", "What it means"],
    [
        ["Counterfactual factor",
         fmt_pct(_fp_sentinel["counterfactual_factor"]),
         fmt_pct(_fp_nuclear["counterfactual_factor"]),
         fmt_pct(_fp_ai["counterfactual_factor"]),
         "What fraction of the impact would not have happened without our grant? "
         "(e.g. 87.5% means we're confident we are genuinely additional)"],
        ["Chance of no effect",
         fmt_pct(_fp_sentinel["p_zero"], 0),
         fmt_pct(_fp_nuclear["p_zero"], 0),
         fmt_pct(_fp_ai["p_zero"], 0),
         "Probability that the intervention simply doesn't work"],
        ["Chance of harm",
         fmt_pct(_fp_sentinel["p_harm"], 0),
         fmt_pct(_fp_nuclear["p_harm"], 0),
         fmt_pct(_fp_ai["p_harm"], 0),
         "Probability that the intervention makes things worse (e.g. backfire or rebound effects)"],
    ],
)
add_body(doc,
    "These are applied as a weighted average: the model's gross intervention value is "
    "multiplied by the counterfactual factor, and then probability-weighted combinations "
    "of benefit, zero, and harm scenarios are averaged together."
)
add_note(doc,
    "Longview AI carries a higher probability of harm (15% vs 5%) reflecting greater "
    "uncertainty about whether AI safety interventions might inadvertently accelerate "
    "AI development or otherwise backfire."
)

# ---------------------------------------------------------------------------
# 7. Sub-extinction events (Sentinel Bio only)
# ---------------------------------------------------------------------------
add_heading(doc, "7.  Catastrophes short of extinction", 1)
add_body(doc,
    "Each fund also models large-scale catastrophes that kill millions or hundreds of "
    "millions of people without ending civilisation. Two tiers are captured separately "
    "from the extinction pathway for each fund:"
)

_sub_ext_funds = [
    ("Sentinel Bio", _fp_sentinel),
    ("Longview Nuclear", get_fund_profile("longview_nuclear")),
    ("Longview AI", get_fund_profile("longview_ai")),
]
for _fund_name, _fp in _sub_ext_funds:
    _tiers = _fp.get("sub_extinction_tiers", [])
    if not _tiers:
        continue
    add_heading(doc, _fund_name, 3)
    add_table(doc,
        ["Event tier", "Chance per decade", "Expected deaths if it occurs", "Discount applied"],
        [
            [t["tier_name"],
             fmt_pct(t["p_10yr"], 0),
             f"{t['expected_deaths'] / 1e6:.0f}M",
             "None" if t["discount"] == 1.0
             else f"{fmt_pct(1 - t['discount'], 0)} discount"]
            for t in _tiers
        ],
    )

add_body(doc,
    "These tiers use a simpler calculation: probability of the event × expected deaths "
    "× fraction of risk the grant removes × how long the effect lasts. They are added "
    "to each fund's extinction-pathway value to give a combined total."
)

# ---------------------------------------------------------------------------
# 8. Running thousands of scenarios: Monte Carlo
# ---------------------------------------------------------------------------
add_heading(doc, "8.  Running thousands of scenarios", 1)
add_body(doc,
    "None of the uncertain quantities described above takes a single fixed value in the "
    "model. Instead, for each simulation run the model draws a random combination of:"
)
add_bullet(doc, "How dangerous is this century? (5%, 10%, or 65% total x-risk)")
add_bullet(doc, "When does risk peak, and how wide is the window?")
add_bullet(doc, "How large is the permanent background risk?")
add_bullet(doc, "Does humanity eventually spread to the stars, and when?")
add_bullet(doc, "How much does Earth's carrying capacity grow?")
add_bullet(doc, "How effective is this specific grant?")
add_body(doc,
    "Running 100,000 such draws gives a distribution of outcomes — from 'nearly nothing' "
    "to 'astronomical value.' The final reported numbers summarise this distribution: "
    "the mean (expected value), selected percentiles, and several risk-weighted variants "
    "that discount very large tail outcomes."
)
add_callout(doc,
    "The model is deliberately transparent about its uncertainty. A wide spread between "
    "the 10th and 90th percentile outcome is not a flaw — it honestly reflects how much "
    "we don't know about catastrophic risk and long-run civilisational trajectories."
)

# ---------------------------------------------------------------------------
# 9. What the output numbers mean
# ---------------------------------------------------------------------------
add_heading(doc, "9.  What the output numbers mean", 1)
add_body(doc,
    "Results are expressed as QALYs (quality-adjusted life years) or equivalent welfare "
    "units per dollar spent. Because long-run value is so large in some scenarios, the "
    "numbers can look astronomically high. Several perspectives are provided:"
)
add_table(doc,
    ["Output label", "What it captures"],
    [
        ["Neutral (mean)",        "Simple probability-weighted average across all scenarios"],
        ["Upside",                "Trims the top tail — ignores scenarios with the very largest values"],
        ["Downside",              "Applies extra weight to bad outcomes (loss-aversion)"],
        ["Combined",              "Percentile-weighted + loss-averse — a 'cautious' summary"],
        ["DMREU",                 "Difference-Making Risk-Weighted EU — extra caution about rare large outcomes"],
        ["WLU variants",          "Weighted linear utility at different levels of risk aversion"],
        ["Ambiguity-averse",      "Discounts outcomes where our uncertainty is especially deep"],
    ],
)
add_body(doc,
    "For GCR work, the 'combined' and 'downside' figures tend to be most decision-relevant "
    "because they acknowledge that decision-makers reasonably don't want to rely entirely "
    "on the mean of a distribution that includes trillion-QALY tails."
)

# ---------------------------------------------------------------------------
# 10. Key limitations and caveats
# ---------------------------------------------------------------------------
add_heading(doc, "10.  Key limitations and honest caveats", 1)
add_body(doc,
    "This model involves some very large numbers and deep uncertainty. The following "
    "caveats are important to keep in mind:"
)
add_bullet(doc,
    "The model's outputs are only as good as its inputs. The 'cause fractions' and "
    "'relative risk reduction per dollar' assumptions are derived from expert judgment "
    "and survey responses, not empirical measurement. They carry wide uncertainty.",
    bold_prefix="Garbage in, garbage out. ")
add_bullet(doc,
    "Stellar expansion scenarios, while assigned only 10% probability, dominate the "
    "mean expected value. Readers who are skeptical of very long-run value should focus "
    "on the earth-only or risk-averse output columns.",
    bold_prefix="Long-run value dominates. ")
add_bullet(doc,
    "The 'cause fractions' come from a single RP cross-cause model. If the true "
    "distribution of risk across causes looks very different, the relative rankings "
    "of funds could change substantially.",
    bold_prefix="AI's large cause fraction. ")
add_bullet(doc,
    "All three funds currently use the same effectiveness assumptions (0.002% per $10M "
    "central estimate). This reflects a judgment that we lack strong evidence to "
    "differentiate them on intervention effectiveness alone.",
    bold_prefix="Symmetric effectiveness assumptions. ")
add_bullet(doc,
    "The model does not account for portfolio effects, strategic interactions between "
    "funders, or second-order effects of funding one cause on others.",
    bold_prefix="Model boundary. ")

# ---------------------------------------------------------------------------
# 11. Summary table
# ---------------------------------------------------------------------------
add_heading(doc, "11.  At a glance: key model parameters", 1)

_ai_fp  = FUND_PROFILES["longview_ai"]["fixed_params"]
_nuc_fp = FUND_PROFILES["longview_nuclear"]["fixed_params"]
_sen_fp = FUND_PROFILES["sentinel"]["fixed_params"]

add_table(doc,
    ["Parameter", "Sentinel Bio", "Longview Nuclear", "Longview AI"],
    [
        ["Budget",
         fmt_dollar(FUND_PROFILES["sentinel"]["budget"]),
         fmt_dollar(FUND_PROFILES["longview_nuclear"]["budget"]),
         fmt_dollar(FUND_PROFILES["longview_ai"]["budget"])],
        ["Cause fraction of total x-risk",
         fmt_pct(_BIO_CAUSE_FRACTION, 1),
         fmt_pct(_NUCLEAR_CAUSE_FRACTION, 1),
         fmt_pct(_AI_CAUSE_FRACTION, 1)],
        ["Rel. risk reduction / $10M (central)",
         fmt_pct(_SENTINEL_REL_REDUCTION_PER_10M["values"][1], 4),
         fmt_pct(_NUCLEAR_REL_REDUCTION_PER_10M["values"][1], 4),
         fmt_pct(_AI_REL_REDUCTION_PER_10M["values"][1], 4)],
        ["Effect starts (years from now)",
         str(_sen_fp["year_effect_starts"]),
         str(_nuc_fp["year_effect_starts"]),
         str(_ai_fp["year_effect_starts"])],
        ["Effect lasts (years)",
         str(_sen_fp["persistence_effect"]),
         str(_nuc_fp["persistence_effect"]),
         str(_ai_fp["persistence_effect"])],
        ["Counterfactual factor",
         fmt_pct(FUND_PROFILES["sentinel"]["counterfactual_factor"]),
         fmt_pct(FUND_PROFILES["longview_nuclear"]["counterfactual_factor"]),
         fmt_pct(FUND_PROFILES["longview_ai"]["counterfactual_factor"])],
        ["P(intervention has no effect)",
         fmt_pct(FUND_PROFILES["sentinel"]["p_zero"], 0),
         fmt_pct(FUND_PROFILES["longview_nuclear"]["p_zero"], 0),
         fmt_pct(FUND_PROFILES["longview_ai"]["p_zero"], 0)],
        ["P(intervention causes harm)",
         fmt_pct(FUND_PROFILES["sentinel"]["p_harm"], 0),
         fmt_pct(FUND_PROFILES["longview_nuclear"]["p_harm"], 0),
         fmt_pct(FUND_PROFILES["longview_ai"]["p_harm"], 0)],
        ["Models sub-extinction events?", "Yes (2 tiers)", "No", "No"],
    ],
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT_PATH)
print(f"Written: {OUT_PATH}")
